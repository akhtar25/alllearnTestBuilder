import base64
# from qrReader import *
import csv
import datetime as dt
import hashlib
import hmac
import json
import logging
import os
import random
import re
import string
import urllib
from calendar import monthrange
from datetime import date
from datetime import timedelta
from io import BytesIO
from logging.handlers import RotatingFileHandler
from random import randint
from urllib.parse import quote, urlparse, parse_qs
from urllib.request import urlopen

# import barCode
import boto3
# from pandas import DataFrame
import numpy as np
import pandas as pd
import plotly
import requests as rq
from algoliasearch.search_client import SearchClient
from docx import Document
from docx.shared import Inches
from elasticsearch import Elasticsearch
from flask import Flask, Markup, render_template, request, flash, redirect, url_for, Response, session, g, jsonify
from flask_api import status, exceptions
from flask_login import current_user, login_user, login_required
from flask_migrate import Migrate
from flask_moment import Moment
# import matplotlib.pyplot as plt
from flask_talisman import Talisman, ALLOW_FROM
from google.auth.transport import requests
from google.oauth2 import id_token
from pytz import timezone
from sqlalchemy import text
from tzlocal import get_localzone

from Accounts.accounts import accounts
from Student_TC.student_tc import student_tc
from applicationDB import *
from config import Config
from forms import ContentManager, LeaderBoardQueryForm, EditProfileForm, ResetPasswordRequestForm, ResetPasswordForm, \
    TestBuilderQueryForm, SchoolRegistrationForm, addEventForm, QuestionBuilderQueryForm, SingleStudentRegistration, \
    SchoolTeacherForm, feedbackReportForm, testPerformanceForm, studentPerformanceForm, QuestionBankQueryForm, \
    studentDirectoryForm, promoteStudentForm, PostForm, createSubscriptionForm, ClassRegisterForm, postJobForm, \
    AddLiveClassForm, SearchForm
from miscFunctions import subjects, topics, subjectPerformance, chapters
from send_email import new_teacher_invitation, new_applicant_for_job, application_processed, job_posted_email, \
    send_notification_email, welcome_email, send_password_reset_email, user_access_request_email, \
    user_school_access_request_email, access_granted_email, new_school_reg_email, performance_report_email, \
    test_report_email, notificationEmail, notificationHelplineEmail
from teacher_register.teacher_register import teacher_register
from payment.payment import payment
from job_post.job_post import job_post
from registration.registration import registration
from fee_details.fee_details import fee_details
from syllabus_mod.syllabus_mod import syllabus_mod
from ques_bank.ques_bank import ques_bank
from test_builder.test_builder import test_builder
from student_survey.student_survey import student_survey
from attendance.attendance import attendance
from student_profile.student_profile import student_profile
from inventory.inventory import inventory
from time_table.time_table import time_table
from dashboard.dashboard import dashboard
from school_details.school_details import school_details
from subject.subject import subject
from course.course import course
from new_task.new_task import new_task
from topic_generate.topic_generate import topic_generate
from whatsapp_bot.whatsapp_bot import whatsapp_bot

app = Flask(__name__)
app.register_blueprint(accounts)
app.register_blueprint(student_tc)
app.register_blueprint(teacher_register)
app.register_blueprint(payment)
app.register_blueprint(job_post)
app.register_blueprint(registration)
app.register_blueprint(fee_details)
app.register_blueprint(syllabus_mod)
app.register_blueprint(ques_bank)
app.register_blueprint(test_builder)
app.register_blueprint(student_survey)
app.register_blueprint(attendance)
app.register_blueprint(student_profile)
app.register_blueprint(inventory)
app.register_blueprint(time_table)
app.register_blueprint(dashboard)
app.register_blueprint(school_details)
app.register_blueprint(subject)
app.register_blueprint(course)
app.register_blueprint(new_task)
app.register_blueprint(topic_generate)
app.register_blueprint(whatsapp_bot)

# End
talisman = Talisman(app, content_security_policy=None)


client = SearchClient.create('RVHAVJXK1B', '58e63c83b4126c21f2e994f1d4e89439')
index = client.init_index('prd_course')

app.config.from_object(Config)
db.init_app(app)
migrate = Migrate(app, db)
login_manager.init_app(app)
moment = Moment(app)


app.elasticsearch = Elasticsearch([app.config['ELASTICSEARCH_URL']]) \
        if app.config['ELASTICSEARCH_URL'] else None

if not app.debug and not app.testing:
    if app.config['LOG_TO_STDOUT']:
        stream_handler = logging.StreamHandler()
        stream_handler.setLevel(logging.INFO)
        app.logger.addHandler(stream_handler)        
    else:
        dateVal= datetime.today()
        if not os.path.exists('logs'):
            os.mkdir('logs')
        file_handler = RotatingFileHandler(
            'logs/allLearn'+str(dateVal).replace(' ','').replace(':','').replace('.','')+'.log', maxBytes=10240, backupCount=10)
        file_handler.setFormatter(
            logging.Formatter(
                '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
            ))
        file_handler.setLevel(logging.INFO)
        app.logger.addHandler(file_handler)

    app.logger.setLevel(logging.INFO)
    app.logger.info('allLearn startup')


@app.before_request
def before_request():
    if current_user.is_authenticated:
        current_user.last_seen = datetime.now()
        db.session.commit()
    g.search_form = SearchForm()
    #scheme = request.headers.get('X-Forwarded-Proto')        
    #if scheme and scheme == 'http' and request.url.startswith('http://'):        
    #    url = request.url.replace('http://', 'https://', 1)
    #    code = 301
    #    return redirect(url, code=code)
    
    

#helper methods
def schoolNameVal():
    if current_user.is_authenticated:
        teacher_id = ''
        if current_user.user_type==134 or current_user.user_type==234:
            teacher_id = StudentProfile.query.filter_by(user_id=current_user.id).first()
        else:
            teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()        
        if teacher_id != None:
            print(teacher_id.school_id)
            school_name_row=SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
            if school_name_row!=None:
                name=school_name_row.school_name            
                return name
            else:
                return None            
        else:
            return None
    else:
        return None


notes = {
    0: 'do the shopping',
    1: 'build the codez',
    2: 'paint the door',
}

def note_repr(key):
    return {
        'url': request.host_url.rstrip('/') + url_for('notes_detail', key=key),
        'text': notes[key]
    }

regex = '^[a-z0-9]+[\._]?[a-z0-9]+[@]\w+[.]\w{2,3}$'
def check(email):
    if re.search(regex,email):
        return 'Y'
    else:
        return 'N'

def send_sms(number,message):
    url = 'https://www.fast2sms.com/dev/bulkV2'
    params = {
        'authorization':'TvQr5N7IgSt3JenP6XAjiLMHfohCKBpUqmduVkWsD0G8b24zyxoIOg7ABHhZ9JFQv2CXW5wkiSYqpdDR',
        'sender_id':'FSTSMS',
        'message':message,
        'route':'p',
        'numbers':number
    }

    headers = {
    'cache-control': "no-cache"
    }
    response = rq.request("GET", url, headers=headers, params=params)
    dic = response.json()
    print('dic',dic)


@app.route('/robots.txt')
@app.route('/sitemap.xml')
def static_from_root():
    return send_from_directory(app.static_folder, request.path[1:])



#Route to verify google sign in token
@app.route('/gTokenSignin',methods=["GET","POST"])
def gTokenSignin():
    try:
        idtoken = request.form.get('idtoken')
        # Specify the CLIENT_ID of the app that accesses the backend:
        idinfo = id_token.verify_oauth2_token(idtoken, requests.Request(), app.config['GOOGLE_CLIENT_ID'])
        # ID token is valid. Get the user's Google Account ID from the decoded token.
        userid = idinfo['sub']
        print('This is the email: ')
        print(idinfo["email"])
        print('#############')
                
        #section to create new user
        chkUserData = User.query.filter_by(email=str(idinfo["email"])).first()
        if chkUserData==None:
            user = User(username=idinfo["email"], email=idinfo["email"], user_type='253', access_status='145', 
                first_name = idinfo["given_name"],last_name= idinfo["family_name"], last_modified_date = datetime.today(),
                user_avatar = idinfo["picture"],school_id=1, login_type=244)
            #user.set_password(form.password.data)
            db.session.add(user)
            db.session.commit()
            #flash('Congratulations! You\'re now a registered user!')
            #if a teacher has already been added during school registration then simply add the new user's id to it's teacher profile value        
            checkTeacherProf = TeacherProfile.query.filter_by(email=idinfo["email"]).first()
            #if a student has already been added during school registration then simply add the new user's id to it's student profile value
            checkStudentProf = StudentProfile.query.filter_by(email=idinfo["email"]).first()
    
            if checkTeacherProf!=None:
                checkTeacherProf.user_id = user.id
                db.session.commit()        
            elif checkStudentProf != None:
                checkStudentProf.user_id = user.id
                db.session.commit()
            else:
                pass
        #end of section
          
        #section to create session and auto login

        #endof section

        return '0'
        # // These six fields are included in all Google ID Tokens.
        # "iss": "https://accounts.google.com",
        # "sub": "110169484474386276334",
        # "azp": "1008719970978-hb24n2dstb40o45d4feuo2ukqmcc6381.apps.googleusercontent.com",
        # "aud": "1008719970978-hb24n2dstb40o45d4feuo2ukqmcc6381.apps.googleusercontent.com",
        # "iat": "1433978353",
        # "exp": "1433981953",
        #
        # // These seven fields are only included when the user has granted the "profile" and
        # // "email" OAuth scopes to the application.
        # "email": "testuser@gmail.com",
        # "email_verified": "true",
        # "name" : "Test User",
        # "picture": "https://lh4.googleusercontent.com/-kYgzyAWpZzJ/ABCDEFGHI/AAAJKLMNOP/tIXL9Ir44LE/s99-c/photo.jpg",
        # "given_name": "Test",
        # "family_name": "User",
        # "locale": "en"
        #}
    except ValueError:
        # Invalid token
        return '1'
        



#@register.filter
#def month_name(month_number):
#    return calendar.month_name[month_number]

@app.route("/api", methods=['GET', 'POST'])
def notes_list():
    """
    List or create notes.
    """
    if request.method == 'POST':
        note = str(request.data.get('text', ''))
        idx = max(notes.keys()) + 1
        notes[idx] = note
        return note_repr(idx), status.HTTP_201_CREATED

    # request.method == 'GET'
    return [note_repr(idx) for idx in sorted(notes.keys())]


#@app.route("/api/leaderBoard/<int:schoolID>",methods=['GET','PUT','DELETE'])
#def leaderboardAPI(schoolID):
#    return leaderboardContent(schoolID)

@property
def serialize(self):
 return {
    'student_id': self.student_id,
    'full_name': self.full_name,
    #'class_val': self.class_val,
    #'section': self.section,
    #'sponsored_status': self.sponsored_status,
    #'sponsored_on': self.sponsored_on,
    #'sponsored_amount': self.sponsored_amount,
    #'sponsored_till': self.sponsored_till,
    #'profile_picture': self.profile_picture,
    #'perf_avg': self.perf_avg,
}



@app.route("/api/studentList/<int:schoolID>/", methods=['GET', 'PUT', 'DELETE'])
def studentListAPI(schoolID):
    #studentList = StudentProfile.query.filter_by(school_id=schoolID).all()
    studentListQuery = "select sp.student_id as student_id, sp.full_name as full_name , cs.class_val as class_val , cs.section as section , sponsored_status as spnsored_status, sponsored_on as sponsoted_on, sponsored_amount as sponsored_amount, sponsored_till as sponsored_till, "
    studentListQuery = studentListQuery + "sp.profile_picture as profile_picture,  CAST ( round( avg(pd.student_score),2) as Varchar) as perf_avg "
    studentListQuery = studentListQuery + "from student_profile sp "
    studentListQuery = studentListQuery + "inner join class_section cs on "
    studentListQuery = studentListQuery + "cs.class_sec_id  = sp.class_sec_id and "
    studentListQuery = studentListQuery + "sp.school_id  = "+str(schoolID)
    studentListQuery = studentListQuery + " inner join performance_detail pd on "
    studentListQuery = studentListQuery + "pd.student_id  = sp.student_id "
    studentListQuery = studentListQuery + "group by sp.student_id , sp.full_name , cs.class_val , cs.section , sp.profile_picture , sponsored_status , sponsored_on , sponsored_amount , sponsored_till "
    studentListQuery = studentListQuery + "order by perf_avg desc"

    studentListData = db.session.execute(text(studentListQuery)).fetchall()
    
    # the two lines below can also be used to send data but without the flask api library
    #resp = jsonify({'result': [dict(row) for row in studentListData]})
    #resp.status_code = 200

    return {'result': [dict(row) for row in studentListData]}
    #return [(str(idx.student_id)+','+str(idx.first_name)+','+str(idx.last_name)+',' +str(idx.sponsored_status)) for idx in studentList]





@app.route("/api/<int:key>/", methods=['GET', 'PUT', 'DELETE'])
def notes_detail(key):
    """
    Retrieve, update or delete note instances.
    """
    if request.method == 'PUT':
        note = str(request.data.get('text', ''))
        notes[key] = note
        return note_repr(key)

    elif request.method == 'DELETE':
        notes.pop(key, None)
        return '', status.HTTP_204_NO_CONTENT

    # request.method == 'GET'
    if key not in notes:
        raise exceptions.NotFound()
    return note_repr(key)
##########################End of test section



def leaderboardContent(qclass_val):
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    query = "select  school,class as class_val,section,studentid,student_name,profile_pic,subjectid,test_count,marks from fn_performance_leaderboard_detail_v1("+str(teacher_id.school_id)+")"
    if qclass_val=='dashboard':
        
        where = " where marks is not null order by marks"
        query = query + where
        print('Query inside leaderboardContent:'+str(query))
    else:
        if qclass_val!='' and qclass_val is not None and str(qclass_val)!='None':
            where = " where class='"+str(qclass_val)+"' order by marks desc"
        else:
            where = " where marks is not null order by marks desc"
        query = query + where
        print('Query inside leaderboardContent:'+str(query))    
    leaderbrd_row = db.session.execute(text(query)).fetchall()
    try:
        df = pd.DataFrame(leaderbrd_row,columns=['school','class_val','section','studentid','student_name','profile_pic','subjectid','test_count','marks'])
        
        leaderbrd_pivot = pd.pivot_table(df,index=['studentid','profile_pic','student_name','class_val','section']
                        , columns='subjectid', values='marks'
                        ,aggfunc='sum').reset_index()
        leaderbrd_pivot = leaderbrd_pivot.rename_axis(None).rename_axis(None, axis=1)
        col_list= list(leaderbrd_pivot)
        col_list.remove('studentid')
        col_list.remove('student_name')
        col_list.remove('class_val')
        col_list.remove('section')
        col_list.remove('profile_pic')
        leaderbrd_pivot['total_marks%'] = leaderbrd_pivot[col_list].mean(axis=1)

        # leaderbrd_pivot = leaderbrd_pivot.groupby(['studentid','student_name','class_val','section']).agg()
        df2 = pd.DataFrame(leaderbrd_row,columns=['school','class_val','section','studentid','student_name','profile_pic','subjectid','test_count','marks'])
        leaderbrd_pivot2 = pd.pivot_table(df2,index=['studentid','profile_pic','student_name','class_val','section']
                        , columns='subjectid', values='test_count'
                        ,aggfunc='sum').reset_index()
        leaderbrd_pivot2 = leaderbrd_pivot2.rename_axis(None).rename_axis(None, axis=1)
        col_list2= list(leaderbrd_pivot2)
        col_list2.remove('studentid')
        col_list2.remove('student_name')
        col_list2.remove('class_val')
        col_list2.remove('section')
        col_list2.remove('profile_pic')
        leaderbrd_pivot2['total_tests'] = leaderbrd_pivot2[col_list2].sum(axis=1)
        result = pd.merge(leaderbrd_pivot,leaderbrd_pivot2,on=('studentid','student_name','class_val','section','profile_pic'))
    except:
        result = 1222
    return result  


def stateList():
    with open('stateList.txt', 'r') as f:
        stateListVal = f.readlines()
        stateListVal = str(stateListVal).split(',')
        return stateListVal


def cityList():
    with open('cityList.txt', 'r') as f:
        cityListVal = f.readlines()
        cityListVal = str(cityListVal)
        cityListVal = cityListVal.replace('[','').replace(']','').replace('\'','').replace(',',':null,')
        cityListVal = cityListVal.split(',')
        cityListVal[-1]=cityListVal[-1]+ ':null'
        cityListDict = dict(item.split(':') for item in cityListVal)
        #cityListVal = cityListVal.split(',')
        return cityListDict


def classSecCheck():
    teacherProfile = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #print('#######this is teacher profile val '+ str(teacherProfile.teacher_id))
    #print('#######this is current user '+ str(current_user.id))
    if teacherProfile==None:
        return 'N'
    else:
        classSecRow = ClassSection.query.filter_by(school_id=teacherProfile.school_id).all()
        #print(classSecRow)
        if len(classSecRow)==0:
            print('returning N')
            return 'N'            
        else:
            return 'Y'


@app.route('/practiceTest',methods=["GET","POST"])
def practiceTest():    
    if request.method=="POST":
        board = request.form.get('board')
        class_val = request.form.get('class_val')
        email = current_user.email     
        #print("board:"+ str(board))
        #print("class_val:" +str(class_val))
        schoolRow = SchoolProfile.query.filter_by(school_name="AllLearn "+str(board)).first()
        #print("school id:" +str(schoolRow.school_id))
        classSectionRow = ClassSection.query.filter_by(school_id=schoolRow.school_id, class_val=str(class_val)).first()
        #print("class sec id:"+ str(classSectionRow.class_sec_id))
        checkStudTable= StudentProfile.query.filter_by(email=current_user.email).first()
        if checkStudTable==None or checkStudTable=="":
            studentDataAdd = StudentProfile(first_name=current_user.first_name,last_name=current_user.last_name,full_name=current_user.first_name +" " + current_user.last_name,
                school_id=schoolRow.school_id,class_sec_id=classSectionRow.class_sec_id,
                phone=current_user.phone,school_adm_number="prac_"+ str(current_user.id), user_id=current_user.id,
                roll_number=000,last_modified_date=datetime.today(), email=current_user.email,points=10)
            db.session.add(studentDataAdd)
            
            #updating the public.user table
            current_user.user_type=234
            current_user.access_status=145
            current_user.school_id=schoolRow.school_id
            current_user.last_modified_date=datetime.today()
            db.session.commit() 
        else:
            studentDataAdd = StudentProfile.query.filter_by(id=current_user.id).first()

        #This section is to update an anonymously taken test with new student id    
        #try:
        #if session.get('anonUser'):
        #    splitVals = str(session['anonUser']).split('_')
        #    respSessionID = splitVals[1]
        #    session['anonUser'] = False
        #    print("Resp Session ID of older test: "+str(respSessionID))
        #    ###Section to update the test results of anon user with the new student id in resp capture
        #    respUPDQuery = "Update response_capture set student_id=" + str(studentDataAdd.student_id)
        #    respUPDQuery = respUPDQuery + " , class_sec_id=" + str(studentDataAdd.class_sec_id) 
        #    respUPDQuery = respUPDQuery + " where resp_session_id=\'"+ str(respSessionID) + "\'"
        #    print(str(respUPDQuery))
        #    db.session.execute(text(respUPDQuery))
                
        #except:
        #    print('error occurred. response cap not updated.')
        #    pass
        db.session.commit()        
        print('New entry made into the student table')

    if ("school.alllearn" in str(request.url)):
            print('#######this is the request url: '+ str(request.url))
            return redirect(url_for('index'))
    if current_user.is_anonymous:            
        studentProfile = StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()  #staging anonymous username f        
    else:
        studentProfile = StudentProfile.query.filter_by(user_id=current_user.id).first()
    if studentProfile==None:
        studentData=""      
        testHistory=""  
        perfRows=""
        testCount = 0
        leaderboardData = ""
        class_val=""
        questionsAnswered=0
        print('#####Student data is null')
    else:        
        studentDataQuery = "with temptable as "
        studentDataQuery = studentDataQuery + " (with total_marks_cte as ( "
        studentDataQuery = studentDataQuery + " select sum(suggested_weightage) as total_weightage, count(*) as num_of_questions  from question_details where question_id in "
        studentDataQuery = studentDataQuery + " (select question_id from response_capture rc2 where student_id ="+ str(studentProfile.student_id)+") ) "
        studentDataQuery = studentDataQuery + " select distinct sp.roll_number, sp.full_name, sp.student_id, "
        studentDataQuery = studentDataQuery + " SUM(CASE WHEN rc.is_correct='Y' THEN qd.suggested_weightage ELSE 0 end) AS  points_scored , "
        studentDataQuery = studentDataQuery + " total_marks_cte.total_weightage "
        studentDataQuery = studentDataQuery + " from response_capture rc inner join student_profile sp on "
        studentDataQuery = studentDataQuery + " rc.student_id=sp.student_id "
        studentDataQuery = studentDataQuery + " inner join question_details qd on "
        studentDataQuery = studentDataQuery + " qd.question_id=rc.question_id, total_marks_cte         "
        studentDataQuery = studentDataQuery + " group by sp.roll_number, sp.full_name, sp.student_id, total_marks_cte.total_weightage ) "
        studentDataQuery = studentDataQuery + " ,temp2 as (select count(distinct resp_session_id) as tests_taken "
        studentDataQuery = studentDataQuery + " , student_id from response_capture group by student_id) "
        studentDataQuery = studentDataQuery + " select *from temptable inner join temp2 on temptable.student_id =temp2.student_id and temp2.student_id =" + str(studentProfile.student_id)
        studentData = db.session.execute(text(studentDataQuery)).first()

        
        #getting class_val for the student
        classDataQuery = "select *from class_section cs where class_sec_id="+ str(studentProfile.class_sec_id)
        classData = db.session.execute(classDataQuery).first()
        class_val = classData.class_val

        #performanceQuery = "select *from fn_leaderboard_responsecapture () where student_id='"+str(studentProfile.student_id)+ "'"
        performanceQuery = "select *from fn_leaderboard_responsecapture() where student_id ='"+str(studentProfile.student_id)+"' order by class_val desc" # and class_val='"+str(class_val)+"'"
        perfRows = db.session.execute(text(performanceQuery)).fetchall()

        testCountQuery = "select sum(test_taken) as tests_taken from fn_leaderboard_ResponseCapture() where student_id='"+str(studentProfile.student_id)+ "'"
        testCount = db.session.execute(text(testCountQuery)).first()

        #section for test history query
        testHistoryQuery = "SELECT *FROM fn_student_performance_response_capture("+str(studentProfile.student_id)+")"
        testHistory = db.session.execute(testHistoryQuery).fetchall()
        ##end of test history query

        #other recent test takers
        #recentTestTakersQuery = "select *from fn_student_test_taken_details() order by test_taken_on desc limit 10"
        #recentTestTakers = db.session.execute(recentTestTakersQuery).fetchall()        
        #Questions answered
        questionsAnsweredQuery = "select count(*) as qanswered from response_capture where student_id="+str(studentProfile.student_id)
        questionsAnswered = db.session.execute(questionsAnsweredQuery).first()
        #leaderboard data
        leaderboardQuery = "SELECT student_id,first_name,SUBJECT,student_score,class_val"
        leaderboardQuery = leaderboardQuery + " FROM fn_leaderboard_ResponseCapture() M "
        leaderboardQuery = leaderboardQuery + " WHERE M.student_score = "
        leaderboardQuery = leaderboardQuery + " (SELECT MAX(MM.student_score) FROM fn_leaderboard_ResponseCapture() MM where class_val='"+str(class_val)+"' GROUP BY "
        leaderboardQuery = leaderboardQuery + " MM.SUBJECT having  MM.SUBJECT =M.subject)"
        leaderboardQuery = leaderboardQuery + " order by student_score desc"
        leaderboardData = db.session.execute(leaderboardQuery).fetchall()
    if studentData!=None or studentData!="":
        try:
            avg_performance = round(((studentData.points_scored/studentData.total_weightage) *100),2)
        except:
            avg_performance=0
    else:
        avg_performance = 0
    
    meta_val = "Preparing for exams need not be difficult. Take free mock tests anytime you want with allLearn. CBSE | IIT JEE | NEET | Competitive Exams"
    return render_template('/practiceTest.html',studentData=studentData, disconn=1,
        studentProfile=studentProfile, avg_performance=avg_performance,testHistory=testHistory,perfRows=perfRows,testCount=testCount,
        leaderboardData=leaderboardData,class_val=class_val,questionsAnswered=questionsAnswered,title='Take unlimited free practice tests anytime',
        meta_val=meta_val)


@app.route('/normal')
def normal():
    return 'Normal'

@app.route('/embeddable')
@talisman(frame_options=ALLOW_FROM, frame_options_allow_from='*')
def embeddable():
    return 'Embeddable'

@app.route("/loaderio-ad2552628971ece0389988c13933a170/")
def performanceTestLoaderFunction():
    return render_template("loaderio-ad2552628971ece0389988c13933a170.html")

# @app.route("/account/")
# @login_required
# def account():
#     return render_template('account.html')

@app.route('/sign-s3')
def sign_s3():
    S3_BUCKET = os.environ.get('S3_BUCKET_NAME')
    #S3_BUCKET = "alllearndatabucketv2"
    file_name = request.args.get('file-name')
    print(file_name)    
    file_type = request.args.get('file-type')
    print(file_type)
    #if file_type=='image/png' or file_type=='image/jpeg':
    #   file_type_folder='images'
    #s3 = boto3.client('s3')
    s3 = boto3.client('s3', region_name='ap-south-1')
    folder_name=request.args.get('folder')
    print('FolderName:'+folder_name)
    # folder_url=signs3Folder(folder_name,file_type)
    folder_url = folder_name
    print('folder_url:'+str(folder_url))
    print(s3)

    presigned_post = s3.generate_presigned_post(
      Bucket = S3_BUCKET,
      Key = str(folder_url)+"/"+str(file_name),
      Fields = {"acl": "public-read", "Content-Type": file_type},
      Conditions = [
        {"acl": "public-read"},
        {"Content-Type": file_type}
      ],
      ExpiresIn = 3600
    )
   
    
    return json.dumps({
      'data': presigned_post,
      'url': 'https://%s.s3.amazonaws.com/%s/%s' % (S3_BUCKET,folder_url,file_name)
    })

@app.route('/s3api',methods=['GET','POST'])
def s3api():
    if request.method == 'POST':
        print('inside s3 api')
        S3_BUCKET = os.environ.get('S3_BUCKET_NAME')
        jsonData = request.json
        # jsonData = {'contact': {'fields': {'age_group': {'inserted_at': '2021-01-25T06:36:45.002400Z', 'label': 'Age Group', 'type': 'string', 'value': '19 or above'}, 'name': {'inserted_at': '2021-01-25T06:35:49.876654Z', 'label': 'Name', 'type': 'string', 'value': 'hi'}}, 'name': 'Zaheen', 'phone': '918802362259'}, 'results': {}, 'custom_key': 'custom_value'}
        print('jsonData:')
        print(jsonData)
        
        userData = json.dumps(jsonData)
        user = json.loads(userData)
        conList = []
        paramList = []
        for con in user['contact'].values():
            conList.append(con)
        for data in user['results'].values():
            paramList.append(data)
        print(conList)
        print(paramList)
        contactNo = conList[2][-10:]
        print(contactNo)
        S3_BUCKET = "alllearndatabucketv2"
        file_name = paramList[0]
        print('fileName:'+str(file_name))
        s3 = boto3.client('s3', region_name='ap-south-1')
        presigned_post = s3.generate_presigned_post(
        Bucket = S3_BUCKET,
        Key = str('coronaDoc')+"/"+str(file_name),
        Fields = {"acl": "public-read"},
        Conditions = [
        {"acl": "public-read"}
        ],
        ExpiresIn = 3600
        )
   
        folder_url = 'corona'
        return json.dumps({
            'data': presigned_post,
            'url': 'https://%s.s3.amazonaws.com/%s/%s' % (S3_BUCKET,folder_url,file_name)
        })


# @app.route("/submit_form/", methods = ["POST"])
# @login_required
# def submit_form():
#     #teacherProfile = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #teacherProfile.teacher_name = request.form["full-name"]
#     #teacherProfile.profile_picture = request.form["avatar-url"]
#     #db.session.commit()
#     #flash('DB values updated')
#     return redirect(url_for('account'))


# @app.route('/reset_password_request', methods=['GET', 'POST'])
# def reset_password_request():
#     if current_user.is_authenticated:
#         return redirect(url_for('index'))
#     form = ResetPasswordRequestForm()
#     if form.validate_on_submit():
#         user = User.query.filter_by(email=form.email.data).first()
#         if user:
#             send_password_reset_email(user)
#         flash('Check your email for the instructions to reset your password')
#         return redirect(url_for('login'))
#     return render_template('reset_password_request.html', title='Reset Password', form=form)


# @app.route('/reset_password/<token>', methods=['GET', 'POST'])
# def reset_password(token):
#     if current_user.is_authenticated:
#         return redirect(url_for('index'))
#     user = User.verify_reset_password_token(token)
#     if not user:
#         return redirect(url_for('index'))
#     form = ResetPasswordForm()
#     if form.validate_on_submit():
#         user.set_password(form.password.data)
#         db.session.commit()
#         flash('Your password has been reset.')
#         return redirect(url_for('login'))
#     return render_template('reset_password_page.html', form=form)

# @app.route('/inReviewSchool')
# def inReviewSchool():
#     print('In review school:'+str(current_user.user_type))
#     return render_template('inReviewSchool.html', disconn = 1)

# @app.route('/schoolProfile')
# @login_required
# def schoolProfile():
#     print('User Type Id:'+str(current_user.user_type))
#     studentDetails = StudentProfile.query.filter_by(user_id = current_user.id).first()
#     if current_user.user_type==134:
#         teacherRow=StudentProfile.query.filter_by(user_id=current_user.id).first()
#     else:
#         teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     registeredStudentCount = db.session.execute(text("select count(*) from student_profile where school_id ='"+str(teacherRow.school_id)+"'")).first()
#     registeredTeacherCount = db.session.execute(text("select count(*) from teacher_profile where school_id ='"+str(teacherRow.school_id)+"'")).first()
#     allTeachers = TeacherProfile.query.filter_by(school_id=teacherRow.school_id).all()
#     classSectionRows = ClassSection.query.filter_by(school_id=teacherRow.school_id).all()
#     schoolProfileRow = SchoolProfile.query.filter_by(school_id = teacherRow.school_id).first()
#     addressRow = Address.query.filter_by(address_id = schoolProfileRow.address_id).first()
#     subscriptionRow = SubscriptionDetail.query.filter_by(sub_id = schoolProfileRow.sub_id).first()
#     value=0
#     #if current_user.user_type==134:
#     #    value=1
#     indic='DashBoard'
#     return render_template('schoolProfile.html',indic=indic,title='School Profile', teacherRow=teacherRow, registeredStudentCount=registeredStudentCount, registeredTeacherCount=registeredTeacherCount,allTeachers=allTeachers,classSectionRows=classSectionRows, schoolProfileRow=schoolProfileRow,addressRow=addressRow,subscriptionRow=subscriptionRow,disconn=value,user_type_val=str(current_user.user_type),studentDetails=studentDetails)




# @app.route('/schoolRegistration', methods=['GET','POST'])
# @login_required
# def schoolRegistration():
#     #queries for subcription 
#     fromImpact = request.args.get('fromImpact')
#     subscriptionRow = SubscriptionDetail.query.filter_by(archive_status='N').order_by(SubscriptionDetail.sub_duration_months).all()    
#     distinctSubsQuery = db.session.execute(text("select distinct group_name, sub_desc, student_limit, teacher_limit, test_limit from subscription_detail where archive_status='N' order by student_limit ")).fetchall()

#     S3_BUCKET = os.environ.get('S3_BUCKET_NAME')
#     form = SchoolRegistrationForm()
#     form.board.choices=[(str(i.description), str(i.description)) for i in MessageDetails.query.with_entities(MessageDetails.description).distinct().filter_by(category='Board').all()]
#     if form.validate_on_submit():
#         user = User.query.filter_by(id=current_user.id).first()
#         user.user_type = 71
#         user.access_status = 145
#         db.session.commit()
#         selected_sub_id = request.form.get('selected_sub_id')        
#         address_id=Address.query.filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#         if address_id is None:
#             address_data=Address(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data,country=form.country.data)
#             db.session.add(address_data)
#             address_id=db.session.query(Address).filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#         board_id=MessageDetails.query.filter_by(description=form.board.data).first()
#         school_picture=request.files['school_image']
#         school_picture_name=request.form['file-input'] 

#         school=SchoolProfile(school_name=form.schoolName.data,board_id=board_id.msg_id,address_id=address_id.address_id,registered_date=dt.datetime.now(), last_modified_date = dt.datetime.now(), sub_id=selected_sub_id,how_to_reach=form.how_to_reach.data,is_verified='N')
#         db.session.add(school)
#         school_id=db.session.query(SchoolProfile).filter_by(school_name=form.schoolName.data,address_id=address_id.address_id).first()
#         if school_picture_name!='':
#             school = SchoolProfile.query.get(school_id.school_id)
#             school.school_picture = 'https://'+ S3_BUCKET + '.s3.amazonaws.com/school_data/school_id_' + str(school_id.school_id) + '/school_profile/' + school_picture_name
#             client = boto3.client('s3', region_name='ap-south-1')
#             client.upload_fileobj(school_picture , os.environ.get('S3_BUCKET_NAME'), 'school_data/school_id_'+ str(school_id.school_id) + '/school_profile/' + school_picture_name,ExtraArgs={'ACL':'public-read'})
       
#         teacher=TeacherProfile(school_id=school.school_id,email=current_user.email,user_id=current_user.id, designation=147, registration_date=dt.datetime.now(), last_modified_date=dt.datetime.now(), phone=current_user.phone, device_preference=78 )
#         db.session.add(teacher)
#         db.session.commit()
#         newTeacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         newSchool = SchoolProfile.query.filter_by(school_id=school_id.school_id).first() 
#         session['school_logo'] = newSchool.school_logo 
#         session['schoolPicture'] = newSchool.school_picture   
#         # Session variable code start
#         session['schoolName'] = schoolNameVal()
#         session['userType'] = current_user.user_type
#         session['username'] = current_user.username
        
#         #print('user name')
#         #print(session['username'])
#         school_id = ''
#         #print('user type')
#         #print(session['userType'])
#         session['studentId'] = ''
#         # if session['userType']==71:
#         #     school_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         # elif session['userType']==134:
#         #     school_id = StudentProfile.query.filter_by(user_id=current_user.id).first()
#         #     session['studentId'] = school_id.student_id
#         # else:
#         #     school_id = User.query.filter_by(id=current_user.id).first()
#         # school_pro = SchoolProfile.query.filter_by(school_id=school_id.school_id).first()
#         # session['school_logo'] = ''
#         # if school_pro:
#         #     session['school_logo'] = school_pro.school_logo
#         #     session['schoolPicture'] = school_pro.school_picture
#         query = "select user_type,md.module_name,description, module_url, module_type from module_detail md inner join module_access ma on md.module_id = ma.module_id where user_type = '"+str(current_user.user_type)+"' and ma.is_archived = 'N' and md.is_archived = 'N' order by module_type"
#         moduleDetRow = db.session.execute(query).fetchall()
#         #print('School profile')
#         # print(session['schoolPicture'])
#         # det_list = [1,2,3,4,5]
#         session['moduleDet'] = []
#         detList = session['moduleDet']
        
#         for det in moduleDetRow:
#             eachList = []
#             print(det.module_name)
#             print(det.module_url)
#             eachList.append(det.module_name)
#             eachList.append(det.module_url)
#             eachList.append(det.module_type)
#             # detList.append(str(det.module_name)+":"+str(det.module_url)+":"+str(det.module_type))
#             detList.append(eachList)
#         session['moduleDet'] = detList
#         # End code   
#         newSchool.school_admin = newTeacherRow.teacher_id

#         db.session.commit()
#         flash('School Registered Successfully!')
#         new_school_reg_email(form.schoolName.data)
#         fromSchoolRegistration = True
       
#         subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#         teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         board = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#         boardRows = MessageDetails.query.filter_by(msg_id=board.board_id).first()
#         school_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#         classValues = "SELECT class_val,sum(class_sec_id) as s FROM class_section cs where school_id = '"+str(teacher_id.school_id)+"' group by class_val order by s"
#         classValues = db.session.execute(text(classValues)).fetchall()
#         classValuesGeneral = "SELECT class_val,sum(class_sec_id) as s FROM class_section cs group by class_val order by s"
#         classValuesGeneral = db.session.execute(text(classValuesGeneral)).fetchall()
#         subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#         bookName = BookDetails.query.all()
#         chapterNum = Topic.query.distinct().all()
#         topicId = Topic.query.all()
#         generalBoardId = SchoolProfile.query.with_entities(SchoolProfile.board_id).filter_by(school_id=teacher_id.school_id).first()
#         generalBoard = MessageDetails.query.filter_by(msg_id=generalBoardId.board_id).first()
#         fromSchoolRegistration = True
#         schoolData = SchoolProfile.query.filter_by(school_admin=newTeacherRow.teacher_id).first()
#         print('current user id:'+str(current_user.id))
#         print('schoolData:'+str(schoolData))
#         print('schoolData.is_verified'+str(schoolData.is_verified))
#         if schoolData:
#             print('if schoolData exist')
#             if schoolData.is_verified == 'N':
#                 print('if schoolData.is_verified is N')
#                 userTableDetails = User.query.filter_by(id=current_user.id).first()
#                 adminEmail=db.session.execute(text("select t2.email,t2.teacher_name,t1.school_name,t3.username from school_profile t1 inner join teacher_profile t2 on t1.school_admin=t2.teacher_id inner join public.user t3 on t2.email=t3.email where t1.school_id='"+str(schoolData.school_id)+"'")).first()
#                 user_school_access_request_email(adminEmail.email,adminEmail.teacher_name, adminEmail.school_name, userTableDetails.first_name+ ''+userTableDetails.last_name, adminEmail.username, userTableDetails.user_type)
#                 return redirect(url_for('inReviewSchool'))
        
#         return render_template('syllabus.html',generalBoard=generalBoard,boardRowsId = boardRows.msg_id , boardRows=boardRows.description,subjectValues=subjectValues,school_name=school_id.school_name,classValues=classValues,classValuesGeneral=classValuesGeneral,bookName=bookName,chapterNum=chapterNum,topicId=topicId,fromSchoolRegistration=fromSchoolRegistration,user_type_val=str(current_user.user_type))
#     return render_template('schoolRegistration.html',fromImpact=fromImpact,disconn = 1,form=form, subscriptionRow=subscriptionRow, distinctSubsQuery=distinctSubsQuery)

@app.route('/admin')
@login_required
def admin():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    query = "select count(*) from public.user where user_type='161'"
    query2 = "SELECT count(*) FROM public.user WHERE last_seen >=current_date - 30;"
    count = db.session.execute(text(query)).fetchall()
    user_type_val = current_user.user_type
    min_user_count = db.session.execute(text(query2)).first()
    schoolDetails = SchoolProfile.query.all()
    teacherDetails = TeacherProfile.query.all()
    schoolCount = "select count(*) from school_profile"
    school_count = db.session.execute(text(schoolCount)).first()
    print(school_count[0])
    teacherCount = "select count(*) from teacher_profile"
    teacher_count = db.session.execute(text(teacherCount)).first()
    studentCount = "select count(*) from student_profile"
    student_count = db.session.execute(text(studentCount)).first()
    userCount = "select count(*) from public.user"
    user_count = db.session.execute(text(userCount)).first()

    userTypeCount = "select user_type,count(*) as user_count,description from public.user inner join message_detail on msg_id=user_type group by user_type,description"
    user_type_count = db.session.execute(text(userTypeCount)).fetchall()
    schoolreg1 = "SELECT count(*) FROM school_profile WHERE registered_date <=current_date - 30;"
    regSchool1 = db.session.execute(text(schoolreg1)).first()
    schoolreg2 = "SELECT count(*) FROM school_profile WHERE registered_date >=current_date - 30;"
    regSchool2 = db.session.execute(text(schoolreg2)).first()

    teacherreg1 = "SELECT count(*) FROM teacher_profile WHERE registration_date <=current_date - 30;"
    regTeacher1 = db.session.execute(text(teacherreg1)).first()
    teacherreg2 = "SELECT count(*) FROM teacher_profile WHERE registration_date >=current_date - 30;"
    regTeacher2 = db.session.execute(text(teacherreg2)).first()

    studentreg1 = "SELECT count(*) FROM student_profile WHERE registration_date <=current_date - 30;"
    regStudent1 = db.session.execute(text(studentreg1)).first()
    studentreg2 = "SELECT count(*) FROM student_profile WHERE registration_date >=current_date - 30;"
    regStudent2 = db.session.execute(text(studentreg2)).first()
    print(regTeacher2[0])
    print(regTeacher1[0])
    try:
        perSchool = float((int(regSchool2[0])*100)/int(regSchool1[0]))
        perSchool = round(perSchool,2)
    except:
        perSchool = 0
    try: 
        perTeacher = float((int(regTeacher2[0])*100)/int(regTeacher1[0]))
        perTeacher = round(perTeacher,2)
    except:
        perTeacher = 0
    try:
        perStudent = float((int(regStudent2[0])*100)/int(regStudent1[0]))
        perStudent = round(perStudent,2)
    except:
        perStudent = 0
    # perSchool=''
    num = ''
    num2 = ''
    for c in count:
        num = c.count
    # for c2 in count2:
    #     num2 = c2.count
    print('Count'+str(num))
    print('Count2:'+str(num2))
    return render_template('admin.html',count=num,user_type_val=str(current_user.user_type),schoolDetails=schoolDetails,school_count=school_count,teacher_count=teacher_count,student_count=student_count,teacherDetails=teacherDetails,user_count=user_count,min_user_count=min_user_count,user_type_count=user_type_count,perSchool=perSchool,perTeacher=perTeacher,perStudent=perStudent)


@app.route('/promoteStudent',methods=['POST','GET'])
@login_required
def promoteStudent():
    form = promoteStudentForm()
    teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    available_class=ClassSection.query.with_entities(ClassSection.class_val,ClassSection.section).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()
    class_list=[(str(i.class_val)+"-"+str(i.section),str(i.class_val)+"-"+str(i.section)) for i in available_class]
    form.class_section1.choices = class_list
    form.class_section2.choices = class_list
    studentListQuery = "select sp.student_id,sp.school_id ,full_name , section ,class_val ,section from student_profile sp inner join class_section cs on sp.class_sec_id =cs.class_sec_id where sp.school_id='"+str(teacher_id.school_id)+"'"
    studentList = db.session.execute(studentListQuery).fetchall()
    if request.method == 'POST':
        #print('Inside post request')    
        flash('Student promoted successfully')
        classSecBefore = form.class_section1.data
        classSecAfter = form.class_section2.data
        #print('Class Section after')
        #print(classSecAfter)
        resClassSection = classSecAfter.split("-")
        classAfter = resClassSection[0]
        #print(classAfter)
        sectionAfter = resClassSection[1]
        #print(sectionAfter)
        value = request.form.getlist('checkbox')
        for i in range(len(value)):
            #print(value[i])
            studentPromote = StudentProfile.query.filter_by(student_id=str(value[i])).first()
            #print(studentPromote.full_name)
            classSection = ClassSection.query.with_entities(ClassSection.class_sec_id).filter_by(class_val=str(classAfter),section=str(sectionAfter),school_id=str(teacher_id.school_id)).first()
            studentPromote.class_sec_id = classSection.class_sec_id
            db.session.commit()
            studentClassSec = StudentClassSecDetail.query.filter_by(student_id=str(value[i])).first()
            if studentClassSec:
                studentClassSec.is_current = 'N'
                db.session.commit()
            classSecAdd = StudentClassSecDetail(student_id=str(value[i]),
            class_sec_id=classSection.class_sec_id,class_val=str(classAfter),
            section=str(sectionAfter),is_current='Y',last_modified_date=datetime.now(),promotion_date=datetime.now())
            db.session.add(classSecAdd)
            db.session.commit()
        indic='DashBoard'
        return render_template('promoteStudent.html',indic=indic,title='Promote Student',form=form,studentList=studentList,user_type_val=str(current_user.user_type))
    else:
        indic='DashBoard'
        return render_template('promoteStudent.html',indic=indic,title='Promote Student',form=form,studentList=studentList,user_type_val=str(current_user.user_type))
      
# @app.route('/classRegistration', methods=['GET','POST'])
# @login_required
# def classRegistration():
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     classSectionRows = ClassSection.query.filter_by(school_id=teacherRow.school_id).all()

#     form = ClassRegisterForm()
#     #if form.validate_on_submit():
#     if request.method == 'POST':
#         #print('passed validation')
#         class_val=request.form.getlist('class_val')
#         class_section=request.form.getlist('section')
#         student_count=request.form.getlist('student_count')

#         for i in range(len(class_val)):
#             #print('there is a range')
#             class_data=ClassSection(class_val=int(class_val[i]),section=str(class_section[i]).upper(),student_count=int(student_count[i]),school_id=teacherRow.school_id)
#             db.session.add(class_data)
        
#         db.session.commit()
#         #adding records to topic tracker while registering school
        
#         classSecRows = ClassSection.query.filter_by(school_id=teacherRow.school_id).all()

#         topicTrackerRows = "select distinct class_sec_id from topic_tracker where school_id='"+str(teacherRow.school_id)+"'"

#         classSecNotInTopicTracker = db.session.execute(text(topicTrackerRows)).fetchall()

#         for i in range(len(class_val)):
#             class_id = ClassSection.query.with_entities(ClassSection.class_sec_id).filter_by(school_id=teacherRow.school_id,class_val=class_val[i]).first()
#             if class_id.class_sec_id not in classSecNotInTopicTracker: 
#                 insertRow = "insert into topic_tracker (subject_id, class_sec_id, is_covered, topic_id, school_id, reteach_count, last_modified_date) (select subject_id, '"+str(class_id.class_sec_id)+"', 'N', topic_id, '"+str(teacherRow.school_id)+"', 0,current_date from Topic_detail where class_val="+str(class_val[i])+")"
#                 db.session.execute(text(insertRow))
#         db.session.commit()

#         flash('Classes added successfully!')
#     return render_template('classRegistration.html', classSectionRows=classSectionRows,form=form)    
    



@app.route('/teacherDirectory',methods=['GET','POST'])
@login_required
def teacherDirectory():
    school_name_val = schoolNameVal()
    
    if school_name_val ==None:
        print('did we reach here')
        return redirect(url_for('disconnectedAccount'))
    else:
        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
        #section for payroll report
        payrollReportQuery = "select to_date(concat('01-',month,'-',year),'DD-MM-YYYY') as period, sum(calc_salary) as salary_spend, count(*) teacher_count, (avg(days_present) / avg(days_in_month))*100 as avg_productivity from teacher_payroll_detail where school_id= "+ str(teacher_id.school_id)
        payrollReportQuery= payrollReportQuery+" group  by month, year "
        payrollReportData = db.session.execute(payrollReportQuery).fetchall()
        #end of payroll report section
        allTeachers = TeacherProfile.query.filter_by(school_id = teacher_id.school_id).all()
        available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher_id.school_id).all()
        class_list=[('select','Select')]
        section_list=[]
        for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all():
            class_list.append((str(i.class_val), "Class "+str(i.class_val)))
        for i in available_section:
            section_list.append((i.section,i.section))
        form=SchoolTeacherForm()
        form.teacher_subject.choices = [(str(i.msg_id), str(i.description)) for i in MessageDetails.query.with_entities(MessageDetails.msg_id,MessageDetails.description).distinct().filter_by(category='Subject').all()]
        form.class_teacher.choices = class_list
        form.class_teacher_section.choices = section_list
        if request.method=='POST':
            teacher_name=request.form.getlist('teacher_name')
            teacher_subject=request.form.getlist('teacher_subject')
            teacher_class=request.form.getlist('class_teacher')
            teacher_class_section=request.form.getlist('class_teacher_section')
            teacher_email=request.form.getlist('teacher_email')
            print(teacher_name)
            for i in range(len(teacher_name)):
                if teacher_class[i]!='select':
                    class_sec_id=ClassSection.query.filter_by(class_val=str(teacher_class[i]),section=teacher_class_section[i]).first()
                    teacher_data=TeacherProfile(teacher_name=teacher_name[i],school_id=teacher_id.school_id,class_sec_id=class_sec_id.class_sec_id,email=teacher_email[i],subject_id=int(teacher_subject[i]),last_modified_date= datetime.now(), registration_date = datetime.now())
                    db.session.add(teacher_data)
                else:
                    teacher_data=TeacherProfile(teacher_name=teacher_name[i],school_id=teacher_id.school_id,email=teacher_email[i],subject_id=int(teacher_subject[i]),last_modified_date= datetime.now(), registration_date = datetime.now(), device_preference=195)
                    db.session.add(teacher_data)
                    #send email to the teachers here
                new_teacher_invitation(teacher_email[i],teacher_name[i],school_name_val, str(teacher_id.teacher_name))
            db.session.commit()
            flash('Successful registration !')   
        indic='DashBoard'         
        return render_template('teacherDirectory.html',indic=indic,form=form, payrollReportData=payrollReportData,allTeachers=allTeachers,user_type_val=str(current_user.user_type))

# New Section added to manage feeDetail
# @app.route('/feeMonthData')
# def feeMonthData():
#     qmonth = request.args.get('month')
#     qyear = request.args.get('year')
#     class_val = request.args.get('class_val')
#     section = request.args.get('section')
#     print('inside Summary Box route')
#     print(class_val)
#     print(section)
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ''
#     if class_val!=None:
#         class_sec_id = ClassSection.query.filter_by(class_val=class_val,section=section,school_id=teacherDataRow.school_id).first()
#     print(qmonth+ ' '+qyear)
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #days in month
#     daysInMonth = monthrange(int(qyear),int(qmonth))
#     daysInMonth = int(daysInMonth[1])
#     feeDetail = ''
#     if class_val=='None' or class_val=='':
#         print('if class is None')
#         paid_fees = "select sum(fee_paid_amount) as collected_fee from fee_detail where school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         paid_fees = db.session.execute(text(paid_fees)).first()
#         paid_student_count = "select count(*) as no_of_paid_students from fee_detail where fee_amount=fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         paid_student_count = db.session.execute(text(paid_student_count)).first()
#         classSec_ids = FeeClassSecDetail.query.filter_by(school_id=teacherDataRow.school_id).all()
#         unpaid_students_count = 0
#         for class_sec_id in classSec_ids:
#             unpaid_students = "select count(*) as no_of_unpaid_students from student_profile sp where sp.student_id not in (select student_id from fee_detail where school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"') and sp.school_id='"+str(teacherDataRow.school_id)+"' and sp.class_sec_id='"+str(class_sec_id.class_sec_id)+"'"
#             unpaid_students = db.session.execute(text(unpaid_students)).first()
#             unpaid_students_count = unpaid_students_count + unpaid_students.no_of_unpaid_students
#         partially_paid_students = "select count(*) as partially_paid_students from fee_Detail where fee_amount>fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         partially_paid_students = db.session.execute(text(partially_paid_students)).first()
#         class_sec_ids = ClassSection.query.filter_by(school_id=teacherDataRow.school_id).all()
#         total_unpaid_students = 0
#         if partially_paid_students:
#             total_unpaid_students = int(unpaid_students_count) + int(partially_paid_students.partially_paid_students)
        
#         total_unpaid_fee = 0
#         for class_sec_id in class_sec_ids:
#             unpaid_students = "select count(*) as no_of_unpaid_students from student_profile sp where sp.student_id not in (select student_id from fee_detail where school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"') and sp.school_id='"+str(teacherDataRow.school_id)+"' and sp.class_sec_id='"+str(class_sec_id.class_sec_id)+"'"
#             unpaid_students = db.session.execute(text(unpaid_students)).first()
#             fee_amount = FeeClassSecDetail.query.filter_by(class_sec_id=class_sec_id.class_sec_id,school_id=teacherDataRow.school_id).first()
#             unpaid_students_fee = 0
#             if unpaid_students and fee_amount:
#                 unpaid_students_fee = int(unpaid_students.no_of_unpaid_students) * int(fee_amount.amount)
#             partially_paid_fee = "select sum(outstanding_amount) as pending_amount from fee_detail where fee_amount>fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#             partially_paid_fee = db.session.execute(text(partially_paid_fee)).first()
#             if partially_paid_fee:
#                 print('partially paid fee:'+str(partially_paid_fee.pending_amount))
#             if partially_paid_fee.pending_amount:
#                 total_unpaid_fee = total_unpaid_fee + unpaid_students_fee + partially_paid_fee.pending_amount
#             else:
#                 total_unpaid_fee = total_unpaid_fee + unpaid_students_fee
#     else:
#         print('if class is not None')
#         paid_fees = "select sum(fee_paid_amount) as collected_fee from fee_detail where school_id='"+str(teacherDataRow.school_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         paid_fees = db.session.execute(text(paid_fees)).first()
#         paid_student_count = "select count(*) as no_of_paid_students from fee_detail where fee_amount=fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         paid_student_count = db.session.execute(text(paid_student_count)).first()
#         unpaid_students = "select count(*) as no_of_unpaid_students from student_profile sp where sp.student_id not in (select student_id from fee_detail where school_id='"+str(teacherDataRow.school_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"') and sp.school_id='"+str(teacherDataRow.school_id)+"' and sp.class_sec_id='"+str(class_sec_id.class_sec_id)+"'"
#         unpaid_students = db.session.execute(text(unpaid_students)).first()
#         partially_paid_students = "select count(*) as partially_paid_students from fee_Detail where fee_amount>fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         partially_paid_students = db.session.execute(text(partially_paid_students)).first()
#         total_unpaid_students = 0
#         if unpaid_students and partially_paid_students:
#             total_unpaid_students = int(unpaid_students.no_of_unpaid_students) + int(partially_paid_students.partially_paid_students)
#         fee_amount = FeeClassSecDetail.query.filter_by(class_sec_id=class_sec_id.class_sec_id,school_id=teacherDataRow.school_id).first()
        
#         unpaid_students_fee = 0
#         if unpaid_students and fee_amount:
#             unpaid_students_fee = int(unpaid_students.no_of_unpaid_students) * int(fee_amount.amount)
#         partially_paid_fee = "select sum(outstanding_amount) as pending_amount from fee_detail where fee_amount>fee_paid_amount and school_id='"+str(teacherDataRow.school_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and month='"+str(qmonth)+"' and year='"+str(qyear)+"'"
#         partially_paid_fee = db.session.execute(text(partially_paid_fee)).first()
#         total_unpaid_fee = 0
#         if partially_paid_fee:
#             print('partially paid fee:'+str(partially_paid_fee.pending_amount))
#             if partially_paid_fee.pending_amount:
#                 total_unpaid_fee = unpaid_students_fee + partially_paid_fee.pending_amount
#             else:
#                 total_unpaid_fee = unpaid_students_fee
#     return render_template('_summaryBox.html',paid_fees=paid_fees.collected_fee,paid_student_count=paid_student_count.no_of_paid_students,total_unpaid_students=total_unpaid_students,total_unpaid_fee=total_unpaid_fee)

# New Section added to manage fee status
# @app.route('/feeStatusDetail')
# def feeStatusDetail():
#     qmonth = request.args.get('month')
#     qyear = request.args.get('year')
#     class_val = request.args.get('class_val')
#     section = request.args.get('section')
    
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,section=section,school_id=teacherDataRow.school_id).first()
#     print(qmonth+ ' '+qyear)
#     #days in month
#     daysInMonth = monthrange(int(qyear),int(qmonth))
#     daysInMonth = int(daysInMonth[1])
#     feeStatusDataQuery = "select sp.student_id as student_id, sp.profile_picture as profile_picture, sp.full_name as student_name,sp.roll_number, fd.fee_amount as fee_amount,fd.fee_paid_amount as paid_amount, fd.outstanding_amount as rem_amount, fd.paid_status as paid_status,fd.delay_reason"
#     feeStatusDataQuery = feeStatusDataQuery + " from student_profile  sp left join "
#     feeStatusDataQuery = feeStatusDataQuery + "fee_detail fd on fd.student_id=sp.student_id "
#     feeStatusDataQuery = feeStatusDataQuery + " and fd.month = "+str(qmonth) + " and fd.year = "+ str(qyear) + " where sp.school_id=" + str(teacherDataRow.school_id) + " and sp.class_sec_id='"+str(class_sec_id.class_sec_id)+"' order by paid_status asc"
#     feeStatusDataRows = db.session.execute(text(feeStatusDataQuery)).fetchall()
#     print(str(len(feeStatusDataRows)))
#     sections = ClassSection.query.filter_by(school_id=teacherDataRow.school_id,class_val=class_val).all()
#     total_amt = ''
#     amount = FeeClassSecDetail.query.filter_by(class_sec_id=class_sec_id.class_sec_id,school_id=teacherDataRow.school_id).first()
#     if amount:
#         total_amt = amount.amount
#     print('Total amount:'+str(total_amt))
#     return render_template('_feeStatusTable.html',total_amt=total_amt,feeStatusDataRows=feeStatusDataRows,qmonth=qmonth,qyear=qyear,class_val=class_val,section=section)
# #New Section added to manage payroll
# @app.route('/payrollMonthData')
# def payrollMonthData():
#     qmonth = request.args.get('month')
#     qyear = request.args.get('year')
#     print(qmonth+ ' '+qyear)
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #days in month
#     daysInMonth = monthrange(int(qyear),int(qmonth))
#     daysInMonth = int(daysInMonth[1])
#     #temporary query
#     payrollDataQuery = "select tp.teacher_id as teacher_id, tp.profile_picture as profile_picture, tp.teacher_name as teacher_name, tp.curr_salary as curr_salary,tpd.days_present as days_present, tpd.calc_salary, tpd.paid_status as paid_status"
#     payrollDataQuery = payrollDataQuery + " from teacher_profile  tp left join "
#     payrollDataQuery = payrollDataQuery + "teacher_payroll_detail tpd on tpd.teacher_id=tp.teacher_id "
#     payrollDataQuery = payrollDataQuery + " and tpd.month = "+str(qmonth) + " and tpd.year = "+ str(qyear) + " where tp.school_id=" + str(teacherDataRow.school_id) + " order by paid_status asc"
#     payrollDataRows = db.session.execute(text(payrollDataQuery)).fetchall()
#     print(str(len(payrollDataRows)))
#     return render_template('_payrollMonthData.html',daysInMonth=daysInMonth, payrollDataRows=payrollDataRows, qmonth=qmonth, qyear = qyear)

# @app.route('/updateFeeData', methods=['GET','POST'])
# def updateFeeData():
#     teacherDetailRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
#     qmonth = request.form.get('qmonth')
#     qyear = request.form.get('qyear')
#     total_amt = request.args.get('total_amt')
#     total_amt = total_amt.strip()
#     qclass_val = request.form.get('qclass_val')
#     qsection = request.form.get('qsection')
#     print('inside updateFeeData')
#     print('Total Fee Amount:'+str(total_amt))
#     class_sec_id = ClassSection.query.filter_by(class_val=qclass_val,section=qsection,school_id=teacherDetailRow.school_id).first()
#     student_id_list = request.form.getlist('student_id')
#     paid_amount_list = request.form.getlist('paid_amount')
#     rem_amount_list = request.form.getlist('rem_amount')
#     # validation when rem_amount is negative
#     for i in range(len(rem_amount_list)-1):
#         print('inside re_amount_list')
#         print(i)
#         print(rem_amount_list[i])
#         if rem_amount_list[i]:
#             if int(rem_amount_list[i])<0:
#                 return jsonify(['1'])
#     # End
#     delay_reason_list = request.form.getlist('delay_reason')
#     count_list = []
#     for i in range(len(paid_amount_list)):
#         if paid_amount_list[i]:
#             print('counter:'+str(i))
#             print('paid amount:'+str(paid_amount_list[i]))
#             count_list.append(i)
#     # print(paid_amount_list)
#     print('count_list length:'+str(len(count_list)))
#     for i in range(len(count_list)):
#         print('inside for loop')
#         print(count_list[i])
#         print(student_id_list[count_list[i]])
#         if paid_amount_list[count_list[i]]:
#             indivFeeRecord = FeeDetail.query.filter_by(student_id=student_id_list[count_list[i]], month=qmonth, year=qyear).first()
#             if indivFeeRecord and indivFeeRecord.outstanding_amount!=0:
#                 print('if record already exist:'+str(paid_amount_list[count_list[i]]))
#                 indivFeeRecord.fee_amount = total_amt
#                 indivFeeRecord.fee_paid_amount = paid_amount_list[count_list[i]]
#                 indivFeeRecord.outstanding_amount = rem_amount_list[count_list[i]]
#                 indivFeeRecord.delay_reason = delay_reason_list[count_list[i]]
#                 print('pending amount:'+str(rem_amount_list[count_list[i]]))
#                 if rem_amount_list[count_list[i]]==0 or rem_amount_list[count_list[i]]=='0':
#                     indivFeeRecord.paid_status = 'Y'
#                 else:
#                     indivFeeRecord.paid_status = 'N'
#             elif indivFeeRecord==None or indivFeeRecord=='':
#                 print('Adding new values:'+str(paid_amount_list[count_list[i]]))
#                 print('Paid Amount:'+paid_amount_list[count_list[i]])
#                 print('Total Amount:'+total_amt)
#                 if float(paid_amount_list[count_list[i]])==float(total_amt):
#                     print('if paid amount equal to total amount')
#                     feeInsert=FeeDetail(school_id=teacherDetailRow.school_id,student_id=student_id_list[count_list[i]],fee_amount = total_amt,
#                     class_sec_id=class_sec_id.class_sec_id,payment_date=datetime.today(),fee_paid_amount = paid_amount_list[count_list[i]],outstanding_amount=rem_amount_list[count_list[i]],month=qmonth,year=qyear
#                     ,paid_status='Y',delay_reason=delay_reason_list[count_list[i]],last_modified_date=datetime.today())
#                 else:
#                     print('if paid amount is less than total amount')
#                     feeInsert=FeeDetail(school_id=teacherDetailRow.school_id,student_id=student_id_list[count_list[i]],fee_amount = total_amt,
#                     class_sec_id=class_sec_id.class_sec_id,payment_date=datetime.today(),fee_paid_amount = paid_amount_list[count_list[i]],outstanding_amount=rem_amount_list[count_list[i]],month=qmonth,year=qyear
#                     ,paid_status='N',delay_reason=delay_reason_list[count_list[i]],last_modified_date=datetime.today())
#                 db.session.add(feeInsert)
#     db.session.commit()
#     return jsonify(['0'])

# @app.route('/updatePayrollData', methods=['GET','POST'])
# def updatePayrollData():
#     teacherDetailRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
#     teacher_id_list = request.form.getlist('teacher_id')    
#     current_salary_list = request.form.getlist('currentSalaryInput')
#     days_count_list = request.form.getlist('dayCountInput')
#     days_present_list = request.form.getlist('days_present')
#     calc_salary_list = request.form.getlist('calcSalaryInput')
#     #paid_status_list = request.form.getlist('paid_status')
#     has_changed_list = request.form.getlist('hasChanged')
#     qmonth = request.form.get('qmonth')
#     qyear = request.form.get('qyear')
#     #print(teacher_id_list)
#     #print(current_salary_list)
#     #print(days_count_list)
#     #print(days_present_list)
#     #print(calc_salary_list)
#     #print(has_changed_list)
#     #print(qmonth)
#     #print(qyear)
#     ##print(paid_status_list)
#     #print("#########")
#     for i in range(len(has_changed_list)):
#         print("This is the value of i "+ str(i))
#         if has_changed_list[i]=='Y':   
#             print ('Something has changed')
#             #if (paid_status_list[i]):
#             #    paidValue = 'Y'
#             #else:
#             #    paidValue='N'
#             indivPayrollRecord = TeacherPayrollDetail.query.filter_by(teacher_id=teacher_id_list[i], month=qmonth, year=qyear).first()
#             if indivPayrollRecord==None:
#                 print('Adding new values')
#                 payrollInsert=TeacherPayrollDetail(teacher_id=teacher_id_list[i],total_salary=current_salary_list[i],month=qmonth,
#                     year=qyear,days_in_month=days_count_list[i],days_present = days_present_list[i], calc_salary = calc_salary_list[i], paid_status='Y',
#                     last_modified_date=datetime.today(), school_id = teacherDetailRow.school_id)
#                 db.session.add(payrollInsert)
#             else:
#                 if indivPayrollRecord.calc_salary!=calc_salary_list[i]:
#                     print('Updating exiting values')
#                     indivPayrollRecord.days_present = days_present_list[i]
#                     indivPayrollRecord.calc_salary= calc_salary_list[i]
#                     indivPayrollRecord.paid_status = 'Y'
#                     indivPayrollRecord.last_modified_date = datetime.today()

#     db.session.commit()
#     return jsonify(['0'])

#End of section for payroll



# @app.route('/bulkStudReg')
# def bulkStudReg():
#     return render_template('_bulkStudReg.html')


# @app.route('/singleStudReg')
# def singleStudReg():
#     student_id = request.args.get('student_id')
#     print('Inside single student Registration:'+str(student_id))
#     if student_id=='':
#         teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher_id.school_id).all()
#         section_list=[(i.section,i.section) for i in available_section]
#         form=SingleStudentRegistration()
#         form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
#         form.section.choices= section_list
#         studentDetailRow = []
#         guardianDetail1 = []
#         guardianDetail2 =[]
#         return render_template('_singleStudReg.html',form=form,student_id=student_id,studentDetailRow=studentDetailRow,guardianDetail1=guardianDetail1,guardianDetail2=guardianDetail2)
#     else:
#         teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher_id.school_id).all()
#         section_list=[(i.section,i.section) for i in available_section]
#         form=SingleStudentRegistration()
#         form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
#         form.section.choices= section_list
#         guardianDetail2 = ''
#         query = "select sp.first_name,sp.last_name,sp.profile_picture,md.description as gender,date(sp.dob) as dob,sp.phone,ad.address_1,ad.address_2,ad.locality,ad.city,ad.state,ad.country,ad.pin,cs.class_val,cs.section, sp.roll_number, sp.school_adm_number from student_profile sp "                 
#         query = query + "inner join message_detail md on md.msg_id=sp.gender "
#         query = query + "left join class_section cs on cs.class_sec_id=sp.class_sec_id " 
#         query = query + "left join address_detail ad on ad.address_id=sp.address_id "
#         query = query + "where sp.student_id='"+str(student_id)+"'"
#         # query = query + "left join guardian_profile gp on gp.student_id=sp.student_id "
#         # query = query + "inner join message_detail md2 on md2.msg_id=gp.relation where sp.student_id='"+str(student_id)+"'"
#         studentDetailRow = db.session.execute(text(query)).first()
#         queryGuardian1 = "select gp.guardian_id,gp.first_name,gp.last_name,gp.email,gp.phone,m1.description as relation from guardian_profile gp inner join message_detail m1 on m1.msg_id=gp.relation where student_id='"+str(student_id)+"'"
#         guardianDetail1 = db.session.execute(text(queryGuardian1)).first()
#         #print('Guardain Detail1 :')
#         #print(guardianDetail1)
#         guardianDetail2 = ''
#         if guardianDetail1!=None:
#             #print('If guardian Detail1 is not empty')
#             queryGuardian2 = "select gp.guardian_id,gp.first_name,gp.last_name,gp.email,gp.phone,m1.description as relation from guardian_profile gp inner join message_detail m1 on m1.msg_id=gp.relation where student_id='"+str(student_id)+"' and guardian_id!='"+str(guardianDetail1.guardian_id)+"'"
#             guardianDetail2 = db.session.execute(text(queryGuardian2)).first()
#         # print(guardianDetail1)
#         # print(guardianDetail2)
#         #print('Name:'+str(studentDetailRow.first_name))
#         #print('Gender:'+str(studentDetailRow.class_val))
#         return render_template('_singleStudReg.html',form=form,student_id=student_id,studentDetailRow=studentDetailRow,guardianDetail1=guardianDetail1,guardianDetail2=guardianDetail2)


# @app.route('/studentRegistration', methods=['GET','POST'])
# @login_required
# def studentRegistration(): 
#     studId = request.args.get('student_id') 
#     form=SingleStudentRegistration()
#     if request.method=='POST':
#         print('Inside Student Registration')
#         if form.submit.data:            
#             studentId = request.form['tag']
#             print('Student Id:'+str(studentId))
#             if studentId:                
#                 print('Inside Student update when student id is not empty')
#                 student_id = request.form['tag']
#                 teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#                 class_sec=ClassSection.query.filter_by(class_val=str(form.class_val.data),section=form.section.data,school_id=teacher_id.school_id).first()
#                 gender=MessageDetails.query.filter_by(description=form.gender.data).first()
#                 address_id=Address.query.filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#                 if address_id is None:
#                     address_data=Address(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data,country=form.country.data)
#                     db.session.add(address_data)
#                     address_id=db.session.query(Address).filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#                 studentDetails = StudentProfile.query.filter_by(student_id=student_id).first()
#                 studentClassSec = StudentClassSecDetail.query.filter_by(student_id=student_id).first()
#                 studProfile = "update student_profile set profile_picture='"+str(request.form['profile_image'])+"' where student_id='"+student_id+"'"
#                 #print('Query:'+str(studProfile))

#                 profileImg = db.session.execute(text(studProfile))
#                 #print(studentDetails)
#                 #if request.form['birthdate']:
#                     #print('DOB:'+str(request.form['birthdate']))
#                 #print('Student id:'+str(student_id))
#                 #print('First Name:'+str(studentDetails.first_name))
#                 #print('Image url:'+str(request.form['profile_image']))
#                 studentDetails.first_name=form.first_name.data
#                 studentDetails.last_name=form.last_name.data
#                 studentDetails.gender=gender.msg_id
#                 if request.form['birthdate']=='':
#                     studentDetails.dob=None
#                 else:
#                     studentDetails.dob=request.form['birthdate']
#                 studentDetails.phone=form.phone.data
#                 studentDetails.address_id=address_id.address_id
#                 studentDetails.profile_image=request.form['profile_image']
#                 studentDetails.class_sec_id=class_sec.class_sec_id
#                 studentDetails.roll_number=int(form.roll_number.data)
#                 studentDetails.school_adm_number=form.school_admn_no.data
#                 studentDetails.full_name=form.first_name.data +" " + form.last_name.data
#                 if studentClassSec!=None and studentClassSec!="":
#                     studentClassSec.class_sec_id = class_sec.class_sec_id
#                     studentClassSec.class_val = str(form.class_val.data)
#                     studentClassSec.section = form.section.data
#                 else:
#                     studentClassSecAdd = StudentClassSecDetail(student_id=student_id, class_sec_id=class_sec.class_sec_id, 
#                         class_val=str(form.class_val.data), section=form.section.data, is_current='Y', last_modified_date=datetime.today())
#                     db.session.add(studentClassSecAdd)
#                 db.session.commit()
#                 first_name=request.form.getlist('guardian_first_name')
#                 last_name=request.form.getlist('guardian_last_name')
#                 phone=request.form.getlist('guardian_phone')
#                 email=request.form.getlist('guardian_email')
#                 relation=request.form.getlist('relation')
#                 gId = ''
#                 for i in range(len(first_name)):
#                     if i==0:
#                         relation_id=MessageDetails.query.filter_by(description=relation[i]).first()
#                         # query = "select first_name,last_name,full_name,relation,email,phone from guardian_profile where student_id='"+str(student_id)+"' limit 1"
#                         # guardianData = db.session.execute(text(query))
#                         guardianData = GuardianProfile.query.filter_by(student_id=student_id).first()
#                         # print('Query:'+query)
#                         # print(guardianData.first_name)
#                         if guardianData:
#                             guardianData.first_name = first_name[i]
#                             guardianData.last_name = last_name[i]
#                             guardianData.full_name = first_name[i] + ' ' + last_name[i]
#                             guardianData.relation = relation_id.msg_id
#                             guardianData.phone = phone[i]
#                             guardianData.email = email[i]
#                             gId = guardianData.guardian_id
#                             # gId = int(gId)+1 
#                             print('Gid:'+str(guardianData.guardian_id))
#                             print('Gid:'+str(gId))
#                             print('Guardian First Name:'+str(first_name[i]))
#                             db.session.commit()
#                         else:
#                             relation_id=MessageDetails.query.filter_by(description=relation[i]).first()
#                             guardian_data=GuardianProfile(first_name=first_name[i],last_name=last_name[i],full_name=first_name[i] + ' ' + last_name[i],relation=relation_id.msg_id,
#                             email=email[i],phone=phone[i],student_id=student_id)
#                             db.session.add(guardian_data)    
#                     if i==1:
#                         query = "select *from guardian_profile where student_id='"+str(student_id)+"' and guardian_id!='"+str(gId)+"'"
#                         print('Query:'+str(query))
#                         guardian_id = db.session.execute(text(query)).first()
                        
#                         if guardian_id:
#                             guarId = guardian_id.guardian_id
#                             relation_id=MessageDetails.query.filter_by(description=relation[i]).first()
#                             guardianData = GuardianProfile.query.filter_by(student_id=student_id,guardian_id=guarId).first()
#                             print('Second guardian first name:'+str(guardianData.first_name))
#                             guardianData.first_name = first_name[i]
#                             guardianData.last_name = last_name[i]
#                             guardianData.full_name = first_name[i] + ' ' + last_name[i]
#                             guardianData.relation = relation_id.msg_id
#                             guardianData.phone = phone[i]
#                             guardianData.email = email[i]
#                             print(guardianData)
#                             print('Second Guardian First Name:'+str(first_name[i]))
#                         else:
#                             relation_id=MessageDetails.query.filter_by(description=relation[i]).first()
#                             guardian_data=GuardianProfile(first_name=first_name[i],last_name=last_name[i],full_name=first_name[i] + ' ' + last_name[i],relation=relation_id.msg_id,
#                             email=email[i],phone=phone[i],student_id=student_id)
#                             db.session.add(guardian_data)
#                 # guardian_data=GuardianProfile(first_name=first_name[i],last_name=last_name[i],full_name=first_name[i] + ' ' + last_name[i],relation=relation_id.msg_id,
#                 # email=email[i],phone=phone[i],student_id=student_data.student_id)
#                 db.session.commit()
#                 flash(Markup('Data Updated Successfully! Go to <a href="/studentProfile"> Student Directory</a>?'))
#                 indic='DashBoard'
#                 return render_template('studentRegistration.html',indic=indic,title='Student Registration',studentId=student_id,user_type_val=str(current_user.user_type))


#             else:
#                 print('Inside Student Registration when student id is empty')
#                 address_id=Address.query.filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#                 if address_id is None:
#                     address_data=Address(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data,country=form.country.data)
#                     db.session.add(address_data)
#                     address_id=db.session.query(Address).filter_by(address_1=form.address1.data,address_2=form.address2.data,locality=form.locality.data,city=form.city.data,state=form.state.data,pin=form.pincode.data).first()
#                 teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#                 print('Print Form Data:'+form.section.data)
                
#                 class_sec=ClassSection.query.filter_by(class_val=str(form.class_val.data),section=form.section.data,school_id=teacher_id.school_id).first()
#                 gender=MessageDetails.query.filter_by(description=form.gender.data).first()
#                 print('Section Id:'+str(class_sec.class_sec_id))
#                 if request.form['birthdate']:
#                     student=StudentProfile(first_name=form.first_name.data,last_name=form.last_name.data,full_name=form.first_name.data +" " + form.last_name.data,
#                     school_id=teacher_id.school_id,class_sec_id=class_sec.class_sec_id,gender=gender.msg_id,
#                     dob=request.form['birthdate'],phone=form.phone.data,profile_picture=request.form['profile_image'],address_id=address_id.address_id,school_adm_number=form.school_admn_no.data,
#                     roll_number=int(form.roll_number.data))
#                 else:
#                     student=StudentProfile(first_name=form.first_name.data,last_name=form.last_name.data,full_name=form.first_name.data +" " + form.last_name.data,
#                     school_id=teacher_id.school_id,class_sec_id=class_sec.class_sec_id,gender=gender.msg_id,
#                     phone=form.phone.data,profile_picture=request.form['profile_image'],address_id=address_id.address_id,school_adm_number=form.school_admn_no.data,
#                     roll_number=int(form.roll_number.data))
#                 #print('Query:'+student)
#                 db.session.add(student)
#                 student_data=db.session.query(StudentProfile).filter_by(school_adm_number=form.school_admn_no.data).first()
#                 for i in range(4):
#                     if i==0:
#                         option='A'
#                         qr_link='https:er5ft/api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + form.roll_number.data + '-' + student_data.first_name + '@' + option
#                     elif i==1:
#                         option='B'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + form.roll_number.data + '-' + student_data.first_name + '@' + option
#                     elif i==2:
#                         option='C'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + form.roll_number.data + '-' + student_data.first_name + '@' + option
#                     else:
#                         option='D'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + form.roll_number.data + '-' + student_data.first_name + '@' + option
#                     student_qr_data=studentQROptions(student_id=student_data.student_id,option=option,qr_link=qr_link)
#                     db.session.add(student_qr_data)
#                 first_name=request.form.getlist('guardian_first_name')
#                 last_name=request.form.getlist('guardian_last_name')
#                 phone=request.form.getlist('guardian_phone')
#                 email=request.form.getlist('guardian_email')
#                 relation=request.form.getlist('relation')
#                 print('Insert into ClassSec table')
#                 student_class_sec = StudentClassSecDetail(student_id=student_data.student_id, class_sec_id=class_sec.class_sec_id,
#                 class_val=str(form.class_val.data), section=form.section.data, is_current='Y',last_modified_date=datetime.today()) 
#                 db.session.add(student_class_sec)
#                 for i in range(len(first_name)):
#                     relation_id=MessageDetails.query.filter_by(description=relation[i]).first()
#                     guardian_id = GuardianProfile.query.filter_by(email=email[i]).first()
#                     if guardian_id:
#                         print('If guardian already exist')
#                         if guardian_id.student_id=='':
#                             guardian_id.student_id = student_data.student_id
#                             guardian_id.relation = relation_id.msg_id
#                             print('If Id is empty')
#                         else:
#                             print('skip')
#                             guardian_data=GuardianProfile(first_name=first_name[i],last_name=last_name[i],full_name=first_name[i] + ' ' + last_name[i],relation=relation_id.msg_id,
#                             email=email[i],phone=phone[i],user_id=guardian_id.user_id,student_id=student_data.student_id)
#                             db.session.add(guardian_data)
#                     else:
#                         guardian_data=GuardianProfile(first_name=first_name[i],last_name=last_name[i],full_name=first_name[i] + ' ' + last_name[i],relation=relation_id.msg_id,
#                         email=email[i],phone=phone[i],student_id=student_data.student_id)
#                         db.session.add(guardian_data)
#                         print('If guardian does not exist')
#                 db.session.commit()
#                 flash('Successful upload !')
#                 indic='DashBoard'
#                 return render_template('studentRegistration.html',indic=indic,title='Student Registration',user_type_val=str(current_user.user_type))

#         else:
#             teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#             print('School_id:'+str(teacher_id.school_id))
#             csv_file=request.files['file-input']
#             df1=pd.read_csv(csv_file)
#             df1=df1.replace(np.nan, '', regex=True)
#             print(df1)
#             for index ,row in df1.iterrows():
#                 if row['first_name']=='' and row['gender']=='' and row['dob']=='' and row['phone'] and row['address_1'] and row['locality'] and row['city'] and row['state'] and row['country'] and row['pin'] and row['class_val'] and row['section'] and row['roll_number'] and row['school_adm_number'] and row['guardian1_first_name'] and row['guardian1_email'] and row['guardian1_phone'] and row['guardian1_relation']:
#                     message = Markup("<h5>(a)Enter first name <br/>(b)Enter gender<br/> (c)Enter date of birth <br/> (d)Enter phone number<br/> (e)Enter address 1 <br/>(f)Enter locality<br/> (g)Enter city<br/> (h)Enter state <br/>(i)Enter country<br/>(j)Enter pin code<br/> (k)Enter class<br/> (l)Enter section<br/> (m)Enter roll number<br/> (n)Enter school admission number<br/>(o)Enter guardian first name<br/> (p)Enter guardian email<br/>(q)Enter guardian phone<br/> (r)Enter guardian relation </h5>")
#                     flash(message)
#                     return render_template('studentRegistration.html')
#                 address_data=Address(address_1=row['address_1'],address_2=row['address_2'],locality=row['locality'],city=row['city'],state=row['state'],pin=str(row['pin']),country=row['country'])
#                 db.session.add(address_data)
#                 address_id=db.session.query(Address).filter_by(address_1=row['address_1'],address_2=row['address_2'],locality=row['locality'],city=row['city'],state=row['state'],pin=str(row['pin']),country=row['country']).first()
#                 teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#                 class_sec=ClassSection.query.filter_by(class_val=str(row['class_val']),section=row['section'],school_id=teacher_id.school_id).first()
#                 gender=MessageDetails.query.filter_by(description=row['gender']).first()
#                 date = row['dob']
#                 li = date.split('/',3)
#                 print('Date'+str(row['dob']))
#                 if int(li[1])>12 or int(li[0])>31:
#                     flash('Invalid Date formate use dd/mm/yyyy')
#                     return render_template('studentRegistration.html')
#                 if row['dob']!='':
#                     date=dt.datetime.strptime(row['dob'], '%d/%m/%Y')
#                 else:
#                     date=''
#                 student=StudentProfile(first_name=row['first_name'],last_name=row['last_name'],full_name=row['first_name'] +" " + row['last_name'],
#                 school_id=teacher_id.school_id,class_sec_id=class_sec.class_sec_id,gender=gender.msg_id,
#                 dob=date,phone=row['phone'],profile_picture=request.form['reference-url'+str(index+1)],address_id=address_id.address_id,school_adm_number=str(row['school_adm_number']),
#                 roll_number=int(row['roll_number']))
#                 db.session.add(student)
#                 student_data=db.session.query(StudentProfile).filter_by(school_adm_number=str(row['school_adm_number'])).first()
#                 print('Insert into ClassSec table')
#                 student_class_sec = StudentClassSecDetail(student_id=student_data.student_id, class_sec_id=class_sec.class_sec_id,
#                 class_val=str(row['class_val']), section=row['section'], is_current='Y',last_modified_date=datetime.today()) 
#                 db.session.add(student_class_sec)
#                 for i in range(4):
#                     if i==0:
#                         option='A'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + str(row['roll_number']) + '-' + student_data.first_name + '@' + option
#                     elif i==1:
#                         option='B'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + str(row['roll_number']) + '-' + student_data.first_name + '@' + option
#                     elif i==2:
#                         option='C'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + str(row['roll_number']) + '-' + student_data.first_name + '@' + option
#                     else:
#                         option='D'
#                         qr_link='https://api.qrserver.com/v1/create-qr-code/?size=150x150&data=' + str(student_data.student_id) + '-' + str(row['roll_number']) + '-' + student_data.first_name + '@' + option
#                     student_qr_data=studentQROptions(student_id=student_data.student_id,option=option,qr_link=qr_link)
#                     db.session.add(student_qr_data)
#                 for i in range(2):
#                     relation_id=MessageDetails.query.filter_by(description=row['guardian'+str(i+1)+'_relation']).first()
#                     print('Inside range')
#                     if relation_id is not None:
#                         print('If relation id is not empty')
#                         guardian_data=GuardianProfile(first_name=row['guardian'+str(i+1)+'_first_name'],last_name=row['guardian'+str(i+1)+'_last_name'],full_name=row['guardian'+str(i+1)+'_first_name'] + ' ' + row['guardian'+str(i+1)+'_last_name'],relation=relation_id.msg_id,
#                         email=row['guardian'+str(i+1)+'_email'],phone=row['guardian'+str(i+1)+'_phone'],student_id=student_data.student_id)
                    
#                     db.session.add(guardian_data)
#                     db.session.commit()
                
            
#             flash('Successful upload !')
#             indic='DashBoard'
#             return render_template('studentRegistration.html',indic=indic,user_type_val=str(current_user.user_type))
#     if studId!='':
#         print('inside if Student Id:'+str(studId))
#         indic='DashBoard'
#         return render_template('studentRegistration.html',indic=indic,studentId=studId,user_type_val=str(current_user.user_type))
#     indic='DashBoard'
#     return render_template('studentRegistration.html',indic=indic,user_type_val=str(current_user.user_type))


'''camera section'''

@app.route('/video_feed')
def video_feed(): 
    cam=VideoCamera()
    #cam.is_record=True
    return Response(cam.gen(), mimetype='multipart/x-mixed-replace; boundary=frame')


@app.route('/video_feed_stop')
def video_feed_stop(): 
    #cam.is_record=False
    cam=VideoCamera()
    return Response(cam.closeCam())

'''camera section ends'''

'''new cam section'''


@app.route('/ScanBooks',methods=['GET', 'POST'])
def ScanBooks():
    print ("We're here!")
    return render_template('ScanBook.html',title='Scan Page')


@app.route('/testingOtherVideo',methods=['GET', 'POST'])
def testingOtherVideo():
    print ("We're here!")
    return render_template('testingOtherVideo.html',title='Test Page')

'''
'''
@app.route('/edit_profile', methods=['GET', 'POST'])
@login_required
def edit_profile():
    #teacher= TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #cityList()

    #cityJSON = json.dumps(cityList())
    #stateJSON = json.dumps(stateList())

    form = EditProfileForm(current_user.username)
    if form.validate_on_submit():        
        current_user.about_me = form.about_me.data        
        current_user.first_name= form.first_name.data
        current_user.last_name= form.last_name.data                
        current_user.education = form.education.data
        current_user.experience = form.experience.data
        current_user.phone=form.phone.data
        #current_user.address=form.address.data
        current_user.city=form.city.data
        current_user.state=form.state.data
        current_user.resume=form.resume.data
        current_user.willing_to_travel = form.willing_to_travel.data
        ##
        db.session.commit()
        flash('Your changes have been saved')        
        if current_user.user_type==161:
            return redirect(url_for('job_post.openJobs'))

    elif request.method == 'GET':        
        form.about_me.data = current_user.about_me
        form.first_name.data = current_user.first_name
        form.last_name.data = current_user.last_name
        form.phone.data = current_user.phone
        form.education.data = current_user.education
        form.experience.data = current_user.experience
        #form.address.data = current_user.address
        form.city.data = current_user.city
        form.state.data = current_user.state
        form.resume.data = current_user.resume
        form.intro_link.data = current_user.intro_link
        
        print('Inside edit profile')
        print(current_user.about_me)
        print(current_user.first_name)
    return render_template(
        'edit_profile.html', title='Edit Profile', form=form,user_type_val=str(current_user.user_type), willing_to_travel=current_user.willing_to_travel)

# @app.route('/',methods=["GET","POST"])
# @app.route('/index')
# @app.route('/dashboard')
# @login_required 
# def index():
#     #print('Inside index')
#     #print("########This is the request url: "+str(request.url))
#     print('current_user.id:'+str(current_user.id))
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     schoolData = ''
#     if teacherData:
#         schoolData = SchoolProfile.query.filter_by(school_admin=teacherData.teacher_id).first()
#     if schoolData:
#         if schoolData.is_verified == 'N':
#             return redirect(url_for('inReviewSchool'))
#     checkUser = User.query.filter_by(id=current_user.id).first()
#     if checkUser:
#         if checkUser.access_status == 143:
#             return redirect(url_for('disconnectedAccount'))
#     user = User.query.filter_by(username=current_user.username).first_or_404()        
#     school_name_val = schoolNameVal()
#     #print('User Type Value:'+str(user.user_type))
#     teacher_id = TeacherProfile.query.filter_by(user_id=user.id).first() 
    
#     school_id = SchoolProfile.query.filter_by(school_name=school_name_val).first()
#     print('school_name_val:',school_name_val)
#     if user.user_type==71:
#         classExist = ClassSection.query.filter_by(school_id=school_id.school_id).first()
#         #print('Insert new school')
#         #print(classExist)
#         if classExist==None:
#             fromSchoolRegistration = True
       
#             subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#             board = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#             boardRows = MessageDetails.query.filter_by(msg_id=board.board_id).first()
#             school_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#             classValues = "SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id = '"+str(teacher_id.school_id)+"' GROUP BY class_val order by s"
#             classValues = db.session.execute(text(classValues)).fetchall()
#             classValuesGeneral = "SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs GROUP BY class_val order by s"
#             classValuesGeneral = db.session.execute(text(classValuesGeneral)).fetchall()
#             subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#             bookName = BookDetails.query.all()
#             chapterNum = Topic.query.distinct().all()
#             topicId = Topic.query.all()
#             generalBoardId = SchoolProfile.query.filter_by(school_id = teacher_id.school_id).first()
#             #print('teacher and board ids')
#             #print(teacher_id.school_id)
#             #print(generalBoardId.board_id)
#             generalBoard = MessageDetails.query.filter_by(msg_id=generalBoardId.board_id).first()
#             fromSchoolRegistration = True
#             return render_template('syllabus.html',title='Syllabus',generalBoard=generalBoard,boardRowsId = boardRows.msg_id , boardRows=boardRows.description,subjectValues=subjectValues,school_name=school_id.school_name,classValues=classValues,classValuesGeneral=classValuesGeneral,bookName=bookName,chapterNum=chapterNum,topicId=topicId,fromSchoolRegistration=fromSchoolRegistration)
#     #if user.user_type==135:
#     #    return redirect(url_for('admin'))
#     if user.user_type==234:
#     #or ("prep.alllearn" in str(request.url)) or ("alllearnprep" in str(request.url))
#         return redirect(url_for('practiceTest'))
#     if user.user_type==253:
#         return redirect(url_for('courseHome'))
#     if user.user_type==72:
#         #print('Inside guardian')
#         return redirect(url_for('disconnectedAccount'))
#     if user.user_type=='161':
#         return redirect(url_for('job_post.openJobs'))
#     if user.user_type==134 and user.access_status==145:        
#         return redirect(url_for('disconnectedAccount'))

#     teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    
#     classSecCheckVal = classSecCheck()

#     if school_name_val ==None:
#         #print('did we reach here')
#         return redirect(url_for('disconnectedAccount'))
#     else:
#     #####Fetch school perf graph information##########
#         performanceQuery = "select * from fn_class_performance("+str(teacher.school_id)+") order by perf_date"
#         performanceRows = db.session.execute(text(performanceQuery)).fetchall()
#         if len(performanceRows)>0:
#             df = pd.DataFrame( [[ij for ij in i] for i in performanceRows])
#             df.rename(columns={0: 'Date', 1: 'Class_1', 2: 'Class_2', 3: 'Class_3', 4:'Class_4',
#                 5:'Class_5', 6:'Class_6', 7:'Class_7', 8:'Class_8', 9:'Class_9', 10:'Class_10'}, inplace=True)
#             #print(df)
#             dateRange = list(df['Date'])
#             class1Data= list(df['Class_1'])
#             class2Data= list(df['Class_2'])
#             class3Data= list(df['Class_3'])
#             class4Data= list(df['Class_4'])
#             class5Data= list(df['Class_5'])
#             class6Data= list(df['Class_6'])
#             class7Data= list(df['Class_7'])
#             class8Data= list(df['Class_8'])
#             class9Data= list(df['Class_9'])
#             class10Data= list(df['Class_10'])
#             #print(dateRange)
#             ##Class 1
#             graphData = [dict(
#                 data1=[dict(y=class1Data,x=dateRange,type='scatter', name='Class 1')],
#                 data2=[dict(y=class2Data,x=dateRange,type='scatter', name='Class 2')],
#                 data3=[dict(y=class3Data,x=dateRange,type='scatter', name='Class 3')],
#                 data4=[dict(y=class4Data,x=dateRange,type='scatter', name='Class 4')],
#                 data5=[dict(y=class5Data,x=dateRange,type='scatter', name='Class 5')],
#                 data6=[dict(y=class6Data,x=dateRange,type='scatter', name='Class 6')],
#                 data7=[dict(y=class7Data,x=dateRange,type='scatter', name='Class 7')],
#                 data8=[dict(y=class8Data,x=dateRange,type='scatter', name='Class 8')],
#                 data9=[dict(y=class9Data,x=dateRange,type='scatter', name='Class 9')],
#                 data10=[dict(y=class10Data,x=dateRange,type='scatter', name='Class 10')]
#                 )]        
#             #print(graphData)

#             graphJSON = json.dumps(graphData, cls=plotly.utils.PlotlyJSONEncoder)
#         else:
#             graphJSON="1"
#     #####Fetch Top Students infor##########        
#         # topStudentsQuery = "select *from fn_monthly_top_students("+str(teacher.school_id)+",8)"
#         qclass_val = 'dashboard'
#         topStudentsRows = ''
#         leaderBoardData = leaderboardContent(qclass_val)
#         # print('leaderBoard Data:'+str(leaderBoardData))
#         # Convert dataframe to a list
#         data = []
#         #print(type(leaderBoardData))
#         column_names = ["a", "b", "c"]
#         datafr = pd.DataFrame(columns = column_names)
#         if type(leaderBoardData)==type(datafr):
#             #print('if data is not empty')
#             df1 = leaderBoardData[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
#             df2 = leaderBoardData.drop(['profile_pic', 'student_name','class_val','section','total_marks%','total_tests'], axis=1)
#             leaderBoard = pd.merge(df1,df2,on=('studentid'))
                
#             d = leaderBoard[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
#             df3 = leaderBoard.drop(['studentid'],axis=1)
            
#             df1.rename(columns = {'profile_pic':'Profile Picture'}, inplace = True)
#             df1.rename(columns = {'student_name':'Student'}, inplace = True)
#             df1.rename(columns = {'class_val':'Class'}, inplace = True)
#             df1.rename(columns = {'section':'Section'}, inplace = True)
#             df1.rename(columns = {'total_marks%':'Total Marks'}, inplace = True)
#             df1.rename(columns = {'total_tests':'Total Tests'}, inplace = True)
#             header = [df1.columns.values.tolist()]
#             headerAll = [df3.columns.values.tolist()]
#             colAll = ''
#             subjHeader = [df2.columns.values.tolist()]
#             columnNames = ''
#             col = ''
#             subColumn = ''
#             for subhead in subjHeader:
#                 subColumn = subhead
#             for h in header:
#                 columnNames = h
#             for headAll in headerAll: 
#                 colAll = headAll
#             n= int(len(subColumn)/2)
#             ndf = df2.drop(['studentid'],axis=1)
#             newDF = ndf.iloc[:,0:n]
#             new1DF = ndf.iloc[:,n:]
                
#             df5 = pd.concat([newDF, new1DF], axis=1)
#             DFW = df5[list(sum(zip(newDF.columns, new1DF.columns), ()))]
            
            
#             dat = pd.concat([d,DFW], axis=1)
                
#             dat = dat.sort_values('total_marks%',ascending=False)  
            
#             subHeader = ''
#             i=1
#             for row in dat.values.tolist():
#                 if i<9:
#                     data.append(row)
#                 i=i+1
#         form  = promoteStudentForm() 
#         available_class=ClassSection.query.with_entities(ClassSection.class_val,ClassSection.section).distinct().order_by(ClassSection.class_val,ClassSection.section).filter_by(school_id=teacher.school_id).all()
#         class_list=[(str(i.class_val)+"-"+str(i.section),str(i.class_val)+"-"+str(i.section)) for i in available_class]
        
#         form.class_section1.choices = class_list 
#         form.class_section2.choices = class_list 
        
#         EventDetailRows = EventDetail.query.filter_by(school_id=teacher.school_id).all()
#     #####Fetch Course Completion infor##########    
#         topicToCoverQuery = "select *from fn_topic_tracker_overall("+str(teacher.school_id)+") order by class, section"
#         topicToCoverDetails = db.session.execute(text(topicToCoverQuery)).fetchall()
#         #print(topicToCoverDetails)

#     ##################Fetch Job post details################################
#         jobPosts = JobDetail.query.filter_by(school_id=teacher.school_id).order_by(JobDetail.posted_on.desc()).all()
#         teacherCount = "select count(*) from teacher_profile tp where school_id = '"+str(teacher.school_id)+"'"
#         teacherCount = db.session.execute(teacherCount).first()
#         studentCount = "select count(*) from student_profile sp where school_id = '"+str(teacher.school_id)+"'"
#         studentCount = db.session.execute(studentCount).first()
#         testCount = "select (select count(distinct upload_id) from result_upload ru where school_id = '"+str(teacher.school_id)+"') + "
#         testCount = testCount + "(select count(distinct resp_session_id) from response_capture rc2 where school_id = '"+str(teacher.school_id)+"') as SumCount"
#         #print(testCount)
#         testCount = db.session.execute(testCount).first()
#         lastWeekTestCount = "select (select count(distinct upload_id) from result_upload ru where school_id = '"+str(teacher.school_id)+"' and last_modified_date >=current_date - 7) + "
#         lastWeekTestCount = lastWeekTestCount + "(select count(distinct resp_session_id) from response_capture rc2 where school_id = '"+str(teacher.school_id)+"' and last_modified_date >=current_date - 7) as SumCount "
#         #print(lastWeekTestCount)
#         lastWeekTestCount = db.session.execute(lastWeekTestCount).first()
#         #print('user type value')
#         #print(session['moduleDet'])
#         query = "select user_type,md.module_name,description, module_url from module_detail md inner join module_access ma on md.module_id = ma.module_id where user_type = '"+str(current_user.user_type)+"'"
#         moduleDetRow = db.session.execute(query).fetchall()
#         return render_template('dashboard.html',form=form,title='Home Page',school_id=teacher.school_id, jobPosts=jobPosts,
#             graphJSON=graphJSON, classSecCheckVal=classSecCheckVal,topicToCoverDetails = topicToCoverDetails, EventDetailRows = EventDetailRows, topStudentsRows = data,teacherCount=teacherCount,studentCount=studentCount,testCount=testCount,lastWeekTestCount=lastWeekTestCount)

# @app.route('/performanceChart',methods=['GET','POST'])
# def performanceChart():
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     query = "Select * from fn_overall_performance_summary('"+str(teacher_id.school_id)+"') where class='All'and section='All' and subject='All'"
    
#     resultSet = db.session.execute(text(query)).fetchall()
    
#     resultArray = []
#     if resultSet:
#         for result in resultSet:
#             Array = {}
#             Array['avg_score'] = str(round(result.avg_score,2))
#             Array['highest_mark'] = str(result.highest_mark)
#             Array['lowest_mark'] = str(result.lowest_mark)
#             Array['no_of_students_above_90'] = str(result.no_of_students_above_90)
#             Array['no_of_students_80_90'] = str(result.no_of_students_80_90)
#             Array['no_of_students_70_80'] = str(result.no_of_students_70_80)
#             Array['no_of_students_50_70'] = str(result.no_of_students_50_70)
#             Array['no_of_students_below_50'] = str(result.no_of_students_below_50)
#             Array['no_of_students_cross_50'] = str(result.no_of_students_cross_50)
#             resultArray.append(Array)
#         return {'result' : resultArray} 
#     else:
#         return jsonify(["NA"])

# @app.route('/performanceBarChart',methods=['GET','POST'])
# def performanceBarChart():
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_v = request.args.get('class_val')
#     section = request.args.get('section')
#     classSection = ClassSection.query.with_entities(ClassSection.class_sec_id).filter_by(class_val=class_v,section=section,school_id=str(teacher_id.school_id)).first()
#     subject = "select distinct subject_id from topic_detail where class_val= '"+str(class_v)+"'"
#     totalStudent = "select count(*) from student_profile where class_sec_id='"+str(classSection.class_sec_id)+"' and school_id='"+str(teacher_id.school_id)+"'"
#     #print(totalStudent)
#     totalStudent = db.session.execute(totalStudent).first()
#     subject_id = db.session.execute(subject).fetchall()
#     performance_array = []
#     for sub in subject_id:
#         pass_count = "select count(*) from student_profile sp where student_id in (select studentid from fn_performance_leaderboard_detail_v1('"+str(teacher_id.school_id)+"') pd where class ='"+str(class_v)+"' and section='"+str(section)+"' and subjectid='"+str(sub.subject_id)+"' and marks>50)"
#         fail_count = "select count(*) from student_profile sp where student_id in (select studentid from fn_performance_leaderboard_detail_v1('"+str(teacher_id.school_id)+"') pd where class ='"+str(class_v)+"' and section='"+str(section)+"' and subjectid='"+str(sub.subject_id)+"' and marks<=50)"
#         #print('pass and fail count:')
#         #print(pass_count)
#         #print(fail_count)
#         passStudents = db.session.execute(pass_count).first()
#         failStudents = db.session.execute(fail_count).first()
#         presentStudents = passStudents[0] + failStudents[0]
#         absentStudents = totalStudent[0] - presentStudents
#         #print(absentStudents)
#         if absentStudents==totalStudent[0]:
#             absentStudents = 0
#         #print((passStudents[0]))
#         #print((failStudents[0]))
#         Array = {}
#         Array['pass_count'] = str(passStudents[0])
#         Array['fail_count'] = str(failStudents[0])
#         Array['absent_students'] = str(absentStudents)
#         subjectName = MessageDetails.query.with_entities(MessageDetails.description).filter_by(msg_id=sub.subject_id).first()
#         Array['description'] = str(subjectName.description)
#         performance_array.append(Array)
#     return {'performance':performance_array}
    
    

@app.route('/disconnectedAccount')
@login_required
def disconnectedAccount():    
    print('Inside disconnected Account')
    userDetailRow=User.query.filter_by(username=current_user.username).first()
    teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()

    #added under the change for practice test module
    boardRows = MessageDetails.query.with_entities(MessageDetails.description).distinct().filter_by(category='Board').all()
    classRows = BoardClassSubject.query.with_entities(BoardClassSubject.class_val).distinct().order_by(BoardClassSubject.class_val.desc()).all()

    if userDetailRow.user_type==72:
        print('Inside Guardian condition')
        return redirect(url_for('guardianDashboard'))
    if teacher==None and userDetailRow.user_type!=161 and userDetailRow.user_type!=134:
        return render_template('disconnectedAccount.html', title='Disconnected Account', disconn = 1, userDetailRow=userDetailRow, boardRows=boardRows,classRows=classRows)
    elif userDetailRow.user_type==161:
        return redirect(url_for('job_post.openJobs'))
    elif userDetailRow.user_type==134 and userDetailRow.access_status==145:
        return redirect(url_for('studentDashboard'))
    else:
        print('Inside else')
        return redirect(url_for('dashboard.index'))




# @app.route('/postJob',methods=['POST','GET'])
# @login_required
# def postJob():
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     schoolCityQuery = "select city from address_detail where address_id =(select address_id from school_profile where school_id ="+str(teacherRow.school_id)+")"
#     schoolCity = db.session.execute(text(schoolCityQuery)).first()
#     form = postJobForm()
    
#     availableCategories=MessageDetails.query.filter_by(category='Job Category').all()
#     availableJobTypes=MessageDetails.query.filter_by(category='Job Type').all()
#     availableStayOptions=MessageDetails.query.filter_by(category='Stay Option').all()
#     availableFoodOptions=MessageDetails.query.filter_by(category='Food Option').all()
#     availableTeachingTermOption=MessageDetails.query.filter_by(category='Teaching Term Option').all()

#     form.category.choices = [(str(i.description),str(i.description)) for i in availableCategories]
#     form.job_type.choices = [(str(i.description),str(i.description)) for i in availableJobTypes]
#     form.stay.choices = [(str(i.description),str(i.description)) for i in availableStayOptions]
#     form.food.choices = [(str(i.description),str(i.description)) for i in availableFoodOptions]
#     form.term.choices = [(str(i.description),str(i.description)) for i in availableTeachingTermOption]


#     if request.method == 'POST' and form.validate():
#         jobData=JobDetail(category=form.category.data,
#             posted_by =teacherRow.teacher_id,school_id=teacherRow.school_id,description=form.description.data,min_pay=form.min_pay.data,max_pay=form.max_pay.data,
#             start_date=form.start_date.data,subject=form.subject.data, 
#             classes= form.classes.data, language= form.language.data,timings= form.timings.data,stay= form.stay.data, 
#             fooding= form.food.data,term= form.term.data,status='Open',num_of_openings=form.num_of_openings.data,city =schoolCity.city,
#             job_type =form.job_type.data,posted_on = datetime.today(),last_modified_date= datetime.today())
#         db.session.add(jobData)
#         db.session.commit()
#         flash('New job posted created!')
#         try:
#             job_posted_email(teacherRow.email,teacherRow.teacher_name,form.category.data)
#         except:
#             pass
#     else:
#         #flash('Please fix the errors to submit')
#         for fieldName, errorMessages in form.errors.items():
#             for err in errorMessages:
#                 print(err)
#     indic='DashBoard'
#     return render_template('postJob.html',indic=indic,title='Post Job',form=form,classSecCheckVal=classSecCheck())


# @app.route('/openJobs')
# def openJobs():
#     page=request.args.get('page',0, type=int)    
#     first_login = request.args.get('first_login','0').strip()
#     jobTermOptions = MessageDetails.query.filter_by(category='Teaching Term Option').all()
#     jobTypeOptions = MessageDetails.query.filter_by(category='Job Type').all()
#     # print('User value in openJobs:'+str(user_type_val))
#     if first_login=='1':
        
#         print('this is the first login section')
#         userRecord = User.query.filter_by(id=current_user.id).first() 
#         userRecord.user_type= '161'
#         db.session.commit()
#         flash('Please complete your profile before applying for jobs')
#         return redirect('edit_profile')
#     else:
#         print('first login not registered')
#         if current_user.is_anonymous:
#             return render_template('openJobs.html',title='Look for Jobs',first_login=first_login,jobTermOptions=jobTermOptions,jobTypeOptions=jobTypeOptions)
#         else:
#             return render_template('openJobs.html',title='Look for Jobs',first_login=first_login,jobTermOptions=jobTermOptions,jobTypeOptions=jobTypeOptions,user_type_val=str(current_user.user_type))


# @app.route('/openJobsFilteredList')
# def openJobsFilteredList():
#     page=request.args.get('page',0, type=int)
#     recordsOnPage = 5
#     offsetVal = page *recordsOnPage
    
#     whereClause = ""
#     qjob_term = request.args.get('job_term') #all /short term / long term
#     qjob_type = request.args.get('job_type') #all /part time/ full time
#     qcity =  request.args.get('city')       # all/ home city
   
#     print("qterm is "+str(qjob_term))
#     print("qtype is "+str(qjob_type))
#     print("qcity is "+str(qcity))

#     whJobTerm=''
#     whJobType=''
#     whCity=''

#     if qjob_term=='All' or qjob_term==None or qjob_term=='':
#         whJobTerm=None
#     else:
#         whJobTerm=" t1.term=\'"+str(qjob_term)+"\'"
#         whereClause = 'where ' + whJobTerm

    
#     if qjob_type=='All' or qjob_type==None or qjob_type=='':
#         whJobType=None
#     else:
#         whJobType=" t1.job_type=\'"+str(qjob_type)+"\'"
#         if whereClause=='':
#             whereClause = 'where '+whJobType
#         else:
#             whereClause =  whereClause + ' and '+whJobType
    
#     if qcity=='All' or qcity==None or qcity=='':
#         whCity=None
#     else:
#         whCity=" t1.city=\'"+ str(qcity)+"\'"
#         if whereClause=='':
#             whereClause = 'where '+whCity
#         else:
#             whereClause = whereClause + ' and '+whCity
    
#     print('this is the where clause' + whereClause)
#     #if whJobTerm!=None and whJobType!=None and whCity!=None:
#     #    whereClause = "where " + whJobTerm + "and "+whJobType + "and "+whCity


#     #teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     openJobsQuery = "select school_picture, school_name, t2.school_id, min_pay, max_pay, t1.city, t1.category, t1.job_type,t1.term, t1.subject,t1.posted_on, t1.job_id "
#     openJobsQuery = openJobsQuery + "from job_detail t1 inner join school_profile t2 on t1.school_id=t2.school_id and t1.status='Open' " + whereClause 
#     openJobsQuery = openJobsQuery + " order by t1.posted_on desc "
#     #openJobsQuery = openJobsQuery +" OFFSET "+str(offsetVal)+" ROWS FETCH FIRST "+str(recordsOnPage)+" ROW ONLY; "
#     #openJobsDataRows = db.session.execute(text(openJobsQuery)).fetchall()    
#     openJobsDataRows = db.session.execute(text(openJobsQuery)).fetchall()
    
#     if len(openJobsDataRows)==0:
#         print('returning 1')
#         return jsonify(['1'])
#     else:
#         next_page=page+1

#         if page!=0:
#             prev_page=page-1
#         else:
#             prev_page=None

#         prev_url=None
#         next_url=None


#         if len(openJobsDataRows)==recordsOnPage:
#             next_url = url_for('openJobsFilteredList', page = next_page,job_term=qjob_term, job_type=qjob_type,city=qcity)
#             prev_url = url_for('openJobsFilteredList', page=prev_page,job_term=qjob_term, job_type=qjob_type,city=qcity)
#         elif len(openJobsDataRows)<recordsOnPage:
#             next_url = None
#             if prev_page!=None:
#                 prev_url = url_for('openJobsFilteredList', page=prev_page,job_term=qjob_term, job_type=qjob_type,city=qcity)
#             else:
#                 prev_url==None
#         else:
#             next_url=None
#             prev_url=None
#         return render_template('_jobList.html',openJobsDataRows=openJobsDataRows,next_url=next_url, prev_url=prev_url)



# @app.route('/jobDetail')

# def jobDetail():
#     job_id = request.args.get('job_id')
#     school_id=request.args.get('school_id')  
#     #teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
#     schoolProfileRow = SchoolProfile.query.filter_by(school_id =school_id).first()
#     addressRow = Address.query.filter_by(address_id = schoolProfileRow.address_id).first()    
#     jobDetailRow = JobDetail.query.filter_by(job_id=job_id).first()
#     if current_user.is_anonymous:
#         print('user Anonymous')
#         jobApplicationRow = ''
#     else:
#         print('user exist')
#         jobApplicationRow = JobApplication.query.filter_by(job_id=job_id, applier_user_id=current_user.id).first()
#     if jobApplicationRow:
#         applied=1
#     else:
#         applied=0
#     if current_user.is_anonymous:
#         return render_template('jobDetail.html', title='Job Detail', 
#             schoolProfileRow=schoolProfileRow,addressRow=addressRow,jobDetailRow=jobDetailRow,applied=applied)
#     else:
#         return render_template('jobDetail.html', title='Job Detail', 
#             schoolProfileRow=schoolProfileRow,addressRow=addressRow,jobDetailRow=jobDetailRow,applied=applied,user_type_val=str(current_user.user_type))
    




# @app.route('/jobDetail')

# def jobDetail():
#     job_id = request.args.get('job_id')
#     school_id = ''
#     userData = User.query.filter_by(id=current_user.id).first()
#     givenSchoolId=request.args.get('school_id')  
#     if givenSchoolId:
#         school_id = givenSchoolId
#     else:
#         school_id = userData.school_id
#     print(school_id)
#     #teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
#     schoolProfileRow = SchoolProfile.query.filter_by(school_id =school_id).first()
#     addressRow = Address.query.filter_by(address_id = schoolProfileRow.address_id).first()    
#     jobDetailRow = JobDetail.query.filter_by(job_id=job_id).first()
#     if current_user.is_anonymous:
#         print('user Anonymous')
#         jobApplicationRow = ''
#     else:
#         print('user exist')
#         jobApplicationRow = JobApplication.query.filter_by(job_id=job_id, applier_user_id=current_user.id).first()
#     if jobApplicationRow:
#         applied=1
#     else:
#         applied=0
#     if current_user.is_anonymous:
#         return render_template('jobDetail.html', title='Job Detail', 
#             schoolProfileRow=schoolProfileRow,addressRow=addressRow,jobDetailRow=jobDetailRow,applied=applied)
#     else:
#         return render_template('jobDetail.html', title='Job Detail', 
#             schoolProfileRow=schoolProfileRow,addressRow=addressRow,jobDetailRow=jobDetailRow,applied=applied,user_type_val=str(current_user.user_type))
    



# @app.route('/sendJobApplication',methods=['POST','GET'])
# @login_required
# def sendJobApplication():
#     print('We are in the right place')    
#     if request.method=='POST':
#         job_id_form = request.form.get('job_id_form')
#         available_from=request.form.get("availableFromID")
#         available_till=request.form.get("availableTillID")
#         if available_from=='':
#             available_from=None
#         if available_till=='':
#             available_till=None
#         school_id=request.form.get("school_id")
#         #teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         jobApplyData=JobApplication(applier_user_id=current_user.id, job_id=job_id_form,
#                 applied_on =datetime.today(),status='Applied',school_id=school_id,available_from=available_from,available_till=available_till,
#                 last_modified_date=date.today())
#         db.session.add(jobApplyData)
#         db.session.commit()
#         flash('Job application submitted!')
#         #try:            
#         jobDetailRow = JobDetail.query.filter_by(job_id=job_id_form).first()
#         teacherRow = TeacherProfile.query.filter_by(teacher_id=jobDetailRow.posted_by).first()
#         new_applicant_for_job(teacherRow.email,teacherRow.teacher_name,current_user.first_name + ' '+current_user.last_name,jobDetailRow.category)
#         #except:
#         #    pass
#         return redirect(url_for('openJobs'))

# @app.route('/appliedJobs')  # this page shows all the job posts that the user has applied to
# @login_required
# def appliedJobs():
#     appliedQuery = "select applied_on, t3.school_id,school_name, category, subject, t2.job_id, "
#     appliedQuery = appliedQuery + "t1.status as application_status, t2.status as job_status "
#     appliedQuery = appliedQuery + "from job_application t1 inner join job_detail t2 on "
#     appliedQuery = appliedQuery + "t1.job_id=t2.job_id inner join school_profile t3 on "
#     appliedQuery = appliedQuery + "t3.school_id=t1.school_id where t1.applier_user_id='"+str(current_user.id)+"'"
#     appliedRows = db.session.execute(text(appliedQuery)).fetchall()
#     return render_template('appliedJobs.html',title='Applied jobs', user_type_val=str(current_user.user_type),appliedRows=appliedRows)


# @app.route('/jobApplications')  # this page shows all the applications received by the job poster for any specifc job post
# @login_required
# def jobApplications():
#     teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     jobidDet=request.args.get('job_id')

#     job_id = ''
#     if jobidDet:
#         job_id = jobidDet
#     else:
#         jobDet = JobApplication.query.filter_by(school_id=teacher.school_id).first()
#         if jobDet:
#             job_id = jobDet.job_id
#         else:
#             job_id = 3
#     #jobApplications = JobApplication.query.filter_by(school_id=teacher.school_id).order_by(JobApplication.applied_on.desc()).all()
#     #pending descision
#     jobAppQuery = "select t1.applied_on, t2.first_name, t2.last_name, t2.username,t1.applier_user_id,t1.job_id, "
#     jobAppQuery=jobAppQuery+"t2.city, t1.available_from, t1.available_till, t2.education, t2.experience from "
#     jobAppQuery=jobAppQuery+"job_application t1 inner join public.user t2 on t1.applier_user_id=t2.id inner join job_detail t3 on "
#     jobAppQuery=jobAppQuery+" t3.job_id=t1.job_id and t3.school_id='"+str(teacher.school_id)+"' and t1.job_id='"+str(job_id)+"' and t1.status='Applied' order by applied_on desc"
#     jobApplications = db.session.execute(text(jobAppQuery)).fetchall()

#     #hired descision
#     jobAppQueryHired = "select t1.applied_on, t2.first_name, t2.last_name, t2.username,t1.applier_user_id, t1.job_id, "
#     jobAppQueryHired=jobAppQueryHired+"t2.city, t1.available_from, t1.available_till, t2.education, t2.experience from "
#     jobAppQueryHired=jobAppQueryHired+"job_application t1 inner join public.user t2 on t1.applier_user_id=t2.id inner join job_detail t3 on "
#     jobAppQueryHired=jobAppQueryHired+" t3.job_id=t1.job_id and t3.school_id='"+str(teacher.school_id)+"' and t1.job_id='"+str(job_id)+"' and t1.status='Hired' order by applied_on desc"
#     jobApplicationsHired = db.session.execute(text(jobAppQueryHired)).fetchall()

#     #shortlist descision
#     jobAppQueryShortlisted = "select t1.applied_on, t2.first_name, t2.last_name, t2.username,t1.applier_user_id, t1.job_id, "
#     jobAppQueryShortlisted=jobAppQueryShortlisted+"t2.city, t1.available_from, t1.available_till, t2.education, t2.experience from "
#     jobAppQueryShortlisted=jobAppQueryShortlisted+"job_application t1 inner join public.user t2 on t1.applier_user_id=t2.id inner join job_detail t3 on "
#     jobAppQueryShortlisted=jobAppQueryShortlisted+" t3.job_id=t1.job_id and t3.school_id='"+str(teacher.school_id)+"' and t1.job_id='"+str(job_id)+"' and t1.status='Shortlisted' order by applied_on desc"
#     jobApplicationsShortlisted = db.session.execute(text(jobAppQueryShortlisted)).fetchall()

#     #rejected descision
#     jobAppQueryRejected = "select t1.applied_on, t2.first_name, t2.last_name, t2.username,t1.applier_user_id, t1.job_id, "
#     jobAppQueryRejected=jobAppQueryRejected+"t2.city, t1.available_from, t1.available_till, t2.education, t2.experience from "
#     jobAppQueryRejected=jobAppQueryRejected+"job_application t1 inner join public.user t2 on t1.applier_user_id=t2.id inner join job_detail t3 on "
#     jobAppQueryRejected=jobAppQueryRejected+" t3.job_id=t1.job_id and t3.school_id='"+str(teacher.school_id)+"' and t1.job_id='"+str(job_id)+"' and t1.status='Rejected' order by applied_on desc"
#     jobApplicationsRejected = db.session.execute(text(jobAppQueryRejected)).fetchall()
    
#     return render_template('jobApplications.html', classSecCheckVal=classSecCheck(),title='Job Applications',jobApplications=jobApplications, jobApplicationsHired=jobApplicationsHired,jobApplicationsShortlisted= jobApplicationsShortlisted, jobApplicationsRejected = jobApplicationsRejected )


# @app.route('/jobPosts')
# @login_required
# def jobPosts():
#     teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     jobPosts = JobDetail.query.filter_by(school_id=teacher.school_id).order_by(JobDetail.posted_on.desc()).all()
#     return render_template('jobPosts.html',jobPosts=jobPosts, classSecCheckVal=classSecCheck(),school_id=teacher.school_id)

# @app.route('/processApplication')
# @login_required
# def processApplication():
#     applier_user_id = request.args.get('applier_user_id')
#     job_id = request.args.get('job_id')
#     process_type = request.args.get('process_type')
#     #try:
#     jobApplicationRow = JobApplication.query.filter_by(applier_user_id=applier_user_id, job_id=job_id).first()
#     jobDetailRow = JobDetail.query.filter_by(job_id=job_id).first()
#     applierRow = User.query.filter_by(id=applier_user_id).first()
#     schoolRow = SchoolProfile.query.filter_by(school_id=jobApplicationRow.school_id).first()

#     print(process_type)
#     if process_type=='shortlist':
#         jobApplicationRow.status= 'Shortlisted'
#         flash('Application Shortlisted')
#         try:
#             application_processed(applierRow.email,applierRow.first_name + ' '+ applierRow.last_name, schoolRow.school_name,jobDetailRow.category, 'Shortlisted')
#         except:
#             pass
#     elif process_type=='reject':
#         jobApplicationRow.status= 'Rejected'
#         flash('Application Rejected')
#         try:
#             application_processed(applierRow.email,applierRow.first_name + ' '+ applierRow.last_name, schoolRow.school_name,jobDetailRow.category, 'Rejected')
#         except:
#             pass
#     elif process_type =='hire':
#         jobApplicationRow.status= 'Hired'
#         flash('Application Hired')
#         try:
#             application_processed(applierRow.email,applierRow.first_name + ' '+ applierRow.last_name, schoolRow.school_name,jobDetailRow.category, 'Hired')
#         except:
#             pass
#     else:
#         flash('Error processing application idk')
#     db.session.commit()
    
#     return redirect(url_for('jobApplications',job_id=job_id))
#     #except:
#     flash('Error processing application')
#     return redirect(url_for('jobApplications',job_id=job_id))


# @app.route('/submitPost', methods=['GET', 'POST'])
# @login_required
# def submitPost():
#     form = PostForm()
#     if form.validate_on_submit():
#         post = Post(body=form.post.data, author=current_user)
#         db.session.add(post)
#         db.session.commit()
#         flash('Your post is now live!')
#         return redirect(url_for('submitPost'))
#     posts = [{
#         'author': {
#             'username': 'John'
#         },
#         'body': 'Beautiful day in Portland!'
#     }, {
#         'author': {
#             'username': 'Susan'
#         },
#         'body': 'The Avengers movie was so cool!'
#     }]
#     return render_template(
#         "submitPost.html", title='Submit Post', form=form, posts=posts)


@app.route('/explore')
@login_required
def explore():
    #page=request.args.get('page',1, type=int)
    #posts = Post.query.order_by(Post.timestamp.desc()).paginate(page,app.config['POSTS_PER_PAGE'],False)
     #next_url = url_for('explore', page=posts.next_num) \
        #if posts.has_next else None
    #prev_url = url_for('explore', page=posts.prev_num) \
        #if posts.has_prev else None
        #return render_template("index.html", title='Explore', posts=posts.items,
         #                 next_url=next_url, prev_url=prev_url)
    posts = [{
        'author': {
            'username': 'John'
        },
        'body': 'Beautiful day in Portland!'
    }, {
        'author': {
            'username': 'Susan'
        },
        'body': 'The Avengers movie was so cool!'
    }]

    return render_template('explore.html', title='Explore', posts=posts)


#new section for liveClass

@app.route('/archiveLiveClass')
@login_required
def archiveLiveClass():
    live_class_id=request.args.get('live_class_id')
    try:
        liveClassVal = LiveClass.query.filter_by(live_class_id=live_class_id,is_archived='N').first()
        liveClassVal.is_archived='Y'
        db.session.commit()
        return jsonify(['0'])
    except:
        return jsonify(['1'])

# start live class section

@app.route('/liveClass', methods=['GET','POST'])
@login_required
def liveClass():    
    form = AddLiveClassForm()
    user_id = request.args.get('user_id')
    student = ''
    studentDetails = ''
    print('inside live class route')
    if user_id:
        student = StudentProfile.query.filter_by(user_id=user_id).first()
        allLiveClassQuery = "(select t1.class_sec_id, t2.class_val, t2.section "
        allLiveClassQuery = allLiveClassQuery + ", t1.subject_id, t3.description as subject, t1.topic_id, t4.topic_name, start_time,end_time, status, teacher_name, "
        allLiveClassQuery = allLiveClassQuery + " conf_link, t1.school_id "
        allLiveClassQuery = allLiveClassQuery + " from live_class t1 "
        allLiveClassQuery = allLiveClassQuery+ " inner join class_section t2 on t1.class_sec_id = t2.class_sec_id "
        allLiveClassQuery= allLiveClassQuery + " inner join message_detail t3 on t1.subject_id = t3.msg_id "
        allLiveClassQuery= allLiveClassQuery + " inner join topic_detail t4 on t1.topic_id = t4.topic_id where t1.school_id= " +str(student.school_id) + " and t1.class_sec_id= "+str(student.class_sec_id) 
        allLiveClassQuery= allLiveClassQuery + " and end_time > now() order by end_time desc)"  
        allLiveClassQuery= allLiveClassQuery + " union "
        allLiveClassQuery = allLiveClassQuery + "(select t1.class_sec_id, t2.class_val, t2.section "
        allLiveClassQuery = allLiveClassQuery + ", t1.subject_id, t3.description as subject, t1.topic_id, t4.topic_name, start_time,end_time, status, teacher_name, "
        allLiveClassQuery = allLiveClassQuery + " conf_link, t1.school_id "
        allLiveClassQuery = allLiveClassQuery + " from live_class t1 "
        allLiveClassQuery = allLiveClassQuery+ " inner join class_section t2 on t1.class_sec_id = t2.class_sec_id "
        allLiveClassQuery= allLiveClassQuery + " inner join message_detail t3 on t1.subject_id = t3.msg_id "
        allLiveClassQuery= allLiveClassQuery + " inner join topic_detail t4 on t1.topic_id = t4.topic_id where t1.is_private= 'N' and t1.school_id<> " +str(student.school_id) + " and t1.class_sec_id= "+str(student.class_sec_id) 
        allLiveClassQuery= allLiveClassQuery + " and end_time > now() order by end_time desc)"
        print('Query for live class:'+str(allLiveClassQuery))
        try:
            allLiveClasses = db.session.execute(text(allLiveClassQuery)).fetchall()
            print('##########Data:'+str(allLiveClasses))
        except:
            allLiveClasses = "" 
        indic='DashBoard'   
        return render_template('liveClass.html',indic=indic,title='Live Classes',allLiveClasses=allLiveClasses,form=form,current_time=datetime.now(),studentDetails=studentDetails)      
    #allLiveClasses = LiveClass.query.filter_by(is_archived='N').order_by(LiveClass.last_modified_date.desc()).all()
    if current_user.user_type==71 or current_user.user_type==135 or current_user.user_type==139:
        teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
        school_id = teacherData.school_id 
        studentDetails=""
    elif current_user.user_type==134:
        #studentData = StudentProfile.query.filter_by(user_id=current_user.id).first()  
        studentDetails = StudentProfile.query.filter_by(user_id=current_user.id).first()       
        school_id = studentDetails.school_id
    else:        
        return redirect(url_for('dashboard.index'))

    print('##########Data:'+str(school_id))
    allLiveClassQuery = "select t1.class_sec_id, t2.class_val, t2.section "
    allLiveClassQuery = allLiveClassQuery + ", t1.subject_id, t3.description as subject, t1.topic_id, t4.topic_name, start_time,end_time, status, teacher_name, "
    allLiveClassQuery = allLiveClassQuery + " conf_link, t1.school_id "
    allLiveClassQuery = allLiveClassQuery + " from live_class t1 "
    allLiveClassQuery = allLiveClassQuery+ " inner join class_section t2 on t1.class_sec_id = t2.class_sec_id "
    allLiveClassQuery= allLiveClassQuery + " inner join message_detail t3 on t1.subject_id = t3.msg_id "
    allLiveClassQuery= allLiveClassQuery + " inner join topic_detail t4 on t1.topic_id = t4.topic_id where t1.school_id= " +str(school_id) 
    allLiveClassQuery= allLiveClassQuery + " and end_time > now() order by end_time desc"
    allLiveClassQuery= allLiveClassQuery + " union "
    allLiveClassQuery = "select t1.class_sec_id, t2.class_val, t2.section "
    allLiveClassQuery = allLiveClassQuery + ", t1.subject_id, t3.description as subject, t1.topic_id, t4.topic_name, start_time,end_time, status, teacher_name, "
    allLiveClassQuery = allLiveClassQuery + " conf_link, t1.school_id "
    allLiveClassQuery = allLiveClassQuery + " from live_class t1 "
    allLiveClassQuery = allLiveClassQuery+ " inner join class_section t2 on t1.class_sec_id = t2.class_sec_id "
    allLiveClassQuery= allLiveClassQuery + " inner join message_detail t3 on t1.subject_id = t3.msg_id "
    allLiveClassQuery= allLiveClassQuery + " inner join topic_detail t4 on t1.topic_id = t4.topic_id where t1.school_id<> " +str(school_id)+" and is_private='N'" 
    allLiveClassQuery= allLiveClassQuery + " and end_time > now() order by end_time desc"
    
    try:
        allLiveClasses = db.session.execute(allLiveClassQuery).fetchall()
        print('##########Data:'+str(allLiveClasses))
    except:
        allLiveClasses = ""

    #if request.method == 'POST':
    #    schoolNameRow = SchoolProfile.query.filter_by(school_id=current_user.school_id).first()
    #    liveClassData=LiveClass(class_val = form.class_val.data,subject = form.subject.data, book_chapter=form.book_chapter.data, 
    #        start_time = form.start_time.data, end_time = form.end_time.data, status = "Active", teacher_id=current_user.id, 
    #        teacher_name = str(current_user.first_name)+' '+str(current_user.last_name), class_link=form.class_link.data,phone_number = form.phone_number.data, school_id = current_user.school_id,
    #        school_name =schoolNameRow.name ,is_archived = 'N',last_modified_date = dt.datetime.now())        
    #    db.session.add(liveClassData)
    #    db.session.commit()     
    #    #adding records to topic tracker while registering school                         
    #    flash('New class listed successfully!') 
    indic='DashBoard'              
    return render_template('liveClass.html',indic=indic,title='Live Classes',allLiveClasses=allLiveClasses,form=form,user_type_val=str(current_user.user_type),current_time=datetime.now(),studentDetails=studentDetails)    

#end of live class section


#####New section for open class modules
@app.route('/updateSearchIndex/<task>')
@login_required
def updateSearchIndex(task, fromPage="default"):
    #queryTest = "select *from topic_detail where topic_name is not null fetch first 100 rows only"
    indexQuery = "select cd.course_id as \"objectID\", cd.course_id, cd.course_name, cd.description,average_rating , cd.image_url, "
    indexQuery = indexQuery + " tp.teacher_name,ideal_for, string_agg(topic_name::text, ', ') as topics, text(cd.last_modified_date) as last_modified_date"
    indexQuery = indexQuery + " from course_detail cd "
    indexQuery = indexQuery + " inner join course_topics ct"
    indexQuery = indexQuery + " on ct.course_id =cd.course_id "
    indexQuery = indexQuery + " inner join topic_detail td"
    indexQuery = indexQuery + " on td.topic_id =ct.topic_id "
    indexQuery = indexQuery + " inner join teacher_profile tp "
    indexQuery = indexQuery + " on tp.teacher_id =cd.teacher_id "
    indexQuery = indexQuery + " group by "
    indexQuery = indexQuery + " cd.course_id, cd.course_name, cd.description , cd.image_url, tp.teacher_name "
    indexQuery = indexQuery + " having course_status=251"
    queryResult = db.session.execute(indexQuery).fetchall()
    queryKeys = db.session.execute(indexQuery).keys()
    items = [dict((queryKeys[i], value) \
               for i, value in enumerate(row)) for row in queryResult]
    #print(json.dumps(items))
    if task == "view":
        return json.dumps(items, indent=3)
    elif task=="send":        
        try:
            #index = client.init_index('staging_COURSE')
            index = client.init_index('prd_course')
            #batch = json.load(open('contacts.json'))            
            index.set_settings({"customRanking": ["desc(last_modified_date)"]})            	
            index.set_settings({"searchableAttributes": ["topic_name", "course_name", "teacher_name"]})
            index.save_objects(items, {'autoGenerateObjectIDIfNotExist': True})
            if fromPage=="default":
                return json.dumps({"The following data was sent to algolia index successfully":items})
            else:
                print("############# Course Indexed")
                return True
        except:
            print('########## Throwing exception')
            return "Error Sending index data to algolia"

    elif task=="deleteAll":
        filter_val = request.args.get('filter_val')
        
        #index = client.init_index('staging_COURSE')
        index = client.init_index('prd_course')         
        index.delete_by({'filters': 'board_id:1001'})
        #batch = json.load(open('contacts.json'))            
        #index.set_settings({"customRanking": ["asc(last_modified_date)"]})            	
        #index.set_settings({"searchableAttributes": ["topic_name", "course_name", "teacher_name"]})
        #index.save_objects(items, {'autoGenerateObjectIDIfNotExist': True})
        return json.dumps({"All indexes have been deleted from algolia under: ": filter_val})
        #except:
        #    return "Error Sending index data to algolia"



@app.route('/courseHome')
# def courseHome():    
#     if ("school.alllearn" in str(request.url)):
#         print('#######this is the request url: '+ str(request.url))
#         return redirect(url_for('index')) 
#     #print(str(current_user.is_anonymous))
#     upcomingClassData = ""
    
#     #if current_user.is_anonymous==False:
#         #upcomingClassQuery = "select * from vw_course_reminder_everyday where email=" + str(current_user.email)
#         #upcomingClassData = db.session.execute(upcomingClassQuery).fetchall()
#     enrolledCourses = "select ce.COURSE_ID, MAX(ce.LAST_MODIFIED_DATE), cd.course_name, cd.average_rating , cd.description ,cd.image_url, cd.is_archived,cd.course_status, tp.teacher_name,tp.teacher_id from course_enrollment ce "
#     enrolledCourses = enrolledCourses + "inner join course_detail cd on cd.course_id =ce.course_id "
#     enrolledCourses = enrolledCourses + "inner join teacher_profile tp on tp.teacher_id =cd.teacher_id "
#     enrolledCourses = enrolledCourses + "group by ce.course_id,cd.course_name,cd.average_rating , cd.description , cd.image_url,cd.course_status, cd.is_archived, cd.teacher_id, tp.teacher_name, tp.teacher_id "
#     enrolledCourses = enrolledCourses + "having cd.course_status =276 and cd.is_archived ='N' order by max(ce.last_modified_date ) desc limit 8"
#     enrolledCourses = db.session.execute(text(enrolledCourses)).fetchall()
#     recentlyAccessed = "select cd.COURSE_ID, MAX(cd.LAST_MODIFIED_DATE), cd.course_name, cd.average_rating , cd.description ,cd.image_url, cd.is_archived,cd.course_status, tp.teacher_name,cd.teacher_id from course_detail cd "
#     recentlyAccessed = recentlyAccessed + "inner join teacher_profile tp on tp.teacher_id =cd.teacher_id "
#     recentlyAccessed = recentlyAccessed + "group by cd.course_id,cd.course_name,cd.average_rating , cd.description , cd.image_url,cd.course_status, cd.is_archived, cd.teacher_id, tp.teacher_name having cd.course_status =276 and cd.is_archived ='N' order by max(cd.last_modified_date ) desc limit 8"
#     recentlyAccessed = db.session.execute(text(recentlyAccessed)).fetchall() 

#     for rate in enrolledCourses:
#         print('Rating:'+str(rate.average_rating))
#         if rate.average_rating:
#             print('rate:'+str(rate.average_rating))

#     idealFor = CourseDetail.query.distinct(CourseDetail.ideal_for).all()
#     idealList = []
#     for ideal in idealFor:
#         print('ideal for:'+str(ideal.ideal_for))
#         if ideal.ideal_for:
#             data = ideal.ideal_for.split(',')
#             for d in data:
#                 if d not in idealList:
#                     idealList.append(d)
#     print('List:'+str(idealList))
#     indic='DashBoard'
#     return render_template('courseHome.html',indic=indic,idealList=idealList,recentlyAccessed=recentlyAccessed,enrolledCourses=enrolledCourses,home=1, upcomingClassData=upcomingClassData)

@app.route('/openLiveClass')
def openLiveClass():
    print('inside openlive class')
    #live_class_id = request.args.get('live_class_id')
    topic_id = request.args.get('topic_id')
    batch_id = request.args.get('batch_id')
    course_id = request.args.get('course_id')
    print('topicID:'+str(topic_id))
    print('batch_id'+str(batch_id))
    print('course_id:'+str(course_id))

    videoExist = CourseTopics.query.filter_by(course_id=course_id,topic_id=topic_id).first()
    #topicName = Topic.query.filter_by(topic_id=topic_id).first()
    #courseId = CourseTopics.query.filter_by(topic_id=topic_id).first()
    #courseName = CourseDetail.query.filter_by(course_id=courseId.course_id).first()
    topicDataQuery = "select td.topic_id,ct.video_class_url, cd.course_id,ct.test_id,cd.course_name, td.topic_name, tp.teacher_id, tp.teacher_name, tp.user_id,tp.room_id "
    topicDataQuery = topicDataQuery + " from topic_detail td inner join course_topics ct on  "
    topicDataQuery = topicDataQuery + " ct.topic_id  = td.topic_id inner join course_detail cd on  "
    topicDataQuery = topicDataQuery + " cd.course_id  = ct.course_id  inner join teacher_profile tp on  "
    topicDataQuery = topicDataQuery + " tp.teacher_id  = cd.teacher_id where td.topic_id  =" + str(topic_id)
    topicDataQuery = topicDataQuery + " and cd.course_id=" + str(course_id)
    topicData = db.session.execute(topicDataQuery).first()

    if topicData== None:
        flash('No relevant course and topic found!')
        return redirect(url_for('course.courseDetail',course_id=course_id))
    
    #batchTestData = BatchTest.query.filter_by(batch_id=batch_id, is_archived='N', is_curren)

    notesList = TopicNotes.query.filter_by(topic_id=topic_id).all()
    comments = "select u.username,c.comment,c.last_modified_date from comments c inner join public.user u on u.id=c.user_id where c.topic_id = '"+str(topic_id)+"' and comment<>' ' and comment is not null  "
    comments = db.session.execute(text(comments)).fetchall()
    lenComm = 0
    if comments:
        lenComm = len(comments)
    #print('comment length:'+str(lenComm))
    #print('courseID:'+str(courseId.course_id))
    rating = CourseDetail.query.filter_by(course_id=topicData.course_id,is_archived='N').first()
    #if rating:
    #    print('Star rating:'+str(rating.average_rating))
    listTopics = "select td.topic_name,td.topic_id from topic_detail td inner join course_topics ct on td.topic_id=ct.topic_id where ct.course_id='"+str(topicData.course_id)+"' and ct.is_archived='N' and ct.topic_id <> '"+str(topic_id)+"' "
    listTopics = db.session.execute(text(listTopics)).fetchall()

    #updating table to say ongoing class
    enrolled=''
    ongoing='N'
    if batch_id!="" and batch_id!=None:
        courseBatchData = CourseBatch.query.filter_by(batch_id = batch_id,is_archived='N').first()
        courseBatchData.is_ongoing = 'Y'
        courseBatchData.ongoing_topic_id = topicData.topic_id
        courseBatchData.last_modified_date = datetime.today()
        db.session.commit()        
    else:
        #checking if a student is seeing the page and then seeing the batch id they're allocated to
        if current_user.is_anonymous==False:
            courseEnrollmentData = CourseEnrollment.query.filter_by(is_archived='N',course_id=topicData.course_id, student_user_id=current_user.id).first()        
            if courseEnrollmentData!=None and courseEnrollmentData!="":
                enrolled='Y'
                batch_id=courseEnrollmentData.batch_id
                courseBatchStudData = CourseBatch.query.filter_by(batch_id = batch_id,is_archived='N',is_ongoing='Y').first()
                if courseBatchStudData!=None and courseBatchStudData!="":
                    ongoing='Y'
        else:
            enrolled='N'
            batch_id=""

    pageTitle = str(topicData.topic_name)+' '+str(topicData.course_name)
    print('$#$$$$$$$$$$$$$'+str(batch_id))
    return render_template('openLiveClass.html',listTopics=listTopics,rating=rating,topicData=topicData
        ,lenComm=lenComm,ongoing=ongoing,title=pageTitle,meta_val=pageTitle,classVideo=videoExist.video_class_url,comments=comments,notesList=notesList,batch_id=batch_id,enrolled=enrolled)

# @app.route('/courseDetail')
# def courseDetail():
#     live_class_id = request.args.get('live_class_id')
#     course_id = request.args.get('course_id')
#     courseDet = CourseDetail.query.filter_by(course_id=course_id).first()
#     teacher = TeacherProfile.query.filter_by(teacher_id=courseDet.teacher_id).first()
#     teacherUser = User.query.filter_by(id=teacher.user_id).first()

#     upcomingDate = "SELECT * FROM course_batch WHERE batch_start_date > NOW() and course_id='"+str(course_id)+"' ORDER BY batch_start_date LIMIT 1"
#     upcomingDate = db.session.execute(text(upcomingDate)).first()
#     checkEnrollment = ''
#     if upcomingDate and current_user.is_authenticated :
#         checkEnrollment = CourseEnrollment.query.filter_by(is_archived='N',course_id=course_id,student_user_id=current_user.id).first()

#     idealFor = courseDet.ideal_for.split(",")
    
#     levelId = courseDet.difficulty_level
#     level = MessageDetails.query.filter_by(msg_id=levelId,category='Difficulty Level').first()
#     #rating = CourseDetail.query.filter_by(course_id=course_id,is_archived='N').first()
#     #if rating:
#     #    print('Star rating:'+str(rating.average_rating))
#     comments = "select u.username,cr.comment,cr.last_modified_date from course_review cr inner join public.user u on u.id=cr.user_id where cr.course_id = '"+str(course_id)+ "' and cr.comment <> ' '"
#     #print(comments)
#     comments = db.session.execute(text(comments)).fetchall()
#     lenComment = len(comments)
#     #print(comments)
#     otherCourses = "select *from course_detail cd where cd.course_id <> '"+str(course_id)+"' and cd.teacher_id='"+str(teacher.teacher_id)+"' and cd.course_status =276"
#     otherCourses = db.session.execute(text(otherCourses)).fetchall()


#     pageTitle = courseDet.course_name
#     return render_template('courseDetail.html',
#         lenComment=lenComment,comments=comments,otherCourses=otherCourses,level=level,
#         idealFor=idealFor,upcomingDate=upcomingDate,
#         courseDet=courseDet,meta_val=pageTitle,title=pageTitle,teacherUser=teacherUser,checkEnrollment=checkEnrollment,course_id=course_id,teacher=teacher)


# @app.route('/courseTopicDetail')
# def courseTopicDetail():
#     course_id = request.args.get('course_id')
#     topicDet = "select case when count(tq.question_id) >0 then count(tq.question_id )"
#     topicDet = topicDet + " else 0 end as no_of_questions , topic_name,ct.video_class_url, ct.topic_id, course_id from course_topics ct "
#     topicDet = topicDet + " inner join topic_detail td on td.topic_id =ct.topic_id and ct.course_id =" + str(course_id)
#     topicDet = topicDet + " left join test_questions tq on tq.test_id =ct.test_id "    
#     topicDet = topicDet + " where ct.is_archived ='N' group by  topic_name, ct.topic_id, course_id,ct.video_class_url"
#     topicDet = topicDet + " order by topic_id asc "
#     #print(topicDet)
#     topicDet = db.session.execute(text(topicDet)).fetchall()
#     return render_template('_courseTopicDetail.html', topicDet=topicDet,course_id=course_id)

@app.route('/fetchClassVideo',methods=['GET','POST'])
def fetchClassVideo():
    topic_id = request.args.get('topic_id')
    class_video = CourseTopics.query.filter_by(topic_id=topic_id).first()
    return jsonify(class_video.video_class_url)


@app.route('/courseBatchDetail')
def courseBatchDetail():
    course_id = request.args.get('course_id')
    teacher_user_id = request.args.get('teacher_user_id')
    #teacher = TeacherProfile.query.filter_by(teacher_id=courseDet.teacher_id).first()
    teacherUser = User.query.filter_by(id=teacher_user_id).first()
    courseBatchData = " select cb.batch_id ,cb.batch_end_date ,cb.batch_start_date ,cb.days_of_week ,cb.student_limit, cb.students_enrolled , "
    courseBatchData = courseBatchData + " cb.course_batch_fee,ce.student_user_id from course_batch cb left join course_enrollment ce on ce.batch_id = cb.batch_id "
    courseBatchData = courseBatchData + " and cb.course_id=ce.course_id "
    if current_user.is_anonymous==False:
        courseBatchData = courseBatchData + " and ce.student_user_id="+str(current_user.id) 
    courseBatchData = courseBatchData + " where cb.course_id = '"+str(course_id)+"' and cb.is_archived='N' "    
    courseBatchData = courseBatchData + " and cb.batch_end_date > NOW() "
    courseBatchData = courseBatchData + " order by cb.batch_start_date desc"
    print('Query:'+str(courseBatchData))
    courseBatchData = db.session.execute(text(courseBatchData)).fetchall()
    return render_template('_courseBatchDetail.html', courseBatchData=courseBatchData,teacherUser=teacherUser
        ,course_id=course_id)

@app.route('/studTakeQuiz', methods=['GET','POST'])
def studTakeQuiz():
    #course_id={{topicData.course_id}}&batch_id={{batch_id}}&topic_id={{topicData.topic_id}}",
    course_id = request.args.get('course_id')
    batch_id = request.args.get('batch_id')
    topic_id = request.args.get('topic_id')
    batchTestData = BatchTest.query.filter_by(batch_id=batch_id, topic_id=topic_id, is_current='Y').first()
    if batchTestData:
        return jsonify([batchTestData.resp_session_id])
    else:
        return jsonify(['1'])


@app.route('/addComments',methods=['GET','POST'])
def addComments():
    topicId = request.form.get('topicId')
    remark = request.form.get('remark')
    if remark:
        addComment = Comments(comment=remark,topic_id=topicId,user_id=current_user.id,is_archived='N',last_modified_date=datetime.now())
        db.session.add(addComment)
        db.session.commit()
    fetchComments = "select u.username,c.comment,c.last_modified_date from comments c inner join public.user u on u.id=c.user_id where c.topic_id = '"+str(topicId)+"'"
    fetchComments = db.session.execute(text(fetchComments)).fetchall()
    lenComm = len(fetchComments)
    print('length:'+str(lenComm))
    commentList = []
    for comment in fetchComments:
        commentList.append(str(comment.username)+":"+str(comment.comment)+":"+str(comment.last_modified_date.strftime('%d %B %Y'))+":"+str(lenComm))
    return jsonify(commentList)

@app.route('/addTopicReview',methods=['GET','POST'])
def addTopicReview():
    course_id = request.args.get('course_id')
    revComment = request.args.get('revComment')
    print('courseId:'+str(course_id))
    reviewRate = CourseReview.query.filter_by(course_id=course_id,is_archived='N',user_id=current_user.id).first()
    reviewRate.comment = revComment
    db.session.commit()
    return jsonify("1")

@app.route('/addReviewComment',methods=['GET','POST'])
def addReviewComment():
    course_id = request.args.get('course_id')
    revComment = request.args.get('revComment')
    starRating = request.args.get('starRating')
    
    print('courseId:'+str(course_id))
    reviewRate = CourseReview.query.filter_by(course_id=str(course_id),is_archived='N',user_id=current_user.id).first()
    if reviewRate==None or revComment=="":
        reviewDataAdd = CourseReview(course_id=str(course_id), star_rating = starRating, comment = str(revComment),
            is_archived = 'N', last_modified_date=datetime.today(), user_id=current_user.id)
        db.session.add(reviewDataAdd )
    else:
        reviewRate.comment = revComment
    db.session.commit()
    fetchReview = "select u.username,cr.comment,cr.last_modified_date from course_review cr inner join public.user u on u.id=cr.user_id where cr.course_id = '"+str(course_id)+"' and cr.comment <> ' '"
    fetchReview = db.session.execute(text(fetchReview)).fetchall()
    lenComm = len(fetchReview)
    print('length:'+str(lenComm))
    reviewList = []
    for review in fetchReview:
        reviewList.append(str(review.username)+":"+str(review.comment)+":"+str(review.last_modified_date.strftime('%d %B %Y'))+":"+str(lenComm))
    return jsonify(reviewList)

@app.route('/addRatingReview',methods=['GET','POST'])
def addRatingReview():
    course_id = request.args.get('course_id')
    rating = request.args.get('rating')
    dataExist = CourseReview.query.filter_by(course_id=course_id,is_archived='N',user_id=current_user.id).first()
    if dataExist:
        dataExist.star_rating = rating
        dataExist.comment = ' '
        db.session.commit()
        courseRevDet = CourseReview.query.filter_by(course_id=course_id,is_archived='N').all()
        aveRat = 0
        for rate in courseRevDet:
            aveRat = aveRat + rate.star_rating
        aveRat = aveRat / len(courseRevDet)
        print('course_id:'+str(course_id))
        print('Average rating:'+str(aveRat))
        courseData = CourseDetail.query.filter_by(course_id=course_id).first()
        courseData.average_rating = aveRat
        db.session.commit()
        return jsonify("1")
    print('Rating:'+str(rating))
    courseRev = CourseReview(course_id=course_id,star_rating=rating,comment=' ',is_archived='N',last_modified_date=datetime.now(),user_id=current_user.id)
    db.session.add(courseRev)
    db.session.commit()
    courseRevDet = CourseReview.query.filter_by(course_id=course_id,is_archived='N').all()
    aveRat = 0
    for rate in courseRevDet:
        aveRat = aveRat + rate.star_rating
    aveRat = aveRat / len(courseRevDet)
    print('Average rating:'+str(aveRat))
    courseData = CourseDetail.query.filter_by(course_id=course_id).first()
    courseData.average_rating = aveRat
    db.session.commit()
    return jsonify("1")

@app.route('/addReview',methods=['GET','POST'])
def addReview():
    course_id = request.args.get('course_id')
    rating = request.args.get('rating')
    dataExist = CourseReview.query.filter_by(course_id=course_id,is_archived='N',user_id=current_user.id).first()
    if dataExist:
        dataExist.star_rating = rating
        dataExist.comment = ' '
        db.session.commit()
        courseRevDet = CourseReview.query.filter_by(course_id=course_id,is_archived='N').all()
        aveRat = 0
        for rate in courseRevDet:
            aveRat = aveRat + rate.star_rating
        aveRat = aveRat / len(courseRevDet)
        print('course_id:'+str(course_id))
        print('Average rating:'+str(aveRat))
        courseData = CourseDetail.query.filter_by(course_id=course_id).first()
        courseData.average_rating = aveRat
        db.session.commit()
        fetchReview = "select u.username,cr.comment,cr.last_modified_date from course_review cr inner join public.user u on u.id=cr.user_id where cr.course_id = '"+str(course_id)+"' and cr.comment <> ' '"
        fetchReview = db.session.execute(text(fetchReview)).fetchall()
        lenComm = len(fetchReview)
        print('length:'+str(lenComm))
        reviewList = []
        for review in fetchReview:
            reviewList.append(str(review.username)+":"+str(review.comment)+":"+str(review.last_modified_date.strftime('%d %B %Y'))+":"+str(lenComm))
        return jsonify(reviewList)
    print('Rating:'+str(rating))
    courseRev = CourseReview(course_id=course_id,star_rating=rating,comment=' ',is_archived='N',last_modified_date=datetime.now(),user_id=current_user.id)
    db.session.add(courseRev)
    db.session.commit()
    courseRevDet = CourseReview.query.filter_by(course_id=course_id,is_archived='N').all()
    aveRat = 0
    for rate in courseRevDet:
        aveRat = aveRat + rate.star_rating
    aveRat = aveRat / len(courseRevDet)
    print('Average rating:'+str(aveRat))
    courseData = CourseDetail.query.filter_by(course_id=course_id).first()
    courseData.average_rating = aveRat
    db.session.commit()
    fetchReview = "select u.username,cr.comment,cr.last_modified_date from course_review cr inner join public.user u on u.id=cr.user_id where cr.course_id = '"+str(course_id)+"' and cr.comment <> ' '"
    fetchReview = db.session.execute(text(fetchReview)).fetchall()
    lenComm = len(fetchReview)
    print('length:'+str(lenComm))
    reviewList = []
    for review in fetchReview:
        reviewList.append(str(review.username)+":"+str(review.comment)+":"+str(review.last_modified_date.strftime('%d %B %Y'))+":"+str(lenComm))
    return jsonify(reviewList)

@app.route('/batchTopicList',methods=['GET','POST'])
def batchTopicList():
    batch_id = request.args.get('batch_id')
    courseId = CourseBatch.query.filter_by(batch_id=batch_id).first()
    topicIds = CourseTopics.query.filter_by(course_id=courseId.course_id,is_archived='N').all()
    topics = []
    for topicId in topicIds:
        topicName = Topic.query.filter_by(topic_id=topicId.topic_id).first()
        topics.append(str(topicName.topic_name)+':'+str(topicName.topic_id))
    if topics:
        return jsonify(topics)
    else:
        return ""

@app.route('/tutorDashboard',methods=['GET','POST'])
def tutorDashboard():
    tutor_id = request.args.get('tutor_id')
    user_id = request.args.get('user_id')
    if user_id:
        TeacherId = TeacherProfile.query.filter_by(user_id=user_id).first()
        
    else:
        TeacherId = TeacherProfile.query.filter_by(teacher_id=tutor_id).first()
    
    if TeacherId==None:
        return redirect(url_for('course.courseHome'))
    else:
        tutor_id=TeacherId.teacher_id
    user = User.query.filter_by(id=TeacherId.user_id).first()
    # courseDet = CourseDetail.query.filter_by(teacher_id=TeacherId.teacher_id).all()
    courseDet = "select count(*) as no_of_topic,cd.course_name,md.description as desc,cd.description,cd.image_url,cd.course_id from course_detail cd left join course_topics ct on ct.course_id=cd.course_id inner join message_detail md on md.msg_id = cd.course_status where cd.teacher_id='"+str(tutor_id)+"' and cd.is_archived='N' group by course_name,md.description,cd.description,cd.image_url,cd.course_id order by cd.course_id desc"
    
    courseDet = db.session.execute(text(courseDet)).fetchall()
    print(courseDet)
    feeType = MessageDetails.query.filter_by(category='Fee Type').all()
    return render_template('tutorDashboard.html',user=user,tutor_id=tutor_id,feeType=feeType,courseDet=courseDet,teacher_name = user.first_name+' '+user.last_name,profile_pic = user.user_avatar,email = user.email,students_taught=TeacherId.students_taught,courses_created=TeacherId.courses_created)

@app.route('/fetchBatch',methods=['GET','POST'])
def fetchBatch():
    print('inside fetch Batch')
    courseId = request.args.get('courseId')
    courseName = CourseDetail.query.filter_by(course_id=courseId).first()
    fetchBatch = "select cb.batch_start_date,cb.batch_end_date,cb.batch_start_time,cb.batch_end_time,cb.days_of_week,cb.student_limit,cb.students_enrolled,cb.course_batch_fee,cb.total_fee_received,cb.fee_type,md.description "
    fetchBatch = fetchBatch + "from course_batch cb inner join message_detail md on cb.fee_type=md.msg_id where cb.course_id='"+str(courseId)+"' and cb.is_archived='N' order by batch_start_date desc"
    print('Query:'+str(fetchBatch))
    fetchBatch = db.session.execute(text(fetchBatch)).fetchall()
    return render_template('_batchDetail.html',fetchBatch=fetchBatch,courseName=courseName)

@app.route('/createBatch',methods=['GET','POST'])
def createBatch():
    startDate = request.form.get('startDate')
    EndDate = request.form.get('EndDate')
    startTime = request.args.get('startTime')
    endTime = request.args.get('endTime')
    days = request.form.getlist('Days')
    studentLimit = request.form.get('studentLimit')
    batchFee = request.form.get('batchFee')
    coId = request.form.get('coId')
    print('New courseId:'+str(coId))
    if batchFee==None:
        batchFee=0
    enrolledStudents = request.form.get('enrolledStudents')
    feeReceived = request.form.get('feeReceived')
    courseId = request.args.get('courseId')
    print('previous courseID:'+str(courseId))
    selectType = request.form.get('selectType')
    print('startDate:'+str(startDate))
    print('selectType:'+str(selectType))
    print('EndDate:'+str(EndDate))
    print('startTime:'+str(startTime))
    print('endTime:'+str(endTime))
    print('days:'+str(days))
    print('studentLimit:'+str(studentLimit))
    print('batchFee:'+str(batchFee))
    print('enrolledStudents:'+str(enrolledStudents))
    print('feeReceived:'+str(feeReceived))
    print('courseId:'+str(courseId))
    dayString = ''
    i=1
    for day in days:
        print('Day:'+str(day))
        dayS =  str(day)
        if i==len(days):
            dayString = dayString + dayS
        else:
            dayString = dayString + dayS + ','
        i=i+1
    print('StringDays:'+str(dayString))
    if startDate and EndDate and startTime and endTime and days and studentLimit and enrolledStudents and selectType:
        createBatch = CourseBatch(course_id=courseId,batch_start_date=startDate,
        batch_end_date=EndDate,batch_start_time=startTime,batch_end_time=endTime,
        days_of_week=dayString,student_limit=studentLimit,course_batch_fee=float(batchFee),
        students_enrolled=enrolledStudents,total_fee_received=0,fee_type=selectType,
        is_archived='N',last_modified_date=datetime.now())
        db.session.add(createBatch)
        db.session.commit()
        return jsonify("1")
    else:
        return ""

# @app.route('/teacherRegistration')
# def teacherRegistration():
#     print('inside teacher Registration')
#     School = "select *from school_profile sp where school_name not like '%_school' order by school_id desc"
#     School = db.session.execute(text(School)).fetchall()
#     NewSchool = "select *from school_profile order by school_id desc"
#     NewSchool = db.session.execute(text(NewSchool)).fetchall()
#     reviewStatus = "select *from teacher_profile where user_id='"+str(current_user.id)+"' "
#     reviewStatus = db.session.execute(text(reviewStatus)).first()
#     teacher_id = request.args.get('teacher_id')
#     print(teacher_id)
#     teacherDetail = ''
#     bankDetail = ''
#     if teacher_id:
#         teacher = TeacherProfile.query.filter_by(teacher_id = teacher_id).first()
#         #for school in School:
#         #    print('School name:'+str(school.school_name))
#         teacherDetail = User.query.filter_by(id = teacher.user_id).first()
#         school_id = SchoolProfile.query.filter_by(school_id=current_user.school_id).first() 
#         email = str(current_user.email).lower().replace(' ','_')
#         vendorId = str(email)+str('_school_')+str(teacher.school_id)+str('_1')
#         print(vendorId)
#         bankDetail = BankDetail.query.filter_by(vendor_id=vendorId).first()
#         print('details')
#         print(bankDetail)
#         print(teacherDetail.about_me)
#         if reviewStatus:
#             return render_template('teacherRegistration.html',School=School,NewSchool=NewSchool,reviewStatus=reviewStatus.review_status,bankDetail=bankDetail,teacherDetail=teacherDetail,teacher_id=teacher_id)
#         else:
#             return render_template('teacherRegistration.html',School=School,NewSchool=NewSchool,bankDetail=bankDetail,teacherDetail=teacherDetail,teacher_id=teacher_id)
    
#     if reviewStatus:
#         print('if not registered as teacher')
#         return render_template('teacherRegistration.html',School=School,NewSchool=NewSchool,reviewStatus=reviewStatus.review_status,bankDetail=bankDetail,teacherDetail=teacherDetail,teacher_id=teacher_id)
#     else:
#         print('if not registered as teacher')
#         return render_template('teacherRegistration.html',School=School,NewSchool=NewSchool,bankDetail=bankDetail,teacherDetail=teacherDetail,teacher_id=teacher_id)
# @app.route('/teacherRegForm',methods=['GET','POST'])
# def teacherRegForm():
#     bankName =request.form.get('bankName')
#     accountHolderName = request.form.get('accountHoldername')
#     accountNo = request.form.get('accountNumber')
#     IfscCode = request.form.get('ifscCode')
#     selectSchool = request.form.get('selectSchool')
#     selectedSchool = request.form.get('NewSchool')
    # if selectSchool==None:
    # print('select school id')
    # print(selectSchool)
    # if selectSchool=='None':
    #     selectSchool = selectedSchool
    # user_avatar = request.form.get('imageUrl')
    # about_me = request.form.get('about_me')
    # schoolName = str(current_user.email)+"_school"
    # current_user.about_me = about_me
    # teacher_id = request.args.get('teacher_id')
    # print('Teacher_id:'+str(teacher_id))
    # if teacher_id=='None':
    #     teacher_id = ''
    #     print(teacher_id)
    # if teacher_id:
    #     user_id = TeacherProfile.query.filter_by(teacher_id=teacher_id).first()
    #     schoolIds = user_id.school_id
    #     user = User.query.filter_by(id=user_id.user_id).first()
    #     user.user_avatar = user_avatar
    #     user.about_me = about_me
    #     db.session.commit()
    #     school = ''
    #     if selectSchool:
    #         school = SchoolProfile.query.filter_by(school_id=selectSchool).first()
    #     else:
    #         schoolName = str(current_user.id)+str('_school')
    #         board  = MessageDetails.query.filter_by(category='Board',description='Other').first()
    #         school = SchoolProfile(school_name = schoolName,registered_date=datetime.now(),last_modified_date=datetime.now(),board_id=board.msg_id,school_admin=teacher_id,school_type='individual')
    #         db.session.add(school)
    #         db.session.commit()
    #         school = SchoolProfile.query.filter_by(school_name = schoolName).first()
    #     user_id.school_id = school.school_id
    #     user_id.review_status = '273'
    #     db.session.commit()
    #     print('previous school id:'+str(schoolIds))
        
    #     vendor = str(current_user.email).lower().replace(' ','_')
    #     vendorId = str(vendor) +'_school_'+ str(schoolIds) + '_1'
    #     ven = str(vendor)+'_school_'+str(school.school_id)+'_1'
        
    #     print('current school id:'+str(school.school_id))
    #     current_user.school_id = school.school_id
    #     print(vendorId)
    #     bankIdExist = BankDetail.query.filter_by(vendor_id=ven).first()
    #     if bankIdExist:
    #         bankIdExist.account_num = accountNo
    #         bankIdExist.ifsc = IfscCode
    #         bankIdExist.bank_name = bankName
    #         bankIdExist.account_name = accountHolderName
    #         db.session.commit()
    #     else:
    #         bankDet = BankDetail(account_num=accountNo,ifsc=IfscCode,vendor_id=ven,bank_name=bankName,account_name=accountHolderName,school_id=school.school_id,is_archived='N')
    #         db.session.add(bankDet)
    #         db.session.commit()
    #         school.current_vendor_id = ven
    #         db.session.commit()
    #     reviewStatus = "select *from teacher_profile where user_id='"+str(current_user.id)+" '"
    #     reviewStatus = db.session.execute(text(reviewStatus)).first()
    #     print('Review status:'+str(reviewStatus.review_status))
    #     return jsonify(reviewStatus.review_status)
    # print('School Id:'+str(selectSchool))
    # schoolEx = ''
    # if selectSchool:
    #     schoolEx = SchoolProfile.query.filter_by(school_id=selectSchool).first()
    # if user_avatar!=None:
    #     current_user.user_avatar = user_avatar
    # if about_me!=None:
    #     current_user.about_me = about_me
    # board  = MessageDetails.query.filter_by(category='Board',description='Other').first()
    # checkTeacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    # print('Teacher:'+str(checkTeacher))
    # if checkTeacher==None:
    #     ## Adding new school record
    #     print('if teacher is none')
    #     schoolAdd = ''
    #     if selectSchool==None:
    #         print('if school id is none')
    #         schoolAdd = SchoolProfile(school_name=schoolName,board_id=board.msg_id,registered_date=datetime.now(),school_type='individual',last_modified_date=datetime.now())
    #         db.session.add(schoolAdd)
    #         current_user.school_id = schoolAdd.school_id
    #         db.session.commit()
    #         exSchool = SchoolProfile.query.filter_by(school_name=schoolName,school_type='individual',board_id=board.msg_id).first()
    #         schoolEx = exSchool
    #     ##Adding new teacher record
    #     print(schoolEx.school_id)
    #     teacherAdd = TeacherProfile(teacher_name=str(current_user.first_name)+' '+str(current_user.last_name),school_id=schoolEx.school_id,registration_date=datetime.now(),email=current_user.email,phone=current_user.phone,user_id=current_user.id,device_preference='195',last_modified_date=datetime.now())
    #     db.session.add(teacherAdd)
    #     db.session.commit()
    #     #Updating school admin with the new teacher ID
    #     #teacherEx = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #     checkTeacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #     schoolEx.school_admin = checkTeacher.teacher_id
    #     #generating vendor id
    #     vendorId = str(schoolName).lower().replace(' ','_')
    #     vendorId = str(vendorId) +'_'+ str(schoolEx.school_id) + '_1'
    #     #schoolAdd.curr_vendor_id = vendorId
    #     current_user.school_id = schoolEx.school_id
    #     db.session.commit()
    #     reviewId = MessageDetails.query.filter_by(description='Inreview',category='Review Status').first()
    #     reviewInsert = "update teacher_profile set review_status='"+str(reviewId.msg_id)+"' where user_id='"+str(current_user.id)+"' "
    #     #Send sms to tech team to follow up
    #     reviewInsert = db.session.execute(text(reviewInsert))
    #     message = "New tutor has been registered in the database. Please setup contact them on email: "+ current_user.email 
    #     message = message + " and phone: " + current_user.phone + " for review and pg setup"
    #     phoneList = "9008500227,9910368828"        
    #     #calling SMS function        
    #     smsResponse = sendSMS(message, phoneList)
    #     print("This is the sms send response: " + smsResponse)

    #     ##Section to bank details
    #     print('Bank Details')
    #     print(accountNo)
    #     print(IfscCode)
    #     print(bankName)
    #     print(accountHolderName)
    #     print(vendorId)
    #     bankDetailAdd = BankDetail(account_num = accountNo, ifsc=IfscCode, bank_name=bankName, account_name=accountHolderName, 
    #         school_id=schoolEx.school_id, vendor_id=vendorId,
    #         is_archived='N')
    #     db.session.add(bankDetailAdd)
    #     db.session.commit()
    #     schoolEx.curr_vendor_id = vendorId
    #     db.session.commit()
    #     ##Section to add vendor with payment gateway


    #     ## Section to create a roomid  for the teacher
    # print("####This is the current room id: " + str(checkTeacher.room_id))
    # if checkTeacher.room_id==None:            
    #     roomResponse = roomCreation()
    #     roomResponseJson = roomResponse.json()
    #     print("New room ID created: " +str(roomResponseJson["url"]))
    #     checkTeacher.room_id = str(roomResponseJson["url"])
    #     db.session.commit()
    # reviewStatus = "select *from teacher_profile where user_id='"+str(current_user.id)+" '"
    # reviewStatus = db.session.execute(text(reviewStatus)).first()
    # print('Review status:'+str(reviewStatus.review_status))
    # return jsonify(reviewStatus.review_status)


@app.route('/getOnlineClassLink',methods=['GET','POST'])
def getOnlineClassLink():
    if request.method == 'POST':
        jsonExamData = request.json 
        # jsonExamData = {"contact":{"phone":"9008262739"},"result":{"data":"1"}}       
        data = json.dumps(jsonExamData)
        response = json.loads(data)
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for val in response['results'].values():
            paramList.append(val)
        for data in response['contact'].values():
            conList.append(data)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        # Start
        selClass = paramList[11]
        classDet = ClassSection.query.filter_by(class_val=selClass,school_id=teacher_id.school_id).first()
        subId  = paramList[15]
        extractChapterQuery = "select td.topic_id ,td.topic_name ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            chap = str(c)+str('-')+str(chapterDet.topic_id)+str('-')+str(chapterDet.topic_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        selChapter = ''
        for chapterName in chapterDetList:
            num = chapterName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(chapterName)
                selChapter = chapterName.split('-')[1]
                print('selChapter:'+str(selChapter))
        selChapter = selChapter.strip()
        print('Topic ID:'+str(selChapter))
        selSubject = paramList[12]
        # End
        if teacher_id.room_id==None:            
            roomResponse = roomCreation()
            roomResponseJson = roomResponse.json()
            print("New room ID created: " +str(roomResponseJson["url"]))
            teacher_id.room_id = str(roomResponseJson["url"])
            db.session.commit()
        link = url_for('classDelivery',class_sec_id=classDet.class_sec_id,subject_id=subId,topic_id=selChapter,retake='N',_external=True)
        OnlineClassLink = str('Online class link:\n')+ str(teacher_id.room_id)+str("\n")
        OnlineClassLink = OnlineClassLink + str("Book Link:\n")+str(link)
        return jsonify({'onlineClassLink':OnlineClassLink})

##Helper function
def roomCreation():
    dailycourl = "https://api.daily.co/v1/rooms"
    payload = {"properties": {
            "max_participants": 50,
            "enable_screenshare": True,
            "enable_chat": True,
            "start_video_off": True,
            "start_audio_off": True,
            "owner_only_broadcast": True,
            "eject_at_room_exp": True,
            "eject_after_elapsed": 4
        }}
    headers = {
        "content-type": "application/json",
        "authorization": "Bearer 157accd1531abddc376517325567f744168d28be21a18cc1bec3e356695259b4"
    }
    response = rq.request("POST", dailycourl, json=payload, headers=headers)
    #print("##########Room Creation response: "+response.text)
    return response

##Helper function
def sendSMS(message, phoneList):
    apiPath = "http://173.212.233.109/app/smsapisr/index.php?key=35EF8379A04DB8&"
    apiPath = apiPath + "campaign=9967&routeid=6&type=text&"
    apiPath = apiPath + "contacts="+str(phoneList) +"&senderid=GLOBAL&"
    apiPath = apiPath + "msg="+ quote(message)
    print(apiPath)
    ##Sending message here
    try:
        r = requests.post(apiPath)                    
        returnData = str(r.text)
        print(returnData)
        
        if "SMS-SHOOT-ID" in returnData:
            return "SMS sent successfully"
        else:
            return "Error sending sms"
        #    for val in studentPhones:
        #        commTransAdd = CommunicationTransaction(comm_id=commDataAdd.comm_id, student_id=val.student_id,last_modified_date = datetime.today())
        #        db.session.add(commTransAdd)
        #    commDataAdd.status=232
        #    db.session.commit()                                        
        #    return jsonify(['0'])
        #else:
        #    return jsonify(['1'])                    
    except:
        return "Exception in sending sms"

#def vendorAddition():

    #Request {"vendorId" : "VEN343", "name" : "343 Industries", "phone" : 9900034300, "email" : "three43@gmail.com", "commission" : 34.3, "bankAccount" : 30004343400003, "accountHolder" : "TFT", "ifsc" : "HDFC0000343", "address1" : "Diamond District", "address2" : "Indranagar", "city" : "Bengaluru", "state" : "Karnataka", "panNo" : "AAAPL1234C", "aadharNo" : "499118665246", "gstin" : "22AAAAA0000A1Z5", "pincode" : 560071}
    #Response {"status":"SUCCESS", "subCode":"200", "message":"Vendor added successfully"}


@app.route('/addCourse')                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                     
def addCourse():
    course_id = request.args.get('course_id')
    if course_id=='':
        courId = CourseDetail(course_name='Untitled Course',course_status=275,is_archived='N',last_modified_date=datetime.now())
        db.session.add(courId)
        db.session.commit()
        course_id = courId.course_id
    return redirect(url_for('course.editCourse',course_id=course_id))

# @app.route('/editCourse')
# def editCourse():
#     print('inside editCourse')
#     course_category = MessageDetails.query.filter_by(category='Course Category').first()
#     desc = course_category.description.split(',')
#     course_id=request.args.get('course_id')
#     teacherIdExist = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     print('userID:'+str(current_user.id))
#     print('Teacher:'+str(teacherIdExist))
#     reviewStatus = "select *from teacher_profile where user_id='"+str(current_user.id)+"'"
#     reviewStatus = db.session.execute(text(reviewStatus)).first()
#     if reviewStatus:
#         print('REview status:'+str(reviewStatus.review_status))
#         if reviewStatus.review_status==273:
#             print('review status Inreview')
#             return redirect(url_for('teacherRegistration'))
#     if teacherIdExist==None:
#         return redirect(url_for('teacherRegistration'))
#     else:
#         print('Description:'+str(desc))
#         print('course_id:'+str(course_id))
#         if course_id:
#             courseDet = CourseDetail.query.filter_by(course_id=course_id).first()
#             levelId = MessageDetails.query.filter_by(category='Difficulty Level',msg_id=courseDet.difficulty_level).first()
#             courseNotes = TopicNotes.query.filter_by(course_id=course_id).first()
#             # topicDet = "select count(*) as no_of_questions,td.topic_name,td.topic_id,ct.course_id from course_topics ct "
#             # topicDet = topicDet + "inner join topic_detail td on ct.topic_id=td.topic_id "
#             # topicDet = topicDet + "left join test_questions tq on ct.test_id = tq.test_id "
#             # topicDet = topicDet + "where ct.course_id = '"+str(course_id)+"' and tq.is_archived='N' and ct.is_archived='N' group by td.topic_name,td.topic_id,ct.course_id "
            
#             topicL = []
#             topicsID = CourseTopics.query.filter_by(course_id=course_id,is_archived='N').all()
#             for topicId in topicsID:
#                 topicList = []
#                 topic_name = Topic.query.filter_by(topic_id=topicId.topic_id).first()
#                 quesNo = TestQuestions.query.filter_by(test_id=topicId.test_id,is_archived='N').all()
#                 questionNo = len(quesNo)
#                 topicList.append(topic_name.topic_name)
#                 topicList.append(questionNo)
#                 topicList.append(topicId.topic_id)
#                 notes = TopicNotes.query.filter_by(topic_id=topicId.topic_id,is_archived='N').first()
#                 recording = "select *from course_topics where course_id='"+str(course_id)+"' and topic_id='"+str(topicId.topic_id)+"' and video_class_url<>'' order by topic_id asc"
#                 recording = db.session.execute(text(recording)).first()
#                 checkNotes = ''
#                 checkRec = ''
#                 if notes:
#                     checkNotes = notes.notes_name
#                 if recording:
#                     checkRec = recording.video_class_url
#                 topicList.append(checkNotes)
#                 topicList.append(checkRec)
#                 print(topicList)
#                 topicL.append(topicList)
#             print(topicL)
#             for topic in topicL:
#                 print(topic[0])
#                 print(topic[1])
#                 print(topic[2])
#             # topicDet = db.session.execute(text(topicDet)).fetchall()
#             idealFor = ''
#             if courseDet:
#                 idealFor = courseDet.ideal_for
#             levelId = ''
#             if levelId:
#                 levelId = levelId.description
#             print('Description:'+str(courseDet.description))
#             status = 1
#             return render_template('editCourse.html',status=status,levelId=levelId,idealFor=idealFor,desc=desc,courseDet=courseDet,course_id=course_id,topicDet=topicL)
#         else:
#             levelId = ''
#             return render_template('editCourse.html',levelId=levelId,desc=desc,course_id=course_id)
    


# #clip = (VideoFileClip("frozen_trailer.mp4")
# #        .subclip((1,22.65),(1,23.2))
# #        .resize(0.3))
# #clip.write_gif("use_your_head.gif")



# @app.route('/searchTopic',methods=['GET','POST'])
# def searchTopic():
#     topic = request.args.get('topic')
#     courseId = request.args.get('courseId')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     coursesId = CourseDetail.query.filter_by(teacher_id=teacher_id.teacher_id).all()
#     topicArray = []
#     for course_id in coursesId:
#         if str(course_id.course_id)!=str(courseId):
#             TopicIds = "select td.topic_id,td.topic_name from topic_detail td inner join course_topics ct on td.topic_id = ct.topic_id where td.topic_name like '"+str(topic)+"%'  and ct.course_id ='"+str(course_id.course_id)+"' and ct.is_archived='N'"
#             TopicIds = db.session.execute(text(TopicIds)).fetchall()
#             for top in TopicIds:
#                 print('Topic:'+str(top))
#                 topicArray.append(str(top.topic_id)+':'+str(top.topic_name)+':'+str(course_id.course_id))
#     if topicArray:
#         return jsonify([topicArray])
#     else:
#         return ""

# @app.route('/fetchQues',methods=['GET','POST'])
# def fetchQues():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#     print('inside fetchQues')
#     courseId = request.args.get('courseId')
#     print('courseId:'+str(courseId))
#     topics = CourseTopics.query.filter_by(course_id=courseId,is_archived='N').order_by(CourseTopics.topic_id.desc()).all()
#     topicsDet = []
#     for topic in topics:
#         topicName = Topic.query.filter_by(topic_id=topic.topic_id).first()
#         quesIds = TestQuestions.query.filter_by(test_id=topic.test_id,is_archived='N').all()
#         quesNo = len(quesIds)
#         notes = TopicNotes.query.filter_by(topic_id=topic.topic_id,is_archived='N').first()
#         checkNotes = ''
#         checkRec = ''
#         if notes:
#             checkNotes = notes.notes_name
#         recording = "select *from course_topics where course_id='"+str(courseId)+"' and topic_id='"+str(topic.topic_id)+"' and video_class_url<>'' order by topic_id"
#         recording = db.session.execute(text(recording)).first()
#         if recording:
#             checkRec = recording.video_class_url
#         topic = topicName.topic_name.replace(",","/")
#         topicsDet.append(str(topic)+':'+str(quesNo)+':'+str(topicName.topic_id)+':'+str(checkNotes)+':'+str(checkRec))
#     if topicsDet:
#         return jsonify(topicsDet)
#     else:
#         return ""

# @app.route('/fetchRecording',methods=['GET','POST'])
# def fetchRecording():
#     print('inside fetchRecording')
#     topic_id = request.args.get('topic_id')
#     courseId = request.args.get('courseId')
#     record = CourseTopics.query.filter_by(course_id=courseId,topic_id=topic_id).first()
#     return jsonify(record.video_class_url)

# @app.route('/deleteNotes',methods=['GET','POST'])
# def deleteNotes():
#     notes_id = request.args.get('notes_id')
#     notes = TopicNotes.query.filter_by(tn_id=notes_id).first()
#     notes.is_archived = 'Y'
#     db.session.commit()
#     return jsonify(notes.topic_id)


# @app.route('/fetchNotes',methods=['GET','POST'])
# def fetchNotes():
#     print('inside fetchNotes')
#     topic_id = request.args.get('topic_id')
#     notes = TopicNotes.query.filter_by(topic_id=topic_id,is_archived='N').all()
#     notesData = []
#     for note in notes:
#         NewNotes = note.notes_name.replace(",","/")
#         notesData.append(str(NewNotes)+'!'+str(note.notes_url)+'!'+str(note.tn_id))
#     print(notesData)
#     if notesData:
#         return jsonify(notesData)
#     else:
#         return ""

# @app.route('/fetchRemQues',methods=['GET','POST'])
# def fetchRemQues():
#     quesIdList = request.get_json()
#     quesArray = []
#     for qId in quesIdList:
#         print('Question Id:'+str(qId))
#         quesObj = {}    
#         quesName = QuestionDetails.query.filter_by(question_id=qId).first()
#         quesObj['quesName'] = quesName.question_description
#         print('Ques:'+str(quesName.question_description))
#         quesOptions = QuestionOptions.query.filter_by(question_id=qId).all()
#         i=0
#         opt1=''
#         opt2=''
#         opt3=''
#         opt4=''
#         for option in quesOptions:
#             if i==0:
#                 opt1 = option.option_desc
#             elif i==1:
#                 opt2 = option.option_desc
#             elif i==2:
#                 opt3 = option.option_desc
#             else:
#                 opt4 = option.option_desc
#             i=i+1
#             print('quesOptions:'+str(option.option_desc))
#         quesArray.append(str(quesName.question_description)+':'+str(opt1)+':'+str(opt2)+':'+str(opt3)+':'+str(opt4)+':'+str(qId))
#     print('quesArray:')
#     print(quesArray)
#     if quesArray:
#         return jsonify(quesArray)
#     else:
#         return ""

# @app.route('/topicName',methods=['GET','POST'])
# def topicName():
#     print('inside topicName')
#     topicId = request.args.get('topic_id')
#     topicName = Topic.query.filter_by(topic_id=topicId).first()
#     topic_name = topicName.topic_name
#     print('topic name:'+topic_name)

#     return jsonify(topic_name)

# @app.route('/fetchTopicQues',methods=['GET','POST'])
# def fetchTopicsQues():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#     print('inside fetchTopics')
#     topic_id = request.args.get('topic_id')
#     courseId = request.args.get('courseId')
#     print('topic_id:'+str(topic_id))
#     topics = CourseTopics.query.filter_by(topic_id=topic_id,is_archived='N').first()
#     topicName = Topic.query.filter_by(topic_id=topics.topic_id).first()
#     print('Topic name:'+str(topicName.topic_name))
#     quesIds = TestQuestions.query.filter_by(test_id=topics.test_id,is_archived='N').all()
#     # topicNotes = TopicNotes.query.filter_by(topic_id=topics.topic_id).first()
#     # NotesName = topicNotes.notes_name
#     # NotesUrl = topicNotes.notes_url
#     quesArray = []
#     NewTopicName = ''
#     if len(quesIds)==0:
#         NewTopicName = topicName.topic_name.replace(",","/")
#         quesArray.append(str('')+':'+str(NewTopicName)+':'+str('')+':'+str('')+':'+str('')+':'+str('')+':'+str('')+':'+str(topicName.topic_id))
#     for qId in quesIds:
#         quesObj = {}    
#         quesName = QuestionDetails.query.filter_by(question_id=qId.question_id).first()
#         quesObj['quesName'] = quesName.question_description
#         quesObj['topic_name'] = topicName.topic_name
#         print('Ques:'+str(quesName.question_description))
#         print('topicName:'+str(topicName.topic_name))
#         quesOptions = QuestionOptions.query.filter_by(question_id=qId.question_id).all()
#         i=0
#         opt1=''
#         opt2=''
#         opt3=''
#         opt4=''
#         for option in quesOptions:
#             if i==0:
#                 opt1 = option.option_desc
#             elif i==1:
#                 opt2 = option.option_desc
#             elif i==2:
#                 opt3 = option.option_desc
#             else:
#                 opt4 = option.option_desc
#             i=i+1
#             print('quesOptions:'+str(option.option_desc))
#         NewTopicName = topicName.topic_name.replace(",","/")
#         quesArray.append(str(quesName.question_description)+':'+str(NewTopicName)+':'+str(opt1)+':'+str(opt2)+':'+str(opt3)+':'+str(opt4)+':'+str(qId.question_id)+':'+str(topicName.topic_id))
#     print('quesArray:')
#     print(quesArray)
#     if quesArray:
#         return jsonify(quesArray)
#     else:
#         return ""

# @app.route('/deleteTopic',methods=['GET','POST'])
# def deleteTopic():
#     print('inside deleteTopic')
#     topicId = request.args.get('topicId')
#     print('Topic id:'+str(topicId))
#     course_id = request.args.get('course_id')
#     courseTopic = CourseTopics.query.filter_by(topic_id=topicId,course_id=course_id,is_archived='N').first()
#     courseTopic.is_archived = 'Y'
#     db.session.commit()
#     notes = "update topic_notes set is_archived='Y' where topic_id='"+str(topicId)+"' and is_archived='N'"
#     notes = db.session.execute(text(notes))
#     db.session.commit()
#     return jsonify("1")

# @app.route('/updateCourseTopic',methods=['GET','POST'])
# def updateCourseTopic():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#     print('inside updateCourseTopic')
#     topicId = request.args.get('topicId')
#     courseId = request.args.get('courseId')
#     topicName = request.args.get('topicName')
#     print('Topic Id:'+str(topicId))
#     print('courseId:'+str(courseId))
#     quesIds = request.get_json()
#     print(quesIds)
#     topicDet = Topic.query.filter_by(topic_id=topicId).first()
#     topicDet.topic_name = topicName
#     db.session.commit()
#     testId = CourseTopics.query.filter_by(course_id=courseId,topic_id=topicId).first()
#     # totalQId = TestQuestions.query.filter_by(test_id=testId.test_id).all()
#     deleteAll = "update test_questions set is_archived='Y' where test_id='"+str(testId.test_id)+"' "
#     deleteAll = db.session.execute(text(deleteAll))
#     # print('Total Question Ids:'+str(totalQId))
#     print('Total not deleted Ques Ids:'+str(quesIds))
#     print(quesIds)
#     print('Length of Ques Ids:'+str(len(quesIds)))
#     total_marks = 10*int(len(quesIds))
#     print('Total marks:'+str(total_marks))
#     testDet = TestDetails.query.filter_by(test_id=testId.test_id).first()
#     # db.session.add(testDet)
#     testDet.total_marks = total_marks
#     db.session.commit()
#         # testId = "select max(test_id) as test_id from test_details"
#         # testId = db.session.execute(text(testId)).first()
#     # courseTopic = ''
#     # if courseId:
#     #     courseTopic = CourseTopics(course_id=courseId,topic_id=topicDet.topic_id,test_id=testDet.test_id,is_archived='N',last_modified_date=datetime.now())
#     #     db.session.add(courseTopic)
#     # else:
#     #     courseTopic = CourseTopics(topic_id=topicDet.topic_id,test_id=testDet.test_id,is_archived='N',last_modified_date=datetime.now())
#     #     db.session.add(courseTopic)
#     # db.session.commit()
#     if len(quesIds)!=0:
#         for quesId in quesIds:
#             print('QuesID:'+str(quesId))
#             if quesId!='[' or quesId!=']':
#                 testQues = TestQuestions(test_id=testDet.test_id,question_id=quesId,is_archived='N',last_modified_date=datetime.now())
#                 db.session.add(testQues)
#                 db.session.commit()
#             # quesDet = QuestionDetails.query.filter_by(question_id=quesId).first()
#             # quesDet.topic_id=topicDet.topic_id
            
    
#     return jsonify("1")



# @app.route('/addCourseTopic',methods=['GET','POST'])
# def addCourseTopic():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#     print('inside addCourseTopic')
#     topicName = request.args.get('topicName')
#     courseId = request.args.get('courseId')
#     topicId = request.args.get('topicId')
#     courId = request.args.get('courId')
#     print('My course ID:'+str(courseId))
#     print('selected topic id:'+str(topicId))
#     print('courseId of selected topic:'+str(courId))
    
    

#     if courId:
#         quesIds = request.get_json()
#         total_marks = 10*len(quesIds)
#         print('Total marks:'+str(total_marks))
#         testDet = TestDetails(board_id=board.board_id,school_id=teacherData.school_id,test_type='Practice Test',total_marks=total_marks,teacher_id=teacherData.teacher_id,date_of_creation=datetime.now(),last_modified_date=datetime.now())
#         db.session.add(testDet)
#         db.session.commit()
#         testId = "select max(test_id) as test_id from test_details"
#         testId = db.session.execute(text(testId)).first()

#         if courId:
#             myTestId = CourseTopics.query.filter_by(course_id=courseId,topic_id=topicId,is_archived='N').first()
#             testID = CourseTopics.query.filter_by(course_id=courId,topic_id=topicId).first()
#             questionIds = TestQuestions.query.filter_by(test_id=testID.test_id).all()
#             print(questionIds)
#             for q in questionIds:
#                 print('Question ID:'+str(q.question_id))
#                 print('test id in which question stored:'+str(testId.test_id))
#                 Questions = TestQuestions(test_id=testId.test_id,question_id=q.question_id,is_archived='N',last_modified_date=datetime.now())
#                 db.session.add(Questions)
#                 db.session.commit()
#         courseTopic = ''
#         if courseId:
#             courseTopic = CourseTopics(course_id=courseId,topic_id=topicId,test_id=testId.test_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(courseTopic)
#         else:
#             courseTopic = CourseTopics(topic_id=topicId,test_id=testId.test_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(courseTopic)
#         db.session.commit()
#         for quesId in quesIds:
#             print('QuesID:'+str(quesId))
#             testQues = TestQuestions(test_id=testId.test_id,question_id=quesId,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(testQues)
        
#             quesDet = QuestionDetails.query.filter_by(question_id=quesId).first()
#             quesDet.topic_id=topicId
#             db.session.commit()
#         return jsonify(topicId)
#     else:
#         print('Topic name:'+str(topicName))
#         print('courseId:'+str(courseId))
#         quesIds = request.get_json()
#         print(quesIds)
#         # courseId = CourseDetail.query.filter_by(course_name=courseName,teacher_id=teacherData.teacher_id,school_id=teacherData.school_id).first()
        
#         # for quesId in quesIds:
#         #     print(quesId)
#         topicDet = Topic(topic_name=topicName,chapter_name=topicName,board_id=board.board_id,teacher_id=teacherData.teacher_id)
#         db.session.add(topicDet)
#         db.session.commit()
#         # topicId = "select max(topic_id) as topic_id from topic_detail"
#         # topicId = db.session.execute(text(topicId)).first()
        
#         topicTr = TopicTracker(school_id=teacherData.school_id,topic_id=topicDet.topic_id,is_covered='N',reteach_count=0,is_archived='N',last_modified_date=datetime.now())
#         db.session.add(topicTr)
#         db.session.commit()
#         total_marks = 10*len(quesIds)
#         print('Total marks:'+str(total_marks))
#         testDet = TestDetails(board_id=board.board_id,school_id=teacherData.school_id,test_type='Practice Test',total_marks=total_marks,teacher_id=teacherData.teacher_id,date_of_creation=datetime.now(),last_modified_date=datetime.now())
#         db.session.add(testDet)
#         db.session.commit()
#         # testId = "select max(test_id) as test_id from test_details"
#         # testId = db.session.execute(text(testId)).first()
#         courseTopic = ''
#         if courseId:
#             courseTopic = CourseTopics(course_id=courseId,topic_id=topicDet.topic_id,test_id=testDet.test_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(courseTopic)
#         else:
#             courseTopic = CourseTopics(topic_id=topicDet.topic_id,test_id=testDet.test_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(courseTopic)
#         db.session.commit()
#         if quesIds:
#             for quesId in quesIds:
#                 testQues = TestQuestions(test_id=testDet.test_id,question_id=quesId,is_archived='N',last_modified_date=datetime.now())
#                 db.session.add(testQues)
                
#                 quesDet = QuestionDetails.query.filter_by(question_id=quesId).first()
#                 quesDet.topic_id=topicDet.topic_id
#                 db.session.commit()
#         return jsonify(topicDet.topic_id)
    
# @app.route('/fetchTickCorrect',methods=['GET','POST'])
# def fetchTickCorrect():
#     print('inside fetchTickCorrect')
#     correctOpt = []
#     topic_id = request.args.get('topic_id')
#     print('TopicId:'+str(topic_id))
#     # quesIdsList = TestQuestions.query.filter_by(topic_id=topic_id).all()
#     quesList  = request.get_json()
#     print('Question List:')
#     print(quesList)
#     for quesId in quesList:
#         print('Question Id:'+str(quesId))
#         corr = QuestionOptions.query.filter_by(question_id=quesId,is_correct='Y').first()
#         correctOpt.append(str(corr.option_desc)+':'+str(quesId))
#     return jsonify(correctOpt)

# @app.route('/addRecording',methods=['GET','POST'])
# def addRecording():
#     topic_id = request.args.get('topic_id')
#     course_id = request.args.get('course_id')
#     videoRec = CourseTopics.query.filter_by(topic_id=topic_id,course_id=course_id).first()
#     recordingURL = request.form.get('recordingURL')
#     print('recording Url:'+str(recordingURL))
#     videoRecordUrl = request.form.get('videoRecordUrl')
#     print('video url:'+str(videoRecordUrl))
    
#     print('video recording url:'+str(videoRecordUrl))
#     if recordingURL:
#         videoRec.video_class_url = recordingURL
#     else:
#         videoRec.video_class_url = videoRecordUrl
#     db.session.commit()
#     return jsonify("1")

# @app.route('/updateNotes',methods=['GET','POST'])
# def updateNotes():
#     topicId = request.args.get('topic_id')
#     notesName = request.form.getlist('notesName')
#     notesURL = request.form.getlist('notesURL')
#     videoNotesUrl = request.form.getlist('videoNotesUrl')
#     print('topicId:'+str(topicId))
#     # print('Notes name:'+str(notesName))  
#     existNotes = "update topic_notes set is_archived='Y' where topic_id='"+str(topicId)+"' "
#     existNotes = db.session.execute(text(existNotes))
#     print('Length of notes url array:'+str(len(notesURL)))
#     for i in range(len(notesName)):
#         print('inside for loop:'+str(i))
#         print('NotesName:'+str(notesName[i]))
#         print('notesUrl:'+str(notesURL[i]))
#         print('videoNotesUrl:'+str(videoNotesUrl[i]))
#         print('index:'+str(i))
#         if i!=0:
#             if notesURL[i]:
#                 print('url not null')
#                 if notesName[i]:
#                     print('notes name not null')
#                     if notesURL[i]:
#                         courseId = CourseTopics.query.filter_by(topic_id=topicId).first()
#                         addNotes = TopicNotes(topic_id=topicId,course_id=courseId.course_id,notes_name=notesName[i],notes_url=notesURL[i],notes_type=226,is_archived='N',last_modified_date=datetime.now())
#                         db.session.add(addNotes)
#                         db.session.commit()
#                     else:
#                         return ""
#                 else:
#                     return ""
#             else:
#                 if notesName[i]: 
#                     if videoNotesUrl[i]:
#                         courseId = CourseTopics.query.filter_by(topic_id=topicId).first()
#                         addNotes = TopicNotes(topic_id=topicId,course_id=courseId.course_id,notes_name=notesName[i],notes_url=videoNotesUrl[i],notes_type=226,is_archived='N',last_modified_date=datetime.now())
#                         db.session.add(addNotes)
#                         db.session.commit()
#                     else:
#                         return ""
#                 else:
#                     return ""
#     return jsonify("1")

# @app.route('/addNotes',methods=['GET','POST'])
# def addNotes():
#     topicId = request.args.get('topic_id')
#     notesName = request.form.getlist('notesName')
#     notesURL = request.form.getlist('notesURL')
#     videoNotesUrl = request.form.getlist('videoNotesUrl')
#     print('topicId:'+str(topicId))
#     print('Notes name:'+str(notesName))
#     for i in range(len(notesName)):
#         print('inside for loop:'+str(i))
#         print('NotesName:'+str(notesName[i]))
#         print('notesUrl:'+str(notesURL[i]))
#         print('videoNotesUrl:'+str(videoNotesUrl[i]))
#         print('index:'+str(i))
#         if notesURL[i]:
#             print('url not null')
#             if notesName[i]:
#                 print('notes name not null')
#                 if notesURL[i]:
#                     courseId = CourseTopics.query.filter_by(topic_id=topicId).first()
#                     addNotes = TopicNotes(topic_id=topicId,course_id=courseId.course_id,notes_name=notesName[i],notes_url=notesURL[i],notes_type=226,is_archived='N',last_modified_date=datetime.now())
#                     db.session.add(addNotes)
#                     db.session.commit()
#                 else:
#                     return ""
#             else:
#                 return ""
#         else:
#             if notesName[i]: 
#                 if videoNotesUrl[i]:
#                     print('inside when notes name and file uploaded')
#                     courseId = CourseTopics.query.filter_by(topic_id=topicId).first()
#                     addNotes = TopicNotes(topic_id=topicId,course_id=courseId.course_id,notes_name=notesName[i],notes_url=videoNotesUrl[i],notes_type=226,is_archived='N',last_modified_date=datetime.now())
#                     db.session.add(addNotes)
#                     db.session.commit()
#                 else:
#                     return ""
#             # return ""
#             else:
#                 return ""
#     return jsonify("1")

# @app.route('/addNewQuestion',methods=['GET','POST'])
# def addNewQuestion():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#     print('inside add question')
#     corr = request.args.get('corr')
#     ques = request.args.get('ques')
#     opt1 = request.args.get('opt1')
#     opt2 = request.args.get('opt2')
#     opt3 = request.args.get('opt3')
#     opt4 = request.args.get('opt4')
#     print('question Desc:'+str(ques))
#     print('option 1:'+str(opt1))
#     print('option 2:'+str(opt2))
#     print('option 3:'+str(opt3))
#     print('option 4:'+str(opt4))
#     print('correct option:'+str(corr))
#     quesCreate = QuestionDetails(board_id=board.board_id,question_description=ques,question_type='MCQ1',suggested_weightage='10',is_private='N',archive_status='N')
#     db.session.add(quesCreate)
#     db.session.commit()
#     for i in range(4):
#         op = request.args.get('opt'+str(i+1))
#         correctOption = ''
#         if str(corr) == str(i+1):
#             correctOption = 'Y'
#         else:
#             correctOption = 'N'
#         option = ''
#         if i==0:
#             option = 'A'
#         elif i==1:
#             option = 'B'
#         elif i==2:
#             option = 'C'
#         else:
#             option = 'D'
#         # ques_det = QuestionDetails.query.filter_by(board_id=board.board_id,question_description=ques,question_type='MCQ1',suggested_weightage='10',is_private='N',archive_status='N').first()
#         options = QuestionOptions(option=option,is_correct=correctOption,option_desc=op,question_id=quesCreate.question_id,weightage='10',last_modified_date=datetime.now())
#         db.session.add(options)
#         db.session.commit()
#     return jsonify(quesCreate.question_id)
    
# @app.route('/fetchQuesList',methods=['GET','POST'])
# def fetchQuesList():
#     print('inside fetchQuesList')
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     quesList = request.get_json()
#     indic = request.args.get('indic')
#     print('questionIdList:')
#     print(quesList)
#     for ques in quesList:
#         print('Question_id:')
#         print(ques)
#     if indic=='1':
#         for quesId in quesList:
#             quesDesc = QuestionDetails.query.filter_by(question_id=quesId).first()
#             print(quesDesc.question_description)
#             quesOptions = QuestionOptions.query.filter_by(question_id=quesId).all()
#             op1=''
#             op2=''
#             op3=''
#             op4=''
#             optionList = []
#             quesDetails = []
#             corrOption = ''
#             for options in quesOptions:
#                 print('Option Desc:'+str(options.option_desc))
#                 optionList.append(options.option_desc)
#                 if options.is_correct=='Y':
#                     corrOption = options.option_desc
#             for i in range(len(optionList)):
#                 if i==0:
#                     op1=optionList[i]
#                 elif i==1:
#                     op2=optionList[i]
#                 elif i==2:
#                     op3=optionList[i]
#                 elif i==3:
#                     op4=optionList[i]
#         print('Question:'+str(quesDesc.question_description)+'op1:'+str(op1)+'op2:'+str(op2)+'op3:'+str(op3)+'op4:'+str(op4))
#         quesDetails.append(str(quesDesc.question_description)+':'+str(op1)+':'+str(op2)+':'+str(op3)+':'+str(op4)+':'+str(corrOption))
#     return jsonify([quesDetails])

# @app.route('/saveAndPublishedCourse',methods=['GET','POST'])
# def saveAndPublishedCourse():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #print('inside saveCourse')
#     course = request.form.get('course')
#     courseId = request.args.get('course_id')
#     description = request.form.get('description')
#     imageUrl = request.form.get('imageUrl')
#     video_url = request.form.get('videoUrl')
#     idealfor = request.args.get('idealfor')
#     level = request.form.get('level')
#     private = request.form.get('private')
#     #print('Course name:'+str(course))
#     #print('courseId:'+str(courseId))
#     #print('description name:'+str(description))
#     #print('Image Url:'+str(imageUrl))
#     #print('Private:'+str(private))
#     course_status = request.args.get('course_status')
#     #print('course status:'+str(course_status))
#     #
#     #print('video_url :'+str(video_url))
#     print('Ideal for:'+str(idealfor))
#     #print('level:'+str(level))
#     updateIndex=False
#     levelId = MessageDetails.query.filter_by(description=level,category='Difficulty Level').first()
#     courseDet = CourseDetail.query.filter_by(course_id=courseId,description=description,summary_url=video_url,
#     teacher_id=teacherData.teacher_id,school_id=teacherData.school_id,ideal_for=idealfor,difficulty_level=levelId.msg_id).first()
#     if courseDet:
#         course_status_id = MessageDetails.query.filter_by(category='Course Status',description=course_status).first()
#         courseDet.course_status=course_status_id.msg_id
#         db.session.commit()
#         print('returning from first')
#         if app.config["MODE"]=="PROD":
#             updateIndex = updateSearchIndex("send","saveCourse")
#         if updateIndex==True:
#             return jsonify("1")
#         else:
#             return jsonify("2")
#     else:
#         course_status_id = MessageDetails.query.filter_by(category='Course Status',description=course_status).first()
#         courseDet = CourseDetail.query.filter_by(course_id=courseId).first()
#         if private:
#             print('if course status is private')
#             courseDet.description=description
#             courseDet.summary_url=video_url
#             courseDet.teacher_id=teacherData.teacher_id
#             courseDet.school_id=teacherData.school_id
#             if idealfor:
#                 courseDet.ideal_for=idealfor
#             courseDet.course_status=course_status_id.msg_id
#             courseDet.is_private='Y'
#             courseDet.image_url = imageUrl
#             courseDet.is_archived = 'N'
#             courseDet.difficulty_level=levelId.msg_id
#         else:
#             print('if course status is public')
#             courseDet.description=description
#             courseDet.summary_url=video_url
#             courseDet.teacher_id=teacherData.teacher_id
#             courseDet.school_id=teacherData.school_id
#             if idealfor:
#                 courseDet.ideal_for=idealfor
#             courseDet.course_status=course_status_id.msg_id
#             courseDet.is_private='N'
#             courseDet.image_url = imageUrl
#             courseDet.is_archived = 'N'
#             courseDet.difficulty_level=levelId.msg_id
#         db.session.commit()
#         updateIndex = updateSearchIndex("send","saveCourse")
#         if updateIndex==True:
#             return jsonify("1")
#         else:
#             return jsonify("2")



# @app.route('/saveCourse',methods=['GET','POST'])
# def saveCourse():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     print('inside saveCourse')
#     course = request.form.get('course')
#     courseId = request.args.get('course_id')
#     description = request.form.get('description')
#     # setDate = request.form.get('setDate')
#     # startTime = request.form.get('startTime')
#     # endTime = request.form.get('endTime')
#     # days = request.form.getlist('Days')
#     imageUrl = ''
#     imgUrl= request.form.get('imageUrl')
#     if imgUrl!=None:
#         imageUrl = imgUrl
#     video_url = request.form.get('videoUrl')
#     idealfor = request.args.get('idealfor')
#     level = request.form.get('level')
#     private = request.form.get('private')
#     print('Course name:'+str(course))
#     print('courseId:'+str(courseId))
#     print('description name:'+str(description))
#     print('Private:'+str(private))
#     print('course Image:'+str(imageUrl))
#     course_status = request.args.get('course_status')
#     print('course status:'+str(course_status))
   
#     print('video_url :'+str(video_url))
#     print('Ideal for:'+str(idealfor))
#     print('level:'+str(level))
#     levelId = MessageDetails.query.filter_by(description=level,category='Difficulty Level').first()
#     course_status_id = MessageDetails.query.filter_by(category='Course Status',description=course_status).first()
#     courseDet = CourseDetail.query.filter_by(course_id=courseId).first()
#     if private:
#         print('if course status is private')
#         courseDet.description=description
#         courseDet.summary_url=video_url
#         courseDet.teacher_id=teacherData.teacher_id
#         courseDet.school_id=teacherData.school_id
#         if idealfor:
#             courseDet.ideal_for=idealfor
#         courseDet.course_status=course_status_id.msg_id
#         courseDet.is_private='Y'
#         courseDet.image_url = imageUrl
#         courseDet.is_archived = 'N'
#         courseDet.difficulty_level=levelId.msg_id
#     else:
#         print('if course status is public')
#         courseDet.description=description
#         courseDet.summary_url=video_url
#         courseDet.teacher_id=teacherData.teacher_id
#         courseDet.school_id=teacherData.school_id
#         if idealfor:
#             courseDet.ideal_for=idealfor
#         courseDet.course_status=course_status_id.msg_id
#         courseDet.is_private='N'
#         courseDet.image_url = imageUrl
#         courseDet.is_archived = 'N'
#         courseDet.difficulty_level=levelId.msg_id
#     db.session.commit()
#     print('course:'+str(course))
#     print('Desc:'+str(description))
#     print('url:'+str(video_url))
#     print('teacher_id:'+str(teacherData.teacher_id))
#     print('school_id:'+str(teacherData.school_id))
#     print('idealfor:'+str(idealfor))
#     print('course_status:'+str(course_status_id.msg_id))    
#     return jsonify("1")

# @app.route('/courseEntry',methods=['GET','POST'])
# def courseEntry():
#     course = request.args.get('course')
#     course_id = "select max(course_id) as course_id from course_detail "
#     course_id = db.session.execute(text(course_id)).first()
#     courseDet = CourseDetail.query.filter_by(course_id=course_id.course_id).first()
#     courseDet.course_name = course
#     db.session.commit()
#     # courseId = "select max(course_id) as course_id from course_detail"
#     # courseId = db.session.execute(text(courseId)).first()
#     return jsonify(courseDet.course_id)

#     # return render_template('editCourse.html')


# @app.route('/myCourses')
# def myCourses():

#     return render_template('myCourses.html')

# @app.route('/paymentForm')
# def paymentForm():    
#     if current_user.is_authenticated:
#         #if current_user.country==None:
#         #    flash('Please update your profile before donating ')
#         #    return jsonify(['2'])

#         #qschool_id = request.args.get('school_id')
#         amount =  request.args.get('amount') 
#         qbatch_id = request.args.get('batch_id')                
#         #amount =              #hard coded value
#         #qbatch_id = 1               #hard coded value

#         if amount=='other':
#             amount = 0
#         #donation_for = request.args.get('donation_for')
#         #if donation_for =='' or donation_for=='undefined':            
#         #    donation_for=24
        
#         courseBatchData = CourseBatch.query.filter_by(batch_id = qbatch_id).first()
#         courseDetailData =CourseDetail.query.filter_by(course_id=courseBatchData.course_id).first()
#         schoolData = SchoolProfile.query.filter_by(school_id=courseDetailData.school_id).first() 
#         print("this is the batch id: "+str(courseBatchData.batch_id))

#         print("this is the course desc: "+str(courseDetailData.description))
#         #New section added to handle different vendors     
#         #if schoolData.curr_sub_charge_type== 41:
#         #    schoolShare = 100-int(schoolData.curr_sub_charge)
#         #    selfShare = schoolData.curr_sub_charge             
#         #else:

#         #every payment amout is to be split in the ratio 97:3 :: Tutor: allLearn
#         schoolShare = 97
#         selfShare = 3
#         vendorData = [
#             {
#                 "vendorId": schoolData.curr_vendor_id,
#                 "commission": 97 #int(schoolShare)
#             }, 
#             {
#                 "vendorId":"SELF",
#                 "commission":3 #int(selfShare)
#             }
#         ]

#         #vendorData=    [{"vendorId":"VENDOR1","commission":30}, {"vendorId":"VENDOR2","commission":40}]        
#         vendorData = json.dumps(vendorData, separators=(',', ':'))
#         print(vendorData)
#         vendorDataEncoded = base64.b64encode(vendorData.encode('utf-8')).decode('utf-8')
#         print(vendorDataEncoded)
#         #end of section

#         note = "Enrollment transaction"   
#         payer_name = current_user.first_name + ' ' + current_user.last_name

#         messageData = "" #MessageDetails.query.filter_by(msg_id=payment_for).first()

#         #Inserting new order and transaction detail in db

#         transactionNewInsert = PaymentTransaction(amount=courseBatchData.course_batch_fee,note=note, 
#             payer_user_id=current_user.id, payer_name=str(payer_name),payer_phone=current_user.phone, payer_email=current_user.email,
#             school_id=courseDetailData.school_id, teacher_id=courseDetailData.teacher_id,batch_id=qbatch_id, trans_type=254, payment_for= 264, tran_status=256, date=datetime.today()) 
#         db.session.add(transactionNewInsert)
#         db.session.commit()

#         #Fetching all required details for the form and signature creation

#         #transactionData = PaymentTransaction.query.filter_by(payer_user_id=current_user.id).order_by(PaymentTransaction.date.desc()).first()
#         transactionData  = transactionNewInsert
#         orderId= str(transactionData.tran_id).zfill(9)
#         print("#######order id: "+str(orderId))
#         currency = transactionData.currency
#         appId= app.config['ALLLEARN_CASHFREE_APP_ID']
#         returnUrl = url_for('paymentResponse',_external=True)
#         notifyUrl = url_for('notifyUrl',_external=True)
#         return render_template('_paymentForm.html',courseDetailData=courseDetailData,courseBatchData=courseBatchData, vendorDataEncoded=vendorDataEncoded,messageData=messageData,notifyUrl=notifyUrl,returnUrl=returnUrl, schoolData=schoolData, appId=appId, orderId = orderId, amount = amount, orderCurrency = currency, orderNote = note, customerName = payer_name)
#     else:
#         flash('Please login to enroll')
#         return jsonify(['1'])

# @app.route('/freeEnrollment')
# def freeEnrollment():
#     if current_user.is_authenticated:
#         batch_id = request.args.get('batch_id')
#         courseBatchData = CourseBatch.query.filter_by(batch_id =batch_id , is_archived='N').first()
#         courseBatchData.students_enrolled = int(courseBatchData.students_enrolled) + 1
#         courseEnrollmentData = CourseEnrollment(course_id= courseBatchData.course_id, 
#             batch_id = batch_id, student_user_id=current_user.id, is_archived='N', 
#             last_modified_date = datetime.today())   
#         db.session.add(courseEnrollmentData)
#         db.session.commit()
#         flash('Course Enrolled')
#         return jsonify(['0'])
#     else:
#         flash('Please login to enroll')
#         return jsonify(['1'])


# @app.route('/request', methods=["POST"])
# def handlerequest():
#     mode = app.config["MODE"] # <-------Change to TEST for test server, PROD for production
#     platformSub = request.args.get('platformSub')
#     if platformSub=="1":
#         postData = {
#             "appId" : request.form['appId'], 
#             "orderId" : request.form['orderId'], 
#             "orderAmount" : request.form['orderAmount'], 
#             "orderCurrency" : request.form['orderCurrency'], 
#             "orderNote" : request.form['orderNote'], 
#             "customerName" : request.form['customerName'], 
#             "customerPhone" : request.form['customerPhone'], 
#             "customerEmail" : request.form['customerEmail'], 
#             "returnUrl" : request.form['returnUrl'], 
#             "notifyUrl" : request.form['notifyUrl'],
#         }
#     else:
#         postData = {
#             "appId" : request.form['appId'], 
#             "orderId" : request.form['orderId'], 
#             "orderAmount" : request.form['orderAmount'], 
#             "orderCurrency" : request.form['orderCurrency'], 
#             "orderNote" : request.form['orderNote'], 
#             "customerName" : request.form['customerName'], 
#             "customerPhone" : request.form['customerPhone'], 
#             "customerEmail" : request.form['customerEmail'], 
#             "returnUrl" : request.form['returnUrl'], 
#             "notifyUrl" : request.form['notifyUrl'],
#             "vendorSplit" : request.form['vendorSplit']
#         }
#     #vendorSplit = request.form['vendorSplit']
#     sortedKeys = sorted(postData)
#     signatureData = ""
#     for key in sortedKeys:
#       signatureData += key+postData[key]
#     message = signatureData.encode('utf-8')
#     #get secret key from config
#     secret = app.config['ALLLEARN_CASHFREE_SECRET_KEY'].encode('utf-8')
#     signature = base64.b64encode(hmac.new(secret,message,digestmod=hashlib.sha256).digest()).decode("utf-8")   
            
#     transactionData = PaymentTransaction.query.filter_by(payer_user_id=current_user.id).order_by(PaymentTransaction.date.desc()).first()
#     transactionData.order_id=postData["orderId"]
#     #transactionData.anonymous_donor = anonymous_donor
#     #transactionData.anonymous_amount = hide_amount
#     transactionData.tran_status = 257 
#     transactionData.request_sign_hash = signature
#     transactionData.amount = postData["orderAmount"]

#     #updating user phone number
#     if current_user.phone==None or current_user.phone=="":
#         userDataUpdate = User.query.filter_by(id=current_user.id).first()
#         userDataUpdate.phone = request.form['customerPhone']
#     db.session.commit()

#     if mode == 'PROD': 
#       url = "https://www.cashfree.com/checkout/post/submit"
#     else: 
#       url = "https://test.cashfree.com/billpay/checkout/post/submit"
#     return render_template('request.html', postData = postData,signature = signature,url = url, platformSub=platformSub)


# #this is the page after response from payment gateway
# @app.route('/paymentResponse', methods=["POST"])
# def paymentResponse():
#     payment = request.args.get('payment')

#     postData = {
#     "orderId" : request.form['orderId'], 
#     "orderAmount" : request.form['orderAmount'], 
#     "referenceId" : request.form['referenceId'], 
#     "txStatus" : request.form['txStatus'], 
#     "paymentMode" : request.form['paymentMode'], 
#     "txMsg" : request.form['txMsg'], 
#     "signature" : request.form['signature'], 
#     "txTime" : request.form['txTime']
#     }

#     signatureData = ""
#     signatureData = postData['orderId'] + postData['orderAmount'] + postData['referenceId'] + postData['txStatus'] + postData['paymentMode'] + postData['txMsg'] + postData['txTime']

#     message = signatureData.encode('utf-8')
#     # get secret key from your config
#     secret = app.config['ALLLEARN_CASHFREE_SECRET_KEY'].encode('utf-8')
#     computedsignature = base64.b64encode(hmac.new(secret,message,digestmod=hashlib.sha256).digest()).decode('utf-8')   
#     print("####this is the txStatus: "+str(postData["txStatus"]))
#     messageData = MessageDetails.query.filter_by(description = postData["txStatus"]).first()

#     #updating response transaction details into the DB
#     transactionData = PaymentTransaction.query.filter_by(order_id=postData["orderId"]).first()
#     currency = transactionData.currency
 
#     transactionData.gateway_ref_id = postData["referenceId"]
#     if transactionData.tran_status!=263:
#         transactionData.tran_status = messageData.msg_id
#     transactionData.payment_mode = postData["paymentMode"]
#     transactionData.tran_msg = postData["txMsg"]
#     transactionData.tran_time = postData["txTime"]
#     transactionData.response_sign_hash = postData["signature"]
#     if postData["signature"]==computedsignature:
#         transactionData.response_sign_check="Matched"
#     else:
#         transactionData.response_sign_check="Not Matched"
#     schoolData = SchoolProfile.query.filter_by(school_id=transactionData.school_id).first()
#     if payment!='sub':
#         #updating school data
#         if transactionData.tran_status==258 or transactionData.tran_status==263:
#             courseBatchData = CourseBatch.query.filter_by(batch_id = transactionData.batch_id, is_archived='N').first()
#             if courseBatchData!=None:
#                 courseBatchData.total_fee_received = int(courseBatchData.total_fee_received)  + int(transactionData.amount)
#                 courseBatchData.students_enrolled = int(courseBatchData.students_enrolled) + 1

#                 courseEnrollmentData = CourseEnrollment(course_id= courseBatchData.course_id, batch_id = transactionData.batch_id, student_user_id=current_user.id, is_archived='N', 
#                     last_modified_date = datetime.today())   
#                 db.session.add(courseEnrollmentData)
#                 courseDataQuery = "select course_id, course_name, tp.teacher_id, teacher_name from course_detail cd"
#                 courseDataQuery = courseDataQuery + " inner join teacher_profile tp on tp.teacher_id=cd.teacher_id"
#                 courseDataQuery = courseDataQuery + " and course_id="+ str(courseBatchData.course_id) 
#                 courseData = db.session.execute(courseDataQuery).first()
#     db.session.commit()

#     return render_template('paymentResponse.html',courseData=courseData, courseBatchData=courseBatchData,transactionData = transactionData,payment=payment,postData=postData,computedsignature=computedsignature, schoolData=schoolData,currency=currency)




# @app.route('/notifyUrl',methods=["POST"])
# def notifyUrl():
#     postData = {
#       "orderId" : request.form['orderId'], 
#       "orderAmount" : request.form['orderAmount'], 
#       "referenceId" : request.form['referenceId'], 
#       "txStatus" : request.form['txStatus'], 
#       "paymentMode" : request.form['paymentMode'], 
#       "txMsg" : request.form['txMsg'], 
#       "txTime" : request.form['txTime'], 
#     }
    
#     transactionData = Transaction.query.filter_by(order_id = postData["orderId"]).first()
#     schoolData = SchoolProfile.query.filter_by(school_id=transactionData.school_id).first()
#     if transactionData!=None:
#         transactionData.tran_status = 263
#         db.session.commit()        
#         #donation_success_email_donor(schoolData.name, transactionData.donor_name,transactionData.donor_email,postData)
#     else:
#         print('############### no transaction detail found')
#     return str(0)


# def requestSignGenerator(appId, orderId, orderAmount, orderCurrency, orderNote, customerName, customerPhone,customerEmail, returnUrl, notifyUrl):    
#     postData = {
#       "appId" : appId,
#       "orderId" : orderId,
#       "orderAmount" : orderAmount,
#       "orderCurrency" : orderCurrency,
#       "orderNote" : orderNote,
#       "customerName" : customerName,
#       "customerPhone" : customerPhone,
#       "customerEmail" : customerEmail,
#       "returnUrl" : returnUrl,
#       "notifyUrl" : notifyUrl
#     }
#     sortedKeys = sorted(postData)
#     signatureData = ""
#     for key in sortedKeys:
#       signatureData += key+postData[key]

#     message = bytes(signatureData,encoding='utf-8')
#     #get secret key from your config
#     secret = bytes(app.config['ALLLEARN_CASHFREE_SECRET_KEY'],encoding='utf-8')
#     signature = base64.b64encode(hmac.new(secret, message,digestmod=hashlib.sha256).digest())
#     return signature


# def verifyResponseSign(receivedResponseSign, postData):
#     signatureData = postData["orderId"] + postData["orderAmount"] + postData["referenceId"] + postData["txStatus"] + postData["paymentMode"] + postData["txMsg"] + postData["txTime"]
#     message = bytes(signatureData).encode('utf-8')
#     #get secret key from your config
#     secret = bytes(app.config['ALLLEARN_CASHFREE_SECRET_KEY']).encode('utf-8')
#     signature = base64.b64encode(hmac.new(secret, message,digestmod=hashlib.sha256).digest())
#     if signature==receivedResponseSign:
#         return True
#     else:
#         return False


##Routes to manage class notifications
@app.route('/classNotification')
def classNotification():    
    return render_template('classNotification.html')

@app.route('/sendClassNotification')
def sendClassNotification():
    upcomingClassStudentsQuery = "select *from public.user"  ##This query needs to be replaced with a new one
    upcomngClassStudentsData = db.session.execute(upcomingClassStudentsQuery).fetchall()
    ##This section to send email notifications
    for val in upcomngClassStudentsData:
        send_notification_email(val.email, val.first_name+ str(' ')+val.last_name, 'Course')
    return jsonify(['0'])
##### end of openClass modules





# @app.route('/register', methods=['GET', 'POST'])
# def register():
#     if current_user.is_authenticated:
#         return redirect(url_for('index'))
    
#     #new arg added for google login
#     #gSigninData = request.args.get

#     #default form submit action
#     form = RegistrationForm() 
#     if form.validate_on_submit():
#         print('Validated form submit')
#         #we're setting the username as email address itself. That way a user won't need to think of a new username to register. 
#         #By default we're setting the user as course taker
#         user = User(username=form.email.data, email=form.email.data, user_type='140', access_status='145', phone=form.phone.data,
#             first_name = form.first_name.data,school_id=1,last_name= form.last_name.data)
#         user.set_password(form.password.data)
#         db.session.add(user)
#         db.session.commit()
#         #if a teacher has already been added during school registration then simply add the new user's id to it's teacher profile value        
#         checkTeacherProf = TeacherProfile.query.filter_by(email=form.email.data).first()
#         #if a student has already been added during school registration then simply add the new user's id to it's student profile value
#         checkStudentProf = StudentProfile.query.filter_by(email=form.email.data).first()

#         if checkTeacherProf!=None:
#             checkTeacherProf.user_id=user.id
#             db.session.commit()        
#         elif checkStudentProf!=None:
#             checkStudentProf.user_id=user.id
#             db.session.commit()
#         else:
#             pass

#         full_name = str(form.first_name.data)+ ' '+str(form.last_name.data)
#         flash('Congratulations '+full_name+', you are now a registered user!')
#         welcome_email(str(form.email.data), full_name)
#         return redirect(url_for('login'))
#     return render_template('register.html', title='Register', form=form)

# @app.route('/logout')
# @login_required
# def logout():
#     logout_user()
#     return redirect(url_for('index'))


@app.route('/teachingApplicantProfile/<user_id>')
@login_required
def teachingApplicantProfile(user_id):
    job_id = request.args.get('job_id')
    user = User.query.filter_by(id=user_id).first_or_404()
    accessingUser = User.query.filter_by(id=current_user.id).first_or_404()
    return render_template('teachingApplicantProfile.html',user=user, user_type_val=str(accessingUser.user_type),job_id=job_id)

@app.route('/user/<username>')
@login_required
def user(username):
    user = User.query.filter_by(username=username).first_or_404() 
    print('current_user.id:'+str(current_user.id))   
    teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    school_name_val = schoolNameVal()        
    disconn = ''
    user_type_val = ''
    if current_user.user_type==72:
        disconn = 1
        user_type_val = current_user.user_type
    if user.user_type==161:        
        return redirect(url_for('teachingApplicantProfile',title='My Profile',user_id=user.id))
    else:
        print('Nope we are not')

    if school_name_val ==None:
        print('did we reach here')
        return redirect(url_for('disconnectedAccount'))
    else:        
        schoolAdminRow = db.session.execute(text("select school_admin from school_profile where school_id ='"+ str(teacher.school_id)+"'")).fetchall()
        print(schoolAdminRow[0][0])
        accessRequestListRows=""
        value=0
        if current_user.user_type==72:
            value=1
        print('schoolAdminRow[0][0]:'+str(schoolAdminRow[0][0]))
        print('teacher.teacher_id:'+str(teacher.teacher_id))

        accessSchoolRequestListRows = ''

        if schoolAdminRow[0][0]==teacher.teacher_id:
            accessReqQuery = "select t1.username, t1.email, t1.phone, t2.description as user_type, t1.about_me, t1.school_id from public.user t1 inner join message_detail t2 on t1.user_type=t2.msg_id where t1.school_id='"+ str(teacher.school_id) +"' and t1.access_status='143'"
            print('accessReqQuery:'+str(accessReqQuery))
            accessRequestListRows = db.session.execute(text(accessReqQuery)).fetchall()
            accessSchoolReqQuery = "select t1.username, t1.email, t1.phone, t2.description as user_type, t1.about_me, t1.school_id from public.user t1 inner join message_detail t2 on t1.user_type=t2.msg_id inner join school_profile sp on t1.school_id = sp.school_id where t1.school_id='"+ str(teacher.school_id) +"' and sp.is_verified='N'"
            print('Query accessSchoolReqQuery:'+str(accessSchoolReqQuery))
            accessSchoolRequestListRows = db.session.execute(text(accessSchoolReqQuery)).fetchall()
        teacherData = "select distinct teacher_name, description as subject_name, cs.class_val, cs.section,cs.class_sec_id from teacher_subject_class tsc "
        teacherData = teacherData + "inner join teacher_profile tp on tsc.teacher_id = tp.teacher_id "
        teacherData = teacherData + "inner join class_section cs on tsc.class_sec_id = cs.class_sec_id "
        teacherData = teacherData + "inner join message_detail md on tsc.subject_id = md.msg_id where tsc.school_id = '"+str(teacher.school_id)+"' and tsc.teacher_id = '"+str(teacher.teacher_id)+"' and tsc.is_archived = 'N' order by cs.class_sec_id"
        teacherData = db.session.execute(text(teacherData)).fetchall()
        indic='DashBoard'
        return render_template('user.html',indic=indic,title='My Profile', classSecCheckVal=classSecCheck(),user=user,teacher=teacher,accessSchoolRequestListRows=accessSchoolRequestListRows,accessRequestListRows=accessRequestListRows, school_id=teacher.school_id,disconn=disconn,user_type_val=str(current_user.user_type),teacherData=teacherData)


# @app.route('/login', methods=['GET', 'POST'])
# def login():
#     print('Inside login')    
#     if current_user.is_authenticated:  
#         print(request.url)    
#         if current_user.user_type=='161':
#             return redirect(url_for('job_post.openJobs'))
#         else:
#             return redirect(url_for('index'))

#     #new section for google login
#     glogin = request.args.get('glogin')
#     gemail = request.args.get('gemail')
#     ##end of new section
    
    # form = LoginForm()
    # print('Validation')
    # print(form.validate_on_submit())
    # session['isGooglelogin'] = ''
    # if form.validate_on_submit() or glogin=="True":
    #     if glogin=="True":
    #         print("###glogin val"+ str(glogin))
    #         print("###email received from page"+ str(gemail))
    #         user=User.query.filter_by(email=gemail).first()   
    #         if user is None:
    #             flash("Email not registered")
    #             print('Email not registered')
    #             return redirect(url_for('login'))
    #     else: 
    #         print('Input data:'+str(form.email.data))
    #         checkEmailValidation = check(form.email.data)
    #         user = ''
    #         if checkEmailValidation == 'Y':
    #             user=User.query.filter_by(email=form.email.data).first() 
    #         else:
    #             Input = form.email.data
    #             print('Type:'+str(type(Input)))
    #             In = Input.upper()
    #             string = 'stud_'
    #             strg = string.upper()
    #             print('Type:'+str(type(strg))+'String:'+str(strg))
    #             if In.find(strg) == 0:
    #                 print('this is student id')
    #                 studentId = Input[5:]
    #                 print('studentId:'+str(studentId))
    #                 studData = StudentProfile.query.filter_by(student_id=studentId).first()
    #                 email = studData.email
    #                 user=User.query.filter_by(email=email).first() 
    #             else:
    #                 print('phone no')
    #                 user=User.query.filter_by(phone=Input).first()

  
#             try:             
#                 if user is None or not user.check_password(form.password.data):        
#                     flash("Invalid email or password")
#                     return redirect(url_for('login'))
#             except:
#                 flash("Invalid email or password")
#                 return redirect(url_for('login'))

#         #logging in the user with flask login
#         try:
#             login_user(user,remember=form.remember_me.data)
#         except:
#             flash("Invalid email or password")
#             return redirect(url_for('login'))

#         next_page = request.args.get('next')
#         print('next_page',next_page)
#         if not next_page or url_parse(next_page).netloc != '':
#             print('if next_page is not empty',next_page)
#             next_page = url_for('index')
        
#         #setting global variables
#         session['classSecVal'] = classSecCheck()
#         session['schoolName'] = schoolNameVal()
        
        # print('user name')
        # #print(session['username'])
        # school_id = ''
        # print('user type')
        # #print(session['userType'])
        # session['studentId'] = ''
        # if current_user.user_type==253:
        #     school_id=1
        # elif current_user.user_type==71:
        #     userProfileData = User.query.filter_by(id=current_user.id).first()
        #     school_id = userProfileData.school_id
        # elif current_user.user_type==134:
        #     studentProfileData = StudentProfile.query.filter_by(user_id=current_user.id).first()
        #     school_id = studentProfileData.school_id            
        #     session['studentId'] = studentProfileData.student_id
        # else:
        #     userData = User.query.filter_by(id=current_user.id).first()
        #     school_id = userData.school_id

        # school_pro = SchoolProfile.query.filter_by(school_id=school_id).first()
        # session['school_logo'] = ''
        # print('school_pro:'+str(school_pro))
        # session['isGooglelogin'] = ''
        # if school_pro:
        #     session['school_logo'] = school_pro.school_logo
        #     session['schoolPicture'] = school_pro.school_picture
        #     session['schoolName'] = school_pro.school_name
        #     session['font'] = school_pro.font
        #     print('session[font]:'+str(session['font']))
        #     session['primary_color'] = school_pro.primary_color
        #     session['isGooglelogin'] = school_pro.google_login
        #     print('session[isGooglelogin]:'+str(session['isGooglelogin']))
        #     print('school_pro.google_login:'+str(school_pro.google_login))
        #     session['show_school_name'] = school_pro.show_school_name
        #     teacherData = TeacherProfile.query.filter_by(teacher_id=school_pro.school_admin).first()
        #     userData = User.query.filter_by(id=teacherData.user_id).first()
        #     session['phone'] = userData.phone
        #     session['email'] = userData.email
        # print(session['primary_color'])
        # query = "select user_type,md.module_name,description, module_url, module_type from module_detail md inner join module_access ma on md.module_id = ma.module_id where user_type = '"+str(current_user.user_type)+"' and ma.is_archived = 'N' and md.is_archived = 'N' order by module_type"
        # print(query)
        # print('Modules')
        # moduleDetRow = db.session.execute(query).fetchall()
        # print('School profile')
        # #print(session['schoolPicture'])
        # # det_list = [1,2,3,4,5]
        # session['moduleDet'] = []
        # detList = session['moduleDet']
        
#         for det in moduleDetRow:
#             eachList = []
#             # print(det.module_name)
#             # print(det.module_url)
#             eachList.append(det.module_name)
#             eachList.append(det.module_url)
#             eachList.append(det.module_type)
#             # detList.append(str(det.module_name)+":"+str(det.module_url)+":"+str(det.module_type))
#             detList.append(eachList)
#         session['moduleDet'] = detList
#         # for each in session['moduleDet']:
#         #     print('module_name'+str(each[0]))
#         #     print('module_url'+str(each[1]))
#         #     print('module_type'+str(each[2]))
#         #print(session['schoolName'])

#         return redirect(next_page)        
#         #return redirect(url_for('index'))
#     # schoolDataQuery = "select *from school_profile"
#     # schoolData = db.session.execute(text(schoolDataQuery)).fetchall()
#     schoolName = ''
#     schoolLogo = ''
#     primaryColor = '' 
#     phone = ''
#     email = ''
#     print('Url:'+str(request.url))
#     subDom = request.url
#     newDom = 'login'
#     print('login:'+str(newDom))
#     newSubDom = subDom.partition(newDom)
#     newSub = newSubDom[0] + newSubDom[1]
#     print('newSubDom:'+str(newSub))
#     schoolDataQuery = "select *from school_profile where sub_domain like '"+str(newSub)+"%'"
#     schoolData = db.session.execute(text(schoolDataQuery)).fetchall()   
#     print(subDom)
#     font=''
#     for row in schoolData:
#         print(row)
#         if row:
#             schoolName = row.school_name
#             schoolLogo = row.school_logo
#             primaryColor = row.primary_color
#             font = row.font
#             print('font:'+str(font))
#             print('primaryColor:'+str(primaryColor))
#             teacherData = TeacherProfile.query.filter_by(teacher_id=row.school_admin).first()
#             userData = User.query.filter_by(id=teacherData.user_id).first()
#             phone = userData.phone
#             email = userData.email
#     print('phone:'+str(phone))
#     print('email:'+str(email))
#     return render_template('login.html',font=font,phone=phone,email=email,primaryColor=primaryColor,schoolName=schoolName,schoolLogo=schoolLogo, title='Sign In', form=form)

@app.route('/success',methods=['POST'])
def success():
    if request.method=='POST':
        email=request.form["email"]
        name=request.form["name"]
        if db.session.query(Survivor).filter(Survivor.sur_email == email).count() == 0:
            #Raw sql example  - db.engine.execute(text("<sql here>")).execution_options(autocommit=True))
            # possibly db.session.execute (text("<sql here>")).execution_options(autocommit=True))
            survivor = Survivor(email, name)
            db.session.add(survivor)
            db.session.commit()
            print(email,name)
            welcome_email(email, name)
            return render_template('newsletterSuccess.html')
        else:
            return render_template('index.html',text='Error: Email already used.')

@app.route('/setFee',methods=['GET','POST'])
def setFee():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
    class_val = request.args.get('class_val')
    section = request.args.get('section')
    total_fee = request.args.get('total_fee')
    classSec_id = ClassSection.query.filter_by(class_val=class_val,section=section,school_id=teacher_id.school_id).first()
    class_sec_id = classSec_id.class_sec_id
    sections = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).all()
    for section in sections:
        insertAmount = FeeClassSecDetail(class_sec_id=section.class_sec_id,class_val=section.class_val,section=section.section,is_current='Y',last_modified_date=datetime.now(),change_date=datetime.now(),amount=total_fee,school_id=teacher_id.school_id)
        db.session.add(insertAmount)
    db.session.commit()
    return jsonify(['0'])

@app.route('/feeManagement')
@login_required
def feeManagement():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
    distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher_id.school_id)+" GROUP BY class_val order by s")).fetchall() 
    classSections=ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
    qclass_val = request.args.get('class_val')
    qsection=request.args.get('section')
    fee = ''
    amount = FeeClassSecDetail.query.filter_by(class_val=qclass_val,section=qsection,school_id=teacher_id.school_id).first()
    print(amount)
    clas=db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher_id.school_id)+" GROUP BY class_val order by s")).first() 
    if amount:
        fee = amount.amount
        print('amount:'+str(fee))
    if qclass_val==None or qclass_val=='':
        qclass_val = clas.class_val
        qsection = 'A'
    indic='DashBoard'
    return render_template('feeManagement.html',indic=indic,qclass_val=qclass_val,qsection=qsection,distinctClasses=distinctClasses,classsections=classSections,fee=fee)


@app.route('/privacyPolicy')
def privacyPolicy():
    return render_template('privacyPolicy.html')

@app.route('/sendHelplineNotification',methods=['GET','POST'])
def sendHelplineNotification():
    if request.method == 'POST':
        jsonData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonData)
        z = json.loads(a)
        conList = []
        paramList = []
        print('data:')
        for con in z['contact'].values():
            conList.append(con)
        for dat in z['results'].values():
            paramList.append(dat)
        medOption = paramList[0]
        fileUrl = paramList[1]
        nameAddress = paramList[2]
        city = paramList[3]
        symptoms = paramList[4]
        services = paramList[5]
        # print('Medicine option:'+str(medOption))
        # print('fileUrl:'+str(fileUrl))
        # print('nameAddress:'+str(nameAddress))
        # print('city:'+str(city))
        # print('symptoms:'+str(symptoms))
        service = ''
        if services == '1':
            service = 'Medicine'
        elif services == '2':
            service = 'Hospital bed'
        elif services == '3':
            service = 'Oxygen cylinder'
        elif services == '4':
            service = 'List of active Vaccination and Corona test ( RT-PCR ) centre'
        elif services == '5':
            service = 'Food'
        elif services == '6':
            service = 'Talk to a doctor'
        elif services == '7':
            service = 'Talk to a coordinator'
        elif services == '8':
            service = 'I want to help volunteer'
        elif services == '9':
            service = 'I can provide meds/food/essential Items'
        
        medicine = ''
        if medOption == '1':
            medicine = 'Remdisivir'
        elif medOption == '2':
            medicine = 'Azithral'
        elif medOption == '3':
            medicine = 'Ivermectin'
        elif medOption == '4':
            medicine = 'Zincovit'
        elif medOption == '5':
            medicine = 'Vitamin C - Celene'
        elif medOption == '6':
            medicine = 'Tocilizumab'
        elif medOption == '7':
            medicine = 'other'
        subject = ''
        if medicine:
            subject = 'Medicine:'+str(medicine)+str('\n')
        if service:
            subject = subject + ';Services:'+str(service)+str('\n')
        fileValue = ''
        f = '@results.uploadedimage.url'
        if fileUrl.find(f) != 0:
            fileValue = fileUrl
            subject = str('\n')+str(subject) + str(';Document:')+str(fileValue)+str('\n')
        address = ''
        a = '@results.nameandaddress'
        if nameAddress.find(a) != 0:
            address = nameAddress
            subject = str('\n')+str(subject) + str(';Name and Address:')+str(address)+str('\n')
        c = ''
        ci = '@results.city' 
        if city.find(ci) != 0:
            c = city
            subject = str('\n')+str(subject) + str(';City:')+str(c)+str('\n')
        sym = ''
        s = '@results.symptoms'
        if symptoms.find(s) != 0:
            sym = symptoms
            if sym == '1':
                s = 'Cough/ Sore throat'
                subject = str('\n')+str(subject) + str(';Symptoms:')+str(s)+str('\n')
            elif sym == '2':
                s = 'Fever'
                subject = str('\n')+str(subject) + str(';Symptoms:')+str(s)+str('\n')
            elif sym == '3':
                s = 'Difficulty breathing'
                subject = str('\n')+str(subject) + str(';Symptoms:')+str(s)+str('\n')
            elif sym == '4':
                s = 'Other'
                subject = str('\n')+str(subject) + str(';Symptoms:')+str(s)+str('\n')
        contactNo = conList[2]
        print('phone:'+str(contactNo))
        print('Subject:'+str(subject))
        email = 'contact@alllearn.in'
        email2 = 'paragsinha+oipkui0jrvwcrso3giqe@boards.trello.com'
        notificationHelplineEmail(email,email2,nameAddress,contactNo,subject)
        return jsonify({'phone':contactNo})                

@app.route('/sendUserNotificationEmail',methods=['POST','GET'])
def sendUserNotificationEmail():
    if request.method == 'POST':
        jsonData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonData)
        z = json.loads(a)
        conList = []
        paramList = []
        print('data:')
        for con in z['contact'].values():
            conList.append(con)
        for dat in z['results'].values():
            paramList.append(dat)
        subject = paramList[0]
        print(conList)
        contactNo = conList[2]
        print('phone:'+str(contactNo))
        name = conList[1]
        print('name:'+str(name))
        email = 'contact@alllearn.in'
        email2 = 'paragsinha+w6uwk6zar1ell7m5oemd@boards.trello.com'
        notificationEmail(email,email2,name,contactNo,subject)
        return jsonify({'phone':contactNo,'name':name})

@app.route('/sendNotificationEmail')
def sendNotificationEmail():
    student_id=request.args.get('student_id')
    resp_session_id = request.args.get('resp_session_id')
    userId = User.query.filter_by(id=current_user.id).first()
    school = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    school_id = school.school_id
    adminEmail=db.session.execute(text("select t2.email,t2.full_name,t1.school_name from school_profile t1 inner join student_profile t2 on t1.school_id=t2.school_id where t1.school_id='"+str(school_id)+"' and t2.student_id='"+str(student_id)+"'")).first()
    testDate = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    respCapture = ResponseCapture.query.filter_by(resp_session_id=resp_session_id).first()
    subject = MessageDetails.query.filter_by(msg_id=respCapture.subject_id).first()
    testType = TestDetails.query.filter_by(test_id=testDate.test_id).first()
    print('Test Type:'+str(testType.test_type))
    print('Response_session_id:'+str(resp_session_id))
    print('Student_id:'+str(student_id))
    if adminEmail!=None:
        test_report_email(adminEmail.email,adminEmail.full_name, adminEmail.school_name,school_id,testDate.last_modified_date,testType.test_type,resp_session_id,student_id,subject.description)
        return jsonify(["0"])
    else:
        return jsonify(["1"])

@app.route('/sendPerformanceReportEmail')
def sendPerformanceReportEmail():
    school_id = request.args.get('school_id')
    schoolAdmin = "select teacher_name as admin from school_profile sp inner join teacher_profile tp on sp.school_admin = tp.teacher_id where sp.school_id='"+str(school_id)+"'"
    schoolAdmin = db.session.execute(schoolAdmin).first()
    score = "Select avg_score from fn_overall_performance_summary("+str(school_id)+") where class='All'and section='All' and subject='All'"
    avg_score = db.session.execute(score).first()
    #####Fetch Top 5 Students info##########        
    # topStudentsQuery = "select *from fn_monthly_top_students("+str(teacher.school_id)+",8)"
    qclass_val = 'dashboard'
    topStudentsRows = ''
    leaderBoardData = leaderboardContent(qclass_val)

    # Convert dataframe to a list

    df1 = leaderBoardData[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
    df2 = leaderBoardData.drop(['profile_pic', 'student_name','class_val','section','total_marks%','total_tests'], axis=1)
    leaderBoard = pd.merge(df1,df2,on=('studentid'))
            
    d = leaderBoard[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
    df3 = leaderBoard.drop(['studentid'],axis=1)
            # print('DF3:')
            # print(df3)
            # print('print new dataframe')
            
    df1.rename(columns = {'profile_pic':'Profile Picture'}, inplace = True)
    df1.rename(columns = {'student_name':'Student'}, inplace = True)
    df1.rename(columns = {'class_val':'Class'}, inplace = True)
    df1.rename(columns = {'section':'Section'}, inplace = True)
    df1.rename(columns = {'total_marks%':'Total Marks'}, inplace = True)
    df1.rename(columns = {'total_tests':'Total Tests'}, inplace = True)
            # print(df1)
            # print('Excluding columns')
            # print(df2)
            # rename(df2)
            # print('LeaderBoard Data:')
            # print(leaderBoardData)
    data = []
            

    header = [df1.columns.values.tolist()]
    headerAll = [df3.columns.values.tolist()]
    colAll = ''
    subjHeader = [df2.columns.values.tolist()]
    columnNames = ''
    col = ''
    subColumn = ''
            # print('Size of dataframe:'+str(len(subjHeader)))
    for subhead in subjHeader:
        subColumn = subhead
                # print('Header with Subject Name')
                # print(subhead)
    for h in header:
        columnNames = h
    for headAll in headerAll: 
        colAll = headAll
            # print(' all header Length:'+str(len(colAll))+'Static length:'+str(len(columnNames))+'sub header length:'+str(len(subColumn)))
    n= int(len(subColumn)/2)
    ndf = df2.drop(['studentid'],axis=1)
    newDF = ndf.iloc[:,0:n]
    new1DF = ndf.iloc[:,n:]
            
    df5 = pd.concat([newDF, new1DF], axis=1)
    DFW = df5[list(sum(zip(newDF.columns, new1DF.columns), ()))]
           
           
    dat = pd.concat([d,DFW], axis=1)
    dat = dat.sort_values('total_marks%',ascending=False)  
    print(dat)        
    subHeader = ''
    i=1
    for row in dat.values.tolist():
        if i<6:
            data.append(row)
        i=i+1

    for d in data:
        print('Print Dataframe Data:')
        print(d[0])
        print(d[1])
        print(d[2])
        print(d[3])
            # print(d[i])

        # End  
    query = "select sum(total) as total_test from "
    query = query + "(select count(*) as total from response_capture rc where last_modified_date >current_date - 7 and school_id='"+str(school_id)+"'"
    query = query + " union all "
    query = query + "select count(*) from (select upload_id from result_upload ru where school_id='"+str(school_id)+"' and last_modified_date >current_date - 7"      
    query = query + " group by ru.upload_id)d ) s"
    test_count = db.session.execute(query).first()
    adminEmail=db.session.execute(text("select t2.email,t2.teacher_name,t1.school_name,t3.username from school_profile t1 inner join teacher_profile t2 on t1.school_admin=t2.teacher_id inner join public.user t3 on t2.email=t3.email where t1.school_id='"+str(school_id)+"'")).first()

    if adminEmail!=None:
        performance_report_email(adminEmail.email,adminEmail.teacher_name, adminEmail.school_name,data,test_count.total_test,avg_score.avg_score,school_id)
        return jsonify(["0"])
    else:
        return jsonify(["1"])
    


@app.route('/requestUserAccess')
def requestUserAccess():
    requestorUsername=request.args.get('username')    
    school_id=request.args.get('school_id')    
    about_me=request.args.get('about_me')    
    quser_type = request.args.get('user_type')        
    adminEmail=db.session.execute(text("select t2.email,t2.teacher_name,t1.school_name,t3.username from school_profile t1 inner join teacher_profile t2 on t1.school_admin=t2.teacher_id inner join public.user t3 on t2.email=t3.email where t1.school_id='"+school_id+"'")).first()
    print(adminEmail)
    print('User Type:'+str(quser_type))
    
    if adminEmail!=None:
        userTableDetails = User.query.filter_by(username=requestorUsername).first()
        userTableDetails.school_id=school_id
        userTableDetails.access_status='143'
        userTableDetails.user_type=quser_type
        userTableDetails.about_me=about_me
        db.session.commit()
        
        user_access_request_email(adminEmail.email,adminEmail.teacher_name, adminEmail.school_name, userTableDetails.first_name+ ''+userTableDetails.last_name, adminEmail.username, quser_type)
        return jsonify(["0"])
    else:
        return jsonify(["1"])



# @app.route('/syllabus')
# @login_required
# def syllabus():
#     fromSchoolRegistration = False
#     subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     boardRows = MessageDetails.query.filter_by(msg_id=board.board_id).first()
#     school_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     classValues = "SELECT class_val,sum(class_sec_id) as s FROM class_section cs where school_id = '"+str(teacher_id.school_id)+"' group by class_val order by s"
#     classValues = db.session.execute(text(classValues)).fetchall()
#     classValuesGeneral = "SELECT class_val,sum(class_sec_id) as s FROM class_section cs group by class_val order by s"
#     classValuesGeneral = db.session.execute(text(classValuesGeneral)).fetchall()
#     subjectValues = MessageDetails.query.filter_by(category='Subject').all()
#     bookName = BookDetails.query.all()
#     chapterNum = Topic.query.distinct().all()
#     topicId = Topic.query.all()
#     generalBoardId = SchoolProfile.query.with_entities(SchoolProfile.board_id).filter_by(school_id=teacher_id.school_id).first()
#     generalBoard = MessageDetails.query.filter_by(msg_id=generalBoardId.board_id).first()
#     for clas in classValues:
#         print('Class value:'+str(clas.class_val))
#     indic='DashBoard'
#     return render_template('syllabus.html',indic=indic,title='Syllabus',generalBoard=generalBoard,boardRowsId = boardRows.msg_id , boardRows=boardRows.description,subjectValues=subjectValues,school_name=school_id.school_name,classValues=classValues,classValuesGeneral=classValuesGeneral,bookName=bookName,chapterNum=chapterNum,topicId=topicId,fromSchoolRegistration=fromSchoolRegistration,user_type_val=str(current_user.user_type))

# @app.route('/addSyllabus',methods=['GET','POST'])
# def addSyllabus():
#     print('inside add syllabus')
#     classes = request.get_json()
#     # class_val = request.args.get('class_val')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     for class_val in classes:
#         classExist = ClassSection.query.filter_by(class_val=class_val,section='A',school_id=teacher_id.school_id).first()
#         if classExist == None:
#             addClass = ClassSection(class_val=class_val,section='A',school_id=teacher_id.school_id,student_count=0,class_teacher=teacher_id.teacher_id,last_modified_date=datetime.now())
#             db.session.add(addClass)
#             db.session.commit()
#     # class_sec_id = ClassSection.query.filter_by(class_val=class_val,section='A',school_id=teacher_id.school_id,student_count=0,class_teacher=teacher_id.teacher_id).first()
#     # for subject in subjects:
#     #     subject_id = MessageDetails.query.filter_by(description=subject,category='Subject').first()
#     #     subjExist = ''
#     #     subjExist = BoardClassSubject.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,school_id=teacher_id.school_id).first()
#     #     print(subjExist)
#     #     if subjExist == None:
#     #         print('is subjExist is null')
#     #         addSubject = BoardClassSubject(board_id=board_id.board_id,class_val=class_val,subject_id=subject_id.msg_id,is_archived='N',school_id=teacher_id.school_id,last_modified_date=datetime.now())
#     #         db.session.add(addSubject)
#     #         db.session.commit()
#         # bookNames = BookDetails.query.distinct(BookDetails.book_name).filter_by(subject_id=subject_id.msg_id,class_val=class_val).all()
#         print('after add subjects')
#         # for book_name in bookNames:
#         #     book_id = BookDetails.query.filter_by(subject_id=subject_id.msg_id,class_val=class_val,book_name=book_name.book_name).first()
#         #     addBook = BoardClassSubjectBooks(school_id=teacher_id.school_id,class_val=class_val,subject_id=subject_id.msg_id,book_id=book_id.book_id,is_archived='N')
#         #     db.session.add(addBook)
#         #     db.session.commit()
#         # insertRow = "insert into topic_tracker (subject_id, class_sec_id, is_covered, topic_id, school_id, reteach_count,is_archived, last_modified_date) (select subject_id, '"+str(class_sec_id.class_sec_id)+"', 'N', topic_id, '"+str(teacher_id.school_id)+"', 0,'N',current_date from Topic_detail where class_val="+str(class_val)+")"
#         # db.session.execute(text(insertRow))
#         # db.session.commit()
#     return ("Syllabus added successfully")


# @app.route('/generalSyllabusClasses')
# def generalSyllabusClasses():
#     board_id=request.args.get('board_id')
#     classArray = []
#     distinctClasses = "SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs GROUP BY class_val order by s"
#     distinctClasses = db.session.execute(text(distinctClasses)).fetchall()
#     for val in distinctClasses:
#         classArray.append(val.class_val)
#     if classArray:
#         return jsonify([classArray])
#     else:
#         return ""

# @app.route('/syllabusClasses')
# @login_required
# def syllabusClasses():
#     board_id=request.args.get('board_id')
#     classSectionArray = []
#     sectionArray = []
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     distinctClasses = "SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id = '"+str(teacher_id.school_id)+"' GROUP BY class_val order by s"
#     distinctClasses = db.session.execute(text(distinctClasses)).fetchall()
#     for val in distinctClasses:
#         #print(val.class_val)
#         sections = ClassSection.query.distinct(ClassSection.section).filter_by(school_id=teacher_id.school_id,class_val=val.class_val).all()
#         # sectionsString = ''
#         sectionsString = '['
#         i=1
#         for section in sections:
#             #print(len(sections))
#             if i<len(sections):
#                 sectionsString = sectionsString + str(section.section)+';'
#             else:
#                 sectionsString = sectionsString + str(section.section)
#             i = i + 1
#         sectionsString = sectionsString + ']'
#         classSectionArray.append(str(val.class_val)+':'+str(sectionsString))
#     if classSectionArray:
#         return jsonify([classSectionArray])
#     else:
#         return ""


# @app.route('/generalSyllabusSubjects',methods=['GET','POST'])
# def generalSyllabusSubjects():
#     board_id=request.args.get('board_id')
#     class_val=request.args.get('class_val')
#     sujectArray=[]
#     subjects = "select distinct description,msg_id from message_detail md inner join topic_detail td on md.msg_id = td.subject_id where td.class_val = '"+str(class_val)+"' order by description"
#     subjects = db.session.execute(text(subjects)).fetchall()
#     for val in subjects:
#         # subject = MessageDetails.query.filter_by(msg_id=val.subject_id).first()
#         sujectArray.append(str(val.msg_id)+":"+str(val.description))
#     if sujectArray:
#         return jsonify([sujectArray])   
#     else:
#         return ""

# @app.route('/syllabusSubjects')
# @login_required
# def syllabusSubjects():
#     board_id=request.args.get('board_id')
#     class_val=request.args.get('class_val')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     distinctSubject = BoardClassSubject.query.filter_by(class_val=class_val,board_id=board_id,school_id=teacher_id.school_id,is_archived='N').all()
#     sujectArray=[]
#     subjects = "select distinct description,msg_id from message_detail md inner join board_class_subject bcs on md.msg_id = bcs.subject_id where bcs.class_val = '"+str(class_val)+"' and school_id='"+str(teacher_id.school_id)+"' and bcs.is_archived= 'N' order by description"
#     subjects = db.session.execute(text(subjects)).fetchall()
#     for val in subjects:
#         # subject = MessageDetails.query.filter_by(msg_id=val.subject_id).first()
#         sujectArray.append(str(val.msg_id)+":"+str(val.description))
#     if sujectArray:
#         return jsonify([sujectArray])   
#     else:
#         return ""

# @app.route('/fetchSubjects',methods=['GET','POST'])
# def fetchSubjects():
#     class_val = request.args.get('class_val')
#     board = request.args.get('board')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     distinctSubject = BoardClassSubject.query.filter_by(class_val=class_val,board_id=board,school_id=teacher_id.school_id,is_archived='N').all()
#     sujectArray=[]
#     subjects = "select distinct description,msg_id from message_detail md inner join board_class_subject bcs on md.msg_id = bcs.subject_id where bcs.class_val = '"+str(class_val)+"' and school_id='"+str(teacher_id.school_id)+"' and bcs.is_archived= 'N' order by description"
#     subjects = db.session.execute(text(subjects)).fetchall()
#     for val in subjects:
#         # subject = MessageDetails.query.filter_by(msg_id=val.subject_id).first()
#         sujectArray.append(str(val.msg_id)+":"+str(val.description))
#     if sujectArray:
#         return jsonify([sujectArray])   
#     else:
#         return "" 

# @app.route('/fetchRemSubjects',methods=['GET','POST'])
# def fetchRemSubjects():
#     print('inside fetchRemSubjects')
#     class_val = request.args.get('class_val')
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher.school_id).first()
#     distinctSubject = BoardClassSubject.query.filter_by(class_val=class_val,board_id=board_id.board_id,school_id=teacher.school_id,is_archived='Y').all()
#     subjectArray=[]
#     generalSubjects = "select distinct msg_id,description from topic_detail td inner join message_detail md on md.msg_id=td.subject_id "
#     generalSubjects = generalSubjects + "where md.msg_id not in (select distinct msg_id from message_detail md "
#     generalSubjects = generalSubjects + "inner join board_class_subject bcs on md.msg_id = bcs.subject_id where bcs.class_val = '"+str(class_val)+"' and school_id='"+str(teacher.school_id)+"' "
#     generalSubjects = generalSubjects + ")  order by description"
#     print('Query: '+str(generalSubjects))
#     generalSubjects = db.session.execute(text(generalSubjects)).fetchall()
#     subjects = "select distinct description,msg_id from message_detail md inner join board_class_subject bcs on md.msg_id = bcs.subject_id where bcs.class_val = '"+str(class_val)+"' and school_id='"+str(teacher.school_id)+"' and bcs.is_archived= 'Y' order by description"
#     print(subjects)
#     subjects = db.session.execute(text(subjects)).fetchall()
#     for val in subjects:
#         # subject = MessageDetails.query.filter_by(msg_id=val.subject_id).first()
#         subjectArray.append(str(val.msg_id)+":"+str(val.description))
#     for val in generalSubjects:
#         subjectArray.append(str(val.msg_id)+":"+str(val.description))
#     if subjectArray:
#         return jsonify([subjectArray])   
#     else:
#         return ""

# @app.route('/addSubject',methods=['GET','POST'])
# def addSubject():
#     subject_id = request.args.get('subject')
#     board_id=request.args.get('board')
#     class_val=request.args.get('class_val')
#     print('Subject:'+str(subject_id))
#     # subject_id = MessageDetails.query.filter_by(description=subjectVal,category='Subject').first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     subjExist = BoardClassSubject.query.filter_by(class_val=class_val,board_id=board_id,subject_id=subject_id,school_id=teacher_id.school_id).first()
#     if subjExist==None:
#         addSubject = BoardClassSubject(class_val=class_val,subject_id=subject_id,school_id=teacher_id.school_id,board_id=board_id,is_archived='N')
#         db.session.add(addSubject)
#         db.session.commit()
#     else:
#         insertSubject = BoardClassSubject.query.filter_by(class_val=class_val,subject_id=subject_id,school_id=teacher_id.school_id,board_id=board_id,is_archived='Y').first()
#         insertSubject.is_archived = 'N'
#         db.session.add(insertSubject)
#         db.session.commit()
#     return ('update data successfully')

# @app.route('/addChapter',methods=['GET','POST'])
# def addChapter():
#     topics=request.get_json()
#     print('inside add Chapter')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     chapterName = request.args.get('chapterName')
    
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     chapter_num = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,chapter_name=chapterName).first()
#     print(topics)
#     print('School id:'+str(teacher_id.school_id))
#     for topic in topics:
#         print('inside for')
#         print(topic)
#         # topic_id = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,topic_name=topic).first()
#         existInTT = TopicTracker.query.filter_by(topic_id=topic,school_id=teacher_id.school_id,class_sec_id=class_sec_id.class_sec_id,subject_id=subject_id.msg_id).first()
        
#         if existInTT:
#             updateTT = "update topic_tracker set is_archived='N' where school_id='"+str(teacher_id.school_id)+"' and subject_id='"+str(subject_id.msg_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and topic_id='"+str(topic)+"'"
#             print(updateTT)
#             updateTT = db.session.execute(text(updateTT))
#         else:
#             insertTT = TopicTracker(subject_id=subject_id.msg_id,class_sec_id=class_sec_id.class_sec_id,is_covered='N',topic_id=topic,school_id=teacher_id.school_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(insertTT)
#         db.session.commit()
#     return ("data updated successfully")

# @app.route('/addBook',methods=['GET','POST'])
# def addBook():
#     book_id = request.args.get('book')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     print("class_val"+str(class_val))
#     print("subject id"+str(subject_id.msg_id))
#     print("book id"+str(book_id))
#     book = BookDetails.query.filter_by(book_id=book_id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id).all()
    
#     for book_id in bookIds:
#         updateBCSB = BoardClassSubjectBooks.query.filter_by(school_id=teacher_id.school_id,class_val=class_val,
#         subject_id=subject_id.msg_id,book_id=book_id.book_id).first()
#         if updateBCSB:
#             updateBCSB.is_archived = 'N'
#         else:
#             addBook = BoardClassSubjectBooks(school_id=teacher_id.school_id,class_val=class_val,subject_id=subject_id.msg_id,book_id=book_id.book_id,is_archived='N',last_modified_date=datetime.now())
#             db.session.add(addBook)
#         db.session.commit()
#     return ("data updated successfully")



# @app.route('/addNewSubject',methods=['GET','POST'])
# def addNewSubject():
#     subject = request.args.get('subject')
#     subject = subject.title()
#     class_val = request.args.get('class_val')
#     board = request.args.get('board')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     insertSubject = MessageDetails(category='Subject',description=subject)
#     db.session.add(insertSubject)
#     db.session.commit()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     insertBCS = BoardClassSubject(class_val=class_val,subject_id=subject_id.msg_id,school_id=teacher_id.school_id,board_id=board,is_archived='N')
#     db.session.add(insertBCS)
#     db.session.commit()
#     return ('New Subject added successfully')

# @app.route('/addNewBook',methods=['GET','POST'])
# def addNewBook():
#     bookName = request.args.get('book')
#     bookLink = request.args.get('bookLink')
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     bookName = bookName.strip()
#     for x in bookName.lower(): 
#         if x in punctuations: 
#             bookName = bookName.replace(x, "") 
#             print(bookName)
#         else:
#             break
#     if bookName==None or bookName=='':
#         return "NA"
#     bookName = bookName.strip()
#     bookLink = bookLink.strip()
#     for x in bookLink.lower(): 
#         if x in punctuations: 
#             bookLink = bookLink.replace(x, "") 
#             print(bookLink)
#         else:
#             break
#     bookLink = bookLink.strip()
#     book = bookName.title()
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     print('class in addNewBook:'+str(class_val))
#     subject_id= MessageDetails.query.filter_by(description=subject).first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     bookExist = BookDetails.query.filter_by(board_id=board_id.board_id,book_name=bookName,class_val=class_val,subject_id=subject_id.msg_id).first()
#     if bookExist==None:
#         if bookLink:
#             insertBook = BookDetails(board_id=board_id.board_id,book_name=book,class_val=class_val,subject_id=subject_id.msg_id,teacher_id=teacher_id.teacher_id,book_link=bookLink,last_modified_date=datetime.now())
#         else:
#             insertBook = BookDetails(board_id=board_id.board_id,book_name=book,class_val=class_val,subject_id=subject_id.msg_id,teacher_id=teacher_id.teacher_id,last_modified_date=datetime.now())
#         db.session.add(insertBook)
#         db.session.commit()
#         book_id = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_name=book).first()
#         insertInBCSB = BoardClassSubjectBooks(school_id=teacher_id.school_id,class_val=class_val,subject_id=subject_id.msg_id,
#         book_id=book_id.book_id,is_archived='N',last_modified_date=datetime.now())
#         db.session.add(insertInBCSB)
#         db.session.commit()
#     return ('New Book added successfully')

# @app.route('/checkForChapter',methods=['GET','POST'])
# def checkForChapter():
#     print('inside checkForChapter')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     chapterNum = request.args.get('chapter_num')
#     bookId = request.args.get('bookId')
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     for x in chapterNum: 
#         if x in punctuations: 
#             chapterNum = chapterNum.replace(x, "") 
#             print(chapterNum)
#             return "NA"
#         else:
#             break
#     chapterName = request.args.get('chapter_name')
#     for x in chapterName.lower(): 
#         if x in punctuations: 
#             chapterName = chapterName.replace(x, "") 
#             print(chapterName)
#             return "NA"
#         else:
#             break
#     subject_id= MessageDetails.query.filter_by(description=subject).first()
#     print('Book Id:'+str(bookId))
#     book = BookDetails.query.filter_by(book_id=bookId).first()
#     bookIds = BookDetails.query.filter_by(class_val=class_val,book_name=book.book_name,subject_id=subject_id.msg_id).all()
#     print('class_val:'+str(class_val)+'subject:'+str(subject_id.msg_id)+'Book name:'+str(book.book_name))
#     # bookIds = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_name=book.book_name).all()
#     k = 0
#     print(book.book_name)
#     for book_id in bookIds:
#         print(str(class_val)+' '+str(subject_id.msg_id)+' '+str(chapterNum)+' '+str(book_id.book_id))
#         topic1 = "select chapter_name,topic_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val='"+str(class_val)+"' and td.subject_id='"+str(subject_id.msg_id)+"' and td.book_id='"+str(book_id.book_id)+"' and tt.is_archived='N' and td.chapter_num='"+str(chapterNum)+"' "
#         topic1 = db.session.execute(text(topic1)).first()
#         topic2 = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,chapter_name=chapterName,book_id=book_id.book_id).first()
#         print('inside for')
#         print(book_id.book_id)
#         print(topic1)
#         if topic1 or topic2:
#             k = 1
#     print(k)
#     if k==1:
#         return ""
#     else:
#         return "1"

# @app.route('/addClassSection',methods=['POST'])
# def addClassSection():
#     print('inside addClassSection')
#     sections=request.get_json()
#     class_val = request.args.get('class_val')
#     print('class values:'+str(class_val))
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     for section in sections:
#         # class_section = class_section.split(':')
#         # class_val = class_section[0]
#         # section = class_section[1]
#         checkClass = ClassSection.query.filter_by(class_val=str(class_val),section=section.upper(),school_id=teacher_id.school_id).first()
#         if checkClass:
#             return ""
#     for section in sections:
#         # print(section)
#         # class_section = class_section.split(':')
#         # class_val = class_section[0]
#         # section = class_section[1]
        
#         print('Class:'+str(class_val)+' Section:'+str(section))
#         class_data=ClassSection(class_val=str(class_val),section=str(section).upper(),student_count=0,school_id=teacher_id.school_id,last_modified_date=datetime.now())
#         db.session.add(class_data)
#         db.session.commit()
    
#     for section in sections:
#         # class_section = class_section.split(':')
#         # class_val = class_section[0]
#         # section = class_section[1]
#         class_id = ClassSection.query.filter_by(class_val=str(class_val),section=section.upper(),school_id=teacher_id.school_id).first()
#         topic_tracker = TopicTracker.query.filter_by(class_sec_id=class_id.class_sec_id,school_id=teacher_id.school_id).first()
#         if topic_tracker:
#             print('data already present')
#         else:
#             print('insert data into topic tracker')
#             insertRow = "insert into topic_tracker (subject_id, class_sec_id, is_covered, topic_id, school_id, reteach_count, last_modified_date) (select subject_id, '"+str(class_id.class_sec_id)+"', 'N', topic_id, '"+str(teacher_id.school_id)+"', 0,current_date from Topic_detail where class_val='"+str(class_val)+"')"
#             db.session.execute(text(insertRow))
#         db.session.commit()   

#     return "success"

# @app.route('/checkForBook',methods=['GET','POST'])
# def checkForBook():
#     book = request.args.get('book')
#     book = book.title()
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     bookName = book.strip()
#     for x in bookName.lower(): 
#         if x in punctuations: 
#             bookName = bookName.replace(x, "") 
#             print(bookName)
#         else:
#             break
#     if bookName==None or bookName=='':
#         return "NA"
#     subject_id = MessageDetails.query.filter_by(category='Subject',description=subject).first()
#     checkBook = BookDetails.query.filter_by(book_name=bookName,class_val=class_val,subject_id=subject_id.msg_id).first()
#     if checkBook:
#         return (book)
#     else:
#         return ""

# @app.route('/checkForSubject',methods=['GET','POST'])
# def checkForSubject():
#     subject = request.args.get('subject')
#     subject = subject.title()
#     print('inside check for subject:'+str(subject))
#     class_val = request.args.get('class_val')
#     board = request.args.get('board')
#     checkSubject = MessageDetails.query.filter_by(category='Subject',description=subject).first()
#     if checkSubject:
#         return (subject)
#     else:
#         return ""

# @app.route('/checkforClassSection',methods=['GET','POST'])
# def checkforClassSection():
#     sections=request.get_json()
#     class_val = request.args.get('class_val')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     print('inside checkforClassSection')
#     print(sections)
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     class_val = class_val.strip()
#     print('before remove punc class:'+str(class_val))
#     if class_val==None or class_val=='':
#         print('if clas_val is none')
#         return "NB"
    
#     for x in class_val: 
#         if x in punctuations: 
#             class_val = class_val.replace(x, "") 
#             print('after remove punc class:'+str(class_val))
#             return "NA"
#         else:
#             break
#     for section in sections:
#         # class_section = class_section.split(':')
#         # class_val = class_section[0]
#         # section = class_section[1]
#         punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#         section = section.strip()
#         for x in section.lower(): 
#             if x in punctuations: 
#                 section = section.replace(x, "") 
#                 print(section)
#                 return "NA"
#             else:
#                 break
#         if section==None or section=='':
#             print('if section is none')
#             return "NB"
#         print('class_val:'+class_val)
#         print('section:'+section.upper())
#         checkClass = ClassSection.query.filter_by(class_val=str(class_val),section=section.upper(),school_id=teacher_id.school_id).first()
#         if checkClass:
#             return str(class_val)+':'+str(section.upper())
#     return ""


# @app.route('/addNewTopic',methods=['GET','POST'])
# def addNewTopic():
#     print('inside add new topic')
#     topics=request.get_json()
#     book_id = request.args.get('book_id')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     chapter = request.args.get('chapter')
#     chapter_num = request.args.get('chapter_num')
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     subject_id = MessageDetails.query.filter_by(description = subject).first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     book = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_id=book_id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     bookId = "select distinct bd.book_id from book_details bd inner join topic_detail td on td.book_id = bd.book_id where td.subject_id = '"+str(subject_id.msg_id)+"' and td.class_val  = '"+str(class_val)+"' and chapter_num = '"+str(chapter_num)+"' and bd.book_name = '"+str(book.book_name)+"'"
#     bookId = db.session.execute(text(bookId)).first()
#     print(topics)
#     print('Book ID:'+str(bookId.book_id))
#     for topic in topics:
#         print(topic)
#         topic = topic.strip()
#         for x in topic: 
#             if x in punctuations: 
#                 topic = topic.replace(x, "") 
#                 print(topic)
#             else:
#                 break
#         topic = topic.strip()
#         topic = topic.capitalize()
#         if bookId:
#             insertTopic = Topic(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=chapter_num,class_val=class_val,book_id=bookId.book_id,teacher_id=teacher_id.teacher_id)
#         db.session.add(insertTopic)
#         db.session.commit()
#         if bookId:
#             topic_id = Topic.query.filter_by(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=chapter_num,class_val=class_val,book_id=bookId.book_id).first()
#         insertTopicTracker = TopicTracker(subject_id=subject_id.msg_id,class_sec_id=class_sec_id.class_sec_id,is_covered='N',topic_id=topic_id.topic_id,school_id=teacher_id.school_id,is_archived='N',last_modified_date=datetime.now())
#         db.session.add(insertTopicTracker)
#         db.session.commit()
#     return ("Add new Topic")
  
# @app.route('/addNewChapter',methods=['GET','POST'])
# def addNewChapter():
#     print('inside add new chapter')
#     topics=request.get_json()
#     book_id = request.args.get('book_id')
#     print('book_id'+str(book_id))
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     chapter = request.args.get('chapter')
#     chapter = chapter.strip()
#     punctuations = '''!()-[]{};:'"\,<>./?@#$%^&*_|`=+~'''
#     for x in chapter.lower(): 
#         if x in punctuations: 
#             chapter = chapter.replace(x, "") 
#             print(chapter)
#         else:
#             break
#     chapter = chapter.strip()
#     chapter = chapter.capitalize() 
#     chapter_num = request.args.get('chapter_num')
#     chapter_num = chapter_num.strip()
#     for x in chapter_num: 
#         if x in punctuations: 
#             chapter_num = chapter_num.replace(x, "") 
#             print(chapter_num)
#         else:
#             break
#     chapter_num = chapter_num.strip()
#     subject_id = MessageDetails.query.filter_by(description = subject).first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     # bookId = "select distinct bd.book_id from book_details bd inner join topic_detail td on td.book_id = bd.book_id where td.subject_id = '"+str(subject_id.msg_id)+"' and td.class_val  = '"+str(class_val)+"' and chapter_num = '"+str(chapter_num)+"'"
#     # bookId = db.session.execute(text(bookId)).first()
#     print(topics)
#     # print('Book ID:'+str(bookId))
#     maxChapterNum = "select max(chapter_num) from topic_detail td"
#     maxChapterNum = db.session.execute(text(maxChapterNum)).first()
#     print('Max chapter no')
#     print(maxChapterNum[0])
#     maxChapterNum = int(maxChapterNum[0]) + 1
#     for topic in topics:
#         print(topic)
#         topic = topic.strip()
#         for x in topic: 
#             if x in punctuations: 
#                 topic = topic.replace(x, "") 
#                 print(topic)
#             else:
#                 break
#         topic = topic.strip()
#         topic = topic.capitalize()
#         if chapter_num:
#             insertTopic = Topic(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=chapter_num,class_val=class_val,book_id=book_id,teacher_id=teacher_id.teacher_id)
#         else:
#             insertTopic = Topic(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=maxChapterNum,class_val=class_val,book_id=book_id,teacher_id=teacher_id.teacher_id)
#         db.session.add(insertTopic)
#         db.session.commit()
#         if chapter_num:
#             topic_id = Topic.query.filter_by(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=chapter_num,class_val=class_val,book_id=book_id).first()
#         else:
#             topic_id = Topic.query.filter_by(topic_name=topic,chapter_name=chapter,subject_id=subject_id.msg_id,board_id=board_id.board_id,chapter_num=maxChapterNum,class_val=class_val,book_id=book_id).first()
#         insertTopicTracker = TopicTracker(subject_id=subject_id.msg_id,class_sec_id=class_sec_id.class_sec_id,is_covered='N',topic_id=topic_id.topic_id,school_id=teacher_id.school_id,is_archived='N',last_modified_date=datetime.now())
#         db.session.add(insertTopicTracker)
#         db.session.commit()
#     return ("Add new Chapter")


# @app.route('/spellCheckBook',methods=['GET','POST'])
# def spellCheckBook():
#     print('inside spellCheckBox')
#     bookText = request.args.get('bookText')
#     return ""
#     #if bookText=='':
#     #    return ""
#     #spell = SpellChecker()
#     #correct = spell.correction(bookText)
#     #print('correct word:'+str(correct))
#     #if bookText==correct:
#     #    return ""
#     #else:
#     #    print('inside if')
#     #    print(bookText)
#     #    print(correct)
#     #    return correct

# @app.route('/deleteSubject',methods=['GET','POST'])
# def deleteSubject():
#     subject_id = request.args.get('subjectId')
#     class_val = request.args.get('class_val')
#     board = request.args.get('board')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     deleteSubject = BoardClassSubject.query.filter_by(class_val=class_val,school_id=teacher_id.school_id,subject_id=subject_id,board_id=board).first()
#     deleteSubject.is_archived = 'Y'
#     db.session.commit()
#     return ("delete subject successfully")

# @app.route('/deleteBook',methods=['GET','POST'])
# def deleteBook():
#     subject = request.args.get('subject')
#     class_val = request.args.get('class_val')
#     bookId = request.args.get('bookId')
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     book = BookDetails.query.filter_by(book_id=bookId,subject_id=subject_id.msg_id,class_val= class_val).first()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id).all()
#     print('book name:'+str(book.book_name))
#     for book_id in bookIds:
#         print(book_id.book_id)
#         updateBook = BoardClassSubjectBooks.query.filter_by(book_id=book_id.book_id,school_id=teacher_id.school_id,class_val=class_val,subject_id=subject_id.msg_id).first()
#         print(updateBook)
#         updateBook.is_archived = 'Y'
#         db.session.commit()
#     return ("delete book successfully")

# @app.route('/deleteTopics',methods=['GET','POST'])
# def deleteTopics():
#     subject = request.args.get('subject')
#     class_val = request.args.get('class_val')
#     bookId = request.args.get('bookId')
#     chapter_num = request.args.get('chapter_num')
#     topic_id = request.args.get('topic_id')
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     updateTT = "update topic_tracker set is_archived='Y' where school_id='"+str(teacher_id.school_id)+"' and subject_id='"+str(subject_id.msg_id)+"' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and topic_id='"+str(topic_id)+"'"
#     print(updateTT)
#     updateTT = db.session.execute(text(updateTT))
#     db.session.commit()
#     return ("delete topic successfully")

# @app.route('/deleteChapters',methods=['GET','POST'])
# def deleteChapters():
#     subject = request.args.get('subject')
#     bookId = request.args.get('bookId')
#     class_val = request.args.get('class_val')
#     chapter_num = request.args.get('chapter_num')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     book  = BookDetails.query.filter_by(book_id=bookId,class_val=class_val,subject_id=subject_id.msg_id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id).all()
#     book_id = Topic.query.filter_by(subject_id=subject_id.msg_id,class_val=class_val,chapter_num=chapter_num).first()
#     # for book_id in bookIds:
#     #     print('inside for of deleteChapters')
#     print('subID:'+str(subject_id.msg_id)+' class_val:'+str(class_val)+' chapter num:'+str(chapter_num)+' bookName:'+str(book.book_name))
#     topic_ids = "select topic_id from topic_detail td where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id.msg_id)+"' and chapter_num = '"+str(chapter_num)+"' "
#     topic_ids = topic_ids + "and td.book_id in (select book_id from book_details bd2 where book_name = '"+str(book.book_name)+"' and class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id.msg_id)+"')"
#     topic_ids = db.session.execute(text(topic_ids)).fetchall()
#     for topic_id in topic_ids:
#         print('Topic id:'+str(topic_id.topic_id))
#         # updateTT = TopicTracker.query.filter_by(school_id=teacher_id.school_id,subject_id=subject_id.msg_id,class_sec_id=class_sec_id.class_sec_id,topic_id=topic_id.topic_id).all()
#         updateTT = "update topic_tracker set is_archived='Y' where school_id='"+str(teacher_id.school_id)+"' and subject_id='"+str(subject_id.msg_id)+"' and class_sec_id in (select class_sec_id from class_section where class_val='"+str(class_val)+"' and school_id='"+str(teacher_id.school_id)+"') and topic_id='"+str(topic_id.topic_id)+"'"
#         print(updateTT)
#         updateTT = db.session.execute(text(updateTT))
#         db.session.commit()
#     return ("delete chapter successfully")

# @app.route('/generalSyllabusBooks')
# def generalSyllabusBooks():
#     subject_name=request.args.get('subject_name')
#     class_val=request.args.get('class_val')
#     board_id = request.args.get('board_id')
#     subject_id = MessageDetails.query.filter_by(description=subject_name).first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     distinctBooks = "select distinct bd.book_name from book_details bd inner join topic_detail td on "
#     distinctBooks = distinctBooks + "bd.book_id = td.book_id where bd.subject_id='"+str(subject_id.msg_id)+"' and td.class_val = '"+str(class_val)+"' order by bd.book_name"
#     distinctBooks = db.session.execute(text(distinctBooks)).fetchall()
#     bookArray=[]
#     for val in distinctBooks:
#         print(val.book_name)
#         book_id = BookDetails.query.filter_by(book_name=val.book_name).first()
#         bookArray.append(str(book_id.book_id)+':'+str(val.book_name))
#     if bookArray:
#         return jsonify([bookArray])  
#     else:
#         return ""

# @app.route('/syllabusBooks')
# @login_required
# def syllabusBooks():
#     subject_name=request.args.get('subject_name')
#     class_val=request.args.get('class_val')
#     board_id = request.args.get('board_id')
#     subject_id = MessageDetails.query.filter_by(description=subject_name).first()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     distinctBooks = "select distinct bd.book_name from book_details bd inner join board_class_subject_books bcsb on "
#     distinctBooks = distinctBooks + "bd.book_id = bcsb.book_id where bcsb.school_id='"+str(teacher_id.school_id)+"' and bcsb.subject_id='"+str(subject_id.msg_id)+"' and bcsb.class_val = '"+str(class_val)+"' and bcsb.is_archived = 'N' order by bd.book_name"
#     print(distinctBooks)
#     distinctBooks = db.session.execute(text(distinctBooks)).fetchall()
#     bookArray=[]
#     for val in distinctBooks:
#         book_id = BookDetails.query.filter_by(book_name=val.book_name,class_val=class_val).first()
#         print(str(book_id.book_id)+':'+str(val.book_name))
#         bookArray.append(str(book_id.book_id)+':'+str(val.book_name))
#     if bookArray:
#         return jsonify([bookArray])  
#     else:
#         return ""

# @app.route('/fetchRemBooks',methods=['GET','POST'])
# def fetchRemBooks():
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
    
#     distinctBooks = ''
#     distinctBooks = "select distinct book_name from book_details bd where class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id.msg_id)+"' and "
#     distinctBooks = distinctBooks + "book_name not in (select distinct book_name from book_details bd inner join board_class_subject_books bcsb on bd.book_id = bcsb.book_id "
#     distinctBooks = distinctBooks + "where bd.class_val = '"+str(class_val)+"' and bd.subject_id = '"+str(subject_id.msg_id)+"' and bcsb.school_id = '"+str(teacher_id.school_id)+"')"
#     distinctBooks = db.session.execute(text(distinctBooks)).fetchall()
#     distinctBooksInBCSB = "select distinct bd.book_name from book_details bd inner join board_class_subject_books bcsb on "
#     distinctBooksInBCSB = distinctBooksInBCSB + "bd.book_id = bcsb.book_id where bcsb.subject_id='"+str(subject_id.msg_id)+"' and bcsb.class_val = '"+str(class_val)+"' and bcsb.school_id='"+str(teacher_id.school_id)+"' and bcsb.is_archived = 'Y' order by bd.book_name"
#     distinctBooksInBCSB = db.session.execute(text(distinctBooksInBCSB)).fetchall()
#     bookArray=[]
#     for val in distinctBooks:
#         print(val.book_name)
#         book_id = BookDetails.query.filter_by(book_name=val.book_name).first()
#         bookArray.append(str(book_id.book_id)+':'+str(val.book_name))
#     for value in distinctBooksInBCSB:
#         book_id = BookDetails.query.filter_by(book_name=value.book_name).first()
#         bookArray.append(str(book_id.book_id)+':'+str(value.book_name))
#     if bookArray:
#         return jsonify([bookArray])
#     else:
#         return "" 

# @app.route('/selectedChapter',methods=['GET','POST'])
# def selectedChapter():
#     chapterNum = request.args.get('chapterNum')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     book_id = request.args.get('bookId')
#     print('inside selected chapter')
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     book = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_id=book_id).first()
#     # chapter = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,chapter_num=chapterNum,book_id=book_id).first()
#     chapter = "select chapter_num,chapter_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where "
#     chapter = chapter + "td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id.msg_id)+"' and chapter_num = '"+str(chapterNum)+"' and book_name  = '"+str(book.book_name)+"'"
#     print(chapter)
#     chapter = db.session.execute(text(chapter)).first()
#     # chapter = "select chapter_name,chapter_num from topic_detail td inner join book_details bd on "
#     # chapter = chapter + "td.book_id = bd.book_id where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id.msg_id)+"' and chapter_num = '"+str(chapterNum)+"' and book_name ='"+str(book.book_name)+"'"
#     # chapter = db.session.execute(text(chapter)).fetchall()
#     selectedChapterArray = []
#     # for chapt in chapter:
#     selectedChapterArray.append(str(chapter.chapter_name)+':'+str(chapter.chapter_num))
#     return jsonify([selectedChapterArray])


# @app.route('/fetchRemChapters',methods=['GET','POST'])
# def fetchRemChapters():
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     bookId = request.args.get('bookId')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val).first()
#     book = BookDetails.query.filter_by(book_id=bookId).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id).all()
    
#     chapterArray=[]
#     print('Book:'+str(book.book_name)+' class:'+str(class_val)+' subId:'+str(subject_id.msg_id))
    
#     queryChapters = "select distinct chapter_name,chapter_num from topic_detail td where td.subject_id = '"+str(subject_id.msg_id)+"' and td.class_val = '"+str(class_val)+"'  and book_id in (select book_id from book_details bd where book_name = '"+str(book.book_name)+"' and subject_id = '"+str(subject_id.msg_id)+"' and class_val = '"+str(class_val)+"') "
#     queryChapters = queryChapters + "and td.topic_id not in (select td.topic_id from topic_detail td inner join topic_tracker tt on "
#     queryChapters = queryChapters + "td.topic_id = tt.topic_id where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id.msg_id)+"' and tt.school_id = '"+str(teacher_id.school_id)+"') order by chapter_num"
#     print('print chapters from general')
#     print(queryChapters)
#     queryChapters = db.session.execute(text(queryChapters)).fetchall()
    
#             # chapterArray.append(str(chapter.chapter_num)+":"+str(chapter.chapter_name))
#     queryBookDetails = "select distinct chapter_name,chapter_num from topic_detail td inner join topic_tracker tt on "
#     queryBookDetails = queryBookDetails + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id.msg_id)+"' and tt.school_id='"+str(teacher_id.school_id)+"' and td.class_val = '"+str(class_val)+"' and tt.is_archived = 'Y' and td.book_id in "
#     queryBookDetails = queryBookDetails + "(select book_id from book_details bd where class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id.msg_id)+"' and book_name='"+str(book.book_name)+"') order by chapter_num"
#     print('deleted chapters')
#     print(queryBookDetails)
#     queryBookDetails = db.session.execute(text(queryBookDetails)).fetchall()
#     j=1
#     for chapter in queryChapters:
#         chapters = chapter.chapter_name
#         chapters = chapters.replace("'","\'")
#         num = chapter.chapter_num
#         if len(queryChapters)>1:
#             if j==1:
#                 chapters = chapters + "/"
#                 print(chapters)
#             elif j==len(queryChapters):
#                 if len(queryBookDetails)>0:
#                     num = "/"+str(num)
#                     chapters = chapters+"/"
#                 else:
#                     num = "/"+str(num)
#                     print(chapters)
#             else:
#                 num = "/"+str(num)
#                 chapters = chapters+"/"
#                 print(chapters)
#             j=j+1
#         else:
#             if len(queryBookDetails)>0:
#                 chapters = chapters+"/"
#         chapterArray.append(str(num)+":"+str(chapters))
#     i=1
#     for book in queryBookDetails:
#         print('inside for queryBookDetails'+str(len(queryChapters)))
#         chapter = book.chapter_name
#         chapter = chapter.replace("'","\'")
#         num = book.chapter_num
#         if len(queryBookDetails)>1:
#             if i==1:
#                 if len(queryChapters)>0:
#                     print('if queryChapters is not null')
#                     num = "/"+str(num)
#                     chapter = chapter + "/"
#                 else:
#                     chapter = chapter + "/"
#                     print(chapter)
#             elif i==len(queryBookDetails):
#                 num = "/"+str(num)
#                 print(chapter)
#             else:
#                 num = "/"+str(num)
#                 chapter = chapter+"/"
#                 print(chapter)
#             i=i+1
#         else:
#             if len(queryChapters)>0:
#                 num = "/"+str(num)
#         print(chapter)
#         chapterArray.append(str(num)+":"+str(chapter))
#     for ch in chapterArray:
#         print(ch)
#     if chapterArray:
#         return jsonify([chapterArray]) 
#     else:
#         return ""


# @app.route('/fetchBooks',methods=['GET','POST'])
# def fetchBooks():
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     distinctBooks = "select distinct bd.book_name from book_details bd inner join board_class_subject_books bcsb on "
#     distinctBooks = distinctBooks + "bd.book_id = bcsb.book_id where bcsb.school_id='"+str(teacher_id.school_id)+"' and bcsb.subject_id='"+str(subject_id.msg_id)+"' and bcsb.class_val = '"+str(class_val)+"' and bcsb.is_archived = 'N' order by bd.book_name"
#     print(distinctBooks)
#     distinctBooks = db.session.execute(text(distinctBooks)).fetchall()
#     bookArray=[]
#     for val in distinctBooks:
#         print(val.book_name)
#         book_id = BookDetails.query.filter_by(book_name=val.book_name).first()
#         bookArray.append(str(book_id.book_id)+':'+str(val.book_name))
#     if bookArray:
#         return jsonify([bookArray])
#     else:
#         return ""

# @app.route('/generalSyllabusChapters')
# def generalSyllabusChapters():
#     book_id=request.args.get('book_id')
#     class_val=request.args.get('class_val')
#     board_id=request.args.get('board_id')
#     subject_id=request.args.get('subject_id')
#     print('Book id:'+str(book_id))
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val).first()
#     book = BookDetails.query.filter_by(book_id=book_id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id,board_id=board_id).all()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     chapterArray=[]
#     # print('Book:'+str(book.book_name)+' class:'+str(class_val)+' subId:'+str(subject_id)+' boardId:'+str(board_id))
#     queryBookDetails = "select distinct chapter_name,chapter_num from topic_detail td where td.subject_id = '"+str(subject_id)+"' and td.class_val = '"+str(class_val)+"' and td.board_id='"+str(board_id)+"' and td.book_id in "
#     queryBookDetails = queryBookDetails + "(select book_id from book_details bd where class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id)+"' and book_name='"+str(book.book_name)+"') order by chapter_num"
#     print('inside general syllabus chapter:')
#     print(queryBookDetails)
#     queryBookDetails = db.session.execute(text(queryBookDetails)).fetchall()
#     i=1
#     for book in queryBookDetails:
#         chapter = book.chapter_name
#         chapter = chapter.replace("'","\'")
#         num = book.chapter_num
#         if len(queryBookDetails)>1:
#             if i==1:
#                 chapter = chapter + "/"
#                 print(chapter)
#             elif i==len(queryBookDetails):
#                 num = "/"+str(num)
#                 print(chapter)
#             else:
#                 num = "/"+str(num)
#                 chapter = chapter+"/"
#                 print(chapter)
#             i=i+1
#         chapterArray.append(str(num)+":"+str(chapter))
#     if chapterArray:
#         return jsonify([chapterArray]) 
#     else:
#         return ""

# @app.route('/syllabusChapters') 
# @login_required
# def syllabusChapters():
#     book_id=request.args.get('book_id')
#     class_val=request.args.get('class_val')
#     board_id=request.args.get('board_id')
#     subject_id=request.args.get('subject_id')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     book = BookDetails.query.filter_by(book_id=book_id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id,board_id=board_id).all()
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     chapterArray=[]
#     print('Book:'+str(book.book_name)+' class:'+str(class_val)+' subId:'+str(subject_id)+' boardId:'+str(board_id))
#     queryBookDetails = "select distinct chapter_name,chapter_num from topic_detail td inner join topic_tracker tt on "
#     queryBookDetails = queryBookDetails + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id)+"' and tt.school_id='"+str(teacher_id.school_id)+"' and td.class_val = '"+str(class_val)+"' and tt.is_archived = 'N' and td.book_id in "
#     queryBookDetails = queryBookDetails + "(select book_id from book_details bd where class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id)+"' and book_name='"+str(book.book_name)+"') order by chapter_num"
#     # queryBookDetails = queryBookDetails + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id)+"' and tt.school_id='"+str(teacher_id.school_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' "
#     # queryBookDetails = queryBookDetails + "order by chapter_num"
#     print(queryBookDetails)
#     queryBookDetails = db.session.execute(text(queryBookDetails)).fetchall()
#     i=1
#     for book in queryBookDetails:
#         chapter = book.chapter_name
#         chapter = chapter.replace("'","\'")
#         num = book.chapter_num
#         book = Topic.query.filter_by(chapter_name=book.chapter_name,chapter_num=book.chapter_num).first()
#         bookId = book.book_id
#         if len(queryBookDetails)>1:
#             if i==1:
#                 bookId = str(bookId) + "/"
#                 print(chapter)
#             elif i==len(queryBookDetails):
#                 num = "/"+str(num)
#                 print(chapter)
#             else:
#                 num = "/"+str(num)
#                 bookId = str(bookId) + "/"
#                 print(chapter)
#             i=i+1
#         chapterArray.append(str(num)+":"+str(chapter)+';'+str(book.book_id))
#     for chapters in chapterArray:
#         print(chapters)
#     if chapterArray:
#         return jsonify([chapterArray]) 
#     else:
#         return ""

# @app.route('/chapterTopic',methods=['GET','POST'])
# def chapterTopic():
#     print('inside chapterTopic')
#     class_val = request.args.get('class_val')
#     subject = request.args.get('subject')
#     chapter_num = request.args.get('chapter_num')
#     book_id = request.args.get('book_id')
#     print('Book id:'+str(book_id))
#     print('class value:'+str(class_val))
#     try:
#         subject_id = MessageDetails.query.filter_by(description=subject).first()
#         teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#         book = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_id=book_id).first()
#         print('class:'+str(class_val)+' subject_id:'+str(subject_id.msg_id)+' book_id:'+str(book_id))
#         remTopics = "select distinct topic_name ,topic_id from topic_detail td where class_val = '"+str(class_val)+"' and subject_id  = '"+str(subject_id.msg_id)+"' and chapter_num = '"+str(chapter_num)+"' and topic_id not in "
#         remTopics = remTopics + "(select distinct td.topic_id from topic_detail td inner join topic_tracker tt on "
#         remTopics = remTopics + "td.topic_id = tt.topic_id where td.class_val = '"+str(class_val)+"' and "
#         remTopics = remTopics + "td.subject_id = '"+str(subject_id.msg_id)+"' and td.chapter_num = '"+str(chapter_num)+"' and tt.school_id = '"+str(teacher_id.school_id)+"') and book_id in (select book_id from book_details bd where book_name = '"+str(book.book_name)+"' and class_val='"+str(class_val)+"' and subject_id='"+str(subject_id.msg_id)+"') order by topic_id"
#         print('inside Rem topics')
#         print('Rem Topics:'+str(remTopics))
#         remTopics = db.session.execute(text(remTopics)).fetchall()
#         topics = "select distinct td.topic_id,td.topic_name from topic_detail td inner join topic_tracker tt on "
#         topics = topics + "td.topic_id = tt.topic_id where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id.msg_id)+"' and td.chapter_num = '"+str(chapter_num)+"' and tt.is_archived = 'Y' and tt.school_id = '"+str(teacher_id.school_id)+"' and td.book_id in (select book_id from book_details bd where book_name = '"+str(book.book_name)+"') order by td.topic_id"
#         topics = db.session.execute(text(topics)).fetchall()
#     except:
#         return ""
#     topicArray = []
#     i=1
#     for topic in topics:
#         print('for topic list')
#         print(topic.topic_name)
#         print(len(remTopics))
#         topic_name = topic.topic_name
#         topic_name = topic_name.replace("'","\'")
#         topic_id = topic.topic_id
#         if len(topics)>1:
#             if i==1:
#                 topic_name = topic_name + "/"
#             elif i==len(topics):
#                 if len(remTopics)>0:
#                     topic_id = "/"+str(topic_id)
#                     topic_name = topic_name + "/"
#                 else:
#                     topic_id = "/"+str(topic_id)
#             else:
#                 topic_id = "/"+str(topic_id)
#                 topic_name = topic_name+"/"
#             i=i+1
#         else:
#             if len(remTopics)>0:
#                 topic_name = topic_name+"/"
#         topicArray.append(str(topic_id)+":"+str(topic_name))
#         # topicArray.append(str(topic.topic_id)+':'+str(topic.topic_name))
#     j=1
#     for remTopic in remTopics:
#         print('rem list')
#         print(remTopic.topic_name)
#         topic_name = remTopic.topic_name
#         topic_name = remTopic.topic_name.replace("'","\'")
#         topic_id = remTopic.topic_id
#         if len(remTopics)>1:
#             if j==1:
#                 if len(topics)>0:
#                     topic_id = "/"+str(topic_id)
#                     topic_name = topic_name + "/"
#                 else:
#                     topic_name = topic_name + "/"
#             elif j==len(remTopics):
#                 topic_id = "/"+str(topic_id)
#             else:
#                 topic_id = "/"+str(topic_id)
#                 topic_name = topic_name+"/"
#             j=j+1
#         else:
#             if len(topics)>0:
#                 topic_id = "/"+str(topic_id)
#         topicArray.append(str(topic_id)+":"+str(topic_name))
#         # topicArray.append(str(remTopic.topic_id)+':'+str(remTopic.topic_name))
#     for top in topicArray:
#         print(top)
#     if topicArray:
#         return jsonify([topicArray])
#     else:
#         return ""

# @app.route('/fetchChapters',methods=['GET','POST'])
# def fetchChapters():
#     book_id=request.args.get('bookId')
#     print('inside fetchChapters')
#     print('Book Id:'+str(book_id))
#     class_val=request.args.get('class_val')
#     # board_id=request.args.get('board_id')
#     subject=request.args.get('subject')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     book = BookDetails.query.filter_by(book_id=book_id).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id).all()
    
#     chapterArray=[]
#     print('Book:'+str(book.book_name)+' class:'+str(class_val)+' subId:'+str(subject_id.msg_id))
#     queryBookDetails = "select distinct chapter_name,chapter_num from topic_detail td inner join topic_tracker tt on "
#     queryBookDetails = queryBookDetails + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id.msg_id)+"' and tt.school_id='"+str(teacher_id.school_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' and td.book_id in "
#     queryBookDetails = queryBookDetails + "(select book_id from book_details bd where class_val = '"+str(class_val)+"' and subject_id = '"+str(subject_id.msg_id)+"' and book_name='"+str(book.book_name)+"') order by chapter_num"
#     # queryBookDetails = queryBookDetails + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id.msg_id)+"' and tt.school_id='"+str(teacher_id.school_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' "
#     # queryBookDetails = queryBookDetails + "order by chapter_num"
#     print(queryBookDetails)
#     queryBookDetails = db.session.execute(text(queryBookDetails)).fetchall()
#     i=1
#     for book in queryBookDetails:
#         chapter = book.chapter_name
#         chapter = chapter.replace("'","\'")
#         num = book.chapter_num
#         book = Topic.query.filter_by(chapter_name=book.chapter_name,chapter_num=book.chapter_num).first()
#         bookId = book.book_id
#         if len(queryBookDetails)>1:
#             if i==1:
#                 bookId = str(bookId) + "/"
#                 print(chapter)
#             elif i==len(queryBookDetails):
#                 num = "/"+str(num)
#                 print(chapter)
#             else:
#                 num = "/"+str(num)
#                 bookId = str(bookId) + "/"
#                 print(chapter)
#             i=i+1
#         chapterArray.append(str(num)+":"+str(chapter)+';'+str(book.book_id))
#     if chapterArray:
#         return jsonify([chapterArray]) 
#     else:
#         return ""

# @app.route('/fetchTopics',methods=['GET','POST'])
# def fetchTopics():
#     subject=request.args.get('subject')
#     chapter_num=request.args.get('chapter_num')
    
#     class_val = request.args.get('class_val')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     subject_id = MessageDetails.query.filter_by(description=subject).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     bookId = request.args.get('bookId')
#     # chapter_name = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,chapter_num=chapter_num).first()
#     book = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,book_id=bookId).first()
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id.msg_id,board_id=board_id.board_id).all()
#     # topicArray=[]
#     # chapter_name = Topic.query.filter_by(class_val=class_val,subject_id=subject_id.msg_id,chapter_num=chapter_num).first()
#     # queryTopics = "select distinct td.topic_id ,td.topic_name from topic_detail td inner join topic_tracker tt on "
#     # queryTopics = queryTopics + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id.msg_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' and tt.school_id = '"+str(teacher_id.school_id)+"' and td.topic_id in "
#     # queryTopics = queryTopics + "(select topic_id from topic_detail td where subject_id = '"+str(subject_id.msg_id)+"' and class_val = '"+str(class_val)+"' and chapter_name = '"+str(chapter_name.chapter_name)+"') order by td.topic_id"
#     # queryTopics = db.session.execute(text(queryTopics)).fetchall()
#     # for topic in queryTopics:
#     #     topicArray.append(str(topic.topic_id)+":"+str(topic.topic_name))
#     # if topicArray:
#     #     return jsonify([topicArray]) 
#     # else:
#     #     return ""
#     topicArray=[]
#     chapter_name = "select chapter_name,td.book_id from topic_detail td inner join book_details bd on td.book_id = bd.book_id where "
#     chapter_name = chapter_name + "td.subject_id = '"+str(subject_id.msg_id)+"' and td.class_val = '"+str(class_val)+"' and td.chapter_num = '"+str(chapter_num)+"' and bd.book_name = '"+str(book.book_name)+"'"
#     chapter_name = db.session.execute(text(chapter_name)).first()
#     chapterName = chapter_name.chapter_name.replace("'","''")
#     queryTopics = "select distinct td.topic_id ,td.topic_name from topic_detail td inner join topic_tracker tt on "
#     queryTopics = queryTopics + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id.msg_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' and tt.school_id = '"+str(teacher_id.school_id)+"' and td.topic_id in "
#     queryTopics = queryTopics + "(select topic_id from topic_detail td where subject_id = '"+str(subject_id.msg_id)+"' and class_val = '"+str(class_val)+"' and chapter_name = '"+str(chapterName)+"') order by td.topic_id"
#     print('fetch Topic Query:'+str(queryTopics))
#     queryTopics = db.session.execute(text(queryTopics)).fetchall()
#     i=1
#     for topic in queryTopics:
#         topic_name = topic.topic_name
#         topic_name = topic_name.replace("'","\'")
#         topic_id = topic.topic_id
#         if len(queryTopics)>1:
#             if i==1:
#                 topic_name = topic_name + "/"
#                 print(topic_name)
#             elif i==len(queryTopics):
#                 topic_id = "/"+str(topic_id)
#             else:
#                 topic_id = "/"+str(topic_id)
#                 topic_name = topic_name+"/"
#                 print(topic_name)
#             i=i+1
#         topicArray.append(str(topic_id)+":"+str(topic_name))
#         # topicArray.append(str(topic.topic_id)+":"+str(topic.topic_name))
#     if topicArray:
#         return jsonify([topicArray]) 
#     else:
#         return ""

# @app.route('/generalSyllabusTopics',methods=['GET','POST'])
# def generalSyllabusTopics():
#     subject_id=request.args.get('subject_id')
#     board_id=request.args.get('board_id')
#     chapter_num=request.args.get('chapter_num')
#     bookId = request.args.get('selectedBookId')
#     class_val = request.args.get('class_val')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val).first()
    
#     print('BookID:'+str(bookId))
#     book = BookDetails.query.filter_by(book_id=bookId).first()
#     print('book name:')
#     print(book.book_name)
#     bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id,board_id=board_id).all()
#     topicArray=[]
#     chapter_name = "select chapter_name,td.book_id from topic_detail td inner join book_details bd on td.book_id = bd.book_id where "
#     chapter_name = chapter_name + "td.subject_id = '"+str(subject_id)+"' and td.class_val = '"+str(class_val)+"' and td.chapter_num = '"+str(chapter_num)+"' and bd.book_name = '"+str(book.book_name)+"'"
#     chapter_name = db.session.execute(text(chapter_name)).first()
#     # queryTopics = "select distinct td.topic_id ,td.topic_name from topic_detail td inner join topic_tracker tt on "
#     # queryTopics = queryTopics + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and td.topic_id in "
#     # queryTopics = queryTopics + "(select topic_id from topic_detail td where subject_id = '"+str(subject_id)+"' and class_val = '"+str(class_val)+"' and chapter_name = '"+str(chapter_name.chapter_name)+"') order by td.topic_id"
#     chapterName = chapter_name.chapter_name.replace("'","''")
#     queryTopics = "select distinct topic_id, topic_name from topic_detail td where class_val = '"+str(class_val)+"' and board_id = '"+str(board_id)+"' and subject_id = '"+str(subject_id)+"' and chapter_name ='"+str(chapterName)+"' order by topic_id"
#     print('inside generalsyllabustopics')
#     print(queryTopics)
#     queryTopics = db.session.execute(text(queryTopics)).fetchall()
#     i=1
#     for topic in queryTopics:
#         topic_name = topic.topic_name
#         topic_name = topic_name.replace("'","\'")
#         topic_id = topic.topic_id
#         if len(queryTopics)>1:
#             if i==1:
#                 topic_name = topic_name + "/"
#                 print(topic_name)
#             elif i==len(queryTopics):
#                 topic_id = "/"+str(topic_id)
#             else:
#                 topic_id = "/"+str(topic_id)
#                 topic_name = topic_name+"/"
#                 print(topic_name)
#             i=i+1
#         topicArray.append(str(topic_id)+":"+str(topic_name))
#     if topicArray:
#         return jsonify([topicArray]) 
#     else:
#         return ""

# @app.route('/syllabusTopics')
# @login_required
# def syllabusTopics():
#     subject_id=request.args.get('subject_id')
#     board_id=request.args.get('board_id')
#     chapter_num=request.args.get('chapter_num')
#     bookId = request.args.get('selectedBookId')
#     class_val = request.args.get('class_val')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = ClassSection.query.filter_by(class_val=class_val,school_id=teacher_id.school_id).first()
#     print('BookID:'+str(bookId))
#     book = BookDetails.query.filter_by(class_val=class_val,subject_id=subject_id,book_id=bookId).first()
#     # bookIds = BookDetails.query.filter_by(book_name=book.book_name,class_val=class_val,subject_id=subject_id,board_id=board_id).all()
#     topicArray=[]
#     chapter_name = "select chapter_name,td.book_id from topic_detail td inner join book_details bd on td.book_id = bd.book_id where "
#     chapter_name = chapter_name + "td.subject_id = '"+str(subject_id)+"' and td.class_val = '"+str(class_val)+"' and td.chapter_num = '"+str(chapter_num)+"' and bd.book_name = '"+str(book.book_name)+"'"
#     chapter_name = db.session.execute(text(chapter_name)).first()
#     chapterName = chapter_name.chapter_name.replace("'","''")
#     queryTopics = "select distinct td.topic_id ,td.topic_name from topic_detail td inner join topic_tracker tt on "
#     queryTopics = queryTopics + "td.topic_id = tt.topic_id where tt.subject_id = '"+str(subject_id)+"' and tt.class_sec_id = '"+str(class_sec_id.class_sec_id)+"' and tt.is_archived = 'N' and tt.school_id = '"+str(teacher_id.school_id)+"' and td.topic_id in "
#     queryTopics = queryTopics + "(select topic_id from topic_detail td where subject_id = '"+str(subject_id)+"' and class_val = '"+str(class_val)+"' and chapter_name = '"+str(chapterName)+"') order by td.topic_id"
#     print(queryTopics)
#     queryTopics = db.session.execute(text(queryTopics)).fetchall()
#     i=1
#     print(len(queryTopics))
#     for topic in queryTopics:
#         topic_name = topic.topic_name
#         topic_name = topic_name.replace("'","\'")
#         topic_id = topic.topic_id
#         if len(queryTopics)>1:
#             if i==1:
#                 topic_name = topic_name + "/"
#                 print(topic_name)
#             elif i==len(queryTopics):
#                 topic_id = "/"+str(topic_id)
#             else:
#                 topic_id = "/"+str(topic_id)
#                 topic_name = topic_name+"/"
#                 print(topic_name)
#             i=i+1
#         topicArray.append(str(topic_id)+":"+str(topic_name))
#     if topicArray:
#         return jsonify([topicArray]) 
#     else:
#         return ""

# @app.route('/syllabusQuestionsDetails',methods=['GET','POST'])
# def syllabusQuestionsDetails():
#     class_val = request.args.get('class_val')
#     subject_id = request.args.get('subject_id')
#     topic_id = request.args.get('topic_id')
#     chapter_num=request.args.get('chapter_num')
#     print('inside syllabusQuestionsDetails')
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     boardName = MessageDetails.query.filter_by(msg_id=board.board_id).first()
#     # questions = QuestionDetails.query.filter_by(subject_id=subject_id,class_val=class_val,topic_id=topic_id).all()
#     topic_name = Topic.query.filter_by(topic_id=topic_id,subject_id=subject_id,class_val=class_val).first()
#     chapter_name = Topic.query.filter_by(subject_id=subject_id,class_val=class_val,chapter_num=chapter_num).first()
#     subQuestion = "select count(*) from question_details qd where subject_id = '"+str(subject_id)+"' and class_val = '"+str(class_val)+"' and topic_id = '"+str(topic_id)+"' and archive_status = 'N' and question_type = 'Subjective'"
#     subQuestion = db.session.execute(text(subQuestion)).first()
#     objQuestion = "select count(*) from question_details qd where subject_id = '"+str(subject_id)+"' and class_val = '"+str(class_val)+"' and topic_id = '"+str(topic_id)+"' and archive_status = 'N' and question_type = 'MCQ1'"
#     objQuestion = db.session.execute(text(objQuestion)).first()
#     refContent = "select count(*) from content_detail cd where class_val = '"+str(class_val)+"' and archive_status = 'N' and subject_id = '"+str(subject_id)+"' and topic_id = '"+str(topic_id)+"'"
#     refContent = db.session.execute(text(refContent)).first()
#     questionDetailsArray = []
#     print('subQuestion:'+str(subQuestion[0])+' objQuestion:'+str(objQuestion[0])+' refContent:'+str(refContent[0]))
#     # for question in questions:
#     #     questionArray.append(question.question_description)
#     questionDetailsArray.append(str(topic_name.topic_name)+':'+str(subQuestion[0])+':'+str(objQuestion[0])+':'+str(refContent[0])+':'+str(boardName.description)+':'+str(chapter_name.chapter_name))
#     if questionDetailsArray:
#         return jsonify([questionDetailsArray])
#     else:
#         return ""


@app.route('/grantSchoolAdminAccess')
def grantSchoolAdminAccess():
    school_id=request.args.get('school_id')
    teacher_id=request.args.get('teacher_id')
    schoolTableDetails = SchoolProfile.query.filter_by(school_id=school_id).first()
    schoolTableDetails.school_admin=teacher_id
    db.session.commit()
    return jsonify(["String"])

@app.route('/grantSchoolAccess')
def grantSchoolAccess():
    username=request.args.get('username')    
    school_id=request.args.get('school_id')
    school=schoolNameVal()
    print("we're in grant access request. ")
    userTableDetails = User.query.filter_by(username=username).first()
    print("#######User Type: "+ str(userTableDetails.user_type))
    SchoolDetailData = SchoolProfile.query.filter_by(school_id=userTableDetails.school_id).first()
    if SchoolDetailData:
        print('if checkSchoolProfile not none')
        print('checkSchoolProfile.is_veirfied:'+str(SchoolDetailData.is_verified))
        print('checkSchoolProfile.school_id'+str(SchoolDetailData.school_id))
        SchoolDetailData.is_verified = 'Y'
        db.session.commit()
    return jsonify(["String"])


@app.route('/grantUserAccess')
def grantUserAccess():
    username=request.args.get('username')    
    school_id=request.args.get('school_id')
    school=schoolNameVal()
    print("we're in grant access request. ")
    userTableDetails = User.query.filter_by(username=username).first()
    print("#######User Type: "+ str(userTableDetails.user_type))
    userTableDetails.access_status='145'
    userFullName = userTableDetails.first_name + " "+ userTableDetails.last_name
    if userTableDetails.user_type==71:
        print('#########Gotten into 71')
        checkTeacherProfile=TeacherProfile.query.filter_by(user_id=userTableDetails.id).first()
        if checkTeacherProfile:
            print('if checkTeacherProfile not none')
            checkSchoolProfile = SchoolProfile.query.filter_by(school_id=checkTeacherProfile.school_id).first()
            if checkSchoolProfile:
                print('if checkSchoolProfile not none')
                print('checkSchoolProfile.is_veirfied:'+str(checkSchoolProfile.is_verified))
                print('checkSchoolProfile.school_id'+str(checkSchoolProfile.school_id))
                checkSchoolProfile.is_verified = 'Y'
                db.session.commit()
        if checkTeacherProfile==None:
            teacherData=TeacherProfile(teacher_name=userFullName,school_id=school_id, registration_date=datetime.now(), email=userTableDetails.email, phone=userTableDetails.phone, device_preference='195', user_id=userTableDetails.id)
            db.session.add(teacherData)    
            db.session.commit()
    elif userTableDetails.user_type==134:
        print('#########Gotten into 134')
        checkStudentProfile=StudentProfile.query.filter_by(user_id=userTableDetails.id).first()
        if checkStudentProfile==None:
            studentData=StudentProfile(full_name=userFullName,school_id=school_id, registration_date=datetime.now(), last_modified_date=datetime.now(), email=userTableDetails.email, phone=userTableDetails.phone, user_id=userTableDetails.id)
            db.session.add(studentData)    
            db.session.commit()
    elif userTableDetails.user_type==72:
        print('#########Gotten into 72')
        print('Email:'+str(userTableDetails.email))
        checkGuardianProfile=GuardianProfile.query.filter_by(email=userTableDetails.email).all()
        print(checkGuardianProfile)
        if checkGuardianProfile:
            for gprofile in checkGuardianProfile:
                print('If guardian profile is not empty')
                gprofile.user_id=userTableDetails.id
                db.session.commit()
        else:
            print('If guardian profile is empty')
            guardianData=GuardianProfile(full_name=userFullName,first_name=userTableDetails.first_name, last_name=userTableDetails.last_name,email=userTableDetails.email,phone=userTableDetails.phone, user_id=userTableDetails.id)
            db.session.add(guardianData)    
            db.session.commit()
    else:
        print('#########Gotten into else')
        pass
    access_granted_email(userTableDetails.email,userTableDetails.username,school )
    return jsonify(["String"])



# @app.route('/questionBank',methods=['POST','GET'])
# @login_required
# def questionBank():
#     topic_list=None
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     form=QuestionBankQueryForm()
#     form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()]
#     form.subject_name.choices= ''
#     form.chapter_num.choices= ''
#     form.test_type.choices= [(i.description,i.description) for i in MessageDetails.query.filter_by(category='Test type').all()]
#     if request.method=='POST':
#         topic_list="select td.topic_id,td.topic_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val = '"+str(form.class_val.data)+"' and td.subject_id = '"+str(form.subject_name.data)+"' and td.chapter_num='"+str(form.chapter_num.data)+"' and tt.is_archived = 'N'"
#         topic_list = db.session.execute(text(topic_list)).fetchall()
#         subject=MessageDetails.query.filter_by(msg_id=int(form.subject_name.data)).first()
#         session['class_val']=form.class_val.data
#         session['sub_name']=subject.description
#         session['test_type_val']=form.test_type.data
#         session['chapter_num']=form.chapter_num.data  
#         print('Class value:'+str(form.class_val.data))
#         form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(str(form.class_val.data))]
#         print('Class value:'+str(form.class_val.data))
#         form.chapter_num.choices= [(int(i['chapter_num']), str(i['chapter_num'])+' - '+str(i['chapter_name'])) for i in chapters(str(form.class_val.data),int(form.subject_name.data))]
#         indic='DashBoard'
#         return render_template('questionBank.html',indic=indic,title='Question Bank',form=form,topics=topic_list,user_type_val=str(current_user.user_type))
#     indic='DashBoard'
#     return render_template('questionBank.html',indic=indic,title='Question Bank',form=form,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

# @app.route('/visitedQuestions',methods=['GET','POST'])
# def visitedQuestions():
#     retake = request.args.get('retake')
#     questions=[]
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     topicList=request.get_json() 
#     for topic in topicList:
#         print(str(retake)+'Retake')
#         topicFromTracker = TopicTracker.query.filter_by(school_id = teacher_id.school_id, topic_id=int(topic)).first()
#         topicFromTracker.is_covered='Y'
#         if topicFromTracker.reteach_count:
#             topicFromTracker.reteach_count=topicFromTracker.reteach_count+1
#         db.session.commit()

#     return jsonify(['1'])
    
# @app.route('/questionBankQuestions',methods=['GET','POST'])
# def questionBankQuestions():
#     questions=[]
#     topicList=request.get_json()
#     for topic in topicList:
#         # question_Details=QuestionDetails.query.filter_by(QuestionDetails.topic_id == int(topic)).first()
#         # questionList = QuestionDetails.query.join(QuestionOptions, QuestionDetails.question_id==QuestionOptions.question_id).add_columns(QuestionDetails.question_id, QuestionDetails.question_description, QuestionDetails.question_type, QuestionDetails.suggested_weightage).filter(QuestionDetails.topic_id == int(topic)).filter(QuestionOptions.is_correct=='Y').all()
#         questionList = QuestionDetails.query.filter_by(topic_id=int(topic),archive_status='N').order_by(QuestionDetails.question_id).all()
#         questions.append(questionList)
#         for q in questionList:
#             print("Question List"+str(q))    
#     if len(questionList)==0:
#         print('returning 1')
#         return jsonify(['1'])
#     else:
#         print('returning template'+ str(questionList))
#         return render_template('questionBankQuestions.html',questions=questions)

# @app.route('/questionBankFileUpload',methods=['GET','POST'])
# def questionBankFileUpload():
#     #question_list=request.get_json()
#     data=request.get_json()
#     question_list=data[0]
#     count_marks=data[1]
#     document = Document()
#     document.add_heading(schoolNameVal(), 0)
#     document.add_heading('Class '+session.get('class_val',None)+" - "+session.get('test_type_val',None)+" - "+str(session.get('date',None)) , 1)
#     document.add_heading("Subject : "+session.get('sub_name',None),2)
#     document.add_heading("Total Marks : "+str(count_marks),3)
#     p = document.add_paragraph()
#     for question in question_list:
#         data=QuestionDetails.query.filter_by(question_id=int(question), archive_status='N').first()
#         document.add_paragraph(
#             data.question_description, style='List Number'
#         )    
#         options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
#         for option in options:
#             if option.option_desc is not None:
#                 document.add_paragraph(
#                     option.option+". "+option.option_desc)     
#     #document.add_page_break()
#     file_name='S'+'1'+'C'+session.get('class_val',"0")+session.get('sub_name',"0")+session.get('test_type_val',"0")+str(datetime.today().strftime("%d%m%Y"))+'.docx'
#     if not os.path.exists('tempdocx'):
#         os.mkdir('tempdocx')
#     document.save('tempdocx/'+file_name)
#     client = boto3.client('s3', region_name='ap-south-1')
#     client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
#     os.remove('tempdocx/'+file_name)

#     return render_template('testPaperDisplay.html',file_name='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name)

# @app.route('/testBuilder',methods=['POST','GET'])
# @login_required
# def testBuilder():
#     topic_list=None
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     form=TestBuilderQueryForm()
#     print(teacher_id.school_id)
#     form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()]
#     form.subject_name.choices= ''
#     form.chapter_num.choices= ''
#     # [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
#     form.test_type.choices= [(i.description,i.description) for i in MessageDetails.query.filter_by(category='Test type').all()]
#     test_papers = MessageDetails.query.filter_by(category='Test type').all()
#     # print(request.form['class_val'])
#     # print(request.form['subject_id'])
#     available_class = "select distinct class_val from class_section where school_id='"+str(teacher_id.school_id)+"'"
#     available_class = db.session.execute(text(available_class)).fetchall()
#     if request.method=='POST':
#         if request.form['test_date']=='':
#             # flash('Select Date')
#             # form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(int(form.class_val.data))]
#             indic='DashBoard'
#             return render_template('testBuilder.html',indic=indic,form=form)
#         topic_list=Topic.query.filter_by(class_val=str(form.class_val.data),subject_id=int(form.subject_name.data),chapter_num=int(form.chapter_num.data)).all()
#         subject=MessageDetails.query.filter_by(msg_id=int(form.subject_name.data)).first()
#         session['class_val']=form.class_val.data
#         session['date']=request.form['test_date']
#         session['sub_name']=subject.description
#         session['sub_id']=form.subject_name.data
#         session['test_type_val']=form.test_type.data
#         session['chapter_num']=form.chapter_num.data 
#         form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(str(form.class_val.data))]
#         form.chapter_num.choices= [(int(i['chapter_num']), str(i['chapter_num'])+' - '+str(i['chapter_name'])) for i in chapters(str(form.class_val.data),int(form.subject_name.data))]
#         indic='DashBoard'
#         return render_template('testBuilder.html',indic=indic,title='Test Builder',form=form,topics=topic_list,user_type_val=str(current_user.user_type))
#     indic='DashBoard'
#     return render_template('testBuilder.html',indic=indic,title='Test Builder',form=form,available_class=available_class,test_papers=test_papers,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

# @app.route('/filterQuestionsfromTopic',methods=['GET','POST'])
# def filterQuestionsfromTopic():
#     topics = request.get_json()
#     print('topics:'+str(topics))
    
#     for topic in topics:
#         print('topic:'+str(topic))
#     class_val = request.args.get('class_val')
#     subject_id = request.args.get('subject_id')
#     test_type = request.args.get('test_type')
#     print('Class feedback:'+str(test_type))
#     questions = []
#     questionList = ''
#     if topics:
#         print('if topics available')
#         for topic in topics:
#             # if test_type == 'Class Feedback':
#             if class_val!=None:
#                 if subject_id!=None:
#                     if test_type!=None:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                     else:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                 else:
#                     if test_type!=None:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                     else:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()           
#             else:
#                 if subject_id!=None:
#                     if test_type!=None:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                     else:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                 else:
#                     if test_type!=None:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#                     else:
#                         questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#             # else:
#             #     if class_val!=None:
#             #         if subject_id!=None:
#             #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#             #         else:
#             #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
#             #     else:
#             #         if subject_id!=None:
#             #             questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id,topic_id=topic).order_by(QuestionDetails.question_id).all()
#             #         else:
#             #             questionList = QuestionDetails.query.filter_by(archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all() 
#             if questionList:  
#                 questions.append(questionList)
#     else:
#         print('if topics not available')
#         # if test_type == 'Class Feedback':
#         if class_val!=None:
#             print('if class_val available:'+str(class_val))
#             if subject_id!=None:
#                 print('if subject_id available:'+str(subject_id))
#                 if test_type:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#                 else:
#                     print('if test type is empty')
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#             else:
#                 if test_type!=None:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#                 else:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#         else:
#             if subject_id!=None:
#                 if test_type!=None:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#                 else:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#             else:
#                 if test_type!=None:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#                 else:
#                     questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#         # else:
#         #     if class_val!=None:
#         #         if subject_id!=None:
#         #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
#         #         else:
#         #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N').order_by(QuestionDetails.question_id).all()
#         #     else:
#         #         if subject_id!=None:
#         #             questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id).order_by(QuestionDetails.question_id).all()
#         #         else:
#         #             questionList = QuestionDetails.query.filter_by(archive_status='N').order_by(QuestionDetails.question_id).all() 
#         print('QuestionList:'+str(questionList))
#         for ques in questionList:
#             print('inside for of QuestionList')
#             print(ques.question_id)
#         if len(questionList)==0:
#             print('returning 1')
#             return jsonify(['1']) 
#         else:
#             return render_template('testBuilderQuestions.html',questions=questionList)
#     if len(questions)==0:
#         print('returning 1')
#         return jsonify(['1']) 
#     else:
#         return render_template('testBuilderQuestions.html',questions=questions,flagTopic = 'true')


# @app.route('/fetchRequiredQues',methods=['GET','POST'])
# def fetchRequiredQues():
#     class_val = request.args.get('class_val')
#     subject_id = request.args.get('subject_id')
    
#     if class_val!=None:
#         if subject_id!=None:
#             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').all()
#         else:
#             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N').all()
#     else:
#         if subject_id!=None:
#             questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id).all()
#         else:
#             questionList = QuestionDetails.query.filter_by(archive_status='N').all()
#     if len(questionList)==0:
#         print('returning 1')
#         return jsonify(['1'])
#     else:
#         print('returning template'+ str(questionList))
#         return render_template('testBuilderQuestions.html',questions=questionList)

# @app.route('/testBuilderQuestions',methods=['GET','POST'])  
# def testBuilderQuestions():
#     questions=[]
#     topicList=request.get_json()
#     for topic in topicList:
#         # questionList = QuestionDetails.query.join(QuestionOptions, QuestionDetails.question_id==QuestionOptions.question_id).add_columns(QuestionDetails.question_id, QuestionDetails.question_description, QuestionDetails.question_type, QuestionOptions.weightage).filter(QuestionDetails.topic_id == int(topic),QuestionDetails.archive_status=='N' ).filter(QuestionOptions.is_correct=='Y').all()
#         questionList = QuestionDetails.query.filter_by(topic_id = int(topic),archive_status='N').order_by(QuestionDetails.question_id).all()
#         questions.append(questionList)
#     if len(questionList)==0:
#         print('returning 1')
#         return jsonify(['1'])
#     else:
#         print('returning template'+ str(questionList))
#         return render_template('testBuilderQuestions.html',questions=questions)



# @app.route('/testBuilderFileUpload',methods=['GET','POST'])
# def testBuilderFileUpload():
#     class_val = request.args.get('class_val')
#     test_type = request.args.get('test_type')
#     subject_id = request.args.get('subject_id')
#     subject_name = MessageDetails.query.filter_by(msg_id=subject_id).first()
#     date = request.args.get('date')
#     print('class_val:'+str(class_val))
#     print('test_type:'+str(test_type))
#     print('subject_id:'+str(subject_id))
#     print('Date:'+str(date))
#     print('Inside Test builder file upload Test Type value:'+str(test_type))
#     #question_list=request.get_json()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     board_id = SchoolProfile.query.filter_by(school_id = teacher_id.school_id).first()
#     data=request.get_json()
#     question_list=data[0]
#     count_marks=data[1]
    
#     document = Document()
#     print('Date')
#     print(date)
#     document.add_heading(schoolNameVal(), 0)
#     document.add_heading('Class '+str(class_val)+" - "+str(test_type)+" - "+str(date) , 1)
#     document.add_heading("Subject : "+str(subject_name.description),2)
#     document.add_heading("Total Marks : "+str(count_marks),3)
#     p = document.add_paragraph()
#     #For every selected question add the question description
#     for question in question_list:
#         data=QuestionDetails.query.filter_by(question_id=int(question), archive_status='N').first()
#         #for every question add it's options
#         options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
#         #add question desc
#         document.add_paragraph(
#             data.question_description, style='List Number'
#         )    
#     #Add the image associated with the question
#         if data.reference_link!='' and data.reference_link!=None:
#             try:
#                 response = requests.get(data.reference_link, stream=True)
#                 image = BytesIO(response.content)
#                 document.add_picture(image, width=Inches(1.25))
#             except:
#                 pass

#         for option in options:
#             if option.option_desc is not None:
#                 document.add_paragraph(
#                     option.option+". "+option.option_desc)     
#     #document.add_page_break()
#     #naming file here
#     cl = class_val.replace("/","-")
#     file_name=str(teacher_id.school_id)+str(cl)+str(subject_name.description)+str(test_type)+str(datetime.today().strftime("%Y%m%d"))+str(count_marks)+'.docx'
#     file_name = file_name.replace(" ", "")
#     if not os.path.exists('tempdocx'):
#         os.mkdir('tempdocx')
#     document.save('tempdocx/'+file_name)
#     #uploading to s3 bucket
#     client = boto3.client('s3', region_name='ap-south-1')
#     client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
#     #deleting file from temporary location after upload to s3
#     os.remove('tempdocx/'+file_name)

#     ###### Inserting record in the Test Detail table
#     file_name_val='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name

#     # Test create date
#     format = "%Y-%m-%d %H:%M:%S"
#     # Current time in UTC
#     now_utc = datetime.now(timezone('UTC'))
#     print(now_utc.strftime(format))
#     # Convert to local time zone
#     now_local = now_utc.astimezone(get_localzone())
#     print('Date of test creation:'+str(now_local.strftime(format)))
#     # date_utc = date.now(timezone('UTC'))
#     # date_local = date_utc.astimezone(get_localzone())
#     # print('Date of Test:'+str(date_local.strftime(format)))
#     # Test end date
#     testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(count_marks),last_modified_date= datetime.now(),
#         board_id=str(board_id.board_id), subject_id=int(subject_id),class_val=str(class_val),date_of_creation=now_local.strftime(format),
#         date_of_test=str(date), school_id=teacher_id.school_id,test_paper_link=file_name_val, teacher_id=teacher_id.teacher_id)
#     db.session.add(testDetailsUpd)
#     db.session.commit()

#     ##### This section to insert values into test questions table #####
#     #try:
#     createdTestID = TestDetails.query.filter_by(teacher_id=teacher_id.teacher_id).order_by(TestDetails.last_modified_date.desc()).first()
#     for questionVal in question_list:
#         testQuestionInsert= TestQuestions(test_id=createdTestID.test_id, question_id=questionVal, last_modified_date=datetime.now(),is_archived='N')
#         db.session.add(testQuestionInsert)
#     db.session.commit()
#     #except:
#     #    print('error inserting values into the test questions table')
#     #### End of section ####
#     testPaperData= TestDetails.query.filter_by(school_id=teacher_id.school_id,teacher_id=teacher_id.teacher_id).order_by(TestDetails.date_of_creation.desc()).first()
#     sections = ClassSection.query.filter_by(school_id=teacher_id.school_id,class_val=testPaperData.class_val).all()
#     return render_template('testPaperDisplay.html',file_name=file_name_val,testPaperData=testPaperData,sections=sections)

# @app.route('/testPapers')
# @login_required
# def testPapers():
#     indic='DashBoard'
#     return render_template('testPapers.html',indic=indic,title='Test Papers')

# @app.route('/testPaperTable')
# def testPaperTable():
#     paper_count = request.args.get('paper_count')
#     if paper_count=="all":
#         paper_count=500
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
#     #testPaperData= TestDetails.query.filter_by(school_id=teacher_id.school_id).order_by(TestDetails.date_of_creation.desc()).all()
#     testPaperQuery = "select *from test_details where school_id="+ str(teacher_id.school_id)
#     testPaperQuery = testPaperQuery  +" order by date_of_creation desc fetch first "+str(paper_count)+" rows only"
#     print('Query:'+str(testPaperQuery))
#     testPaperData = db.session.execute(testPaperQuery).fetchall()
#     subjectNames=MessageDetails.query.filter_by(category='Subject')
#     return render_template('_testPaperTable.html',testPaperData=testPaperData,subjectNames=subjectNames,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

# @app.route('/findSection',methods=['GET','POST'])
# def findSection():
#     class_val=request.args.get('class_val')
#     school_id = User.query.filter_by(id=current_user.id).first()
#     print('Class:'+str(class_val))
#     print('School Id:'+str(school_id.school_id))
#     sections = ClassSection.query.filter_by(school_id=school_id.school_id,class_val=class_val)
#     sec = []
#     for section in sections:
#         sec.append(section.section)
#     return jsonify(sec)

# @app.route('/getQuestionDetails')
# def getQuestionDetails():
#     print('inside getQuestionDetails')
#     qtest_id = request.args.get('test_id')
#     print('Test Id:'+str(qtest_id)) 
#     getQuestionQuery = "select tq.question_id,qd.question_type,qd.question_description,td.total_marks as weightage from test_questions tq "
#     getQuestionQuery = getQuestionQuery + "inner join question_details qd on tq.question_id = qd.question_id "
#     getQuestionQuery = getQuestionQuery + "inner join test_details td on td.test_id = tq.test_id where td.test_id = '"+str(qtest_id)+"'"
#     testQuestionsDetails = db.session.execute(text(getQuestionQuery)).fetchall()
#     for questions in testQuestionsDetails:
#         print(questions.question_description)
#     totalQues = len(testQuestionsDetails)
#     return render_template('_getQuestionDetails.html',totalQues=totalQues,testQuestionsDetails=testQuestionsDetails)

# @app.route('/getChapterDetails')
# def getChapterDetails():
#     qtest_id=request.args.get('test_id')
#     getChapterQuery = "select distinct topic_name, chapter_name, chapter_num from test_questions tq "
#     getChapterQuery= getChapterQuery+ "inner join question_details qd  on "
#     getChapterQuery= getChapterQuery+ " qd.question_id=tq.question_id inner join "
#     getChapterQuery= getChapterQuery+ " topic_detail td on td.topic_id=qd.topic_id "
#     getChapterQuery= getChapterQuery+ "where tq.test_id='"+str(qtest_id)+"'"

#     getChapterRows = db.session.execute(text(getChapterQuery)).fetchall()

#     return render_template('_getChapterDetails.html', getChapterRows=getChapterRows)

# @app.route('/calendar')
# @login_required
# def calendar():
#     return render_template('calendar.html')

# @app.route('/schoolPerformanceRanking')
# @login_required
# def schoolPerformanceRanking():
#     return render_template('schoolPerformanceRanking.html')

# @app.route('/recommendations')
# @login_required
# def recommendations():
#     return render_template('recommendations.html')


# @app.route('/attendance')
# @login_required
# def attendance():
#     return render_template('attendance.html')

@app.route('/guardianDashboard')
@login_required
def guardianDashboard():
    print('Id:'+str(current_user.id))
    user_type_val = ''
    if current_user.user_type==72:
        user_type_val = 72
        print('User Type:'+str(user_type_val))
    guardian=GuardianProfile.query.filter_by(user_id=current_user.id).all()
    print(guardian)
    print('Id:'+str(current_user.id))
    school= User.query.filter_by(id=current_user.id).first()
    student=[]
    
    for g in guardian:
        if g.student_id!=None:
            query = "select fn.score as marks,fn.strong_2_subjects as subjects,sp.full_name as student,spro.school_name as school,cs.class_val as class,cs.section as section,sp.profile_picture as pic,sp.student_id from student_profile sp "
            query = query + "inner join class_section cs on cs.class_sec_id=sp.class_sec_id "
            query = query + "left join fn_guardian_dashboard_summary('"+str(school.school_id)+"') fn on fn.student_name=sp.full_name "
            query = query + "inner join school_profile spro on spro.school_id = sp.school_id where student_id='"+str(g.student_id)+"' order by marks desc limit 1"
            student_data = db.session.execute(text(query)).first()
            student.append(student_data)
    
    return render_template('guardianDashboard.html',students=student,disconn = 1,user_type_val=user_type_val)

@app.route('/performanceDetails/<student_id>',methods=['POST','GET'])
@login_required
def performanceDetails(student_id):
    user = User.query.filter_by(username=current_user.username).first_or_404()        
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    
    form1=studentPerformanceForm()
    
    available_class=ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher.school_id).all()
    available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher.school_id).all()    
    available_test_type=MessageDetails.query.filter_by(category='Test type').all()
    available_student_list=StudentProfile.query.filter_by(school_id=teacher.school_id).all()


    class_list=[(str(i.class_val), "Class "+str(i.class_val)) for i in available_class]
    section_list=[(i.section,i.section) for i in available_section]    
    test_type_list=[(i.msg_id,i.description) for i in available_test_type]
    student_list=[(i.student_id,i.full_name) for i in available_student_list]

    #selectfield choices
    form1.class_val1.choices = class_list
    form1.section1.choices= section_list    
    form1.test_type1.choices=test_type_list
    form1.student_name1.choices = student_list

    student=StudentProfile.query.filter_by(student_id=student_id).first()
    if request.method=='POST':
        student=StudentProfile.query.filter_by(student_id=student_id).first()
        class_sec=ClassSection.query.filter_by(class_sec_id=student.class_sec_id).first()
        subject=subjectPerformance(class_sec.class_val,class_sec.school_id)
        print(subject)
        date=request.form['performace_date']        


        return render_template('studentPerfDetails.html',date=date,subjects=subject,students=student)
    return render_template('performanceDetails.html',students=student, student_id=student_id,form1=form1)


@app.route('/studentfeedbackreporttemp')
def studentfeedbackreporttemp():
    student_name=request.args.get('student_name')
    return render_template('studentfeedbackreporttemp.html',student_name=student_name)

@app.route('/unpaidStudentsList',methods=["GET","POST"])
def unpaidStudentsList():
    class_val = request.args.get('class_val')
    section = request.args.get('section')
    
    teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    class_sec_id = ClassSection.query.filter_by(class_val=class_val,section=section,school_id=teacherData.school_id).first()
    studentData = "select sp.full_name from student_profile sp where class_sec_id ='"+str(class_sec_id.class_sec_id)+"' and school_id='"+str(teacherData.school_id)+"' and student_id not in (select student_id from fee_detail where paid_status='Y' and class_sec_id='"+str(class_sec_id.class_sec_id)+"' and school_id='"+str(teacherData.school_id)+"')"
    print(studentData)
    studentList = db.session.execute(text(studentData)).fetchall()
    return render_template('_studentList.html',studentList=studentList)

@app.route('/sendFeeSMS',methods=["GET","POST"])
def sendFeeSMS():
    if request.method=="POST":
        commType = request.form.get('commType')
        message = request.form.get('message')
        class_val = request.form.get('qclass_val')
        section = request.form.get('qsection')
        teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
        class_sec = ClassSection.query.filter_by(school_id=teacherData.school_id,class_val=class_val,section=section).first()
        class_sec_id = class_sec.class_sec_id
        if class_sec_id !=None and class_sec_id !="":
            commDataAdd=CommunicationDetail(message = message,status=231 , school_id=teacherData.school_id,             
            teacher_id=teacherData.teacher_id,last_modified_date =datetime.today())
            db.session.add(commDataAdd)
            #db.session.flush()
            db.session.commit()
            getStudNumQuery = "select student_id , phone from student_profile sp where class_sec_id ='"+str(class_sec_id)+"' and school_id='"+str(teacherData.school_id)+"' and student_id not in (select student_id from fee_detail where paid_status='Y' and class_sec_id='"+str(class_sec_id)+"' and school_id='"+str(teacherData.school_id)+"')"            
            studentPhones = db.session.execute(getStudNumQuery).fetchall()
            phoneList =[]
            for phoneRow in studentPhones:
                if phoneRow.phone!=None and phoneRow.phone!='':
                    phoneList.append(phoneRow.phone)
            if studentPhones!=None:
                if commType=='sms':
                    apiPath = "http://173.212.233.109/app/smsapisr/index.php?key=35EF8379A04DB8&"
                    apiPath = apiPath + "campaign=9967&routeid=6&type=text&"
                    apiPath = apiPath + "contacts="+str(phoneList).replace('[','').replace(']','').replace('\'','').replace(' ','')+"&senderid=GLOBAL&"
                    apiPath = apiPath + "msg="+ quote(message)
                    print(apiPath)
                    ##Sending message here
                    try:
                        r = requests.post(apiPath)                    
                        returnData = str(r.text)
                        print(returnData)
                        if "SMS-SHOOT-ID" in returnData:
                            for val in studentPhones:
                                commTransAdd = CommunicationTransaction(comm_id=commDataAdd.comm_id, student_id=val.student_id,last_modified_date = datetime.today())
                                db.session.add(commTransAdd)
                            commDataAdd.status=232
                            db.session.commit()                                        
                            return jsonify(['0'])
                        else:
                            return jsonify(['1'])                    
                    except:
                        return jsonify(['1'])
                    ##Message sent 
    return jsonify(['1']) 

@app.route('/sendSMS',methods=["GET","POST"])
def sendSMS():
        # commType = request.form.get('commType')
    commType = 'sms'
    message = request.args.get('message')
    student_id = request.args.get('student_id')
    print('Message:'+str(message))
    print('student_id:'+str(student_id))
    class_sec = StudentProfile.query.filter_by(student_id=student_id).first()
    class_sec_id = class_sec.class_sec_id
    if class_sec_id !=None and class_sec_id !="":
        teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
            #insert into communication detail table
        commDataAdd=CommunicationDetail(message = message,status=231 , school_id=teacherData.school_id,             
        teacher_id=teacherData.teacher_id,last_modified_date =datetime.today())
        db.session.add(commDataAdd)
            #db.session.flush()
        db.session.commit() 
        getStudNumQuery = "select student_id, phone from student_profile sp where class_sec_id ="+ str(class_sec_id)            
        studentPhones = db.session.execute(getStudNumQuery).fetchall()
        phoneList =[]
        for phoneRow in studentPhones:
            if phoneRow.phone!=None and phoneRow.phone!='':
                phoneList.append(phoneRow.phone)
        if studentPhones!=None:
            if commType=='sms':
                apiPath = "http://173.212.233.109/app/smsapisr/index.php?key=35EF8379A04DB8&"
                apiPath = apiPath + "campaign=9967&routeid=6&type=text&"
                apiPath = apiPath + "contacts="+str(phoneList).replace('[','').replace(']','').replace('\'','').replace(' ','')+"&senderid=GLOBAL&"
                apiPath = apiPath + "msg="+ quote(message)
                print(apiPath)
                    ##Sending message here
                try:
                    r = requests.post(apiPath)                    
                    returnData = str(r.text)
                    print(returnData)
                    if "SMS-SHOOT-ID" in returnData:
                        for val in studentPhones:
                            commTransAdd = CommunicationTransaction(comm_id=commDataAdd.comm_id, student_id=val.student_id,last_modified_date = datetime.today())
                            db.session.add(commTransAdd)
                        commDataAdd.status=232
                        db.session.commit()                                        
                        return jsonify(['0'])
                    else:
                        return jsonify(['1'])                    
                except:
                    return jsonify(['1'])
                    ##Message sent                    
    return jsonify(['1'])

@app.route('/sendComm',methods=["GET","POST"])
def sendComm():
    if request.method=="POST":
        commType = request.form.get('commType')
        message = request.form.get('message')
        class_sec_id = request.form.get('class_sec_id')
        if class_sec_id !=None and class_sec_id !="":
            teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
            #insert into communication detail table
            commDataAdd=CommunicationDetail(message = message,status=231 , school_id=teacherData.school_id,             
            teacher_id=teacherData.teacher_id,last_modified_date =datetime.today())
            db.session.add(commDataAdd)
            #db.session.flush()
            db.session.commit() 
            getStudNumQuery = "select student_id, phone from student_profile sp where class_sec_id ="+ str(class_sec_id)            
            studentPhones = db.session.execute(getStudNumQuery).fetchall()
            phoneList =[]
            for phoneRow in studentPhones:
                if phoneRow.phone!=None and phoneRow.phone!='':
                    phoneList.append(phoneRow.phone)
            if studentPhones!=None:
                if commType=='sms':
                    apiPath = "http://173.212.233.109/app/smsapisr/index.php?key=35EF8379A04DB8&"
                    apiPath = apiPath + "campaign=9967&routeid=6&type=text&"
                    apiPath = apiPath + "contacts="+str(phoneList).replace('[','').replace(']','').replace('\'','').replace(' ','')+"&senderid=GLOBAL&"
                    apiPath = apiPath + "msg="+ quote(message)
                    print(apiPath)
                    ##Sending message here
                    try:
                        r = requests.post(apiPath)                    
                        returnData = str(r.text)
                        print(returnData)
                        if "SMS-SHOOT-ID" in returnData:
                            for val in studentPhones:
                                commTransAdd = CommunicationTransaction(comm_id=commDataAdd.comm_id, student_id=val.student_id,last_modified_date = datetime.today())
                                db.session.add(commTransAdd)
                            commDataAdd.status=232
                            db.session.commit()                                        
                            return jsonify(['0'])
                        else:
                            return jsonify(['1'])                    
                    except:
                        return jsonify(['1'])
                    ##Message sent                    
    return jsonify(['1'])


@app.route('/class')
@login_required
def classCon(): 
    if current_user.is_authenticated:        
        user = User.query.filter_by(username=current_user.username).first_or_404()        
        teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    
        qclass_val = request.args.get('class_val')
        qsection=request.args.get('section')

        #db query

        classSections=ClassSection.query.filter_by(school_id=teacher.school_id).all()
        count = 0
        for section in classSections:
            #print("Class Section:"+section.section)
            #this section is to load the page for the first class section if no query value has been provided
            if count==0:
                getClassVal = section.class_val
                getSection = section.section
                count+=1

        distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher.school_id)+" GROUP BY class_val order by s")).fetchall()
        #if no value has been passed for class and section in query string then use the values fetched from db
        if qclass_val==None:
            qclass_val = getClassVal
            qsection=getSection
            
        selectedClassSection=ClassSection.query.filter_by(school_id=teacher.school_id, class_val=str(qclass_val), section=qsection).order_by(ClassSection.class_val).first()
        topicTrackerQuery = "with cte_total_topics as "
        topicTrackerQuery = topicTrackerQuery + "(select subject_id,  "
        topicTrackerQuery = topicTrackerQuery +"count(is_covered) as total_topics , max(last_modified_Date) as last_updated_date "
        topicTrackerQuery = topicTrackerQuery +"  from topic_tracker where class_sec_id = '"+ str(selectedClassSection.class_sec_id)+"' group by subject_id)  "
        topicTrackerQuery = topicTrackerQuery +"select c1.subject_id,  t2.description as subject_name, c1.last_updated_date, "
        topicTrackerQuery = topicTrackerQuery +"CASE WHEN COUNT(t1.subject_id) <> 0 THEN COUNT(c1.subject_id) ELSE 0 END "
        topicTrackerQuery = topicTrackerQuery +"topics_covered, c1.total_topics  "
        topicTrackerQuery = topicTrackerQuery +"from topic_tracker t1  "
        topicTrackerQuery = topicTrackerQuery +"right outer join cte_total_topics c1  "
        topicTrackerQuery = topicTrackerQuery +"on c1.subject_id=t1.subject_id and class_sec_id= '"+ str(selectedClassSection.class_sec_id)+"'  "
        topicTrackerQuery = topicTrackerQuery +"and t1.is_covered='Y'  "
        topicTrackerQuery = topicTrackerQuery +"inner join   "
        topicTrackerQuery = topicTrackerQuery +"message_detail t2 on   "
        topicTrackerQuery = topicTrackerQuery +"c1.subject_id=t2.msg_id  "
        topicTrackerQuery = topicTrackerQuery +"group by c1.subject_id, t2.description, c1.total_topics,  c1.last_updated_date"                
        topicRows  = db.session.execute(text(topicTrackerQuery)).fetchall()
        #print('this is the number of topicRows' + str(len(topicRows)))

        ##Section for login info
        loginDataQuery = "select *from fn_student_last_login("+str(selectedClassSection.class_sec_id)+")"
        loginData = db.session.execute(loginDataQuery).fetchall()
        ##
        
        #Summary query
        summaryQuery = "with perfCTE as(select round(avg(student_score ),2) as avgClassPerfomance "
        summaryQuery = summaryQuery + " from performance_detail pd where class_sec_id ="+ str(selectedClassSection.class_sec_id)
        summaryQuery = summaryQuery + "), studCountCTE as (select count(student_id) studcount from student_profile "
        summaryQuery = summaryQuery + " sp where class_sec_id =" + str(selectedClassSection.class_sec_id) 
        summaryQuery = summaryQuery + ") select t1.avgclassperfomance as avgclassperfomance, t2.studcount as studcount "
        summaryQuery = summaryQuery +" from perfCTE t1, studCountCTE t2"
        #print(summaryQuery)
        summaryData = db.session.execute(text(summaryQuery)).first()
        # End of query section 
        attendance = "select count(*) from attendance where class_sec_id="+str(selectedClassSection.class_sec_id)+" and school_id="+str(teacher.school_id)+" and is_present='Y'"
        attendance = db.session.execute(text(attendance)).first()
        print('Attendance:'+str(attendance[0]))
        indic='class'
        return render_template('class.html',indic=indic,attendance=attendance, classSecCheckVal=classSecCheck(),classsections=classSections,summaryData=summaryData, 
            qclass_val=qclass_val, qsection=qsection, class_sec_id=selectedClassSection.class_sec_id, distinctClasses=distinctClasses,
            topicRows=topicRows,title='Class', user_type_val=str(current_user.user_type), loginData=loginData)
    else:
        return redirect(url_for('accounts.login'))    

@app.route('/topicList')
def topicList():
    class_sec_id = request.args.get('class_sec_id','1')
    subject_id = request.args.get('subject_id','15')
    class_val = request.args.get('class_val')
    teacher = ''
    if current_user.user_type==134:
        teacher = StudentProfile.query.filter_by(user_id=current_user.id).first()
    else:
        teacher= TeacherProfile.query.filter_by(user_id=current_user.id).first() 
    #topicList = TopicTracker.query.filter_by(subject_id=subject_id, class_sec_id=class_sec_id).all()
    topicListQuery = "select distinct t1.subject_id, t3.description as subject_name, t1.topic_id, t2.topic_name,t1.is_covered, "
    topicListQuery = topicListQuery + "t2.chapter_num, t2.unit_num, t4.book_name from topic_tracker t1 "
    topicListQuery = topicListQuery + "inner join topic_detail t2 on t1.topic_id=t2.topic_id "
    topicListQuery = topicListQuery + "inner join message_detail t3 on t1.subject_id=t3.msg_id "
    topicListQuery = topicListQuery + "inner join book_details t4 on t4.book_id=t2.book_id where "
    topicListQuery = topicListQuery + "t2.book_id in (select bd.book_id from book_details bd inner join topic_detail td on td.book_id = bd.book_id inner join topic_tracker tt on td.topic_id = tt.topic_id where bd.class_val = '"+str(class_val)+"' and bd.subject_id = '"+str(subject_id)+"' and tt.school_id = '"+str(teacher.school_id)+"') and "
    topicListQuery = topicListQuery + "t1.subject_id = '" + subject_id+"' and t1.is_archived='N' and t1.school_id='"+str(teacher.school_id)+"' and t1.class_sec_id='" +class_sec_id+"' order by  t2.chapter_num, is_covered desc"
    print('inside topicList')
    print(topicListQuery)
    topicList= db.session.execute(text(topicListQuery)).fetchall()

    return render_template('_topicList.html', topicList=topicList, class_sec_id=class_sec_id,class_val=class_val)

# @app.route('/setGoogleLogin',methods=['POST','GET'])
# def setGoogleLogin():
#     isgoogleLogin = request.args.get('isGoogleLogin')
#     school_id = request.args.get('school_id')
#     schoolData = SchoolProfile.query.filter_by(school_id=school_id).first()
#     print(isgoogleLogin)
#     schoolData.google_login = isgoogleLogin
#     db.session.commit()
#     session['isGooglelogin'] = isgoogleLogin
#     return jsonify([0])

@app.route('/setSchoolName',methods=['POST','GET'])
def setSchoolName():
    isSchoolName = request.args.get('isSchoolName')
    school_id = request.args.get('school_id')
    print(school_id)
    schoolData = SchoolProfile.query.filter_by(school_id=school_id).first()
    print(isSchoolName)
    schoolData.show_school_name = isSchoolName
    db.session.commit()
    session['show_school_name'] = isSchoolName
    return jsonify([0])

@app.route('/qrSessionScanner')
@login_required
def qrSessionScanner():
    return render_template('qrSessionScanner.html')


@app.route('/qrSessionScannerStudent')
@login_required
def qrSessionScannerStudent():
    
    studentDetails = StudentProfile.query.filter_by(user_id=current_user.id).first()
    
    testHistoryQuery = "SELECT *FROM fn_student_performance_response_capture("+str(studentDetails.student_id)+") order by test_date desc limit 50"
    testHistory = db.session.execute(testHistoryQuery).fetchall()
    
    return render_template('qrSessionScannerStudent.html',user_type_val=str(current_user.user_type),studentDetails=studentDetails,testHistory=testHistory)

@app.route('/fetchStudentTestReport', methods=['GET','POST'])
def fetchStudentTestReport():
    student_id = request.args.get('student_id')
    print('Student Id:'+str(student_id))
    testHistoryQuery = "SELECT *FROM fn_student_performance_response_capture("+str(student_id)+") order by test_date desc limit 50"
    testHistory = db.session.execute(testHistoryQuery).fetchall()
    testResData = []
    for test in testHistory:
        testData = {}
        testData['subject'] = test.subject
        testData['topics'] = test.topics
        testData['test_date'] = test.test_date.strftime('%d %B %Y')
        testData['perf_percentage'] = test.perf_percentage
        testData['resp_session_id'] = test.resp_session_id
        testResData.append(testData)
    print('Test Res Data:')
    print(testResData)
    return jsonify({'testResult':testResData})


@app.route('/viewHomework')
@login_required
def viewHomework():
    return render_template('viewHomework.html')

@app.route('/mobFeedbackCollection', methods=['GET', 'POST'])
def mobQuestionLoader():

    resp_session_id=request.args.get('resp_session_id')
    print('Response Session Id in mobFeedbackCollection:'+str(resp_session_id))
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id=str(resp_session_id)).first()
    if sessionDetailRow:
        if sessionDetailRow.session_status=='80':
            sessionDetailRow.session_status='81'        
            db.session.commit()    
        classSectionRow = ClassSection.query.filter_by(class_sec_id=sessionDetailRow.class_sec_id).first()        
        testDetailRow = TestDetails.query.filter_by(test_id = sessionDetailRow.test_id).first()
        testQuestions = TestQuestions.query.filter_by(test_id=sessionDetailRow.test_id).all()

        if testQuestions!=None:
            questionListSize = len(testQuestions)
        return render_template('mobFeedbackCollection.html',class_val = classSectionRow.class_val, 
            section=classSectionRow.section,questionListSize=questionListSize,
            resp_session_id=str(resp_session_id), questionList=testQuestions, subject_id=testDetailRow.subject_id, test_type=testDetailRow.test_type,disconn=1)
    else:
        flash('This is not a valid id')
        return render_template('qrSessionScanner.html')




@app.route('/startPracticeTest', methods=['GET', 'POST'])
def startPracticeTest():
    #topics = request.get_json()
    #for topic in topics:
    
    topics = request.args.getlist('topics')
    print('topic:'+str(topics))
    topicsList = str(topics).replace('[','').replace(']','').replace('\'','')
    difficulty = request.args.get('difficulty')    
    qcount = request.args.get('qcount')
    subject_id = request.args.get('subject_id')
    class_val = request.args.get('class_val')
    print('difficulty:'+str(difficulty))
    print('qcount:'+str(qcount))
    print('subject_id:'+str(subject_id))
    print('class_val:'+str(class_val))
    #subject_id
    #board_id = request.args.get('board_id')
    #topicList = request.for.getlist('topicList')adsfsdfasdf

    if current_user.is_anonymous:
        print('for anonymous user')
        studentData = StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()
    else:
        print('for student:'+str(current_user.id))
        studentData = StudentProfile.query.filter_by(user_id=current_user.id).first()
    schoolData = SchoolProfile.query.filter_by(school_id = studentData.school_id).first()
    classSecData = ClassSection.query.filter_by(school_id=studentData.school_id, class_val=str(class_val)).first()
    #Collection Questions
    questions = []
    total_marks = 0
    questionIDList = []
    if topics:
        #for topic in topics:
            #questionList = QuestionDetails.query.filter_by(archive_status='N',topic_id=topic,question_type='MCQ1').all()
        questionListQuery = "select *from question_details where archive_status='N' and question_type='MCQ1' and topic_id in (" + str(topicsList) + ")"
        questionListQuery = questionListQuery + " order by random() " #limit " +  str(qcount)
        questionList = db.session.execute(text(questionListQuery)).fetchall()
        if questionList:  
            for q in questionList:
                if len(questions)< int(qcount):
                    questions.append(q)  
    for val in questions:        
        print(str(val))
        total_marks = total_marks + val.suggested_weightage
        questionIDList.append(val.question_id)
    ##Create test
    testDetailsAdd = TestDetails(test_type='238', total_marks=str(total_marks),last_modified_date= datetime.today(),
        board_id=str(schoolData.board_id), subject_id=int(subject_id),class_val=str(class_val),date_of_creation=datetime.today(),
        date_of_test=str(datetime.today()), school_id=studentData.school_id)        
    db.session.add(testDetailsAdd)
    if current_user.is_anonymous==False:
        if studentData.points:
            studentData.points= int(studentData.points) + 1
    db.session.commit()
    print('test_id:'+str(testDetailsAdd.test_id))
    print('Data feed to test details complete')

    ##### This section to insert values into test questions table #####        
    for questionVal in questions:
        testQuestionInsert= TestQuestions(test_id=testDetailsAdd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now())
        db.session.add(testQuestionInsert)
    db.session.commit()
    
    print('Data feed to test questions complete')
    ##Create response session ID
    dateVal= datetime.today().strftime("%d%m%Y%H%M%S")

    responseSessionID = str(subject_id).strip()+ str(dateVal).strip() + str(classSecData.class_sec_id).strip()

    print('resp session id:'+str(responseSessionID))
    print('Response ID generated')

    if current_user.is_anonymous:
        if session.get('anonUser'):
            print('the anon user is true')
            #this section forces a user to take an unsigned test only once
            #flash('Please login to start any further tests') 
            #return jsonify(['2']) 
        else:
            session['anonUser'] = "user_"+ str(responseSessionID) + '_'+str(random.randint(1,100))
            print("This is the value of anon user: " + session['anonUser'])            
    else:
        print('last segment. anon user is false')
        session['anonUser']==False

    ##Create session
    if len(questions) >0:  
        format = "%Y-%m-%d %H:%M:%S"
        # Current time in UTC
        now_utc = datetime.now(timezone('UTC'))
        print(now_utc.strftime(format))
        # Convert to local time zone
        now_local = now_utc.astimezone(get_localzone())
        print(now_local.strftime(format))      
        sessionDetailRowInsert=SessionDetail(resp_session_id=responseSessionID,session_status='80',
            class_sec_id=classSecData.class_sec_id,test_id=testDetailsAdd.test_id, last_modified_date=now_local.strftime(format),correct_marks=10,incorrect_marks=0,test_time=0,total_marks=total_marks )
        db.session.add(sessionDetailRowInsert)
        db.session.commit()        

        print('Data feed to session detail completed')
        ## Start test
        return jsonify([responseSessionID])    
    else:
        return jsonify(['1'])  


@app.route('/feedbackCollectionStudDev', methods=['GET', 'POST'])
@login_required
def feedbackCollectionStudDev():
    resp_session_id=request.args.get('resp_session_id')
    print('inside feedbackCollectionStudDev')
    print('Resp_session_id:'+str(resp_session_id))
    instructionsRows = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    if instructionsRows:
        instructions = instructionsRows.instructions
    else:
        instructions = ''
    # student_id = request.args.get('student_id')
    school_id = request.args.get('school_id')
    school_profile_data = SchoolProfile.query.filter_by(school_id=school_id).first()
    primaryColor = school_profile_data.primary_color
    uploadStatus=request.args.get('uploadStatus')
    resultStatus = request.args.get('resultStatus')
    advance = request.args.get('advance')
    print('upload status:'+str(uploadStatus))
    print('result status:'+str(resultStatus))
    print('advance:'+str(advance))
    # print('Student Id:'+str(student_id))
    # studId = None
    # if student_id!=None:
    #     studId=student_id
    # if current_user.is_anonymous:
    #     print('user not registered')
    # else:
    #     studentDetails = StudentProfile.query.filter_by(user_id=current_user.id).first()
    #     studId = studentDetails.student_id

    # if studId==None:
    #     print('Student Id is null')
    #     return render_template('feedbackCollectionStudDev.html',resp_session_id=str(resp_session_id),studId=studId,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)
    # emailDet = StudentProfile.query.filter_by(student_id=studId).first()
    # user = ''
    # if emailDet:
    #     user = User.query.filter_by(email=emailDet.email).first()
    # if user:
    #     login_user(user,remember='Y')
    #     session['schoolName'] = schoolNameVal()
        
    #     print('user name')
    #     #print(session['username'])
    #     school_id = ''
    #     print('user type')
    #     #print(session['userType'])
    #     session['studentId'] = ''
    #     if current_user.user_type==253:
    #         school_id=1
    #     elif current_user.user_type==71:
    #         teacherProfileData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #         school_id = teacherProfileData.school_id
    #     elif current_user.user_type==134:
    #         studentProfileData = StudentProfile.query.filter_by(user_id=current_user.id).first()
    #         school_id = studentProfileData.school_id            
    #         session['studentId'] = studentProfileData.student_id
    #     else:
    #         userData = User.query.filter_by(id=current_user.id).first()
    #         school_id = userData.school_id

    #     school_pro = SchoolProfile.query.filter_by(school_id=school_id).first()
    #     session['school_logo'] = ''
    #     if school_pro:
    #         session['school_logo'] = school_pro.school_logo
    #         session['schoolPicture'] = school_pro.school_picture
    #     query = "select user_type,md.module_name,description, module_url, module_type from module_detail md inner join module_access ma on md.module_id = ma.module_id where user_type = '"+str(current_user.user_type)+"' and ma.is_archived = 'N' and md.is_archived = 'N' order by module_type"
    #     print(query)
    #     print('Modules')
    #     moduleDetRow = db.session.execute(query).fetchall()
    #     print('School profile')
    #     #print(session['schoolPicture'])
    #     # det_list = [1,2,3,4,5]
    #     session['moduleDet'] = []
    #     detList = session['moduleDet']
        
    #     for det in moduleDetRow:
    #         eachList = []
    #         print(det.module_name)
    #         print(det.module_url)
    #         eachList.append(det.module_name)
    #         eachList.append(det.module_url)
    #         eachList.append(det.module_type)
    #         # detList.append(str(det.module_name)+":"+str(det.module_url)+":"+str(det.module_type))
    #         detList.append(eachList)
    #     session['moduleDet'] = detList
    # else:
    #     flash('please create student account first')
    #     return render_template('feedbackCollectionStudDev.html',resp_session_id=str(resp_session_id),studId=None)

    # print('student_id in feedbackCollectionStudDev:'+str(studId))
    print('Response Session Id:'+str(resp_session_id))
    studentRow = StudentProfile.query.filter_by(user_id=current_user.id).first()
    studId = studentRow.student_id
    classData = ClassSection.query.filter_by(class_sec_id=studentRow.class_sec_id).first()
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id=str(resp_session_id)).first()
    print('Session Detail Row:'+str(sessionDetailRow))
    if sessionDetailRow==None:
        print('Response Session id is wrong')
        flash('please enter correct test id')
        return render_template('qrSessionScannerStudent.html',user_type_val=str(current_user.user_type),studentDetails=studentRow)
    testDet = TestDetails.query.filter_by(test_id=sessionDetailRow.test_id).first()
    print('Test of Class:'+str(testDet.class_val))
    print('Student of class:'+str(classData.class_val))
    responseExist = "select rc.student_id,sd.test_id from response_capture rc inner join session_detail sd on rc.resp_session_id = sd.resp_session_id where sd.resp_session_id='"+str(resp_session_id)+"' and rc.student_id='"+str(studId)+" '"
    responseExist = db.session.execute(text(responseExist)).first()
    
    if responseExist:
        print('Inside if test already attempt')
        flash('Sorry, you have already attempt this test')
        studenId = None
        return render_template('feedbackCollectionStudDev.html',resp_session_id=str(resp_session_id),studId=studenId)
    if((str(testDet.class_val)!=str(classData.class_val)) and (school_id==classData.school_id)):
        print('Inside if classes are same')
        flash('Sorry, you can not attempt this test')
        studenId = None
        return render_template('feedbackCollectionStudDev.html',resp_session_id=str(resp_session_id),studId=studenId)
    if sessionDetailRow!=None:
        print("This is the session status - "+str(sessionDetailRow.session_status))
        if sessionDetailRow.session_status=='80':
            sessionDetailRow.session_status='81'        
            db.session.commit()    
        classSectionRow = ClassSection.query.filter_by(class_sec_id=sessionDetailRow.class_sec_id).first()        
        testDetailRow = TestDetails.query.filter_by(test_id = sessionDetailRow.test_id).first()
        testQuestions = TestQuestions.query.filter_by(test_id=sessionDetailRow.test_id,is_archived='N').all()

        if testQuestions!=None:
            questionListSize = len(testQuestions)
        print('Student ID:'+str(studentRow.student_id))
        return render_template('feedbackCollectionStudDev.html',class_val = classSectionRow.class_val, 
            section=classSectionRow.section,questionListSize=questionListSize,
            resp_session_id=str(resp_session_id),primaryColor=primaryColor, questionList=testQuestions, subject_id=testDetailRow.subject_id, test_type=testDetailRow.test_type,disconn=1,student_id = studId,studentName=studentRow.full_name,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance,instructions=instructions)
    else:
        flash('This is not a valid id or there are no question in this test')
        return redirect('index')
        #return render_template('qrSessionScannerStudent.html',disconn=1)


# @app.route('/updateQuestion')
# def updateQuestion():
#     question_id = request.args.get('question_id')
#     updatedCV = request.args.get('updatedCV')
#     topicId = request.args.get('topicName')
#     subId = request.args.get('subName')
#     qType = request.args.get('qType')
#     qDesc = request.args.get('qDesc')
#     corrans = request.args.get('corrans')
#     weightage = request.args.get('weightage')
#     print('Weightage:'+str(weightage))
#     print('Correct option:'+str(str(corrans)))
#     preview = request.args.get('preview')
#     options = request.args.get('options')
#     op1 = request.args.get('op1')
#     op2 = request.args.get('op2')
#     op3 = request.args.get('op3')
#     op4 = request.args.get('op4')
#     print(op1)
#     print(op2)
#     print(op3)
#     print(op4)
#     form = QuestionBuilderQueryForm()
#     print("Updated class Value+:"+updatedCV)
#     print(str(updatedCV)+" "+str(topicId)+" "+str(subId)+" "+str(qType)+" "+str(qDesc)+" "+str(preview)+" "+str(corrans)+" "+str(weightage))
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
#     form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
#     form.topics.choices=[(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]
#     flag = False
#     # updateQuery = "update question_details t1 set topic_id='" + str(topicId) + "' where question_id='" + question_id + "'"
#     updateQuery = "update question_details set class_val='" + str(updatedCV) +  "',topic_id='"+ str(topicId) + "',subject_id='"+ str(subId) + "',question_type='" + str(qType) + "',question_description='"+ str(qDesc) + "',reference_link='"+ str(preview) +"' where question_id='" + str(question_id) + "'"

#     queryOneExe = db.session.execute(text(updateQuery))
#     updateWeightage = "update question_details set suggested_weightage='" + str(weightage) + "' where question_id='" + str(question_id) + "'" 
#     querytwoExe = db.session.execute(text(updateWeightage))
    

#     option_id_list = QuestionOptions.query.filter_by(question_id=question_id).order_by(QuestionOptions.option_id).all()
#     if corrans:
#         print(option_id_list)
#         i=0
#         opId1=''
#         opId2=''
#         opId3=''
#         opId4=''
#         for opt in option_id_list:
#             if i==0:
#                 opId1 = opt.option_id
#             elif i==1:
#                 opId2 = opt.option_id
#             elif i==2:
#                 opId3 = opt.option_id
#             else:
#                 opId4 = opt.option_id
#             i=i+1
#         print("Options order of id:"+str(opId1)+' '+str(opId2)+' '+str(opId3)+' '+str(opId4))
#         # print(option_id) 
#         # option = "select option_id from question_options where question_id='"+ str(question_id) + "'"
#         # opt = db.session.execute(text(option))
#         # print("Updated options "+str(opt))
#         if opId1 and opId2 and opId3 and opId4:
#             updateOption1 = "update question_options set option_desc='"+str(op1)+"' where option_id='"+ str(opId1) + "'"
#             print(updateOption1)
#             updateOpt1Exe = db.session.execute(text(updateOption1))
#             updateOption2 = "update question_options set option_desc='"+str(op2)+"' where option_id='"+ str(opId2) + "'"
#             print(updateOption2)
#             updateOpt2Exe = db.session.execute(text(updateOption2))
#             updateOption3 = "update question_options set option_desc='"+str(op3)+"' where option_id='"+ str(opId3) + "'"
#             print(updateOption3)
#             updateOpt3Exe = db.session.execute(text(updateOption3))
#             updateOption4 = "update question_options set option_desc='"+str(op4)+"' where option_id='"+ str(opId4) + "'"
#             print(updateOption4)
#             updateOpt4Exe = db.session.execute(text(updateOption4))
#             if str(corrans)!='':
#                 updatequery1 = "update question_options set is_correct='N' where is_correct='Y' and question_id='" +str(question_id)+"'"
#                 print(updatequery1)
#                 update1 = db.session.execute(text(updatequery1))
#                 updateCorrectOption = "update question_options set is_correct='Y' where option_desc='"+str(corrans)+"' and question_id='"+str(question_id)+"'"
#                 print(updateCorrectOption)
#                 updateOp = db.session.execute(text(updateCorrectOption))
#         else:
#             optionlist = []
#             optionlist.append(op1)
#             optionlist.append(op2)
#             optionlist.append(op3)
#             optionlist.append(op4)
#             corrAns = 'Y'
#             for optionDesc in optionlist:
#                 if optionDesc==corrans:
#                     query = "insert into question_options(option_desc,question_id,weightage,is_correct,option) values('"+optionDesc+"','"+question_id+"','"+weightage+"','Y','A')"
#                 else:
#                     query = "insert into question_options(option_desc,question_id,weightage,option) values('"+optionDesc+"','"+question_id+"','"+weightage+"','A')"
#                     db.session.execute(query)

#     print('Inside Update Questions')
#     db.session.commit()
#     print(updateQuery)
#     # updateSecondQuery = "update question_options set weightage='" + str(weightage) +"' where question_id='" + str(question_id) + "'"
#     # querySecondExe = db.session.execute(text(updateSecondQuery)) 
#     # db.session.commit()
#     print("Question Id in update Question:"+question_id)
#     # print(updatedData)
#     return render_template('questionUpload.html', form=form, flag=flag)


# @app.route('/questionOptions')
# def questionOptions():
#     question_id_arg=request.args.get('question_id')
#     questionOptionResults = QuestionOptions.query.filter_by(question_id=question_id_arg).all()
#     questionOptionsList=[]
#     for value in questionOptionResults:
#         print("This is the value: "+str(value))        
#         questionOptionsList.append(value.option+". "+value.option_desc)

#     print(questionOptionsList)

#     return jsonify([questionOptionsList])


# @app.route('/deleteQuestion')
# def deleteQuestion():
#     question_id = request.args.get('question_id')
#     print('Delete Question Id:'+question_id)
#     print("Question Id:-"+question_id)

#     updateQuery = "update question_details set archive_status='Y' where question_id='"+question_id+"'"
#     print(updateQuery)
#     db.session.execute(updateQuery)
#     db.session.commit()
#     return "text" 




# @app.route('/questionDetails')
# def questionDetails():
#     flag = True
#     question_id = request.args.get('question_id')
#     print("Question Id-:"+question_id)
#     form = QuestionBuilderQueryForm()
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
#     form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
#     form.topics.choices=[(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]

#     questionDetailsQuery = "select t2.class_val, t1.question_id, t2.subject_id, t1.reference_link, t1.suggested_weightage, t2.topic_name, t2.topic_id, t1.question_type, t1.question_description, t4.description from question_details t1 "
#     questionDetailsQuery = questionDetailsQuery + "inner join topic_detail t2 on t1.topic_id=t2.topic_id "
#     questionDetailsQuery = questionDetailsQuery + "inner join message_detail t4 on t1.subject_id = t4.msg_id"
#     questionDetailsQuery = questionDetailsQuery + " where t1.question_id ='" + question_id + "' order by t1.question_id"

#     questionUpdateUploadSubjective = db.session.execute(text(questionDetailsQuery)).first()   
#     question_desc = questionUpdateUploadSubjective.question_description.replace('\n', ' ').replace('\r', '')
#     print(questionUpdateUploadSubjective)
#     questionUpdateUpload=questionUpdateUploadSubjective
#     if questionUpdateUpload.question_type=='MCQ1':
       
#         query = "select option_desc from question_options where question_id='" + question_id + "' order by option_id"
#         #print(query)
#         avail_options = db.session.execute(text(query)).fetchall()
#         queryCorrectoption = "select option_desc from question_options where is_correct='Y' and question_id='" + question_id + "'"  
#         #print(queryCorrectoption)
#         correctoption = db.session.execute(text(queryCorrectoption)).fetchall()
#         print(correctoption)
#         correctOption = ''
#         for c in correctoption:
#             print(c.option_desc)
#             correctOption = c.option_desc
#         print('Correct Option:'+correctOption)
#         for q in questionUpdateUploadSubjective:
#             print('this is check for MCQ ' + str(q))
#         for a in avail_options:
#             print(a)
#         print('Correct Option Again:'+correctOption)
#         return render_template('questionUpload.html', question_id=question_id, questionUpdateUpload=questionUpdateUpload, form=form, flag=flag, avail_options=avail_options, correctOption=correctOption,question_desc=question_desc)
#         # return render_template('questionUpload.html',question_id=question_id, questionUpdateUploadSubjective=questionUpdateUploadSubjective,form=form,flag=flag,avail_options=avail_options,correctOption=correctOption)

#     for q in questionUpdateUpload:
#         print('this is check for Subjective ' + str(q))
    
#     return render_template('questionUpload.html', question_id=question_id, questionUpdateUpload=questionUpdateUpload, form=form, flag=flag,question_desc=question_desc)


@app.route('/updateStudentProfile')
def updateStudentProfile():
    first_name = request.args.get('first_name')
    last_name = request.args.get('last_name')
    gender = request.args.get('gender')
    birthdate = request.args.get('birthdate')
    phone = request.args.get('phone')
    address2 = request.args.get('address2')
    address1 = request.args.get('address1')
    locality = request.args.get('locality')
    city = request.args.get('city')
    state = request.args.get('state')
    country = request.args.get('country')
    pincode = request.args.get('pincode')
    class_val = request.args.get('class_val')
    section = request.args.get('section')
    school_admn_no = request.args.get('school_admn_no')
    roll_number = request.args.get('roll_number')
    preview = request.args.get('preview')
    student_id = request.args.get('student_id')
    url = request.args.get('preview')

    # query = "update student_profile set g"
    print('StudentId:'+str(student_id)+"fName:"+str(first_name)+"lName:"+str(last_name)+"gender:"+str(gender)+"bdate:"+str(birthdate)+"phone:"+str(phone)+"add1:"+str(address1)+"add2:"+str(address2)+"loc:"+str(locality)+"city:"+str(city)+"state:"+str(state)+"country:"+str(country)+"pincode:"+str(pincode)+"class:"+str(class_val)+"sect:"+str(section)+"Sch_num:"+str(school_admn_no)+"roll:"+str(roll_number)+"url:"+str(url))
    form=SingleStudentRegistration()
    return render_template('studentRegistration.html',form=form)
# @app.route('/topperListAll')
# def topperListAll():
#     user = User.query.filter_by(username=current_user.username).first_or_404()
#     teacher= TeacherProfile.query.filter_by(user_id=user.id).first() 
#     query = "select  * from fn_performance_leaderboard_detail('"+ str(teacher.school_id)+"') order by marks desc, student_name "
#     #print('Query:'+query)
#     leaderBoardData = db.session.execute(text(query)).fetchall()
#     return render_template('_leaderBoardTable.html',leaderBoardData=leaderBoardData)
def rename(dataframe):
    subject = MessageDetails.query.with_entities(MessageDetails.msg_id,MessageDetails.description).distinct().filter_by(category='Subject').all()
    # df = df.drop(['studentid'],axis=1)
    df = dataframe.columns.values.tolist()
    i=0
    for col in df:
        if i!=0:
            
            c = col.split('_')
            for sub in subject:
                if c[0]==str(sub.msg_id) and c[1]=='x':
                    dataframe.rename(columns = {col:sub.description}, inplace = True)
                elif c[1]=='y':
                    dataframe.rename(columns = {col:'Total Test'}, inplace = True)
                else:
                    print(col)
        i = i +1
    return dataframe
@app.route('/leaderBoard')
@login_required
def leaderBoard():
    form = LeaderBoardQueryForm()
    qclass_val = request.args.get("class_val")
    print('Class:'+str(qclass_val))
    #print('class:'+str(qclass_val))
    if current_user.is_authenticated:        
        user = User.query.filter_by(username=current_user.username).first_or_404()
        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
        distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher_id.school_id)+" GROUP BY class_val order by s")).fetchall()    
        form.subject_name.choices = [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
        class_sec_id=ClassSection.query.filter_by(class_val='1',school_id=teacher_id.school_id).first()
        form.test_type.choices= [(i.description,i.description) for i in MessageDetails.query.filter_by(category='Test type').all()]

        # form.testdate.choices = [(i.exam_date,i.exam_date) for i in ResultUpload.query.filter_by(class_sec_id=class_sec_id.class_sec_id).all()]
        available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher_id.school_id).all()  
        form.section.choices= [(i.section,i.section) for i in available_section]
        

        leaderBoardData = leaderboardContent(qclass_val)
        #print('Type')
        #print(type(leaderBoardData))
        column_names = ["a", "b", "c"]
        datafr = pd.DataFrame(columns = column_names)
        if type(leaderBoardData)!=type(datafr):
            classSecCheckVal=''
            colAll = ''
            columnNames = ''
            qclass_val = ''
            subj = ''
            subColumn = ''
            subHeader = ''
            data = ''
            classSecCheckVal = ''
            indic='leaderBoard'
            return render_template('leaderBoard.html',indic=indic,title='Leaderboard',classSecCheckVal=classSecCheckVal,form=form,distinctClasses=distinctClasses,leaderBoardData=leaderBoardData,colAll=colAll,columnNames=columnNames, qclass_val=qclass_val,subject=subj,subColumn=subColumn,subHeader=subHeader,user_type_val=str(current_user.user_type))

        else:
            df1 = leaderBoardData[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
            df2 = leaderBoardData.drop(['profile_pic', 'student_name','class_val','section','total_marks%','total_tests'], axis=1)
            leaderBoard = pd.merge(df1,df2,on=('studentid'))
            
            d = leaderBoard[['studentid','profile_pic','student_name','class_val','section','total_marks%','total_tests']]
            df3 = leaderBoard.drop(['studentid'],axis=1)
            #print('DF3:')
            #print(df3)
            #print('print new dataframe')
            
            df1.rename(columns = {'profile_pic':'Profile Picture'}, inplace = True)
            df1.rename(columns = {'student_name':'Student'}, inplace = True)
            df1.rename(columns = {'class_val':'Class'}, inplace = True)
            df1.rename(columns = {'section':'Section'}, inplace = True)
            df1.rename(columns = {'total_marks%':'Total Marks'}, inplace = True)
            df1.rename(columns = {'total_tests':'Total Tests'}, inplace = True)
            #print(df1)
            #print('Excluding columns')
            #print(df2)
            ## rename(df2)
            #print('LeaderBoard Data:')
            #print(leaderBoardData)
            data = []
            

            header = [df1.columns.values.tolist()]
            headerAll = [df3.columns.values.tolist()]
            colAll = ''
            subjHeader = [df2.columns.values.tolist()]
            columnNames = ''
            col = ''
            subColumn = ''
            #print('Size of dataframe:'+str(len(subjHeader)))
            for subhead in subjHeader:
                subColumn = subhead
                #print('Header with Subject Name')
                #print(subhead)
            for h in header:
                columnNames = h
            for headAll in headerAll: 
                colAll = headAll
            #print(' all header Length:'+str(len(colAll))+'Static length:'+str(len(columnNames))+'sub header length:'+str(len(subColumn)))
            n= int(len(subColumn)/2)
            ndf = df2.drop(['studentid'],axis=1)
            newDF = ndf.iloc[:,0:n]
            new1DF = ndf.iloc[:,n:]
            
            df5 = pd.concat([newDF, new1DF], axis=1)
            DFW = df5[list(sum(zip(newDF.columns, new1DF.columns), ()))]
            #print('New DF')
            #print(DFW)
            dat = pd.concat([d,DFW], axis=1)
            #print(dat)
            subHeader = ''
            for row in dat.values.tolist():
                data.append(row)
            subH = [DFW.columns.values.tolist()]

            for s in subH:
                subHeader = s
            subHead = [dat.columns.values.tolist()]
            for column in subHead:
                col = column
            subject = MessageDetails.query.with_entities(MessageDetails.msg_id,MessageDetails.description).distinct().filter_by(category='Subject').order_by(MessageDetails.msg_id).all()
            #print(subject)
            subj = []

            #for d in data:
            #    print('In data Student id')
            #    print(data[0])
            
            for sub in subject:
                li = []
                i=0
                for col in subColumn:
                    if i!=len(subColumn)/2:
                        c = col.split('_')
                        #print(c[0])
                        #print(sub.msg_id)
                        if(c[0]==str(sub.msg_id)):
                            #print(c[0])
                            #print(sub.msg_id)                                                        
                            li.append(sub.msg_id)
                            li.append(sub.description)
                            subj.append(li)
                            break
                    i=i+1
                    
            #print('List with Subjects')
            #print(subj)
            #for s in subj:
            #    print(s[0])
            #    print(s[1])
            #print('Inside subjects')
            classSecCheckVal=classSecCheck()
            indic='leaderBoard'
            return render_template('leaderBoard.html',indic=indic,title='Leaderboard',classSecCheckVal=classSecCheckVal,form=form,distinctClasses=distinctClasses,leaderBoardData=data,colAll=colAll,columnNames=columnNames, qclass_val=qclass_val,subject=subj,subColumn=subColumn,subHeader=subHeader,user_type_val=str(current_user.user_type))
    # classSecCheckVal=''
    indic='leaderBoard'
    return render_template('leaderBoard.html',indic=indic,title='Leaderboard',classSecCheckVal=classSecCheckVal,form=form,distinctClasses=distinctClasses,leaderBoardData=data,colAll=colAll,columnNames=columnNames, qclass_val=qclass_val,subject=subj,subColumn=subColumn,subHeader=subHeader,user_type_val=str(current_user.user_type))

@app.route('/classDelivery',methods=['GET','POST'])
@login_required
def classDelivery():
    form = ContentManager()
    public = request.args.get('public')
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
    form.subject_name.choices = ''
    
    form.chapter_num.choices = ''
    
    form.topics.choices = ''
    
    form.content_type.choices = ''
    #if current_user.is_authenticated:        
    user = User.query.filter_by(username=current_user.username).first_or_404()        
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    
    qtopic_id=request.args.get('topic_id')
    qsubject_id=request.args.get('subject_id')
    qclass_sec_id = request.args.get('class_sec_id')
    retake = request.args.get('retake')
  
    contentData = ContentDetail.query.filter_by(topic_id=int(qtopic_id),archive_status='N').all()
    subject_name = MessageDetails.query.filter_by(msg_id=qsubject_id).all()
    subName = ''
    for sub in subject_name:
        subName = sub.description
        break
    q=0
    for content in contentData:            
        q=q+1
    classSections=ClassSection.query.filter_by(school_id=teacher.school_id).order_by(ClassSection.class_val).all()
    
    currClassSecDet = ClassSection.query.filter_by(class_sec_id=qclass_sec_id).first()
    distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher_id.school_id)+" GROUP BY class_val order by s")).fetchall()        
        # end of sidebar        
    #for curr in currClass:        
    #topicTrack = TopicTracker.query.filter_by(class_sec_id=currClass.class_sec_id, subject_id=qsubject_id).first()
    #print ("this is topic Track: " + topicTrack)
    topicDet = Topic.query.filter_by(topic_id=qtopic_id).order_by(Topic.chapter_num).first()
    bookDet= BookDetails.query.filter_by(book_id = topicDet.book_id).first()
    #if retake is true then set is_covered to No
    if retake == 'Y':
        topicFromTracker = TopicTracker.query.filter_by(school_id = teacher.school_id, topic_id=qtopic_id).first()
        topicFromTracker.is_covered='N'
        topicFromTracker.reteach_count=int(topicFromTracker.reteach_count)+1
        db.session.commit()
    
    topicTrackerQuery = "select t1.topic_id, t1.topic_name, t1.chapter_name, t1.chapter_num, " 
    topicTrackerQuery = topicTrackerQuery + " t1.unit_num, t1.book_id, t2.is_covered, t1.subject_id, t2.class_sec_id "
    topicTrackerQuery = topicTrackerQuery + " from "
    topicTrackerQuery = topicTrackerQuery + " topic_detail t1, "
    topicTrackerQuery = topicTrackerQuery + " topic_tracker t2"
    topicTrackerQuery = topicTrackerQuery + " where"
    topicTrackerQuery = topicTrackerQuery + " t1.topic_id=t2.topic_id"
    topicTrackerQuery = topicTrackerQuery + " and t2.class_sec_id = '" + str(qclass_sec_id) + "'"
    topicTrackerQuery = topicTrackerQuery + " and t1.subject_id= '" + str(qsubject_id ) + "' order by chapter_num"
    topicTrackerDetails= db.session.execute(text(topicTrackerQuery)).fetchall()
    # new changes for live classes
    if request.method=='POST':
        print('post method')
        # print(request.form('duration'))
        # print(request.form('conferenceLink'))
        qconf_link = request.args.get('conf_link')
        qduration = request.args.get('duration')
        print(qconf_link)
        print(qduration)
        if qduration!=None or qduration!='':
            print('if duration is not empty')
            print(int(qduration))
            end_time = datetime.now() + timedelta(hours=int(qduration))
            print(end_time)
        else:
            print('if duration is empty') 
            end_time = datetime.now() + timedelta(hours=1)
        print('Start time')
        
        format = "%Y-%m-%d %H:%M:%S"
    # Current time in UTC
        now_utc = datetime.now(timezone('UTC'))
        print(now_utc.strftime(format))
    # Convert to local time zone
        now_local = now_utc.astimezone(get_localzone())
        print(now_local.strftime(format))
        end_utc = end_time.now(timezone('UTC'))
        end_local = end_time.astimezone(get_localzone())
        print(end_local.strftime(format))
        print('end time')
        liveClassData = ''
        if public=='true':
            print('if data is public:'+str(public))
            # liveClassData=LiveClass(class_sec_id = qclass_sec_id,subject_id = qsubject_id, topic_id=qtopic_id, 
            #     start_time = now_local.strftime(format), end_time = end_local.strftime(format), status = "Active", teacher_id=teacher.teacher_id, 
            #     teacher_name = str(current_user.first_name)+' '+str(current_user.last_name), conf_link=str(qconf_link), school_id = teacher.school_id,
            #     is_archived = 'N',is_private='N',last_modified_date = now_local.strftime(format))  
            liveClassData = db.session.execute(text("insert into live_class(class_sec_id,subject_id,topic_id,start_time,end_time,status,teacher_id,teacher_name,conf_link,school_id,is_archived,is_private,last_modified_date) values('"+str(qclass_sec_id)+"','"+str(qsubject_id)+"','"+str(qtopic_id)+"','"+str(now_local.strftime(format))+"','"+str(end_local.strftime(format))+"','Active','"+str(teacher.teacher_id)+"','"+str(current_user.first_name)+' '+str(current_user.last_name)+"','"+str(qconf_link)+"','"+str(teacher.school_id)+"','N','N','"+str(now_local.strftime(format))+"')"))
        else:
            print('if data is not public:'+str(public))
            # liveClassData=LiveClass(class_sec_id = qclass_sec_id,subject_id = qsubject_id, topic_id=qtopic_id, 
            #     start_time = now_local.strftime(format), end_time = end_local.strftime(format), status = "Active", teacher_id=teacher.teacher_id, 
            #     teacher_name = str(current_user.first_name)+' '+str(current_user.last_name), conf_link=str(qconf_link), school_id = teacher.school_id,
            #     is_archived = 'N',is_private='Y',last_modified_date = now_local.strftime(format))    
            liveClassData = db.session.execute(text("insert into live_class(class_sec_id,subject_id,topic_id,start_time,end_time,status,teacher_id,teacher_name,conf_link,school_id,is_archived,is_private,last_modified_date) values('"+str(qclass_sec_id)+"','"+str(qsubject_id)+"','"+str(qtopic_id)+"','"+str(now_local.strftime(format))+"','"+str(end_local.strftime(format))+"','Active','"+str(teacher.teacher_id)+"','"+str(current_user.first_name)+' '+str(current_user.last_name)+"','"+str(qconf_link)+"','"+str(teacher.school_id)+"','N','Y','"+str(now_local.strftime(format))+"')"))
        # db.session.add(liveClassData)
        db.session.commit() 
        print('Before class Delivery')
        indic='DashBoard'
        return render_template('classDelivery.html',indic=indic,title='Class Delivery', classSecCheckVal=classSecCheck(),classsections=classSections, currClassSecDet= currClassSecDet, distinctClasses=distinctClasses,form=form ,topicDet=topicDet ,bookDet=bookDet,topicTrackerDetails=topicTrackerDetails,contentData=contentData,subName=subName,retake=retake,user_type_val=str(current_user.user_type))
    print('Before class Delivery')
    indic='DashBoard'
    return render_template('classDelivery.html',indic=indic,title='Class Delivery', classSecCheckVal=classSecCheck(),classsections=classSections, currClassSecDet= currClassSecDet, distinctClasses=distinctClasses,form=form ,topicDet=topicDet ,bookDet=bookDet,topicTrackerDetails=topicTrackerDetails,contentData=contentData,subName=subName,retake=retake,user_type_val=str(current_user.user_type))



@app.route('/contentManager',methods=['GET','POST'])
@login_required
def contentManager():
    topic_list=None
    user_type_val = current_user.user_type
    formContent = ContentManager()
    studentDetails = StudentProfile.query.filter_by(user_id=current_user.id).first()
    teacher_id = ''
    if user_type_val==134:
        teacher_id = StudentProfile.query.filter_by(user_id=current_user.id).first()
    else:
        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    formContent.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
    formContent.subject_name.choices = ''
    formContent.chapter_num.choices = ''
    formContent.topics.choices = ''    
    formContent.content_type.choices = ''    
    form=QuestionBankQueryForm() # resusing form used in question bank 
    available_class = "select distinct class_val from class_section where school_id='"+str(teacher_id.school_id)+"'"
    available_class = db.session.execute(text(available_class)).fetchall()
    
    if request.method=='POST':
        topic_list=Topic.query.filter_by(class_val=str(form.class_val.data),subject_id=int(form.subject_name.data),chapter_num=int(form.chapter_num.data)).all()
        subject=MessageDetails.query.filter_by(msg_id=int(form.subject_name.data)).first()
        session['class_val']=form.class_val.data
        session['sub_name']=subject.description
        session['test_type_val']=form.test_type.data
        session['chapter_num']=form.chapter_num.data    
        form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(str(form.class_val.data))]
        form.chapter_num.choices= [(int(i['chapter_num']), str(i['chapter_num'])+' - '+str(i['chapter_name'])) for i in chapters(str(form.class_val.data),int(form.subject_name.data))]
        if user_type_val==134:
            return render_template('contentManager.html',title='Content Manager',form=form,formContent=formContent,topics=topic_list,disconn=1,user_type_value=str(current_user.user_type),user_type_val=str(current_user.user_type),studentDetails=studentDetails)
        else:
            indic='selected'
            return render_template('contentManager.html',indic=indic,title='Content Manager',form=form,formContent=formContent,topics=topic_list,user_type_val=str(current_user.user_type),studentDetails=studentDetails)
    if user_type_val==134:
        classVal = ClassSection.query.filter_by(class_sec_id=teacher_id.class_sec_id).first()
        class_values = classVal.class_val
        available_subject = "select distinct subject_id,description as subject_name from board_class_subject bcs inner join message_detail md on bcs.subject_id=md.msg_id where class_val='"+str(class_values)+"' and school_id='"+str(teacher_id.school_id)+"'"
        print('Subject:'+str(available_subject))
        available_subject = db.session.execute(text(available_subject)).fetchall()
        flag='home'
        return render_template('contentManager.html',flag=flag,available_subject=available_subject,class_values=class_values,title='Content Manager',classSecCheckVal=classSecCheck(),form=form,formContent=formContent,disconn=1,user_type_value=str(current_user.user_type),user_type_val=str(current_user.user_type),studentDetails=studentDetails)
    else:
        indic='selected'
        return render_template('contentManager.html',indic=indic,title='Content Manager',classSecCheckVal=classSecCheck(),form=form,formContent=formContent,user_type_val=str(current_user.user_type),studentDetails=studentDetails,available_class=available_class)


@app.route('/loadContent',methods=['GET','POST'])
def loadContent():
    class_val = request.args.get('selected_class_value')
    selected_subject = request.args.get('selected_subject_value')
    selected_chapter = request.args.get('selected_chapter_value')
    selected_topic = request.args.get('selected_topic_value')
    contentName = request.args.get('contentName')
    contentTypeId = request.args.get('contentTypeId')
    contentUrl = request.args.get('contentUrl')
    contentUrl = contentUrl.replace("watch?v=", "embed/")
    print('contentUrl:'+str(contentUrl))
    reference = request.args.get('reference')
    print('Reference:'+str(reference))
    # url = "https://youtube.com/watch?v=TESTURLNOTTOBEUSED"
    reference = reference.replace("watch?v=", "embed/")
    public = request.args.get('public')
    teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    today = date.today()
    d4 = today.strftime("%b-%d-%Y")
    print(d4)
    print('public check value')
    print(public)
    if public=='true':
        if reference!='':
            contentData = ContentDetail(content_name=str(contentName),class_val=str(class_val),subject_id=int(selected_subject),
            topic_id=int(selected_topic),is_private='N',content_type=contentTypeId,school_id=teacher_id.school_id,reference_link=reference,archive_status='N',last_modified_date=d4,uploaded_by=teacher_id.teacher_id)
            db.session.add(contentData)
        else:
            contentData = ContentDetail(content_name=str(contentName),class_val=str(class_val),subject_id=int(selected_subject),
            topic_id=int(selected_topic),is_private='N',school_id=teacher_id.school_id,content_type=contentTypeId,reference_link=contentUrl,archive_status='N',last_modified_date=d4,uploaded_by=teacher_id.teacher_id)
            db.session.add(contentData)
    else:
        if reference!='':
            contentData = ContentDetail(content_name=str(contentName),class_val=str(class_val),subject_id=int(selected_subject),
            topic_id=int(selected_topic),is_private='Y',school_id=teacher_id.school_id,content_type=contentTypeId,reference_link=reference,archive_status='N',last_modified_date=d4,uploaded_by=teacher_id.teacher_id)
            db.session.add(contentData)
        else:
            contentData = ContentDetail(content_name=str(contentName),class_val=str(class_val),subject_id=int(selected_subject),
            topic_id=int(selected_topic),is_private='Y',school_id=teacher_id.school_id,content_type=contentTypeId,reference_link=contentUrl,archive_status='N',last_modified_date=d4,uploaded_by=teacher_id.teacher_id)
            db.session.add(contentData)
    db.session.commit()
    return "Upload"

@app.route('/getContentDetails',methods=['GET','POST'])
def getContentDetails():
    topic_id = request.args.get('topic_id')
    if current_user.user_type==134:
        print('if user is student')
        teacher = StudentProfile.query.filter_by(user_id=current_user.id).first()
        classVal = ClassSection.query.filter_by(class_sec_id=teacher.class_sec_id).first()
        content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
        content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
        content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
        content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where td.topic_id = '"+str(topic_id)+"' and cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(classVal.class_val)+"' and cd.school_id<>'"+str(teacher.school_id)+"' "
        content = content + "union "
        content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
        content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
        content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
        content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where td.topic_id = '"+str(topic_id)+"' and cd.archive_status = 'N' and cd.class_val='"+str(classVal.class_val)+"' and cd.school_id='"+str(teacher.school_id)+"' order by content_id"
        print(content)
        contentDetail = db.session.execute(text(content)).fetchall()
        # if len(contentDetail)==0:
        #     print("No data present in the content manager details")
        #     return jsonify(["NA"])
        # else:
        print(len(contentDetail))
        for c in contentDetail:
            print("Content List"+str(c.content_name))    
            
        return render_template('_topicContentDetails.html',contents=contentDetail)


@app.route('/contentDetails',methods=['GET','POST'])
def contentDetails():
    info = request.args.get('info')
    data = request.args.get('data')
    forDash = request.args.get('forDash')
    teacher = ''
    print(info)
    if current_user.user_type==134:
        print('if user is student')
        teacher = StudentProfile.query.filter_by(user_id=current_user.id).first()
        classVal = ClassSection.query.filter_by(class_sec_id=teacher.class_sec_id).first()
        content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
        content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
        content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
        content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(classVal.class_val)+"' and cd.school_id<>'"+str(teacher.school_id)+"' "
        content = content + "union "
        content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
        content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
        content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
        content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and cd.class_val='"+str(classVal.class_val)+"' and cd.school_id='"+str(teacher.school_id)+"' order by content_id desc limit 5"
        print('query:'+str(content))
        contentDetail = db.session.execute(text(content)).fetchall()
    
        if len(contentDetail)==0:
            print("No data present in the content manager details")
            return jsonify(["NA"])
        else:
            print(len(contentDetail))
            for c in contentDetail:
                print("Content List"+str(c.content_name))    
            
            return render_template('_contentDetails.html',forDash=forDash,contents=contentDetail,info=info,data=data)
    elif current_user.user_type==71:
        print('if user is teacher')
        teacher= TeacherProfile.query.filter_by(user_id=current_user.id).first()
    else:
        print('for other users')
        teacher = User.query.filter_by(id=current_user.id).first()
    content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
    content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
    content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
    content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.school_id<>'"+str(teacher.school_id)+"' "
    content = content + "union "
    content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
    content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
    content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
    content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and cd.school_id='"+str(teacher.school_id)+"' order by content_id desc limit 5"
    print('query:'+str(content))
    contentDetail = db.session.execute(text(content)).fetchall()
    
    if len(contentDetail)==0:
        print("No data present in the content manager details")
        return jsonify(["NA"])
    else:
        print(len(contentDetail))
        for c in contentDetail:
            print("Content List"+str(c.content_name))    
        
        return render_template('_contentDetails.html',forDash=forDash,contents=contentDetail,info=info,data=data)


@app.route('/filterContentfromTopic',methods=['GET','POST'])
def filterContentfromTopic():
    teacher = ''
    if current_user.user_type==71 or current_user.user_type==135 :
        teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    elif current_user.user_type==134:
        teacher = StudentProfile.query.filter_by(user_id=current_user.id).first()
    class_value = request.args.get('class_val')
    subject_id = request.args.get('subject_id')
    print(class_value)
    print(subject_id)
    topic_list = request.get_json()
    for topic in topic_list:
        print(topic)
        if topic:
            print('topic exist')
        else:
            print('topic not exist')
    # contentList = "select *from content_detail cd where is_private = 'N' "
    # contentList = contentList + " and archive_status = 'N' "
    content = ''
    contents = []
    l=1
    for topic in topic_list:
        if class_value != None:
            if subject_id!=None:
                content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' "
                content = content + "union "
                content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' ) order by content_id desc"
            else:
                content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(class_value)+"' and cd.topic_id='"+str(topic)+"' "
                content = content + "union "
                content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.class_val='"+str(class_value)+"' and cd.topic_id='"+str(topic)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.class_val='"+str(class_value)+"' and cd.topic_id='"+str(topic)+"' ) order by content_id desc"
        else:
            if subject_id!=None:
                content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.subject_id= '"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' "
                content = content + "union "
                content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.subject_id= '"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.subject_id= '"+str(subject_id)+"' and cd.topic_id='"+str(topic)+"' ) order by content_id desc"
            else:
                content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.topic_id='"+str(topic)+"' "
                content = content + "union "
                content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
                content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
                content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
                content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.topic_id='"+str(topic)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.topic_id='"+str(topic)+"') order by content_id desc limit 10"
        contentDetail = db.session.execute(text(content)).fetchall()
        if contentDetail:
            contents.append(contentDetail)

    # if contents:
    #     contentDetail = db.session.execute(text(content)).fetchall()
    # else:
    #     return jsonify(["NS"])
    print('length of list:'+str(len(contentDetail)))
    if len(contents)==0:
        print("No data present in the content manager details")
        return jsonify(["NA"])
    else:
        print(len(contents))
        # for c in contentDetail:
        #     print("Content List"+str(c.content_name)) 
        flag = 'true'
        flagTopic = 'true'
        forDash = ''
        return render_template('_contentDetails.html',forDash = forDash,contents=contents,flag=flag,flagTopic=flagTopic)

@app.route('/recentContentDetails',methods=['GET','POST'])
def recentContentDetails():
    teacher = ''
    if current_user.user_type==71 or current_user.user_type==135 :
        teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    elif current_user.user_type==134:
        teacher = StudentProfile.query.filter_by(user_id=current_user.id).first()
    class_value = request.args.get('class_val')
    subject_id = request.args.get('subject_id')
    print(class_value)
    print(subject_id)
    # contentList = "select *from content_detail cd where is_private = 'N' "
    # contentList = contentList + " and archive_status = 'N' "

    if class_value != None:
        if subject_id!=None:
            content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' "
            content = content + "union "
            content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.class_val='"+str(class_value)+"' and cd.subject_id='"+str(subject_id)+"' ) order by content_id desc"
        else:
            content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.class_val='"+str(class_value)+"' "
            content = content + "union "
            content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.class_val='"+str(class_value)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.class_val='"+str(class_value)+"' ) order by content_id desc"
    else:
        if subject_id!=None:
            content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' and cd.subject_id= '"+str(subject_id)+"' "
            content = content + "union "
            content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and cd.subject_id= '"+str(subject_id)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and cd.subject_id= '"+str(subject_id)+"' ) order by content_id desc"
        else:
            content = "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='N' "
            content = content + "union "
            content = content + "select cd.last_modified_date,cd.content_id, cd.content_type,cd.reference_link, cd.content_name,td.topic_name,md.description subject_name, cd.class_val,tp.teacher_name uploaded_by from content_detail cd "
            content = content + "inner join topic_detail td on cd.topic_id = td.topic_id "
            content = content + "inner join message_detail md on md.msg_id = cd.subject_id "
            content = content + "inner join teacher_profile tp on tp.teacher_id = cd.uploaded_by where cd.archive_status = 'N' and is_private='Y' and cd.school_id='"+str(teacher.school_id)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N') order by content_id desc limit 10"

    contentDetail = db.session.execute(text(content)).fetchall()
    
    if len(contentDetail)==0:
        print("No data present in the content manager details")
        return jsonify(["NA"])
    else:
        print(len(contentDetail))
        for c in contentDetail:
            print("Content List"+str(c.content_name)) 
        flag = 'true'
        forDash = ''
        return render_template('_contentDetails.html',forDash = forDash,contents=contentDetail,flag=flag)
    # if contentList:
    #     return render_template('_contentManagerDetails.html',contents=contentList)
    # else:
    #     return jsonify(["NA"])

@app.route('/deleteContentById',methods=['GET','POST'])
def deleteContentById():
    content_id=request.args.get('content_id')
    content = ContentDetail.query.filter_by(content_id=content_id).first()
    content.archive_status = 'Y'
    db.session.commit()
    return jsonify(['1'])
 
@app.route('/contentManagerDetails',methods=['GET','POST'])
def contentManagerDetails():
    print('inside contentManagerDetails')
    contents=[]
    topicList=request.get_json()
    user_type_val = current_user.user_type
    teacher_id = ''
    if user_type_val == 134:
        teacher_id = StudentProfile.query.filter_by(user_id=current_user.id).first()
    else:
        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    for topic in topicList:

        contentList = "select *from content_detail cd where is_private = 'N' and archive_status = 'N' and topic_id='"+str(topic)+"' union select *from content_detail cd2 where is_private = 'Y' and archive_status = 'N' and topic_id='"+str(topic)+"' and school_id = '"+str(teacher_id.school_id)+"' and content_id not in (select content_id from content_detail cd where is_private = 'N' and archive_status = 'N' and topic_id='"+str(topic)+"')"
        print(contentList)
        contentList = db.session.execute(text(contentList)).fetchall()
        
        if len(contentList)!=0:
            contents.append(contentList)
    if len(contents)==0:
        print("No data present in the content manager details")
        return jsonify(["NA"])
    else:
        print(len(contents))
        for c in contents:
            print("Content List"+str(c))    
        return render_template('_contentManagerDetails.html',contents=contents)

# API for existed Test Paper and Test Link Generation
@app.route('/existedTestPaperLinkGenerate',methods=['POST'])
def existedTestPaperLinkGenerate():
    teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    school_id=teacher_id.school_id
    print('SchoolId:',school_id)
    uploadStatus=request.args.get('uploadStatus')
    duration = request.args.get('duration')
    if duration =='':
        duration = 0
    print('Duration:'+str(duration))
    if uploadStatus=='' or uploadStatus==None:
        uploadStatus = 'Y'
    resultStatus = request.args.get('resultStatus')
    if resultStatus=='' or resultStatus==None:
        resultStatus = 'Y'
    instructions = request.args.get('instructions')

    advance = request.args.get('advance')
    if advance=='' or advance==None:
        advance = 'Y'
    weightage = request.args.get('weightage')
    if weightage=='' or weightage==None:
        weightage = 10
    NegMarking = request.args.get('negativeMarking')
    if NegMarking=='' or NegMarking==None:
        NegMarking = 0
    testId = request.args.get('test_id')
    testPaperLinkQuery = TestDetails.query.filter_by(test_id=testId).first()
    test_paper_link = testPaperLinkQuery.test_paper_link
    selectOption = request.args.get('selectOption')
    print('SelectOption:'+str(selectOption))
    currClassSecRow=ClassSection.query.filter_by(school_id=str(testPaperLinkQuery.school_id),class_val=str(testPaperLinkQuery.class_val).strip()).first()
    resp_session_id = str(testPaperLinkQuery.subject_id).strip()+ str(datetime.today().strftime("%d%m%Y%H%M%S")).strip() + str(currClassSecRow.class_sec_id).strip()
    if selectOption=='0':
        print('download test paper')
        return jsonify({'testPaperLink':test_paper_link})
    elif selectOption=='1':
        print('test link generate')
        linkForTeacher=url_for('testLinkWhatsappBoot',resp_session_id=resp_session_id,test_id=testId,weightage=weightage,negativeMarking=NegMarking,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance,instructions=instructions,duration=duration,class_val=testPaperLinkQuery.class_val,section=currClassSecRow.section,subject_id=testPaperLinkQuery.subject_id, _external=True)
        linkForStudent=url_for('feedbackCollectionStudDev',resp_session_id=resp_session_id,school_id=testPaperLinkQuery.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance, _external=True)
        return jsonify({'onlineTestLinkForTeacher':linkForTeacher,'onlineTestLinkForStudent':linkForStudent})
    else:
        print('test link generate and download paper')
        linkForTeacher=url_for('testLinkWhatsappBoot',resp_session_id=resp_session_id,test_id=testId,weightage=weightage,negativeMarking=NegMarking,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance,instructions=instructions,duration=duration,class_val=testPaperLinkQuery.class_val,section=currClassSecRow.section,subject_id=testPaperLinkQuery.subject_id, _external=True)
        linkForStudent=url_for('feedbackCollectionStudDev',resp_session_id=resp_session_id,school_id=testPaperLinkQuery.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance, _external=True)
        return jsonify({'testPaperLink':test_paper_link,'onlineTestLinkForTeacher':linkForTeacher,'onlineTestLinkForStudent':linkForStudent})

# Start API
@app.route('/testApp',methods=['POST'])
def testApp():
    if request.method == 'POST':
        jsonExamData = request.json
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        for data in z['results'].values():
            paramList.append(data)
        for con in z['contact'].values():
            conList.append(con)
        testIDQuery = TestDetails.query.filter_by(test_id=paramList[0]).first()
        subjectQuery = MessageDetails.query.filter_by(msg_id=testIDQuery.subject_id).first()
        userId = User.query.filter_by(phone=conList[0]).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        print('Test ID:'+str(paramList[0]))
        quesIdQuery = TestQuestions.query.filter_by(test_id=paramList[0]).all()
        document = Document()
        document.add_heading(schoolNameVal(), 0)
        document.add_heading('Class '+str(testIDQuery.class_val)+" - "+str(testIDQuery.test_type)+" - "+str(datetime.today().strftime("%d%m%Y%H%M%S")) , 1)
        document.add_heading("Subject : "+str(subjectQuery.description),2)
        document.add_heading("Total Marks : "+str(testIDQuery.total_marks),3)
        p = document.add_paragraph()
        for question in quesIdQuery:
            data=QuestionDetails.query.filter_by(question_id=int(question.question_id), archive_status='N').first()
            options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
            #add question desc
            document.add_paragraph(
                data.question_description, style='List Number'
            )    
            if data.reference_link!='' and data.reference_link!=None:
                try:
                    response = requests.get(data.reference_link, stream=True)
                    image = BytesIO(response.content)
                    document.add_picture(image, width=Inches(1.25))
                except:
                    pass
            for option in options:
                if option.option_desc is not None:
                    document.add_paragraph(
                        option.option+". "+option.option_desc) 
        cl = testIDQuery.class_val.replace("/","-")
        file_name=str(teacher_id.school_id)+str(cl)+str(subjectQuery.description)+str(testIDQuery.test_type)+str(datetime.today().strftime("%Y%m%d"))+str(testIDQuery.total_marks)+'.docx'
   
        if not os.path.exists('tempdocx'):
            os.mkdir('tempdocx')
        document.save('tempdocx/'+file_name.replace(" ", ""))
        #uploading to s3 bucket
        client = boto3.client('s3', region_name='ap-south-1')
        client.upload_file('tempdocx/'+file_name.replace(" ", "") , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name.replace(" ", "")),ExtraArgs={'ACL':'public-read'})
        #deleting file from temporary location after upload to s3
        os.remove('tempdocx/'+file_name.replace(" ", ""))
        file_name_val='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name.replace(" ", "")
        print('Test Created successfully:'+str(file_name_val))
    return jsonify({'fileName':file_name_val})




def insertData(class_sec_id,resp_session_id,question_ids,test_type,total_marks,class_val,teacher_id,school_id):
    with app.app_context():
        print('inside insertData')
        subjId = ''
        topicID = ''
        boardID = ''
        print('question_ids:')
        print(question_ids)
        for det in question_ids:
            subjId = det.subject_id
            topicID = det.topic_id
            boardID = det.board_id
            break
        format = "%Y-%m-%d %H:%M:%S"
        schoolQuery = SchoolProfile.query.filter_by(school_id=school_id).first()
        schoolName = schoolQuery.school_name
        now_utc = datetime.now(timezone('UTC'))
        now_local = now_utc.astimezone(get_localzone())
        print('Date of test creation:'+str(now_local.strftime(format)))
        subjectQuery = MessageDetails.query.filter_by(msg_id=subjId).first()
        document = Document()
        document.add_heading(schoolName, 0)
        document.add_heading('Class '+str(class_val)+" - "+str(test_type)+" - "+str(datetime.today().strftime("%d%m%Y%H%M%S")) , 1)
        document.add_heading("Subject : "+str(subjectQuery.description),2)
        document.add_heading("Total Marks : "+str(total_marks),3)
        p = document.add_paragraph()
        for question in question_ids:
            data=QuestionDetails.query.filter_by(question_id=int(question.question_id), archive_status='N').first()
            options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
            #add question desc
            document.add_paragraph(
                data.question_description, style='List Number'
            )    
            print(data.reference_link)
            if data.reference_link!='' or data.reference_link!=None:
                print('inside threadUse if ')
                print(data.reference_link)
                try:
                    response = requests.get(data.reference_link, stream=True)
                    image = BytesIO(response.content)
                    document.add_picture(image, width=Inches(1.25))
                except:
                    pass
            for option in options:
                if option.option_desc is not None:
                    document.add_paragraph(
                        option.option+". "+option.option_desc) 
        cl = class_val.replace("/","-")
        file_name=str(school_id)+str(cl)+str(subjectQuery.description)+str(test_type)+str(datetime.today().strftime("%Y%m%d"))+str(total_marks)+'.docx'
   
        if not os.path.exists('tempdocx'):
            os.mkdir('tempdocx')
        document.save('tempdocx/'+file_name.replace(" ", ""))
        #uploading to s3 bucket
        client = boto3.client('s3', region_name='ap-south-1')
        client.upload_file('tempdocx/'+file_name.replace(" ", "") , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name.replace(" ", "")),ExtraArgs={'ACL':'public-read'})
        #deleting file from temporary location after upload to s3
        os.remove('tempdocx/'+file_name.replace(" ", ""))
        file_name_val='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name.replace(" ", "")

        testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
            board_id=str(boardID), subject_id=int(subjId),class_val=str(class_val),date_of_creation=now_local.strftime(format),
            date_of_test=datetime.now(),test_paper_link=file_name_val, school_id=school_id, teacher_id=teacher_id)
        db.session.add(testDetailsUpd)
        db.session.commit()
        sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
            test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
        db.session.add(sessionDetailRowInsert)
        for questionVal in question_ids:
            testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
            db.session.add(testQuestionInsert)
        db.session.commit()
        print('after insertData')
        return 'sent Asynchronous data'

# def threadUse(class_sec_id,resp_session_id,question_ids,test_type,total_marks,class_val,teacher_id,school_id):
#     print('Inside threadUse')
#     Thread(target=insertData,args=(class_sec_id,resp_session_id,question_ids,test_type,total_marks,class_val,teacher_id,school_id)).start()
    

# API for New Test Paper Link and Test Link Generation

@app.route('/getLeaderBoardLink',methods=['GET','POST'])
@login_required
def getLeaderBoardLink():
    if request.method == 'POST':
        leaderBoardLink = url_for('leaderBoard', _external=True)
        return jsonify({'leaderboardLink':leaderBoardLink})

@app.route('/getCustomerSupportLink',methods=['GET','POST'])
@login_required
def getCustomerSupportLink():
    if request.method == 'POST':
        customerSupportLink = url_for('help',_external=True)
        return jsonify({'helpLink':customerSupportLink})

@app.route('/getStudentTopicList',methods=['GET','POST'])
def getStudentTopicList():
    if request.method == 'POST':
        print('inside getStudentTopicList')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","subject":"1","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        paramList = []
        conList = []
        for con in dataList['contact'].values():
            conList.append(con)
        for data in dataList['results'].values():
            paramList.append(data)
        contactNo = conList[2][-10:]
        print(contactNo)
        studentData = StudentProfile.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=studentData.user_id).first()
        schoolData = SchoolProfile.query.filter_by(school_id=studentData.school_id).first()
        selClass = paramList[0].strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(studentData.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
        
        print('Subject:')
        selSubject = selSubject.strip()
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            if c==1:
                chap = str('Here???s the full list of chapters:\n')+str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            else:
                chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        msg = 'no topics available'
        if chapterDetList:
            return jsonify({'chapterDetList':chapterDetList,'selClass':selClass,'selSubject':selSubject,'userId':studentData.user_id,'teacher_id':teacher_id.teacher_id,'subId':subId,'schoolId':teacher_id.school_id,'schoolName':schoolData.school_name})
        else:
            return jsonify({'chapterDetList':msg})

@app.route('/getTopicIdList',methods=['POST','GET'])
def getTopicIdList():
    if request.method == 'POST':
        print('inside getTopicList')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","subject":"1","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        conList = []
        selectedOptions = []
        for con in dataList['contact'].values():
            conList.append(con)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        schoolData = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
        # classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        # classList = [] 
        # j=1
        # for classlist in classesListData:
        #     classVal = str(j)+str(' - ')+str(classlist.class_val)
        #     classList.append(classVal)
        #     j=j+1
        for clas in dataList['results'].values():
            selectedOptions.append(clas)
        selClass = ''
        selSubject = ''
        # for className in classList:
        #     num = className.split('-')[0]
        #     print('num:'+str(num))
        #     print('class:'+str(selectedOptions[0]))
        #     if int(num) == int(selectedOptions[0]):
        #         print(className)
        #         selClass = className.split('-')[1]
        #         print('selClass:'+str(selClass))
        print('class')
        selClass = selectedOptions[0].strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedOptions[1]))
            if int(num) == int(selectedOptions[1]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
        
        print('Subject:')
        selSubject = selSubject.strip()
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.topic_name ,td.topic_id ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            if c==1:
                chap = str('Here???s the full list of chapters:\n')+str(c)+str('-')+str(chapterDet.topic_id)+str('-')+str(chapterDet.topic_name)+str("\n")
            else:
                chap = str(c)+str('-')+str(chapterDet.topic_id)+str('-')+str(chapterDet.topic_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        msg = 'no topics available'
        if chapterDetList:
            return jsonify({'chapterDetList':chapterDetList,'selClass':selClass,'selSubject':selSubject,'userId':userId.id,'teacher_id':teacher_id.teacher_id,'subId':subId,'schoolId':teacher_id.school_id,'schoolName':schoolData.school_name})
        else:
            return jsonify({'chapterDetList':msg})

                      

@app.route('/getTopicList',methods=['POST','GET'])
def getTopicList():
    if request.method == 'POST':
        print('inside getTopicList')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","subject":"1","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        conList = []
        selectedOptions = []
        for con in dataList['contact'].values():
            conList.append(con)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        schoolData = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
        # classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        # classList = [] 
        # j=1
        # for classlist in classesListData:
        #     classVal = str(j)+str(' - ')+str(classlist.class_val)
        #     classList.append(classVal)
        #     j=j+1
        for clas in dataList['results'].values():
            selectedOptions.append(clas)
        selClass = ''
        selSubject = ''
        # for className in classList:
        #     num = className.split('-')[0]
        #     print('num:'+str(num))
        #     print('class:'+str(selectedOptions[0]))
        #     if int(num) == int(selectedOptions[0]):
        #         print(className)
        #         selClass = className.split('-')[1]
        #         print('selClass:'+str(selClass))
        print('class')
        selClass = selectedOptions[0].strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedOptions[1]))
            if int(num) == int(selectedOptions[1]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
        
        print('Subject:')
        selSubject = selSubject.strip()
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.topic_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            if c==1:
                chap = str('Here???s the full list of chapters:\n')+str(c)+str('-')+str(chapterDet.topic_name)+str('-')+str(chapterDet.book_name)+str("\n")
            else:
                chap = str(c)+str('-')+str(chapterDet.topic_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        msg = 'no topics available'
        if chapterDetList:
            return jsonify({'chapterDetList':chapterDetList,'selClass':selClass,'selSubject':selSubject,'userId':userId.id,'teacher_id':teacher_id.teacher_id,'subId':subId,'schoolId':teacher_id.school_id,'schoolName':schoolData.school_name})
        else:
            return jsonify({'chapterDetList':msg})

@app.route('/getStudentRequiredData',methods=['GET','POST'])
def getStudentRequiredData():
    if request.method == 'POST':
        print('inside getStudentRequiredData')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        
        conList = []
        selectedOptions = []
        for con in dataList['contact'].values():
            conList.append(con)
        for clas in dataList['results'].values():
            selectedOptions.append(clas)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        studentData = StudentProfile.query.filter_by(phone=contactNo).first()
        schoolData = SchoolProfile.query.filter_by(school_id=studentData.school_id).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=studentData.user_id).first()
        classData = ClassSection.query.filter_by(class_sec_id=studentData.class_sec_id).first()
        selClass = classData.class_val
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(studentData.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        selSubject = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedOptions[1]))
            if int(num) == int(selectedOptions[1]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
                
        print('Subject:')
        selSubject = selSubject.strip()
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))

        return jsonify({'selClass':selClass,'selSubject':selSubject,'userId':studentData.user_id,'teacher_id':teacher_id.teacher_id,'subId':subId,'schoolId':teacher_id.school_id,'schoolName':schoolData.school_name})        


@app.route('/getRequiredData',methods=['GET','POST'])
def getRequiredData():
    if request.method == 'POST':
        print('inside getRequiredData')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        
        conList = []
        selectedOptions = []
        for con in dataList['contact'].values():
            conList.append(con)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        schoolData = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
        classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)
            classList.append(classVal)
            j=j+1
        for clas in dataList['results'].values():
            selectedOptions.append(clas)
        selClass = ''
        selSubject = ''
        for className in classList:
            num = className.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedOptions[0]))
            if int(num) == int(selectedOptions[0]):
                print(className)
                selClass = className.split('-')[1]
                print('selClass:'+str(selClass))
        print('class')
        selClass = selClass.strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedOptions[1]))
            if int(num) == int(selectedOptions[1]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
                
        print('Subject:')
        selSubject = selSubject.strip()
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))

        return jsonify({'selClass':selClass,'selSubject':selSubject,'userId':userId.id,'teacher_id':teacher_id.teacher_id,'subId':subId,'schoolId':teacher_id.school_id,'schoolName':schoolData.school_name})


@app.route('/getSubjectsList',methods=['POST','GET'])
def getSubjectsList():
    if request.method == 'POST':
        print('inside getSubjectsList')
        jsonData = request.json
        # jsonData = {"contact": {"phone":"9008262739" },"results": {"class_val":"4","custom_key": "custom_value"}}
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        
        conList = []
        selectedClassOption = []
        for con in dataList['contact'].values():
            conList.append(con)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)
            classList.append(classVal)
            j=j+1
        for clas in dataList['results'].values():
            selectedClassOption.append(clas)
        selClass = ''
        print('Selected Class option:')
        print(selectedClassOption[0])
        for className in classList:
            num = className.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(selectedClassOption[0]))
            if int(num) == int(selectedClassOption[0]):
                print(className)
                selClass = className.split('-')[1]
                print('selClass:'+str(selClass))
        print('class')
        selClass = selClass.strip()
        print(selClass)
        
        subQuery = "select md.description as subject from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        for subj in subjectData:
            if k==1:
                sub = str('Which Subject?\n')+str(k)+str('-')+str(subj.subject)+str("\n")
            else:
                sub = str(k)+str('-')+str(subj.subject)+str("\n")
            subjectList.append(sub)
            k=k+1
        msg = 'no subjects available'
        if subjectList:
            return jsonify({'subject_list':subjectList,'class_val':selClass}) 
        else:
            return jsonify({'subject_list':msg})

@app.route('/getStudentDetails',methods=['POST','GET'])
def getStudentDetails():
    if request.method == 'POST':
        print('inside getStudentDetails')
        jsonData = request.json
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        selectedStudentOption = []
        for data in dataList['results'].values():
            selectedStudentOption.append(data)
        conList = []
        print('SelectedOption:'+str(selectedStudentOption))
        for con in dataList['contact'].values():
            conList.append(con)
        print(conList[2])
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)+str('-')+str(classlist.section)
            classList.append(classVal)
            j=j+1
        print(classList)
        selClass = ''
        selSection = ''
        for clas in classList:
            option = clas.split('-')[0]
            if int(option) == selectedStudentOption[1]:
                selClass = clas.split('-')[1]
                selSection = clas.split('-')[2]
        selClass = selClass.strip()
        selSection = selSection.strip()
        print('Class:'+str(selClass))
        print('Section:'+str(selSection))
        classSec = ClassSection.query.filter_by(class_val=selClass,section=selSection,school_id=teacher_id.school_id).first()
        classSecId = classSec.class_sec_id
        studentListQuery = StudentProfile.query.filter_by(school_id=teacher_id.school_id,class_sec_id=classSecId).all()
        l=1
        studentList = []
        for student in studentListQuery:
            stud = str(l)+str('-')+str(student.full_name)+str("-")+str(student.student_id)
            studentList.append(stud)
            l=l+1
        selStudentId = ''
        for stud in studentList:
            option = stud.split('-')[0]
            print(option)
            print(selectedStudentOption[0])
            if int(option) == int(selectedStudentOption[0]):
                print(stud)
                selStudentId = stud.split('-')[2]
        print('getStudentDetails student_id:'+str(selStudentId))
        studentDetailLink = url_for('student_profile.studentProfile',student_id=selStudentId, _external=True)
        newLink = ''
        if studentDetailLink:
            newLink = str('Student Detail Link:\n')+str(studentDetailLink)
        if newLink:
            return jsonify({'studentDetailLink':newLink})
        else:
            msg = 'No students available'
            return jsonify({'studentDetailLink':msg})

@app.route('/getStudentsList',methods=['GET','POST'])
def getStudentsList():
    if request.method == 'POST':
        print('inside getStudentsList')
        jsonData = request.json
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        conList = []
        for con in dataList['contact'].values():
            conList.append(con)
        print(conList[2])
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)+str('-')+str(classlist.section)
            classList.append(classVal)
            j=j+1
        print(classList)
        selectedClassOption = ''
        selClass = ''
        selSection = ''
        for data in dataList['results'].values():
            selectedClassOption = data
        for clas in classList:
            option = clas.split('-')[0]
            if int(option) == selectedClassOption:
                selClass = clas.split('-')[1]
                selSection = clas.split('-')[2]
        selClass = selClass.strip()
        selSection = selSection.strip()
        print('Class:'+str(selClass))
        print('Section:'+str(selSection))
        classSec = ClassSection.query.filter_by(class_val=selClass,section=selSection,school_id=teacher_id.school_id).first()
        classSecId = classSec.class_sec_id
        studentListQuery = StudentProfile.query.filter_by(school_id=teacher_id.school_id,class_sec_id=classSecId).all()
        l=1
        studentList = []
        for student in studentListQuery:
            if l==1:
                stud = str("Here's your students list:\n")+str(l)+str('-')+str(student.full_name)+str("\n")
            else:
                stud = str(l)+str('-')+str(student.full_name)+str("\n")
            studentList.append(stud)
            l=l+1
        msg = 'no students available'
        if studentList:
            return jsonify({'studentNewList':studentList})
        else:
            return jsonify({'studentNewList':msg})
    

@app.route('/getClassSectionList',methods=['POST','GET'])
def getClassSectionList():
    if request.method == 'POST':
        print('inside getClassSectionList')
        jsonData = request.json
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        conList = []
        for con in dataList['contact'].values():
            conList.append(con)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = ''
        if userId:
            teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        else:
            Msg = 'you are not a registered user'
            return jsonify({'class_list':Msg})
        classesListData = ''
        if teacher_id:
            classesListData = ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
        else:
            Msg = 'you are not a registered teacher'
            return jsonify({'class_list':Msg})
        
        classList = [] 
        j=1
        for classlist in classesListData:
            if j==1:
                classVal = str('Which class?\n')+str(j)+str(' - ')+str(classlist.class_val)+str('-')+str(classlist.section)+str("\n")
            else:
                classVal = str(j)+str(' - ')+str(classlist.class_val)+str('-')+str(classlist.section)+str("\n")
            classList.append(classVal)
            j=j+1
        print(classList)
        # return jsonify({'class_list':classList}) 
        msg = 'no classes available'
        if classList:
            return jsonify({'class_list':classList})
        else:
            return jsonify({'class_list':msg})

@app.route('/getClassList',methods=['POST','GET'])
def getClassList():
    if request.method == 'POST':
        print('inside getClassList')
        jsonData = request.json
        data = json.dumps(jsonData)
        dataList = json.loads(data)
        print('all data:')
        print(dataList)
        conList = []
        paramList = []
        for con in dataList['contact'].values():
            conList.append(con)
        for data in dataList['results'].values():
            paramList.append(data)
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = ''
        if userId:
            print('you are registered user')
            teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        else:
            Msg = 'you are not a registered user'
            return jsonify({'class_list':Msg})
        classesListData = ''
        if teacher_id:
            classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        else:
            Msg = 'you are not a registered teacher'
            return jsonify({'class_list':Msg})
        
        classList = [] 
        j=1
        for classlist in classesListData:
            if j==1:
                if paramList[0] == '1':
                    classVal = str('Which class?\n')+str(j)+str(' - ')+str(classlist.class_val)+str("\n")
                else:
                    classVal = str('Which class do you want to test?\n')+str(j)+str(' - ')+str(classlist.class_val)+str("\n")
            else:
                classVal = str(j)+str(' - ')+str(classlist.class_val)+str("\n")
            classList.append(classVal)
            j=j+1
        print(classList)
        msg = 'no classes available'
        if classList:
            return jsonify({'class_list':classList})
        else:
            return jsonify({'class_list':msg})

# @app.route('/getReqTopicList',methods=['POST','GET'])
# def getReqTopicList():
#     if request.method == 'POST':
#         jsonData = request.json
#         a = json.dumps(jsonData)
#         data = json.loads(a)
#         for value in data['results'].values():
#             print(value)
#     dataValue = 10
#     return jsonify({'Data':dataValue})

@app.route('/getStudentPerformance',methods=['POST','GET'])
def getStudentPerformance():
    if request.method == 'POST':
        print('inside getStudentPerformance')
        jsonExamData = request.json
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        conList = []
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        studentDetails = StudentProfile.query.filter_by(user_id=userId.id).first()
        emailDet = studentDetails.email
        if emailDet:
            user = User.query.filter_by(email=studentDetails.email).first()
        if user:
            login_user(user,remember='Y')
        link = url_for('student_profile.studentProfile',student_id=studentDetails.student_id,_external=True)
        return jsonify({'studentProfile':link})

# @app.route('/getStudentDashboard',methods=['POST','GET'])
# def getStudentDashboard():
#     if request.method == 'POST':
#         print('inside getStudentDashboard')
#         jsonExamData = request.json
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         conList = []
#         for con in z['contact'].values():
#             conList.append(con)
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         userId = User.query.filter_by(phone=contactNo).first()
#         studentDetails = StudentProfile.query.filter_by(user_id=userId.id).first()
#         emailDet = studentDetails.email
#         if emailDet:
#             user = User.query.filter_by(email=studentDetails.email).first()
#         if user:
#             login_user(user,remember='Y')
#         link = url_for('studentDashboard',student_id=studentDetails.student_id,_external=True)
#         return jsonify({'studentDashboard':link})

@app.route('/getStudentDet',methods=['POST','GET'])
def getStudentDet():
    if request.method == 'POST':
        print('inside getStudentDetails')
        jsonExamData = request.json
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        conList = []
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        studentDetails = StudentProfile.query.filter_by(user_id=userId.id).first()
        if studentDetails:
            msg = 'What do you want to do today?\n1-See dashboard\n2-see performance\n3-Start practice test'
            return jsonify({'studentDetails':msg})
        else:
            msg = 'you are not a registered student'
            return jsonify({'studentDetails':msg})

# @app.route('/checkContact',methods=['POST','GET'])
# def checkContact():
#     if request.method == 'POST':
#         print('inside checkContact')
#         jsonExamData = request.json
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         conList = []
#         for con in z['results'].values():
#             conList.append(con)
#         contactNo = conList[0]
#         print(contactNo)
#         msg = ''
#         checkContact = User.query.filter_by(phone=contactNo).first()
#         print(checkContact)
#         if checkContact:
#             msg = 'Exist'
#             print('Exist')
#             return jsonify({'msg':msg})
#         else:
#             msg = 'New'
#             print('New')
#             return jsonify({'msg':msg})


# @app.route('/checkQuesNo',methods=['GET','POST'])
# def checkQuesNo():
#     if request.method == 'POST':
#         print('inside checkQuesNo')
#         jsonData = request.json
#         print('jsonData:')
#         print(jsonData)
        
#         userData = json.dumps(jsonData)
#         user = json.loads(userData)   
#         paramList = []
#         for con in user['results'].values():
#             paramList.append(con)     
#         quesCount = paramList[0]
#         if int(quesCount) > 15:
#             print('question count greater then 18')
#             return jsonify({'msg':'Greater'})
#         else:
#             print('question count less then 18')
#             return jsonify({'msg':'Less'})    

@app.route('/isSpecialCharacter',methods=['POST','GET'])
def isSpecialCharacter():
    if request.method == 'POST':
        print('inside isSpecialCharacter')
        jsonData = request.json
        # jsonData = {'contact': {'fields': {'age_group': {'inserted_at': '2021-01-25T06:36:45.002400Z', 'label': 'Age Group', 'type': 'string', 'value': '19 or above'}, 'name': {'inserted_at': '2021-01-25T06:35:49.876654Z', 'label': 'Name', 'type': 'string', 'value': 'hi'}}, 'name': 'Zaheen', 'phone': '918802362259'}, 'results': {}, 'custom_key': 'custom_value'}
        print('jsonData:')
        print(jsonData)
        
        userData = json.dumps(jsonData)
        user = json.loads(userData)   
        paramList = []
        for con in user['results'].values():
            paramList.append(con)  
        # name = re.sub('[^a-zA-Z.\d\s]', '', paramList[0])  
        # name = paramList[0].isalpha()
        regex = re.compile('[@_!#$%^&*()<>?/\|}{~:]')
        print('isSpecialChar:')
        # print(name)
        if(regex.search(paramList[0]) == None):
            print('if special character does not exist')
            return jsonify({'name':'Not'})
        else:
            print('if special character present')
            return jsonify({'name':'Char'})    


# @app.route('/registerUser',methods=['POST','GET'])
# def registerUser():
#     if request.method == 'POST':
#         print('inside register user')
#         jsonData = request.json
#         # jsonData = {'contact': {'fields': {'age_group': {'inserted_at': '2021-01-25T06:36:45.002400Z', 'label': 'Age Group', 'type': 'string', 'value': '19 or above'}, 'name': {'inserted_at': '2021-01-25T06:35:49.876654Z', 'label': 'Name', 'type': 'string', 'value': 'hi'}}, 'name': 'Zaheen', 'phone': '918802362259'}, 'results': {}, 'custom_key': 'custom_value'}
#         print('jsonData:')
#         print(jsonData)
        
#         userData = json.dumps(jsonData)
#         user = json.loads(userData)
#         conList = []
#         for con in user['contact'].values():
#             conList.append(con)
#         print(conList)
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         userId = User.query.filter_by(phone=contactNo).first()
#         teacher = ''
#         if userId:
#             teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
#             student_id = StudentProfile.query.filter_by(user_id=userId.id).first()
#             if userId.user_type == 71:
#                 teacher = 'Teacher'
#                 print('User is Teacher')
#                 if teacher_id.school_id:
#                     schoolDet = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#                 else:
#                     return jsonify({'user':'Unregistered'})
#                 if schoolDet.is_verified == 'N':
#                     return jsonify({'user':'Parent'})
#                 lastName = ''
#                 if userId.last_name:
#                     lastName = userId.last_name
#                 return jsonify({'user':teacher,'firstName':str(userId.first_name)+str(' ')+str(lastName)})
                
#             elif userId.user_type == 134:
#                 student = 'Student'
#                 print('user is student')
#                 if student_id.school_id:
#                     schoolDet = SchoolProfile.query.filter_by(school_id=student_id.school_id).first()
#                 else:
#                     return jsonify({'user':'Unregistered'})
#                 if schoolDet.is_verified == 'N':
#                     return jsonify({'user':'Parent'})
#                 lastName = ''
#                 if userId.last_name:
#                     lastName = userId.last_name
#                 return jsonify({'user':student,'firstName':str(userId.first_name)+str(' ')+str(lastName),'studentId':student_id.student_id})
#             else:
#                 parent = 'Parent'
#                 print('user is parent outside if')
#                 if userId.user_type==72:
#                     print('user is parent inside if')
#                     guardianDet = GuardianProfile.query.filter_by(user_id=userId.id).first()
#                     print('guardianDet:')
#                     print(guardianDet)
#                     studentDet = StudentProfile.query.filter_by(student_id=guardianDet.student_id).first()
#                     print('studentDet:')
#                     print(studentDet)
#                     lastName = ''
#                     if userId.last_name:
#                         lastName = userId.last_name
#                     return jsonify({'user':parent,'firstName':str(userId.first_name)+str(' ')+str(lastName),'studentName':studentDet.full_name,'studentId':studentDet.student_id})
#         print('not registered user')
#         return jsonify({'user':'null'})

@app.route('/checkSchoolAddress',methods=['POST','GET'])
def checkSchoolAddress():
    if request.method == 'POST':
        print('inside checkSchoolAddress')
        jsonExamData = request.json
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        for data in z['results'].values():
            paramList.append(data)
        msg = ''
        if paramList[0].upper()==paramList[1].upper() or paramList[0].upper()==paramList[2].upper() or paramList[1].upper()==paramList[2].upper():
            print('Same address')
            return jsonify({'msg':'Same'})
        else:
            print('different address')
            return jsonify({'msg':'Different'})

@app.route('/getUserDetails',methods=['POST','GET'])
def getUserDetails():
    if request.method == 'POST':
        print('inside getUserDetails')
        jsonExamData = request.json
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        conList = []
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        if teacher_id:
            msg = ' What do you want to do today?\n1 - Create or start Online Tests\n2 - Create Online Class Link\n3 - See student profile (Performance report included)\n4 - See Leaderboard\n5 - Customer Support'
            return jsonify({'userDetails':msg})
        else:
            msg = 'you are not a registered teacher'
            return jsonify({'userDetails':msg})

@app.route('/getStudentProfileById',methods=['GET','POST'])
def getStudentProfileById():
    if request.method == 'GET':
        return url_for('student_profile.studentProfile',student_id=743,_external=True)
    if request.method == 'POST':
        print('inside getStudentProfileById')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print('data:')
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        print(paramList[0])
        finalResult = "Here's the link to the student profile:\n"
        studProfLink = url_for('student_profile.studentProfile',student_id=paramList[0],_external=True)
        newRes = str(finalResult) + str(studProfLink)
                
        return jsonify({'studentData':newRes})  

# @app.route('/accessSchool',methods=['GET','POST'])
# def accessSchool():
#     if request.method == 'POST':
#         print('inside accessSchool')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []
#         conList = []
#         print(data)
#         for values in data['results'].values():
#             paramList.append(values)    
#         for con in data['contact'].values():
#             conList.append(con)
#         contactNo = conList[2][-10:]
#         print(contactNo)      
#         schoolDet = SchoolProfile.query.filter_by(school_id=paramList[0]).first()
#         if schoolDet:
#             userDet = User.query.filter_by(phone=contactNo).first()
#             userDet.school_id =  schoolDet.school_id
#             db.session.commit()
#             teacherDet = TeacherProfile.query.filter_by(phone=contactNo).first()
#             studentDet = StudentProfile.query.filter_by(phone=contactNo).first()
#             if teacherDet:
#                 teacherDet.school_id = schoolDet.school_id
#             if studentDet:
#                 studentDet.school_id = schoolDet.school_id
#             db.session.commit()
#             statement = 'Access request sent to the school admin.'
#             return jsonify({'statement':statement})   
#         else:
#             statement = "School id does not exist."
#             return jsonify({'statement':statement})              

@app.route('/accessRegisteredSchool',methods=['GET','POST'])
def accessRegisteredSchool():
    if request.method == 'POST':
        print('inside accessRegisteredSchool')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        schoolDet = SchoolProfile.query.filter_by(school_id=paramList[0]).first()
        checkUser = User.query.filter_by(email=paramList[2]).first()
        statement = ''
        if checkUser:
            statement = 'Mail id already exist.'
            return jsonify({'statement':statement})
        createUser = User(username=paramList[2],school_id=schoolDet.school_id,email=paramList[2],last_seen=datetime.now(),user_type=71,access_status=143,phone=contactNo,last_modified_date=datetime.now(),first_name=paramList[1])
        createUser.set_password(contactNo)
        db.session.add(createUser)
        db.session.commit()
        createTeacher = TeacherProfile(teacher_name=paramList[1],school_id=schoolDet.school_id,designation=148,registration_date=datetime.now(),email=paramList[2],last_modified_date=datetime.now(),user_id=createUser.id,phone=contactNo,device_preference=195)
        db.session.add(createTeacher)
        db.session.commit() 
        schoolDet.school_admin = createTeacher.teacher_id
        db.session.commit()
        statement = 'Access request sent to the school admin.'
        return jsonify({'statement':statement})  
                
# @app.route('/checkMailId',methods=['GET','POST'])
# def checkMailId():
#     if request.method == 'POST':
#         print('inside checkMailID')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []
#         for values in data['results'].values():
#             paramList.append(values) 
#         email = paramList[0]
#         checkUser = User.query.filter_by(email=email).first()
#         statement = ''
#         if checkUser:
#             statement = 'Exist'
#             return jsonify({'statement':statement})
#         else:
#             statement = 'No'
#         return jsonify({'statement':statement})  


@app.route('/insertUserTeacherDetails',methods=['GET','POST'])
def insertUserTeacherDetails():
    if request.method == 'POST':
        print('inside insertUserTeacherDetails')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        checkUser = User.query.filter_by(email=paramList[1]).first()
        statement = ''
        if checkUser:
            statement = 'Mail id already exist.'
            return jsonify({'statement':statement})
        createUser = User(username=paramList[1],email=paramList[1],last_seen=datetime.now(),user_type=71,access_status=143,phone=contactNo,last_modified_date=datetime.now(),first_name=paramList[0])
        createUser.set_password(contactNo)
        db.session.add(createUser)
        db.session.commit()
        createTeacher = TeacherProfile(teacher_name=paramList[0],designation=148,registration_date=datetime.now(),email=paramList[1],last_modified_date=datetime.now(),user_id=createUser.id,phone=contactNo,device_preference=195)
        db.session.add(createTeacher)
        db.session.commit()  
        statement = "What's your school's name?"
        return jsonify({'statement':statement})  

@app.route('/checkSchoolList',methods=['GET','POST'])
def checkSchoolList():
    if request.method == 'POST':
        print('inside schoolList')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        schoolNam = paramList[3].upper()
        schoolDetQuery = "select school_id,school_name from school_profile where INITCAP(school_name) like initcap('%"+str(schoolNam)+"%')"
        print(schoolDetQuery)
        schoolDet = db.session.execute(text(schoolDetQuery)).fetchall()
        data = ''
        i=1
        if len(schoolDet) != 0:
            data = 'Exist'
        else:
            data = 'None'
        # data = data + '\n If your school is not in this list, please type 00'
        print(data)
        return jsonify({'isSchoolList':data})         

@app.route('/schoolList',methods=['GET','POST'])
def schoolList():
    if request.method == 'POST':
        print('inside schoolList')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        schoolNam = paramList[3].upper()
        schoolDetQuery = "select school_id,school_name from school_profile where INITCAP(school_name) like initcap('%"+str(schoolNam)+"%')"
        print(schoolDetQuery)
        schoolDet = db.session.execute(text(schoolDetQuery)).fetchall()
        data = ''
        i=1
        if len(schoolDet) != 0:
            data = 'Please type the school id of your school from the list\n'
            for school in schoolDet:
                data = data + str(school.school_id)+str(' ')+str(school.school_name) + str(' ') + str(i) +str('\n')
                i = i +1
        else:
            data = 'None'
        # data = data + '\n If your school is not in this list, please type 00'
        return jsonify({'schoolNameList':data}) 

@app.route('/registerTeacher',methods=['GET','POST'])
def registerTeacher():
    if request.method == 'POST':
        print('inside registerTeacher')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])  
        userDet = User.query.filter_by(phone=contactNo).first()       
        checkUser = User.query.filter_by(email=paramList[1]).first()
        statement = ''
        if checkUser:
            statement = 'Mail id already exist.Would you like to go back to the main menu?'
            return jsonify({'teacherId':statement}) 
        createUser = User(username=paramList[1],school_id=userDet.school_id,email=paramList[1],last_seen=datetime.now(),user_type=71,access_status=145,phone=paramList[2],last_modified_date=datetime.now(),first_name=paramList[0])
        createUser.set_password(paramList[2])
        db.session.add(createUser)
        db.session.commit()
        createTeacher = TeacherProfile(teacher_name=paramList[0],school_id=userDet.school_id,designation=148,registration_date=datetime.now(),email=paramList[1],last_modified_date=datetime.now(),user_id=createUser.id,phone=paramList[2],device_preference=195)
        db.session.add(createTeacher)
        db.session.commit()
        statement = "Success! Teacher has been successfully registered on the platform. The Teacher Id is "+str(createTeacher.teacher_id)+" The account password is the teacher's phone number. Would you like to go back to the main menu?"
        return jsonify({'teacherId':statement}) 

# @app.route('/checkClass',methods=['GET','POST'])
# def checkClass():
#     if request.method == 'POST':
#         print('inside checkClass')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []  
#         for values in data['results'].values():
#             paramList.append(values) 
#         clas = paramList[0].split('-')[0]
#         section = paramList[0].split('-')[1]
#         print('Class'+str(clas))
#         print('Section'+str(section))
#         classSecId = ClassSection.query.filter_by(class_val=clas,section=section,school_id=userDet.school_id).first()
#         if classSecId == '' or classSecId == None:
#             statement = 'Class does not exist.'
#             return jsonify({'statement':statement})  
#         statement = "What's your school's name?"    
#         return jsonify({'statement':statement})

# @app.route('/studentSubjectList',methods=['GET','POST'])
# def studentSubjectList():
#     if request.method == 'POST':
#         print('inside studentSubjectList')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []
#         conList = []
#         print(data)
#         for values in data['results'].values():
#             paramList.append(values) 
#         for con in data['contact'].values():
#             conList.append(con)
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         studentDet = StudentProfile.query.filter_by(phone=contactNo).first()
#         clasDet = ClassSection.query.filter_by(class_sec_id=studentDet.class_sec_id).first()
#         subQuery = "select md.description as subject from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(studentDet.school_id)+"' and class_val = '"+str(clasDet.class_val)+"'"
#         print(subQuery)
#         subjectData = db.session.execute(text(subQuery)).fetchall()
#         print(subjectData)
#         subjectList = []
#         k=1
#         for subj in subjectData:
#             if k==1:
#                 sub = str('Which Subject?\n')+str(k)+str('-')+str(subj.subject)+str("\n")
#             else:
#                 sub = str(k)+str('-')+str(subj.subject)+str("\n")
#             subjectList.append(sub)
#             k=k+1      
#         msg = 'no subjects available'
#         if subjectList:
#             return jsonify({'subject_list':subjectList,'class_val':clasDet.class_val}) 
#         else:
#             return jsonify({'subject_list':msg})   

# @app.route('/classSectionCheck',methods=['GET','POST'])
# def classSectionCheck():
#     if request.method == 'POST':
#         print('inside classSectionCheck')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []
#         conList = []
#         for values in data['results'].values():
#             paramList.append(values)    
#         for con in data['contact'].values():
#             conList.append(con)
#         contactNo = conList[2][-10:]
#         print(paramList)
#         print(paramList[0])   
#         subString = '-'
#         if subString in paramList[0]:
#             print('Y')
#             return jsonify({'msg':'Y'})
#         else:
#             print('N')
#             return jsonify({'msg':'N'})


@app.route('/registerStudent',methods=['GET','POST'])
def registerStudent():
    if request.method == 'POST':
        print('inside registerNewStudent')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        userDet = User.query.filter_by(phone=contactNo).first()
        checkUser = User.query.filter_by(email=paramList[2]).first()
        statement = ''
        if checkUser:
            statement = 'Mail id already exist. Would you like to go back to the main menu?'
            return jsonify({'studentId':statement})
        createUser = User(username=paramList[2],school_id=userDet.school_id,email=paramList[2],last_seen=datetime.now(),user_type=134,access_status=145,phone=paramList[1],last_modified_date=datetime.now(),first_name=paramList[0])
        createUser.set_password(paramList[1])
        db.session.add(createUser)
        db.session.commit()
        clas = paramList[3].split('-')[0]
        section = paramList[3].split('-')[1].upper()
        print('Class'+str(clas))
        print('Section'+str(section))
        classSecId = ClassSection.query.filter_by(class_val=clas,section=section,school_id=userDet.school_id).first()
        if classSecId == '' or classSecId == None:
            statement = 'Class does not exist.'
            return jsonify({'studentId':statement}) 
        createStudent = StudentProfile(school_id=userDet.school_id,registration_date=datetime.now(),last_modified_date=datetime.now(),class_sec_id=classSecId.class_sec_id,first_name=paramList[0],full_name=paramList[0],email=paramList[2],phone=paramList[1],user_id=createUser.id,is_archived='N')
        db.session.add(createStudent)
        db.session.commit()   
        statement = "Congratulations! Student registered successfully. Your student ID is "+str(createStudent.student_id)+" Your password is your phone number.Would you like to go back to the main menu?"     
        return jsonify({'studentId':statement})

@app.route('/registerNewStudent',methods=['GET','POST'])
def registerNewStudent():
    if request.method == 'POST':
        print('inside registerNewStudent')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[1])
        strg = '-'
        print(paramList[2])
        if paramList[2].find(strg)==0:
            statement = 'invalid class format'
            return jsonify({'studentId':statement})
        clas = paramList[2].split('-')[0]
        section = paramList[2].split('-')[1].upper()
        print('Class'+str(clas))
        print('Section'+str(section))
        classSecId = ClassSection.query.filter_by(class_val=clas,section=section,school_id=paramList[5]).first()
        statement = ''
        print('classSecId:'+str(classSecId))
        if classSecId == '' or classSecId == None:
            statement = 'Class does not exist'
            return jsonify({'studentId':statement})
        createUser = User(username=paramList[1],school_id=paramList[5],email=paramList[1],last_seen=datetime.now(),user_type=134,access_status=143,phone=contactNo,last_modified_date=datetime.now(),first_name=paramList[0])
        createUser.set_password(contactNo)
        db.session.add(createUser)
        db.session.commit()
        createStudent = StudentProfile(school_id=paramList[5],registration_date=datetime.now(),last_modified_date=datetime.now(),class_sec_id=classSecId.class_sec_id,first_name=paramList[0],full_name=paramList[0],email=paramList[1],phone=contactNo,user_id=createUser.id,is_archived='N')
        db.session.add(createStudent)
        db.session.commit()
        statement = "Congratulations! You're registered. Your student ID is "+str(createStudent.student_id)+" Your password is your phone number."
        return jsonify({'studentId':statement})

# @app.route('/unregisterSchoolRegistered',methods=['GET','POST'])
# def unregisterSchoolRegistered():
#     if request.method == 'POST':
#         print('inside unregisterSchoolRegistered')
#         jsonStudentData = request.json
#         newData = json.dumps(jsonStudentData)
#         data = json.loads(newData)
#         paramList = []
#         conList = []
#         print(data)
#         for values in data['results'].values():
#             paramList.append(values)    
#         for con in data['contact'].values():
#             conList.append(con)
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         latitude = paramList[9]
#         longitude = paramList[10]
#         print('latitude:'+str(latitude))
#         print('longitude:'+str(longitude))
#         substring = '.latitude'
#         if substring in latitude:
#             createAddress = Address(address_1=paramList[3],city=paramList[4],state=paramList[5],country='india')
#             db.session.add(createAddress)
#             db.session.commit()
#         else:
#             createAddress = Address(latitude=latitude,longitude=longitude)
#             db.session.add(createAddress)
#             db.session.commit()
#         boardId = ''
#         schoolType = ''
#         if paramList[7] == '1':
#             boardId = 1001
#         elif paramList[7] == '2':
#             boardId = 1002
#         elif paramList[7] == '3':
#             boardId = 1005
#         else:
#             boardId = 1003
#         if paramList[6] == '1':
#             schoolType = 'Affordable private school'
#         elif paramList[6] == '2':
#             schoolType = 'NGO School'
#         elif paramList[6] == '3':
#             schoolType = 'Elite private school'
#         else: 
#             schoolType = 'Other'
#         createTeacher = TeacherProfile.query.filter_by(phone=contactNo).first()
#         createStudent = StudentProfile.query.filter_by(phone=contactNo).first()
#         createUser = User.query.filter_by(phone=contactNo).first()
#         createSchool = SchoolProfile(school_name=paramList[2],registered_date=datetime.now(),last_modified_date=datetime.now(),address_id=createAddress.address_id,board_id=boardId,school_admin=createTeacher.teacher_id,sub_id=2,is_verified='N',school_type=schoolType)
#         db.session.add(createSchool)
#         db.session.commit()
#         print(createSchool.school_id)
#         createUser.school_id = createSchool.school_id
#         db.session.commit()
#         if createTeacher:
#             print('user is teacher School id: '+str(createSchool.school_id))
#             print('school id in teacher table before update:'+str(createTeacher.school_id))
#             createTeacher.school_id = createSchool.school_id
#             db.session.commit()
#             print('school id in teacher table after update:'+str(createTeacher.school_id))
#         if createStudent:
#             print('user is student')
#             createStudent.school_id = createSchool.school_id
#             db.session.commit()
#         db.session.commit()
#         return jsonify({'success':'success'})        


@app.route('/registerSchool',methods=['GET','POST'])
def registerSchool():
    if request.method == 'POST':
        print('inside registerSchool')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print('data:')
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        for param in paramList:
            print(param)
        print(paramList[3])
        latitude = paramList[9]
        longitude = paramList[10]
        print('latitude:'+str(latitude))
        print('longitude:'+str(longitude))
        # geolocator = GoogleV3(api_key='AIzaSyDIUer3-m41C8aHiNlo0mld7aKndhuPqLM')
        # coordinates = str(lattitude)+','+str(longitude)
        # locations = geolocator.reverse(coordinates)
        # if locations:
            # print(locations[0].address)  # select first location
        substring = '.latitude'
        if substring in latitude:
            createAddress = Address(address_1=paramList[3],city=paramList[4],state=paramList[5],country='india')
            db.session.add(createAddress)
            db.session.commit()
        else:
            createAddress = Address(latitude=latitude,longitude=longitude)
            db.session.add(createAddress)
            db.session.commit()
        
        boardId = ''
        schoolType = ''
        if paramList[7] == '1':
            boardId = 1001
        elif paramList[7] == '2':
            boardId = 1002
        elif paramList[7] == '3':
            boardId = 1005
        else:
            boardId = 1003
        if paramList[6] == '1':
            schoolType = 'Affordable private school'
        elif paramList[6] == '2':
            schoolType = 'NGO School'
        elif paramList[6] == '3':
            schoolType = 'Elite private school'
        else: 
            schoolType = 'Other' 
        createTeacher = TeacherProfile.query.filter_by(teacher_name=paramList[0],email=paramList[1],phone=contactNo).first()
        createUser = User.query.filter_by(username=paramList[1],email=paramList[1],phone=contactNo).first()
        createSchool = SchoolProfile(school_name=paramList[2],registered_date=datetime.now(),last_modified_date=datetime.now(),address_id=createAddress.address_id,board_id=boardId,school_admin=createTeacher.teacher_id,sub_id=2,is_verified='N',school_type=schoolType)
        db.session.add(createSchool)
        db.session.commit()
        print(createSchool.school_id)
        createUser.school_id = createSchool.school_id
        db.session.commit()
        createTeacher.school_id = createSchool.school_id
        db.session.commit()
        return jsonify({'success':'success'})




@app.route('/checkStudent',methods=['GET','POST'])
def checkStudent():
    if request.method == 'POST':
        print('inside checkStudent')
        jsonStudentData = request.json
        newData = json.dumps(jsonStudentData)
        data = json.loads(newData)
        paramList = []
        conList = []
        print('data:')
        print(data)
        for values in data['results'].values():
            paramList.append(values)    
        for con in data['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        print(paramList[0])
        studentDataQuery = "select student_id,full_name from student_profile where initcap(full_name) like initcap('%"+str(paramList[0])+"%') and is_archived='N' and school_id='"+str(teacher_id.school_id)+"'"
        print('studentDataQuery:'+str(studentDataQuery))
        studentData = db.session.execute(text(studentDataQuery)).fetchall()
        newRes = ''
        print(studentData)
        print(type(studentData))
        if len(studentData) != 0:
            if len(studentData) == 1:
                print(len(studentData))
                for student in studentData:
                    finalResult = "Here's the link to the student profile:\n"
                    studProfLink = url_for('student_profile.studentProfile',student_id=student.student_id,_external=True)
                    newRes = str(finalResult) + str(studProfLink)
                    
                return jsonify({'studentData':newRes,'flag':'1'}) 
            else:
                print(len(studentData))
                newString = "Multiple students found with similar name.\n Please enter the student ID of the student from the below list:\n"
                i=1
                studData = ''
                for student in studentData:
                    studData = str(student.student_id)+str(' ')+str(student.full_name)+str('\n')+ studData
                    i=i+1
                newRes = newString + studData
                return jsonify({'studentData':newRes,'flag':'More'}) 
        else:
            newRes = 'No student available'
            return jsonify({'studentData':newRes,'flag':'Other'})  

@app.route('/checkrequiredquestions',methods=['POST','GET'])
def checkrequiredquestions():
    if request.method == 'POST':
        print('inside checkrequiredquestions')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
        a = json.dumps(jsonExamData)
      
        z = json.loads(a)
        
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for data in z['results'].values():
            
            paramList.append(data)
        topics = paramList[0].strip()
        topicList = topics.split(',')
        print(topicList[0])
        topic = topicList[0].capitalize()
        print('Topic:'+str(topic))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        p =1
        selClass = paramList[1]
        selSubject = paramList[2]
        for topic in topicList:
            fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"'"
            if p<len(topicList):
                fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
            p=p+1
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        Msg = 'Test Link'
        if len(fetchQuesIds)==0 or fetchQuesIds == '':
            Msg = 'No questions available'
            return jsonify({'msg':Msg})
        return jsonify({'msg':Msg})

@app.route('/getEnteredTopicOnlineClassLink',methods=['POST','GET'])
def getEnteredTopicOnlineClassLink():
    if request.method == 'POST':
        print('inside getStudentEnteredTopicList')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
        a = json.dumps(jsonExamData)
      
        z = json.loads(a)
        
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for data in z['results'].values():
            
            paramList.append(data)
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)
        print(conList[2])
        
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        topics = paramList[0].strip()
        topicList = topics.split(',')
        print(topicList[0])
        topic = topicList[0].capitalize()
        selClass = paramList[10]
        selSubject = paramList[11]
        classDet = ClassSection.query.filter_by(class_val=selClass,school_id=teacher_id.school_id).first()
        subId  = paramList[14]
        p =1
        for topic in topicList:
            fetchQuesIdsQuery = "select td.topic_id,td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"'"
            if p<len(topicList):
                fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
            p=p+1
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        Msg = 'no questions available'
        if len(fetchQuesIds)==0:
            return jsonify({'onlineClassLink':Msg})
        topicID = ''
        for det in fetchQuesIds:
            if det.topic_id:
                topicID = det.topic_id
                break
        if teacher_id.room_id==None:            
            roomResponse = roomCreation()
            roomResponseJson = roomResponse.json()
            print("New room ID created: " +str(roomResponseJson["url"]))
            teacher_id.room_id = str(roomResponseJson["url"])
            db.session.commit()
        link = url_for('classDelivery',class_sec_id=classDet.class_sec_id,subject_id=subId,topic_id=topicID,retake='N',_external=True)
        OnlineClassLink = str('Online class link:\n')+ str(teacher_id.room_id)+str("\n")
        OnlineClassLink = OnlineClassLink + str("Book Link:\n")+str(link)
        print('topicID:'+str(topicID))
        print(OnlineClassLink)
        return jsonify({'onlineClassLink':OnlineClassLink})
        


# @app.route('/getStudentEnteredTopicList',methods=['POST','GET'])
# def getStudentEnteredTopicList():
#     if request.method == 'POST':
#         print('inside getStudentEnteredTopicList')
#         jsonExamData = request.json
#         # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
#         a = json.dumps(jsonExamData)
      
#         z = json.loads(a)
        
#         paramList = []
#         conList = []
#         print('data:')
#         # print(z['result'].class_val)
#         # print(z['result'])
#         for data in z['results'].values():
            
#             paramList.append(data)
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)
#         print(conList[2])
        
#         print('Data Contact')
#         # print(conList[2])
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         studentData = StudentProfile.query.filter_by(phone=contactNo).first()
#         teacher_id = TeacherProfile.query.filter_by(user_id=studentData.user_id).first()
#         classesListData = ClassSection.query.filter_by(class_sec_id=studentData.class_sec_id).first()
#         print('class')
#         selClass = paramList[12]
#         selClass = selClass.strip()
#         print(selClass)

#         print('Subject:')
#         selSubject = paramList[13]
#         selSubject = selSubject.strip()
#         # Start for topic
#         subQuery = MessageDetails.query.filter_by(description=selSubject).first()
#         subId = subQuery.msg_id
#         print(selSubject)
#         print('SubId:'+str(subId))
#         topics = paramList[1].strip()
#         topicList = topics.split(',')
#         print(topicList[0])
#         topic = topicList[0].capitalize()
#         print('Topic:'+str(topic))
#         dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
#         p =1
#         for topic in topicList:
#             fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
#             fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
#             fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
#             if p<len(topicList):
#                 fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
#             p=p+1
#         print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
#         fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
#         Msg = 'no questions available'
#         if len(fetchQuesIds)==0:
#             return jsonify({'onlineTestLink':Msg})
#         listLength = len(fetchQuesIds)
#         count_marks = int(paramList[0]) * int(listLength)
        
#         subjId = ''
#         topicID = ''
#         boardID = ''
#         for det in fetchQuesIds:
#             subjId = det.subject_id
#             topicID = det.topic_id
#             boardID = det.board_id
#             break
#         print('subjId:'+str(subjId))
#         print(fetchQuesIds)
#         # currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher_id.school_id),class_val=str(selClass).strip()).first()
#         resp_session_id = str(subId).strip()+ str(dateVal).strip() + str(randint(10,99)).strip()
#         format = "%Y-%m-%d %H:%M:%S"
#         now_utc = datetime.now(timezone('UTC'))
#         now_local = now_utc.astimezone(get_localzone())
#         print('Date of test creation:'+str(now_local.strftime(format)))

#         # clasVal = selClass.replace('_','@')
#         # testType = paramList[11].replace('_','@')
#         # linkForTeacher=url_for('testLinkWhatsappBot',testType=paramList[11],totalMarks=count_marks,respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=currClassSecRow.section,subjectId=subjId,phone=contactNo, _external=True)
#         # key = '265e29e3968fc62f68da76a373e5af775fa60'
#         # url = urllib.parse.quote(linkForTeacher)
#         # name  = ''
#         # r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
#         # print('New Link')
#         # print(r.text)
#         # print(type(r.text))
#         # linkList = []
#         # jsonLink = json.dumps(r.text)
#         # newData = json.loads(r.text)
#         # print(type(newData))
#         # for linkData in newData['url'].values():
#         #     linkList.append(linkData)
#         # finalLink = linkList[3]
#         # newLink = str('Here is the link to the online test:\n')+finalLink+str('\nDo you want to download the question paper?\n1 - Yes\n2 - No')
#         # print('newLink'+str(newLink))
#         test_type=paramList[11]
#         count = paramList[3]
#         weightage = paramList[0]
#         total_marks = int(count) * int(weightage)
#         class_sec_id = classesListData.class_sec_id
#         print('selected chapter')
#         print(paramList[1])
#         # file_name_val = url_for('question_paper',limit=paramList[3],chapter=paramList[1],schoolName=paramList[18],class_val=selClass,test_type=paramList[11],subject=selSubject,total_marks=count_marks,today=datetime.today().strftime("%d%m%Y%H%M%S"),_external=True)
#         file_name_val = url_for('downloadPaper',test_id='123')
#         return jsonify({'fileName':file_name_val,'selChapter':paramList[1],'boardID':boardID,'resp_session_id':resp_session_id})      



@app.route('/getEnteredTopicList',methods=['POST','GET'])
def getEnteredTopicList():
    if request.method == 'POST':
        print('inside getEnteredTopicList')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
        a = json.dumps(jsonExamData)
      
        z = json.loads(a)
        
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for data in z['results'].values():
            
            paramList.append(data)
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)
        print(conList[2])
        
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        print('class')
        selClass = paramList[12]
        selClass = selClass.strip()
        print(selClass)

        print('Subject:')
        selSubject = paramList[13]
        selSubject = selSubject.strip()
        # Start for topic
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        topics = paramList[1].strip()
        topicList = topics.split(',')
        print(topicList[0])
        topic = topicList[0].capitalize()
        print('Topic:'+str(topic))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        p =1
        for topic in topicList:
            fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
            if p<len(topicList):
                fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
            p=p+1
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        Msg = 'no questions available'
        if len(fetchQuesIds)==0:
            return jsonify({'onlineTestLink':Msg})
        listLength = len(fetchQuesIds)
        count_marks = int(paramList[0]) * int(listLength)
        
        subjId = ''
        topicID = ''
        boardID = ''
        for det in fetchQuesIds:
            subjId = det.subject_id
            topicID = det.topic_id
            boardID = det.board_id
            break
        print('subjId:'+str(subjId))
        print(fetchQuesIds)
        currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher_id.school_id),class_val=str(selClass).strip()).first()
        resp_session_id = str(subId).strip()+ str(dateVal).strip() + str(randint(10,99)).strip()
        format = "%Y-%m-%d %H:%M:%S"
        now_utc = datetime.now(timezone('UTC'))
        now_local = now_utc.astimezone(get_localzone())
        print('Date of test creation:'+str(now_local.strftime(format)))

        # clasVal = selClass.replace('_','@')
        # testType = paramList[11].replace('_','@')
        # linkForTeacher=url_for('testLinkWhatsappBot',testType=paramList[11],totalMarks=count_marks,respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=currClassSecRow.section,subjectId=subjId,phone=contactNo, _external=True)
        # key = '265e29e3968fc62f68da76a373e5af775fa60'
        # url = urllib.parse.quote(linkForTeacher)
        # name  = ''
        # r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
        # print('New Link')
        # print(r.text)
        # print(type(r.text))
        # linkList = []
        # jsonLink = json.dumps(r.text)
        # newData = json.loads(r.text)
        # print(type(newData))
        # for linkData in newData['url'].values():
        #     linkList.append(linkData)
        # finalLink = linkList[3]
        # newLink = str('Here is the link to the online test:\n')+finalLink+str('\nDo you want to download the question paper?\n1 - Yes\n2 - No')
        # print('newLink'+str(newLink))
        test_type=paramList[11]
        count = paramList[3]
        weightage = paramList[0]
        total_marks = int(count) * int(weightage)
        class_sec_id = currClassSecRow.class_sec_id
        print('selected chapter')
        print(paramList[1])
        # file_name_val = url_for('question_paper',limit=paramList[3],chapter=paramList[1],schoolName=paramList[18],class_val=selClass,test_type=paramList[11],subject=selSubject,total_marks=count_marks,today=datetime.today().strftime("%d%m%Y%H%M%S"),_external=True)
        file_name_val = url_for('downloadPaper',test_id='123')
        return jsonify({'fileName':file_name_val,'selChapter':paramList[1],'boardID':boardID,'resp_session_id':resp_session_id})      

# @app.route('/downloadPaper/<test_id>')
# def downloadPaper(test_id):
#     print('inside question paper')
#     school_name = request.args.get('schoolName')
#     class_val = request.args.get('class_val')
#     test_type = request.args.get('test_type')
#     today = request.args.get('today')
#     limit = request.args.get('limit')
#     chapter = request.args.get('chapter')
#     total_marks = request.args.get('total_marks')
#     subject = request.args.get('subject')
#     # fetchQuesIds = request.args.get('fetchQuesIds')
#     topics = chapter.strip()
#     topicList = topics.split(',')
#     dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
#     p =1
#     for topic in topicList:
#         fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id,qd.reference_link "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and qd.question_type='MCQ1' and qd.archive_status='N' and td.class_val='"+str(class_val)+"' and md.description ='"+str(subject)+"' limit '"+str(limit)+"'"
#         if p<len(topicList):
#             fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
#         p=p+1
#     print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
#     fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()

#     myDict = {}
#     options = ''
#     for question in fetchQuesIds:
#         print('QId:'+str(question.question_id))
#         data=QuestionDetails.query.filter_by(question_id=int(question.question_id), archive_status='N').first()
#         print(data)
#         options=QuestionOptions.query.filter_by(question_id=data.question_id).all()    
#         newOpt = [] 
#         for option in options:
#             newOpt.append(option.option_desc)
#         myDict[question.question_id] = newOpt
#     # myDict['1'] = [1,2,3,4]
#     print(myDict)
#     return render_template('questionPaper.html',myDict=myDict,school_name=school_name,class_val=class_val,test_type=test_type,today=today,total_marks=total_marks,subject=subject,fetchQuesIds=fetchQuesIds)


@app.route('/enteredTopicTestDet',methods=['POST','GET'])
def enteredTopicTestDet():
    if request.method == 'POST':
        print('insert enteredTopicTestDet')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "8802362259"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
                
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)

        print(conList[2])

        return jsonify({'success':'success'})  

# @app.route('/addStudentEnteredTopicTestDet',methods=['GET','POST'])
# def addStudentEnteredTopicTestDet():
#     if request.method == 'POST':
#         print('inside addStudentEnteredTopicTestDet')
#         print('insert addEnteredTopicTestDet')
#         jsonExamData = request.json
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         paramList = []
#         conList = []
#         print('data:')
#         print(z)
#         for data in z['results'].values():
                
#             paramList.append(data)
#             print('data:'+str(data))
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)

#         print(conList[2])
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#             # 
#         print('Data Contact')
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         studentData = StudentProfile.query.filter_by(phone=contactNo).first()
#         userId = paramList[14]
#         teacher_id = paramList[15]
            
#         selClass = paramList[12]
            
#         selSubject = paramList[13]
#         subId = paramList[16]
#         print(selSubject)
#         print('SubId:'+str(subId))
#         topics = paramList[1].strip()
#         topicList = topics.split(',')
#         print(topicList[0])
#         topic = topicList[0].capitalize()
#         print('Topic:'+str(topic))
#         dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
#         p =1
#         for topic in topicList:
#             fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
#             fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
#             fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and qd.archive_status='N' and qd.question_type='MCQ1' and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
#             if p<len(topicList):
#                 fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
#             p=p+1
#         print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
#         fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
#         msg = 'No questions available'
#         if len(fetchQuesIds)==0 or fetchQuesIds=='':
#             return jsonify({'testId':msg})
#         listLength = len(fetchQuesIds)
#         total_marks = int(paramList[0]) * int(listLength)
#         boardID = paramList[20]
#         test_type = paramList[11]
#         subjId = paramList[16]
#         class_val = paramList[4]
#         file_name_val = paramList[21]
#         school_id = paramList[17]
#         teacher_id = paramList[15]
#         resp_session_id = paramList[22]
#         format = "%Y-%m-%d %H:%M:%S"
#         now_utc = datetime.now(timezone('UTC'))
#         now_local = now_utc.astimezone(get_localzone())
#         print('Date of test creation:'+str(now_local.strftime(format)))
#         classDet = ClassSection.query.filter_by(class_val=selClass,school_id=studentData.school_id).first()
#         class_sec_id = classDet.class_sec_id
#         testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
#             board_id=str(boardID), subject_id=int(subjId),class_val=str(selClass),date_of_creation=now_local.strftime(format),
#             date_of_test=datetime.now(), school_id=studentData.school_id, teacher_id=teacher_id)
#         db.session.add(testDetailsUpd)
#         db.session.commit()
#         file_name_val = url_for('downloadPaper',test_id=testDetailsUpd.test_id,_external=True)
#         testDet = TestDetails.query.filter_by(test_id=testDetailsUpd.test_id).first()
#         testDet.test_paper_link = file_name_val
#         db.session.commit()
#         sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
#             test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
#         db.session.add(sessionDetailRowInsert)
#         for questionVal in fetchQuesIds:
#             testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
#             db.session.add(testQuestionInsert)
#         db.session.commit()
#         testId = testDetailsUpd.test_id 
#         return jsonify({'testId':testId,'section':classDet.section,'total_marks':total_marks})   


# @app.route('/addEnteredTopicTestDet',methods=['GET','POST'])
# def addEnteredTopicTestDet():
#     if request.method == 'POST':
#         print('insert addEnteredTopicTestDet')
#         jsonExamData = request.json
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         paramList = []
#         conList = []
#         print('data:')
#         print(z)
#         for data in z['results'].values():
                
#             paramList.append(data)
#             print('data:'+str(data))
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)

#         print(conList[2])
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#             # 
#         print('Data Contact')
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         # studentData = StudentProfile.query.filter_by(phone=contactNo).first()
#         userId = paramList[14]
#         teacher_id = paramList[15]
            
#         selClass = paramList[12]
            
        selSubject = paramList[13]
        # subId = paramList[16]
        # print(selSubject)
        # print('SubId:'+str(subId))
        # topics = paramList[1].strip()
        # topicList = topics.split(',')
        # print(topicList[0])
        # topic = topicList[0].capitalize()
        # print('Topic:'+str(topic))
        # dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        # p =1
        # for topic in topicList:
        #     fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
        #     fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
        #     fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and qd.archive_status='N' and qd.question_type='MCQ1' and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
        #     if p<len(topicList):
        #         fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
        #     p=p+1
        # print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        # fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        # msg = 'No questions available'
        # if len(fetchQuesIds)==0 or fetchQuesIds=='':
        #     return jsonify({'testId':msg})
        # listLength = len(fetchQuesIds)
        # total_marks = int(paramList[0]) * int(listLength)
        # boardID = paramList[20]
        # test_type = paramList[11]
        # subjId = paramList[16]
        # class_val = paramList[4]
        # file_name_val = paramList[21]
        # school_id = paramList[17]
        # teacher_id = paramList[15]
        # resp_session_id = paramList[22]
        # format = "%Y-%m-%d %H:%M:%S"
        # now_utc = datetime.now(timezone('UTC'))
        # now_local = now_utc.astimezone(get_localzone())
        # print('Date of test creation:'+str(now_local.strftime(format)))
        # classDet = ClassSection.query.filter_by(class_val=selClass,school_id=school_id).first()
        # class_sec_id = classDet.class_sec_id
        # testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
        #     board_id=str(boardID), subject_id=int(subjId),class_val=str(selClass),date_of_creation=now_local.strftime(format),
        #     date_of_test=datetime.now(), school_id=school_id, teacher_id=teacher_id)
        # db.session.add(testDetailsUpd)
        # db.session.commit()
        # file_name_val = url_for('downloadPaper',test_id=testDetailsUpd.test_id,_external=True)
        # testDet = TestDetails.query.filter_by(test_id=testDetailsUpd.test_id).first()
        # testDet.test_paper_link = file_name_val
        # db.session.commit()
        # sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
        #     test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
        # db.session.add(sessionDetailRowInsert)
        # for questionVal in fetchQuesIds:
        #     testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
        #     db.session.add(testQuestionInsert)
        # db.session.commit()
        # testId = testDetailsUpd.test_id 
        # return jsonify({'testId':testId,'section':classDet.section,'total_marks':total_marks})   

@app.route('/checkQuestions',methods=['GET','POST'])
def checkQuestions():
    if request.method == 'POST':
        print('inside checkQuestions')
        jsonData = request.json
        a = json.dumps(jsonData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
                
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)

        print(conList[2])
        print('Testing for topic')
        # print(type(paramList[1]))
        # print(int(paramList[1]))
            # 
        print('Data Contact')
        contactNo = conList[2][-10:]
        print(contactNo)
        selClass = paramList[11]
        subId  = paramList[15]
        extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        selChapter = ''
        for chapterName in chapterDetList:
            num = chapterName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(chapterName)
                selChapter = chapterName.split('-')[1]
                print('selChapter:'+str(selChapter))
        selChapter = selChapter.strip()
        print('Chapter'+str(selChapter))
        selSubject = paramList[12]
        
        fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.archive_status='N' and qd.question_type='MCQ1' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"'"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        msg = 'Test Link'
        if len(fetchQuesIds)==0 or fetchQuesIds=='':
            msg = 'No questions available'
            return jsonify({'msg':msg})
        return jsonify({'msg':msg})

# @app.route('/addStudentTestDet',methods=['GET','POST'])
# def addStudentTestDet():
#     if request.method == 'POST':
#         print('inside addStudentTestDet')        
#         jsonExamData = request.json
#             # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         paramList = []
#         conList = []
#         print('data:')
#         print(z)
#         for data in z['results'].values():
                
#             paramList.append(data)
#             print('data:'+str(data))
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)

#         print(conList[2])
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#             # 
#         print('Data Contact')
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         studentData = StudentProfile.query.filter_by(phone=contactNo).first()
#         teacherData = TeacherProfile.query.filter_by(user_id=studentData.user_id).first()
#         schoolData = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#         teacher_id = teacherData.teacher_id
#         classData = ClassSection.query.filter_by(class_sec_id=studentData.class_sec_id).first()
#         selClass = classData.class_val
#         selSubject = paramList[13]
#         subId = paramList[16]
#         selChapter = paramList[19]
#         print('Chapter'+str(selChapter))
#         dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
#         fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.archive_status='N' and qd.question_type='MCQ1' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
#         print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
#         fetchQuesIds = db.session.execute(fet # subId = paramList[16]
        # print(selSubject)
        # print('SubId:'+str(subId))
        # topics = paramList[1].strip()
        # topicList = topics.split(',')
        # print(topicList[0])
        # topic = topicList[0].capitalize()
        # print('Topic:'+str(topic))
        # dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        # p =1
        # for topic in topicList:
        #     fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
        #     fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
        #     fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and qd.archive_status='N' and qd.question_type='MCQ1' and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
        #     if p<len(topicList):
        #         fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
        #     p=p+1
        # print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        # fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        # msg = 'No questions available'
        # if len(fetchQuesIds)==0 or fetchQuesIds=='':
        #     return jsonify({'testId':msg})
        # listLength = len(fetchQuesIds)
        # total_marks = int(paramList[0]) * int(listLength)
        # boardID = paramList[20]
        # test_type = paramList[11]
        # subjId = paramList[16]
        # class_val = paramList[4]
        # file_name_val = paramList[21]
        # school_id = paramList[17]
        # teacher_id = paramList[15]
        # resp_session_id = paramList[22]
        # format = "%Y-%m-%d %H:%M:%S"
        # now_utc = datetime.now(timezone('UTC'))
        # now_local = now_utc.astimezone(get_localzone())
        # print('Date of test creation:'+str(now_local.strftime(format)))
        # classDet = ClassSection.query.filter_by(class_val=selClass,school_id=school_id).first()
        # class_sec_id = classDet.class_sec_id
        # testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
        #     board_id=str(boardID), subject_id=int(subjId),class_val=str(selClass),date_of_creation=now_local.strftime(format),
        #     date_of_test=datetime.now(), school_id=school_id, teacher_id=teacher_id)
        # db.session.add(testDetailsUpd)
        # db.session.commit()
        # file_name_val = url_for('downloadPaper',test_id=testDetailsUpd.test_id,_external=True)
        # testDet = TestDetails.query.filter_by(test_id=testDetailsUpd.test_id).first()
        # testDet.test_paper_link = file_name_val
        # db.session.commit()
        # sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
        #     test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
        # db.session.add(sessionDetailRowInsert)
        # for questionVal in fetchQuesIds:
        #     testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
        #     db.session.add(testQuestionInsert)
        # db.session.commit()
        # testId = testDetailsUpd.test_id 
        # return jsonify({'testId':testId,'section':classDet.section,'total_marks':total_marks})   
# chQuesIdsQuery).fetchall()
#         msg = 'Finally, how many questions?'
#         if len(fetchQuesIds)==0 or fetchQuesIds=='':
#             msg = 'No questions available'
#             return jsonify({'msg':msg})
#         return jsonify({'msg':msg})

# @app.route('/addStudentTestDet',methods=['GET','POST'])
# def addStudentTestDet():
#     if request.method == 'POST':
#         print('inside addStudentTestDet')        
#         jsonExamData = request.json
#             # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
#         a = json.dumps(jsonExamData)
#         z = json.loads(a)
#         paramList = []
#         conList = []
#         print('data:')
#         print(z)
#         for data in z['results'].values():
                
#             paramList.append(data)
#             print('data:'+str(data))
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)

#         print(conList[2])
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#             # 
#         print('Data Contact')
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         studentData = StudentProfile.query.filter_by(phone=contactNo).first()
#         teacherData = TeacherProfile.query.filter_by(user_id=studentData.user_id).first()
#         schoolData = SchoolProfile.query.filter_by(school_id=teacherData.school_id).first()
#         teacher_id = teacherData.teacher_id
#         classData = ClassSection.query.filter_by(class_sec_id=studentData.class_sec_id).first()
#         selClass = classData.class_val
#         selSubject = paramList[13]
#         subId = paramList[16]
#         selChapter = paramList[19]
#         print('Chapter'+str(selChapter))
#         dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
#         fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
#         fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.archive_status='N' and qd.question_type='MCQ1' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit '"+str(paramList[3])+"'"
#         print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
#         fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
#         msg = 'No questions available'
#         if len(fetchQuesIds)==0 or fetchQuesIds=='':
#             return jsonify({'testId':msg})
#         listLength = len(fetchQuesIds)
#         total_marks = int(paramList[0]) * int(listLength)
#         boardID = schoolData.board_id
#         test_type = paramList[11]
#         subjId = paramList[16]                
#         file_name_val = paramList[21]
#         school_id = paramList[17]
#         teacher_id = paramList[15]
#         resp_session_id = paramList[22]        
#         format = "%Y-%m-%d %H:%M:%S"
#         now_utc = datetime.now(timezone('UTC'))
#         now_local = now_utc.astimezone(get_localzone())
#         print('Date of test creation:'+str(now_local.strftime(format)))
#         # classDet = ClassSection.query.filter_by(class_val=class_val,school_id=school_id).first()
#         # class_sec_id = classDet.class_sec_id
#         testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
#             board_id=str(boardID), subject_id=int(subjId),class_val=str(selClass),date_of_creation=now_local.strftime(format),
#             date_of_test=datetime.now(), school_id=teacherData.school_id, teacher_id=teacher_id)
#         db.session.add(testDetailsUpd)
#         db.session.commit()
#         file_name_val = url_for('downloadPaper',test_id=testDetailsUpd.test_id,_external=True)
#         testDet = TestDetails.query.filter_by(test_id=testDetailsUpd.test_id).first()
#         testDet.test_paper_link = file_name_val
#         db.session.commit()
#         sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
#             test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=classData.class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
#         db.session.add(sessionDetailRowInsert)
#         for questionVal in fetchQuesIds:
#             testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
#             db.session.add(testQuestionInsert)
#         db.session.commit()
#         testId = testDetailsUpd.test_id 
#         return jsonify({'testId':testId,'section':classData.section,'total_marks':total_marks})   


@app.route('/addTestDet',methods=['GET','POST'])
def addTestDet():
    if request.method == 'POST':
        print('insert addTestDet')
        jsonExamData = request.json
            # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
                
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)

        print(conList[2])
        print('Testing for topic')
        # print(type(paramList[1]))
        # print(int(paramList[1]))
            # 
        print('Data Contact')
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = paramList[14]
        teacher_id = paramList[15]
            
        selClass = paramList[12]
            
        selSubject = paramList[13]
        subId = paramList[16]
        print(selSubject)
        print('SubId:'+str(subId))
        selChapter = paramList[19]
        print('Chapter'+str(selChapter))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.archive_status='N' and qd.question_type='MCQ1' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        
        msg = 'No questions available'
        if len(fetchQuesIds)==0 or fetchQuesIds=='':
            return jsonify({'testId':msg})
        listLength = len(fetchQuesIds)
        total_marks = int(paramList[0]) * int(listLength)
        boardID = paramList[20]
        for ques in fetchQuesIds:
            boardID = ques.board_id
            break
        test_type = paramList[11]
        subjId = paramList[16]
        class_val = paramList[12]
        file_name_val = paramList[21]
        school_id = paramList[17]
        teacher_id = paramList[15]
        resp_session_id = paramList[22]
        format = "%Y-%m-%d %H:%M:%S"
        now_utc = datetime.now(timezone('UTC'))
        now_local = now_utc.astimezone(get_localzone())
        print('Date of test creation:'+str(now_local.strftime(format)))
        classDet = ClassSection.query.filter_by(class_val=class_val,school_id=school_id).first()
        class_sec_id = classDet.class_sec_id
        testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(total_marks),last_modified_date= datetime.now(),
            board_id=str(boardID), subject_id=int(subjId),class_val=str(class_val),date_of_creation=now_local.strftime(format),
            date_of_test=datetime.now(), school_id=school_id, teacher_id=teacher_id)
        db.session.add(testDetailsUpd)
        db.session.commit()
        file_name_val = url_for('downloadPaper',test_id=testDetailsUpd.test_id,_external=True)
        testDet = TestDetails.query.filter_by(test_id=testDetailsUpd.test_id).first()
        testDet.test_paper_link = file_name_val
        db.session.commit()
        sessionDetailRowInsert=SessionDetail(resp_session_id=resp_session_id,session_status='80',teacher_id= teacher_id,
            test_id=str(testDetailsUpd.test_id).strip(),class_sec_id=class_sec_id,correct_marks=10,incorrect_marks=0, test_time=0,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)))
        db.session.add(sessionDetailRowInsert)
        for questionVal in fetchQuesIds:
            testQuestionInsert= TestQuestions(test_id=testDetailsUpd.test_id, question_id=questionVal.question_id, last_modified_date=datetime.now(),is_archived='N')
            db.session.add(testQuestionInsert)
        db.session.commit()
        testId = testDetailsUpd.test_id 
        return jsonify({'testId':testId,'section':classDet.section,'total_marks':total_marks})   

@app.route('/addTopicTestPaperNewDet',methods=['GET','POST'])
def addTopicTestPaperNewDet():
    if request.method == 'POST':
        print('inside addEnteredTopicTestDet')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
            
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = paramList[14]
        teacher_id = paramList[15]
            
        selClass = paramList[12]
            
        selSubject = paramList[13]
        subId = paramList[16]
        total_marks = paramList[25]
        section = paramList[24]
        testId = paramList[23]
        resp_session_id = paramList[22]
        print(selSubject)
        print('SubId:'+str(subId))
        topics = paramList[1].strip()
        topicList = topics.split(',')
        print(topicList[0])
        topic = topicList[0].capitalize()
        print('Topic:'+str(topic))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        p =1
        for topic in topicList:
            fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "from question_details qd inner join topic_detail td on qd.topic_id = td.topic_id inner join message_detail md on md.msg_id = td.subject_id "
            fetchQuesIdsQuery = fetchQuesIdsQuery + "where initcap(td.topic_name) like initcap('%"+str(topic.capitalize())+"%') and qd.question_type='MCQ1' and qd.archive_status='N' and td.class_val='"+str(selClass)+"' and md.description ='"+str(selSubject)+"' limit 5"
            if p<len(topicList):
                fetchQuesIdsQuery = fetchQuesIdsQuery + "union "
            p=p+1
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        clasVal = selClass.replace('_','@')
        testType = paramList[11].replace('_','@')
        linkForTeacher=url_for('testLinkWhatsappBot',testType=str(testType),totalMarks=str(total_marks),respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=section,subjectId=subId,phone=contactNo, _external=True)
        key = '265e29e3968fc62f68da76a373e5af775fa60'
        url = urllib.parse.quote(linkForTeacher)
        name  = ''
        r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
        print('New Link')
        print(r.text)
        print(type(r.text))
        linkList = []
        jsonLink = json.dumps(r.text)
        newData = json.loads(r.text)
        print(type(newData))
        for linkData in newData['url'].values():
            linkList.append(linkData)
        finalLink = linkList[3]
        newLink = finalLink
        return jsonify({'onlineTestLink':newLink,'testId':testId})        


@app.route('/generateNewTestLink',methods=['GET','POST'])
def generateNewTestLink():
    if request.method == 'POST':
        print('inside generateNewTestLink')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
            
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = paramList[14]
        teacher_id = paramList[15]
            
        selClass = paramList[12]
            
        selSubject = paramList[13]
        subId = paramList[16]
        total_marks = paramList[25]
        section = paramList[24]
        testId = paramList[23]
        resp_session_id = paramList[22]
        print(selSubject)
        print('SubId:'+str(subId))
        selChapter = paramList[19]
        print('Chapter'+str(selChapter))
        fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.question_type='MCQ1' and qd.archive_status='N' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        clasVal = selClass.replace('_','@')
        testType = paramList[11].replace('_','@')
        linkForTeacher=url_for('testLinkWhatsappBot',testType=str(testType),totalMarks=str(total_marks),respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=section,subjectId=subId,phone=contactNo, _external=True)
        key = '265e29e3968fc62f68da76a373e5af775fa60'
        url = urllib.parse.quote(linkForTeacher)
        name  = ''
        r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
        print('New Link')
        print(r.text)
        print(type(r.text))
        linkList = []
        jsonLink = json.dumps(r.text)
        newData = json.loads(r.text)
        print(type(newData))
        for linkData in newData['url'].values():
            linkList.append(linkData)
        finalLink = linkList[3]
        newLink = finalLink
        return jsonify({'onlineTestLink':newLink,'testId':testId})  

@app.route('/enteredTopicTestPaperNewLink',methods=['GET','POST'])
def enteredTopicTestPaperNewLink():
    if request.method == 'POST':
        print('inside enteredTopicTestPaperNewLink')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
            
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        testId = paramList[0]
        testPaperDet = TestDetails.query.filter_by(test_id=testId).first()
        test_paper = testPaperDet.test_paper_link
        return jsonify({'testPaper':test_paper})               


@app.route('/getTestPaperLinkNew',methods=['GET','POST'])
def getTestPaperLinkNew():
    if request.method == 'POST':
        print('inside getTestPaperLink')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
            
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        contactNo = conList[2][-10:]
        print(contactNo)
        testId = paramList[0]
        testPaperDet = TestDetails.query.filter_by(test_id=testId).first()
        test_paper = testPaperDet.test_paper_link
        return jsonify({'testPaper':test_paper})                  


@app.route('/insertTestData',methods=['GET','POST'])
def insertTestData():
    if request.method == 'POST':
        print('inside insertTestData')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        a = json.dumps(jsonExamData)
        z = json.loads(a)
        paramList = []
        conList = []
        print('data:')
        print(z)
        for data in z['results'].values():
            
            paramList.append(data)
            print('data:'+str(data))
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)

        print(conList[2])
        print('Testing for topic')
        # print(type(paramList[1]))
        # print(int(paramList[1]))
        # 
        print('Data Contact')
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = paramList[14]
        teacher_id = paramList[15]
        
        selClass = paramList[12]
        
        selSubject = paramList[13]
        subId = paramList[16]
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        selChapter = ''
        for chapterName in chapterDetList:
            num = chapterName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(chapterName)
                selChapter = chapterName.split('-')[1]
                print('selChapter:'+str(selChapter))
        selChapter = selChapter.strip()
        print('Chapter'+str(selChapter))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        fetchQuesIdsQuery = "select qd.question_id,qd.topic_id,qd.board_id,qd.subject_id,qd.question_description,qo.option_desc,qd.reference_link from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join question_options qo on qd.question_id=qo.question_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(selChapter)+"%' and qd.question_type='MCQ1' and qd.archive_status='N' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        msg = 'no questions available'
        print('fetchQuesIds:'+str(fetchQuesIds))
        if len(fetchQuesIds)==0 or fetchQuesIds=='':
            return jsonify({'onlineTestLink':msg})
        listLength = len(fetchQuesIds)
        count_marks = int(paramList[0]) * int(listLength)
        
        subjId = ''
        topicID = ''
        boardID = ''
        for det in fetchQuesIds:
            subjId = det.subject_id
            topicID = det.topic_id
            boardID = det.board_id
            break
        print('subjId:'+str(subjId))
        print(fetchQuesIds)
        # currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher_id.school_id),class_val=str(selClass).strip()).first()
        resp_session_id = str(subId).strip()+ str(dateVal).strip() + str(randint(10,99)).strip()
        print('inside insertData')
        school_id = paramList[17]
        format = "%Y-%m-%d %H:%M:%S"
        # schoolQuery = SchoolProfile.query.filter_by(school_id=school_id).first()
        schoolName = paramList[18]
        now_utc = datetime.now(timezone('UTC'))
        now_local = now_utc.astimezone(get_localzone())
        print('Date of test creation:'+str(now_local.strftime(format)))
        # subjectQuery = MessageDetails.query.filter_by(msg_id=subjId).first()
        # document = Document()
        # document.add_heading(schoolName, 0)
        # document.add_heading('Class '+str(selClass)+" - "+str(paramList[11])+" - "+str(datetime.today().strftime("%d%m%Y%H%M%S")) , 1)
        # document.add_heading("Subject : "+str(selSubject),2)
        # document.add_heading("Total Marks : "+str(count_marks),3)
        # p = document.add_paragraph()
        # for question in fetchQuesIds:
        #     data=QuestionDetails.query.filter_by(question_id=int(question.question_id), archive_status='N').first()
        #     options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
        #     document.add_paragraph(
        #         data.question_description, style='List Number'
        #     )    
        #     print(data.reference_link)
        #     if data.reference_link!='' or data.reference_link!=None:

        #         print('inside threadUse if ')
        #         print(data.reference_link)
        #         try:
        #             response = requests.get(data.reference_link, stream=True)
        #             image = BytesIO(response.content)
        #             document.add_picture(image, width=Inches(1.25))
        #         except:
        #             pass
        #     for option in options:
        #         if option.option_desc is not None:
        #             document.add_paragraph(
        #                 option.option+". "+option.option_desc) 
        # cl = selClass.replace("/","-")
        # file_name=str(school_id)+str(cl)+str(selSubject)+str(paramList[11])+str(datetime.today().strftime("%Y%m%d"))+str(count_marks)+'.docx'
   
        # if not os.path.exists('tempdocx'):
        #     os.mkdir('tempdocx')
        # document.save('tempdocx/'+file_name.replace(" ", ""))
        # client = boto3.client('s3', region_name='ap-south-1')
        # client.upload_file('tempdocx/'+file_name.replace(" ", "") , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name.replace(" ", "")),ExtraArgs={'ACL':'public-read'})
        # os.remove('tempdocx/'+file_name.replace(" ", ""))
        # file_name_val='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name.replace(" ", "")
        # file_name_val = url_for('test',limit=paramList[3],chapter=selChapter,schoolName=schoolName,class_val=selClass,test_type=paramList[11],subject=selSubject,total_marks=count_marks,today=datetime.today().strftime("%d%m%Y%H%M%S"),_external=True)
        print(url_for('downloadPaper',test_id='123',_external=True))
        file_name_val = url_for('downloadPaper',test_id='123',_external=True)
        print(file_name_val)
        return jsonify({'fileName':file_name_val,'selChapter':selChapter,'boardID':boardID,'resp_session_id':resp_session_id})      

@app.route('/SampleRoute')
def SampleRoute():
    print(url_for('downloadPaper',test_id='123',_external=True))
    url = url_for('downloadPaper',test_id='123',_external=True)
    return jsonify({'url':url})

@app.route('/downloadPaper/<test_id>')
def downloadPaper(test_id):
    print('inside question paper')
    print('test_id:'+str(test_id))
    sessionDet = SessionDetail.query.filter_by(test_id=test_id).first()
    testDet = TestDetails.query.filter_by(test_id=test_id).first()
    teacher = TeacherProfile.query.filter_by(teacher_id=testDet.teacher_id).first()
    schoolDet = SchoolProfile.query.filter_by(school_id=teacher.school_id).first()
    school_name = schoolDet.school_name
    class_val = testDet.class_val
    test_type = testDet.test_type
    today = datetime.today().strftime("%d%m%Y%H%M%S")
    subjectQuery = MessageDetails.query.filter_by(msg_id=testDet.subject_id).first()
    subject = subjectQuery.description
    # school_name = request.args.get('schoolName')
    # class_val = request.args.get('class_val')
    # test_type = request.args.get('test_type')
    # today = request.args.get('today')
    # limit = request.args.get('limit')
    # chapter = request.args.get('chapter')
    total_marks = sessionDet.total_marks
    # subject = request.args.get('subject')
    # fetchQuesIds = request.args.get('fetchQuesIds')
    # fetchQuesIdsQuery = "select qd.question_id,qd.question_description,qd.reference_link from question_details qd "
    # fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
    # fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
    # fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name like '%"+str(chapter)+"%' and qd.question_type='MCQ1' and qd.archive_status='N' and md.description = '"+str(subject)+"' and td.class_val = '"+str(class_val)+"' limit '"+str(limit)+"'"
    # print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
    # fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
    fetchQuesIdsQuery = "select qd.question_id, qd.reference_link,qd.question_description from question_details qd inner join test_questions tq on qd.question_id = tq.question_id where test_id  = '"+str(test_id)+"'"
    fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()

    print('fetchQuesIds:')
    print(fetchQuesIds)
    myDict = {}
    options = ''
    for question in fetchQuesIds:
        data=QuestionDetails.query.filter_by(question_id=int(question.question_id), archive_status='N').first()
        options=QuestionOptions.query.filter_by(question_id=data.question_id).all()    
        newOpt = [] 
        for option in options:
            newOpt.append(option.option_desc)
        myDict[question.question_id] = newOpt
    # myDict['1'] = [1,2,3,4]
    print(myDict)
    return render_template('questionPaper.html',myDict=myDict,school_name=school_name,class_val=class_val,test_type=test_type,today=today,total_marks=total_marks,subject=subject,fetchQuesIds=fetchQuesIds)

# @app.route('/newTestLinkGenerate',methods=['POST','GET'])
# def newTestLinkGenerate():
#     if request.method == 'POST':
#         print('newTestLinkGenerate')
#         jsonExamData = request.json
#         # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
#         a = json.dumps(jsonExamData)
      
#         z = json.loads(a)
        
        
#         paramList = []
#         conList = []
#         print('data:')
#         # print(z['result'].class_val)
#         # print(z['result'])
#         for data in z['results'].values():
            
#             paramList.append(data)
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)
#         print(conList[2])
#         # Test for topic
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#         # 
#         print('Data Contact')
#         # print(conList[2])
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         userId = User.query.filter_by(phone=contactNo).first()
#         teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
#         classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
#         classList = [] 
#         j=1
#         for classlist in classesListData:
#             classVal = str(j)+str(' - ')+str(classlist.class_val)
#             classList.append(classVal)
#             j=j+1
        
#         selClass = ''
#         print('Selected Class option:')
#         print(paramList[4])
#         for className in classList:
#             num = className.split('-')[0]
#             print('num:'+str(num))
#             print('class:'+str(paramList[4]))
#             if int(num) == int(paramList[4]):
#                 print(className)
#                 selClass = className.split('-')[1]
#                 print('selClass:'+str(selClass))
#         print('class')
#         selClass = selClass.strip()
#         print(selClass)
#         subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
#         print(subQuery)
#         subjectData = db.session.execute(text(subQuery)).fetchall()
#         print(subjectData)
#         subjectList = []
#         k=1
#         subId = ''
#         for subj in subjectData:
#             sub = str(k)+str('-')+str(subj.subject)
#             subjectList.append(sub)
#             k=k+1
#         for subjectName in subjectList:
#             num = subjectName.split('-')[0]
#             print('num:'+str(num))
#             print('class:'+str(paramList[2]))
#             if int(num) == int(paramList[2]):
#                 print(subjectName)
#                 selSubject = subjectName.split('-')[1]
#                 print('selSubject:'+str(selSubject))
                
        # print('Subject:')
        # selSubject = selSubject.strip()
        # # Start for topic
        # subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        # subId = subQuery.msg_id
        # print(selSubject)
        # print('SubId:'+str(subId))
        # extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        # print('Query:'+str(extractChapterQuery))
        # extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        # print(extractChapterData)
        # c=1
        # chapterDetList = []
        # for chapterDet in extractChapterData:
        #     chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
        #     chapterDetList.append(chap)
        #     c=c+1
        # selChapter = ''
        # for chapterName in chapterDetList:
        #     num = chapterName.split('-')[0]
        #     print('num:'+str(num))
        #     print('class:'+str(paramList[1]))
        #     if int(num) == int(paramList[1]):
        #         print(chapterName)
        #         selChapter = chapterName.split('-')[1]
        #         print('selChapter:'+str(selChapter))
        # #End topic
        # selChapter = selChapter.strip()
        # print('Chapter'+str(selChapter))
        # dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        # fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
        # fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        # fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        # fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name = '"+str(selChapter)+"' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
        # print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        # fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        # msg = 'no questions available'
        # print('fetchQuesIds:'+str(fetchQuesIds))
        # if len(fetchQuesIds)==0 or fetchQuesIds=='':
        #     return jsonify({'onlineTestLink':msg})
        # listLength = len(fetchQuesIds)
        # count_marks = int(paramList[0]) * int(listLength)
        
#         subjId = ''
#         topicID = ''
#         boardID = ''
#         for det in fetchQuesIds:
#             subjId = det.subject_id
#             topicID = det.topic_id
#             boardID = det.board_id
#             break
#         print('subjId:'+str(subjId))
#         print(fetchQuesIds)
#         currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher_id.school_id),class_val=str(selClass).strip()).first()
#         resp_session_id = str(subId).strip()+ str(dateVal).strip() + str(randint(10,99)).strip()
#         # threadUse(currClassSecRow.class_sec_id,resp_session_id,fetchQuesIds,paramList[11],count_marks,selClass,teacher_id.teacher_id,teacher_id.school_id)
#         # task = exampleData.delay(10,20)
#         task = insertData.delay(currClassSecRow.class_sec_id,resp_session_id,fetchQuesIds,paramList[11],count_marks,selClass,teacher_id.teacher_id,teacher_id.school_id)
#         clasVal = selClass.replace('_','@')
#         testType = paramList[11].replace('_','@')
#         linkForTeacher=url_for('testLinkWhatsappBot',testType=paramList[11],totalMarks=count_marks,respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=currClassSecRow.section,subjectId=subjId,phone=contactNo, _external=True)
#         key = '265e29e3968fc62f68da76a373e5af775fa60'
#         url = urllib.parse.quote(linkForTeacher)
#         name  = ''
#         r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
#         print('New Link')
#         print(r.text)
#         print(type(r.text))
#         linkList = []
#         jsonLink = json.dumps(r.text)
#         newData = json.loads(r.text)
#         print(type(newData))
#         for linkData in newData['url'].values():
#             linkList.append(linkData)
#         finalLink = linkList[3]
#         newLink = str('Here is the link to the online test:\n')+finalLink+str('\nDo you want to download the question paper?\n1 - Yes\n2 - No')
#         print('newLink'+str(newLink))
#         return jsonify({'onlineTestLink':newLink})

@app.route('/getNewUrl',methods=['POST','GET'])
def getNewUrl():
    jsonData = {"url":{"status":7,"fullLink":"https:\/\/alllearnreview-pr-229.herokuapp.com\/testLinkWhatsappBot?testType=Class+Feedback&totalMarks=50&respsessionid=3320402202117572449&fetchQuesIds=%281004%2C+10%2C+%27Subjective%27%2C+6046%2C+%27Double+Attack%3A+White+to+Move%27%2C+332%2C+3081%29&fetchQuesIds=%281004%2C+10%2C+%27Subjective%27%2C+6048%2C+%27Double+Attack%3A+White+to+Move%27%2C+332%2C+3081%29&fetchQuesIds=%281004%2C+10%2C+%27Subjective%27%2C+6047%2C+%27Double+Attack%3A+White+to+Move%27%2C+332%2C+3081%29&fetchQuesIds=%281004%2C+10%2C+%27Subjective%27%2C+6051%2C+%27Double+Attack%3A+White+to+Move%27%2C+332%2C+3081%29&fetchQuesIds=%281004%2C+10%2C+%27Subjective%27%2C+6049%2C+%27Double+Attack%3A+White+to+Move%27%2C+332%2C+3081%29&weightage=10&negativeMarking=0&uploadStatus=Y&resultStatus=Y&advance=Y&instructions=&duration=0&classVal=Beginner%40Level%403&section=Abhishek+P&subjectId=332&phone=8802362259","date":"2021-02-04","shortLink":"https:\/\/cutt.ly\/IkkYl2L","title":"allLearn"}}
    linkList = []
    jsonLink = json.dumps(jsonData)
    newData = json.loads(jsonLink)
    print(newData)
    for linkData in newData['url'].values():
        linkList.append(linkData)
    finalLink = linkList[3]
#     return jsonify({'data':finalLink})
# @app.route('/getTestPaperLink',methods=['POST','GET'])
# def getTestPaperLink():
#     if request.method == 'POST':
#         print('inside getTestPaperLink')
#         jsonExamData = request.json
#         # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
#         a = json.dumps(jsonExamData)
      
#         z = json.loads(a)
        
        
#         paramList = []
#         conList = []
#         print('data:')
#         # print(z['result'].class_val)
#         # print(z['result'])
#         for data in z['results'].values():
            
#             paramList.append(data)
#         for con in z['contact'].values():
#             conList.append(con)
#         print(paramList)
#         print(conList[2])
#         # Test for topic
#         print('Testing for topic')
#         # print(type(paramList[1]))
#         # print(int(paramList[1]))
#         # 
#         print('Data Contact')
#         # print(conList[2])
#         contactNo = conList[2][-10:]
#         print(contactNo)
#         userId = User.query.filter_by(phone=contactNo).first()
#         teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
#         classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
#         classList = [] 
#         j=1
#         for classlist in classesListData:
#             classVal = str(j)+str(' - ')+str(classlist.class_val)
#             classList.append(classVal)
#             j=j+1
        
#         selClass = ''
#         print('Selected Class option:')
#         print(paramList[4])
#         for className in classList:
#             num = className.split('-')[0]
#             print('num:'+str(num))
#             print('class:'+str(paramList[4]))
#             if int(num) == int(paramList[4]):
#                 print(className)
#                 selClass = className.split('-')[1]
#                 print('selClass:'+str(selClass))
#         print('class')
#         selClass = selClass.strip()
#         print(selClass)
#         subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
#         print(subQuery)
#         subjectData = db.session.execute(text(subQuery)).fetchall()
#         print(subjectData)
#         subjectList = []
#         k=1
#         subId = ''
#         for subj in subjectData:
#             sub = str(k)+str('-')+str(subj.subject)
#             subjectList.append(sub)
#             k=k+1
#         for subjectName in subjectList:
#             num = subjectName.split('-')[0]
#             print('num:'+str(num))
#             print('class:'+str(paramList[2]))
#             if int(num) == int(paramList[2]):
#                 print(subjectName)
#                 selSubject = subjectName.split('-')[1]
#                 print('selSubject:'+str(selSubject))
                
    print('Subject:')
    selSubject = selSubject.strip()
    # Start for topic
    subQuery = MessageDetails.query.filter_by(description=selSubject).first()
    subId = subQuery.msg_id
    print(selSubject)
    print('SubId:'+str(subId))
    extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
    print('Query:'+str(extractChapterQuery))
    extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
    print(extractChapterData)
    c=1
    chapterDetList = []
    for chapterDet in extractChapterData:
        chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
        chapterDetList.append(chap)
        c=c+1
    selChapter = ''
    for chapterName in chapterDetList:
        num = chapterName.split('-')[0]
        print('num:'+str(num))
        print('class:'+str(paramList[1]))
        if int(num) == int(paramList[1]):
            print(chapterName)
            selChapter = chapterName.split('-')[1]
            print('selChapter:'+str(selChapter))
    #End topic
    selChapter = selChapter.strip()
    print('Chapter'+str(selChapter))
    dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
    fetchQuesIdsQuery = "select qd.question_id from question_details qd "
    fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
    fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
    fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name = '"+str(selChapter)+"' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit 5"
    print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))        
    fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
    oldQuesIds = []
    for ques in fetchQuesIds:
        if ques:
            oldQuesIds.append(ques.question_id)
    testPaperQuery = "select test_id,test_paper_link from test_details order by test_id desc limit 1"
    print(testPaperQuery)
    testPaperData = db.session.execute(text(testPaperQuery)).first()
    fetchLastPaperQuestionIds = TestQuestions.query.filter_by(test_id=testPaperData.test_id).all()
    newQuesIds = []
    for ques in fetchLastPaperQuestionIds:
        if ques:
            newQuesIds.append(ques.question_id) 
    testPaperLink = ''
    if  oldQuesIds ==  newQuesIds:   
        testPaperLink = str("Here's the test paper link:\n")+str(testPaperData.test_paper_link)
        print('testPaperLink:'+str(testPaperLink))
    else:
        testPaperLink = 'No testpaper available'
    return jsonify({'TestPaperLink':testPaperLink})

    
# @app.route('/testLinkWhatsappBot', methods=['POST','GET'])
# def testLinkWhatsappBot(): 
#     phone = request.args.get('phone') 
#     user = User.query.filter_by(phone=phone).first()
#     teacher= TeacherProfile.query.filter_by(user_id=user.id).first() 
#     student = StudentProfile.query.filter_by(user_id=user.id).first()
#     subject_id = request.args.get('subjectId')
#     print('inside testlinkwhatsappbot')
#     print(subject_id)
#     subjectQuery = MessageDetails.query.filter_by(msg_id=subject_id).first()
#     subjectName = subjectQuery.description
#     classVal = request.args.get('classVal')
#     emailDet = ''
#     if student:
#       emailDet = StudentProfile.query.filter_by(student_id=student.student_id).first()
#     user = ''
    
#     if emailDet:
#         user = User.query.filter_by(email=teacher.email).first()
#     if user:
#         login_user(user,remember='Y')
#     clasVal = classVal.replace('@','_')
#     respsessionid = request.args.get('respsessionid')
#     testQuery = SessionDetail.query.filter_by(resp_session_id=respsessionid).first()
#     testId = testQuery.test_id
#     section = request.args.get('section')
#     fetchQuesQuery = "select question_id from test_questions where test_id='"+str(testId)+"'"
#     fetchQuesIds = db.session.execute(fetchQuesQuery).fetchall()
#     quesIds = []
#     for fetchIds in fetchQuesIds:
#         quesIds.append(fetchIds.question_id)
#     questions = QuestionDetails.query.filter(QuestionDetails.question_id.in_(quesIds)).all()  
#     for ques in questions:
#         print('question description:')
#         print(ques.question_id)
#         print(ques.question_description)
#     # questions = QuestionDetails.query.filter(QuestionDetails.question_id.in_(fetchQuesIds)).all()
#     questionListSize = len(fetchQuesIds)
#     respsessionid = request.args.get('respsessionid')
#     total_marks = request.args.get('totalMarks')
#     weightage = request.args.get('weightage')
#     test_type = request.args.get('testType')
#     test_type = test_type.replace('@','_')
#     uploadStatus = request.args.get('uploadStatus')
#     resultStatus = request.args.get('resultStatus')
#     advance = request.args.get('advance')
#     print('inside testLinkWhatsappBot')
#     print('Subject Id:'+str(subject_id))
#     studId = None
#     if current_user.is_anonymous:
#         print('user id student')
#         return redirect(url_for('feedbackCollectionStudDev',student_id=studId,resp_session_id=respsessionid,school_id=teacher.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance,_external=True))
#         # return render_template('feedbackCollectionStudDev.html',resp_session_id=str(respsessionid),studId=studId,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)
#     else:
#         print('user is teacher') 
#         url = "http://www.school.alllearn.in/feedbackCollectionStudDev?resp_session_id="+str(respsessionid)+"&school_id="+str(teacher.school_id)
#         responseSessionIDQRCode = "https://api.qrserver.com/v1/create-qr-code/?size=150x150&data="+url
#         return render_template('feedbackCollectionTeachDev.html',classSecCheckVal='Y', subject_id=subject_id, 
#             class_val = clasVal, section = section,questions=questions, questionListSize = questionListSize, resp_session_id = respsessionid,responseSessionIDQRCode=responseSessionIDQRCode,
#             subjectName = subjectName, totalMarks=total_marks,weightage=weightage, 
#             batch_test=0,testType=test_type,school_id=teacher.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)

@app.route('/feedbackCollection', methods=['GET', 'POST'])
@login_required
def feedbackCollection():    
    if request.method=='POST':
        teacher= TeacherProfile.query.filter_by(user_id=current_user.id).first()  
        #classSections=ClassSection.query.filter_by(school_id=teacher.school_id).order_by(ClassSection.class_val).all()  
        #distinctClasses = db.session.execute(text("select distinct class_val, count(class_val) from class_section where school_id="+ str(teacher.school_id)+" group by class_val order by class_val")).fetchall()
        teacherProfile = teacher
        #using today's date to build response session id
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        qtest_id = request.form.get('test_id')
        weightage = request.form.get('weightage')
        NegMarking = request.form.get('negativeMarking')
        # code for upload file status
        uploadStatus = request.form.get('uploadStatus')
        resultStatus = request.form.get('resultStatus')
        advance = request.form.get('advance')
        instructions = request.form.get('instructions')
        dueDate = request.form.get('dueDate')
        print('dueDate:'+str(dueDate))
        print('upload status:'+str(uploadStatus))
        print('resultStatus:'+str(resultStatus))
        print('feeedback Collection instructions:'+str(instructions))
        print('advance:'+str(advance))
        if resultStatus=='Y':
            print('result status is yes')
        else:
            print('result status is no')
        if uploadStatus=='Y':
            print('upload status is yes')
        else:
            print('upload status is no')
        if advance=='Y':
            print('Test type is Advance')
        else:
            print('Test type is basic')
        print('Neg Marks:'+str(NegMarking))
        print(type(NegMarking))
        nMarking = abs(int(NegMarking))
        # Marks = int(nMarking)
        nMark = -nMarking
        duration = request.form.get('duration')
        print('Test Id:'+str(qtest_id))
        print('Duration:'+str(duration))
        if duration!='':
            durTime = duration.split('.')[0]
            duration = int(durTime)
        print('Duration in int:'+str(duration))
        if duration=='':
            print('if duration is null')
            duration=0 
        qclass_val = request.form.get('class_val')
        qsection = request.form.get('section')
        qsubject_id = request.form.get('subject_id')
        batch_test =  request.form.get('batch_test')
        batch_id =  request.form.get('batch_id')    
            
        print("this is the section, class_val and teacher: "+ str(qsection).upper() + ' ' + str(qclass_val).strip() + ' '+ str(teacher.school_id))
        if batch_test=="1":
            print("Entered feedback coll")
            print("batchtest: " + batch_test)
            print("batch_id: " + batch_id)
            courseBatchData = CourseBatch.query.filter_by(batch_id=batch_id).first()
            courseBatchData.is_ongoing='N'
            db.session.commit()
        # if all(v is not None for v in [qtest_id, qclass_val, qsection, qsubject_id]):
        qsection = str(qsection)
        currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher.school_id),class_val=str(qclass_val).strip(),section=str(qsection).strip()).first()

        if currClassSecRow is None and batch_test!="1":
            flash('Class and section value not valid')
            return redirect(url_for('test_builder.testPapers'))
        elif batch_test=="1" and  currClassSecRow is None:
            class_sec_id = 1
            qsubject_id=54
        else:
            class_sec_id = currClassSecRow.class_sec_id
                #qsubject_id
                #pass
            #building response session ID
            #print('This is the class section id found in DB:'+ str(currClassSecRow.class_sec_id))
        responseSessionID = request.args.get('resp_session_id')
        print('Response session id:'+str(responseSessionID))
        if responseSessionID=='' or responseSessionID==None:            
            responseSessionID = str(qsubject_id).strip()+ str(dateVal).strip() + str(class_sec_id).strip()
        subjectQueryRow = MessageDetails.query.filter_by(msg_id=qsubject_id).first()
            
        url = "http://www.school.alllearn.in/feedbackCollectionStudDev?resp_session_id="+str(responseSessionID)+"&school_id="+str(teacher.school_id)
        responseSessionIDQRCode = "https://api.qrserver.com/v1/create-qr-code/?size=150x150&data="+url
    
        questionIDList = TestQuestions.query.filter_by(test_id=qtest_id,is_archived='N').all()              
        if weightage==None or weightage=="":
            qId = TestQuestions.query.filter_by(test_id=qtest_id,is_archived='N').first()
            weightage = QuestionDetails.query.filter_by(question_id=qId.question_id).first()

        print('Inside question id list')
        print(questionIDList)          
        questionListSize = len(questionIDList)

        print('Question list size:'+str(questionListSize))
        total_marks = int(weightage)*questionListSize
            #creating a record in the session detail table  
        if questionListSize !=0:
            sessionDetailRowCheck = SessionDetail.query.filter_by(resp_session_id=responseSessionID).first()
                #print('Date:'+str(dateVal))
            print('##########Response Session ID:'+str(responseSessionID))
                #print('If Question list size is not zero')
                #print(sessionDetailRowCheck)
            if sessionDetailRowCheck==None:
                print('if sessionDetailRowCheck is none')
                    #print(sessionDetailRowCheck)   
                format = "%Y-%m-%d %H:%M:%S"
                    # Current time in UTC
                now_utc = datetime.now(timezone('UTC'))
                print(now_utc.strftime(format))
                    # Convert to local time zone
                now_local = now_utc.astimezone(get_localzone())
                print(now_local.strftime(format))  
                testLink = url_for('feedbackCollectionStudDev',resp_session_id=responseSessionID,school_id=teacherProfile.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance, _external=True)
                print('TestLink:'+str(testLink))            
                sessionDetailRowInsert=SessionDetail(resp_session_id=responseSessionID,session_status='80',teacher_id= teacherProfile.teacher_id,
                    class_sec_id=class_sec_id, test_id=str(qtest_id).strip(),correct_marks=weightage,incorrect_marks=nMark, test_time=duration,total_marks=total_marks, last_modified_date = str(now_local.strftime(format)),instructions=instructions,test_due_date=dueDate,test_link=testLink)
                db.session.add(sessionDetailRowInsert)
                print('Adding to the db')

            if batch_test=="1":
                batchTestInsert = BatchTest(batch_id=request.args.get('batch_id'), topic_id=request.args.get('topic_id'), test_id=request.args.get('test_id'), 
                    resp_session_id=responseSessionID, is_current='Y', is_archived='N', last_modified_date=datetime.today())
                db.session.add(batchTestInsert)
                    #courseBatchData = CourseBatch.query.filter_by(batch_id=batch_id).first()
                    #courseBatchData.is_ongoing='N'
            db.session.commit()

        questionList = []
        for questValue in questionIDList:
            print('Question ID:'+str(questValue.question_id))
            questionList.append(questValue.question_id)
            

        testDetailRow = TestDetails.query.filter_by(test_id=qtest_id).first()
        testType = testDetailRow.test_type
            #testTypeNameRow = MessageDetails.query.filter_by(msg_id=testTypeID).first()


        questions = QuestionDetails.query.filter(QuestionDetails.question_id.in_(questionList)).all()  
        for  question in questions:
            print('Question:'+str(question.question_description))         
        totalMarks = 0
        for eachQuest in questions:
            totalMarks = totalMarks + int(eachQuest.suggested_weightage)
        responseSessionIDQRCode = "https://api.qrserver.com/v1/create-qr-code/?size=150x150&data="+str(responseSessionID)
        if teacherProfile.device_preference==195:
            print('the device preference is as expected:' + str(teacherProfile.device_preference))
            return render_template('feedbackCollectionTeachDev.html',classSecCheckVal=classSecCheck(), subject_id=qsubject_id, 
                class_val = qclass_val, section = qsection,questions=questions, questionListSize = questionListSize, resp_session_id = responseSessionID,responseSessionIDQRCode=responseSessionIDQRCode,
                subjectName = subjectQueryRow.description, totalMarks=total_marks,weightage=weightage, 
                batch_test=batch_test,testType=testType,school_id=testDetailRow.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)
        elif teacherProfile.device_preference==78:
            print('the device preference is not as expected' + str(teacherProfile.device_preference))
            return render_template('feedbackCollection.html',classSecCheckVal=classSecCheck(), subject_id=qsubject_id,classSections = classSections, distinctClasses = distinctClasses, class_val = qclass_val, section = qsection, questionList = questionIDList, questionListSize = questionListSize, resp_session_id = responseSessionID)
        else:
            print('the device preference is external webcame' + str(teacherProfile.device_preference))
            return render_template('feedbackCollectionExternalCam.html',classSecCheckVal=classSecCheck(), responseSessionIDQRCode = responseSessionIDQRCode, resp_session_id = responseSessionID,  subject_id=qsubject_id,classSections = classSections, distinctClasses = distinctClasses,questions=questions , class_val = qclass_val, section = qsection, questionList = questionIDList, questionListSize = questionListSize,qtest_id=qtest_id)

    # else:
    #     return redirect(url_for('classCon'))

@app.route('/markSessionComplete')
@login_required
def markSessionComplete():
    resp_session_id = request.args.get('resp_session_id')
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    if sessionDetailRow!=None:
        sessionDetailRow.session_status='82'
        db.session.commit()
        return jsonify(["0"])
    else:
        return jsonify(["1"])


@app.route('/checkQuestionChange')
@login_required
def checkQuestionChange():
    resp_session_id = request.args.get('resp_session_id')
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    if str(sessionDetailRow.session_status).strip()=='80':
        if sessionDetailRow.load_new_question=='Y':
            return jsonify(["Y"])
        else:
            return jsonify(["N"])
    elif str(sessionDetailRow.session_status).strip()=='82':
        print ("returning FR")
        return jsonify(["FR"])
    else:
        print("We're in returning sessionDetail row's sessionstatus NA")
        return jsonify([str(sessionDetailRow.session_status)+'NA'])
            

@app.route('/loadQuestionExtCam')
@login_required
def loadQuestionExtCam():
    resp_session_id=request.args.get('resp_session_id')
    totalQCount = request.args.get('total')
    qnum= request.args.get('qnum')
    print("This is the complete response session ID received in load quest ext cam"+resp_session_id)
    sessionDetailRow=SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    if sessionDetailRow!=None:
        sessionDetailRow.load_new_question='N'
        db.session.commit()
    current_question_id=sessionDetailRow.current_question
    question = QuestionDetails.query.filter_by(question_id=current_question_id, archive_status='N').first()
    questionOp = QuestionOptions.query.filter_by(question_id=current_question_id).all()
    return render_template('_loadQuestionExtCam.html',question=question, questionOp=questionOp,qnum = qnum,totalQCount = totalQCount)

@app.route('/loadQuestion')
@login_required
def loadQuestion():
    question_id = request.args.get('question_id')
    totalQCount = request.args.get('total')
    qnum= request.args.get('qnum')
    resp_session_id=request.args.get('resp_session_id')
    print('Question Id:'+str(question_id))
    print(resp_session_id)
    question = QuestionDetails.query.filter_by(question_id=question_id, archive_status='N').first()
    questionOp = QuestionOptions.query.filter_by(question_id=question_id).all()
    if resp_session_id!=None:
        respSessionQuestionRow=RespSessionQuestion.query.filter_by(resp_session_id=resp_session_id,question_status='86').first()
        if respSessionQuestionRow!=None:
            respSessionQuestionRow.question_status='87'
            db.session.commit()
        sessionDetRow=SessionDetail.query.filter_by(resp_session_id=str(resp_session_id).strip()).first()        
        sessionDetRow.current_question=question_id
        sessionDetRow.load_new_question='Y'
        db.session.commit()
    #for option in questionOp:
    #    print(option.option_desc)
    return render_template('_question.html',question=question, questionOp=questionOp,qnum = qnum,totalQCount = totalQCount,  )    


# route for fetching next question and updating db for each response from student - tablet assessment process

@app.route('/loadQuestionStud')
def loadQuestionStud():
    question_id = request.args.get('question_id')
    totalQCount = request.args.get('total')
    student_id = request.args.get('student_id')
    uploadStatus = request.args.get('uploadStatus')
    resultStatus = request.args.get('resultStatus')
    print('Student id in student_id:'+str(student_id))
    qnum= request.args.get('qnum')
    print('Question Num:'+str(qnum))
    print('totalQCount:'+str(totalQCount))
    print('question_id:'+str(question_id))
    print('questionId:'+str(question_id))
    btn = request.args.get('btn')
    textAns = request.args.get('textAns')
    url = request.args.get('url')
    ######################################################
    response_option = request.args.get('response_option')
    resp_session_id = request.args.get('resp_session_id')
    subject_id =  request.args.get('subject_id')
    last_q_id =  request.args.get('last_q_id')
    questionDet = QuestionDetails.query.filter_by(question_id=question_id).first()
    print('This is the response session id in: ' + str(resp_session_id) )
    studentRow = ''
    if current_user.is_anonymous and student_id=='':     
        print('if user is anonymous')   
        studentRow=StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()
    else:
        print('if user is student of student id:'+str(student_id))
        studentRow=StudentProfile.query.filter_by(student_id=student_id).first()
    #print('#######this is the current user id'+ str(current_user.id))
    print('student_id:'+str(studentRow.student_id))
    resp_id = str(resp_session_id)
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id = resp_id).first()
    #print('########### Session details have been fetched')
    #print(sessionDetailRow)
    teacherID = sessionDetailRow.teacher_id
    format = "%Y-%m-%d %H:%M:%S"
    # Current time in UTC
    now_utc = datetime.now(timezone('UTC'))
    print(now_utc.strftime(format))
    now_local = now_utc.astimezone(get_localzone())
    print('Time:')
    print(now_local.strftime(format))
    # If Test is submitted
    if btn=='submit' or btn=='timeout':
        currentTestId = sessionDetailRow.test_id
        totQues = "select count(*) as totQues from test_questions tq where test_id = '"+str(sessionDetailRow.test_id)+"'"
        print(totQues)
        totQuesVal = db.session.execute(text(totQues)).first()
        print(totQuesVal)
        totMCQQues = "select count(*) as totMCQQues from test_questions tq inner join question_details qd on tq.question_id = qd.question_id where qd.question_type='MCQ1' and tq.test_id='"+str(sessionDetailRow.test_id)+"'"
        print(totMCQQues)
        totMCQQuesVal = db.session.execute(text(totMCQQues)).first()
        print(totMCQQuesVal)
        print('Total Ques:'+str(totQuesVal[0]))
        print('Total MCQ Ques:'+str(totMCQQuesVal[0]))
        
        if totQuesVal[0] == totMCQQuesVal[0]:
            db.session.execute(text('call sp_performance_detail_load_feedback()'))
            db.session.commit()
            print('After load sp_performance_detail_load_feedback()')
            # return render_template('_feedbackReportIndiv.html',flag='')
        fetchRemQues = "select tq.question_id,qd.question_type from test_questions tq inner join question_details qd on tq.question_id = qd.question_id where tq.question_id not in (select question_id from response_capture rc where resp_session_id = '"+str(resp_session_id)+"' or answer_status='279') and tq.test_id='"+str(sessionDetailRow.test_id)+"'"
        print(fetchRemQues)
        fetchRemQues = db.session.execute(text(fetchRemQues)).fetchall()
        
        for remQues in fetchRemQues:
            print('insert into responsecapture table')
            insertRes = ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
            question_id= remQues.question_id, teacher_id= teacherID,
            class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id,answer_status=240,marks_scored=0,last_modified_date= now_local.strftime(format),question_type=remQues.question_type)
            db.session.add(insertRes)
            db.session.commit()
        if studentRow:
            totalMarksQuery = "select sum(marks_scored) as total_marks, count(*) as num_of_questions from response_capture where student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"' and  answer_status<>'279' and question_type='MCQ1'"
        else:
            totalMarksQuery = "select sum(marks_scored) as total_marks, count(*) as num_of_questions from response_capture where student_user_id="+str(current_user.id)+" and resp_session_id='"+str(resp_session_id)+"' and  answer_status<>'279' and question_type='MCQ1'"
        print('Total Marks Query:'+totalMarksQuery)
        totalQ = "select count(*) as num_of_questions from test_questions where test_id='"+str(sessionDetailRow.test_id)+"' and is_archived='N'"
        print('Query:'+str(totalQ))
        totalQ = db.session.execute(text(totalQ)).first()
        
        print('Total questions:'+str(totalQ.num_of_questions))
        totalMarksVal = db.session.execute(text(totalMarksQuery)).first()
        # neg_marks = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
        # marksScoredQuery =  "select sum(suggested_weightage) as marks_scored, count(*) as correct_ans from question_details where question_id "
        # marksScoredQuery=marksScoredQuery+"in (select distinct question_id from response_capture where is_correct='Y' and "
        # marksScoredQuery=marksScoredQuery+"student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"')"
        # incorrect_ques = "select count(*) as incorrect_ques from response_capture rc where is_correct = 'N' and resp_session_id = '"+str(resp_session_id)+"' and (answer_status=239 or answer_status=241)"
        # print(' Query for incorrect question:'+str(incorrect_ques))
        # incorrect_ques = db.session.execute(text(incorrect_ques)).first()
        marksScoredQuery = "select sum(marks_scored) as marks_scored from response_capture where student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279' and question_type='MCQ1'"
        print('Query for scored marks:'+str(marksScoredQuery))
        marksScoredVal = db.session.execute(text(marksScoredQuery)).first()
        # print('Marks Scored Query:'+marksScoredQuery)
        # print('Marks Scored:'+str(marksScoredVal.marks_scored))
        print('Total Marks:'+str(marksScoredVal.marks_scored))
        correctAns = "select count(*) as correct_ans from response_capture where is_correct='Y' and student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279' and question_type='MCQ1'"
        print('Query for number of correct answer:'+str(correctAns))
        correctAns = db.session.execute(text(correctAns)).first()
        # negative_marks = 0
        marks_scored = 0
        # if neg_marks.incorrect_marks>0:
        #     print('incorrect Ques:'+str(incorrect_ques.incorrect_ques))

        #     negative_marks = int(neg_marks.incorrect_marks) * int(incorrect_ques.incorrect_ques)
        if marksScoredVal.marks_scored!=None:
            print('inside marksscoredval is not empty')
            marks_scored = int(marksScoredVal.marks_scored)
        # if negative_marks>0:
        #     print('Negative Marks:'+str(negative_marks))
        #     marks_scored = int(marks_scored) - int(negative_marks)
        # else:
        #     marks_scored = int(marks_scored)
        try:
            if marks_scored>0:
                marksPercentage = (marks_scored/sessionDetailRow.total_marks) *100
            else:
                marksPercentage = 0
        except:
            marksPercentage=0        
        
        print('Marks Percentage:'+str(marksPercentage))
        
        flag = 1
        if studentRow:
            if studentRow.points!=None and studentRow.points!="":
                studentRow.points = int(studentRow.points) + 1
                db.session.commit()
            return render_template('_feedbackReportIndiv.html',btn=btn,totalQ=totalQ,sessionDetailRow=sessionDetailRow,marksPercentage=marksPercentage,marksScoredVal=marksScoredVal,correctAns=correctAns , marks_scored= marks_scored,totalMarksVal =totalMarksVal, student_id=studentRow.student_id, student_name= studentRow.full_name, resp_session_id = resp_session_id,resultStatus=resultStatus )
        else:
            return render_template('_feedbackReportIndiv.html',btn=btn,totalQ=totalQ,sessionDetailRow=sessionDetailRow,marksPercentage=marksPercentage,marksScoredVal=marksScoredVal,correctAns=correctAns , marks_scored= marks_scored,totalMarksVal =totalMarksVal, student_id=current_user.id, student_name= str(current_user.first_name)+' '+str(current_user.last_name), resp_session_id = resp_session_id,resultStatus=resultStatus )

    # End
    print(studentRow)
    if btn=='next':
        if questionDet.question_type=='Subjective':
            if studentRow!=None:
                print('inside if studentRow exist')
                responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id,question_type='Subjective', subject_id = subject_id, resp_session_id = resp_session_id,answer_status=242,last_modified_date= now_local.strftime(format))
            else:
                print('if student Row not exist')
                responseStudUpdateQuery=ResponseCapture(student_user_id=current_user.id,
                    question_id= last_q_id, teacher_id= teacherID,
                    resp_session_id = resp_session_id,question_type='Subjective',answer_status=242,last_modified_date= now_local.strftime(format))
            print(responseStudUpdateQuery)
            db.session.add(responseStudUpdateQuery)
            db.session.commit()
        else:
            if studentRow!=None:
                print('inside if studentRow exist')
                responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id ,question_type='MCQ1', subject_id = subject_id, resp_session_id = resp_session_id, marks_scored= sessionDetailRow.correct_marks,answer_status=242,last_modified_date= now_local.strftime(format))
            else:
                print('if student Row not exist')
                responseStudUpdateQuery=ResponseCapture(student_user_id=current_user.id,
                    question_id= last_q_id, teacher_id= teacherID,
                    resp_session_id = resp_session_id ,question_type='MCQ1', marks_scored= sessionDetailRow.correct_marks,answer_status=242,last_modified_date= now_local.strftime(format))
            print(responseStudUpdateQuery)
            db.session.add(responseStudUpdateQuery)
            db.session.commit()
    print('qId:'+str(last_q_id))
    checkResponse = ''
    if last_q_id:
        print('inside if qId not empty')
        checkResponse = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,question_id= last_q_id,student_id=studentRow.student_id).first()
        if checkResponse:
            if btn=='submitandnext':
                print('inside submitandnext')
                if questionDet.question_type=='Subjective':
                    checkResponse.response_option = response_option
                    if textAns:
                        checkResponse.answer_type = 334
                    else:
                        checkResponse.answer_type = 335
                else:
                    checkResponse.response_option = response_option
                    checkResponse.answer_type = 336
                checkResponse.answer_status = 239
                db.session.commit()
            if btn=='save':
                print('inside savebtn')
                if questionDet.question_type=='Subjective':
                    checkResponse.response_option = response_option
                    if textAns:
                        checkResponse.answer_type = 334
                    else:
                        checkResponse.answer_type = 335
                else:
                    checkResponse.response_option = response_option
                    checkResponse.answer_type = 336
                checkResponse.answer_status = 239
                db.session.commit()
    print('Response Option:'+str(response_option))
    if response_option!='':
        print('if response_option is not null')
        #print('###############Response option is not null###############. This is it' + str(response_option)+ '-')
        optionCheckRow = QuestionOptions.query.filter_by(question_id=last_q_id, option=response_option).first()                    

        #print('this is optionCheckRow'+ str(optionCheckRow))
        ansCheck = ''
        if (optionCheckRow==None):
            ansCheck='N'
        elif (optionCheckRow.is_correct=='Y'):
            ansCheck='Y'
        else:
            ansCheck='N'
             
        print('Class:'+str(studentRow.class_sec_id))
        resp_weightage = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
        if checkResponse:
            if checkResponse.response_option==None or checkResponse.response_option:
                checkResponse.response_option = response_option
                if ansCheck=='N':
                    checkResponse.is_correct = ansCheck
                    checkResponse.marks_scored = resp_weightage.incorrect_marks
                else:
                    checkResponse.is_correct = ansCheck
                
                db.session.commit()
            print('data already exist')
        else:
            questionType = QuestionDetails.query.filter_by(question_id=last_q_id).first()
            if questionType.question_type=='Subjective':
                print('new data insert in response capture in Subjective type')
                if textAns:
                    responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, response_option=response_option, is_correct = ansCheck, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id,answer_status=240,marks_scored=0,last_modified_date= now_local.strftime(format),answer_type=334,question_type='Subjective')
                else:
                    responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, response_option=response_option, is_correct = ansCheck, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id,answer_status=240,marks_scored=0,last_modified_date= now_local.strftime(format),answer_type=335,question_type='Subjective')
            else:
                print('new data insert in response capture in objective type')
                if ansCheck=='N':
                    responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, response_option=response_option, is_correct = ansCheck, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id, marks_scored= resp_weightage.incorrect_marks,answer_status=240,last_modified_date= now_local.strftime(format),answer_type=336,question_type='MCQ1')
                else:
                    responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
                    question_id= last_q_id, response_option=response_option, is_correct = ansCheck, teacher_id= teacherID,
                    class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id, marks_scored= resp_weightage.correct_marks,answer_status=240,last_modified_date= now_local.strftime(format),answer_type=336,question_type='MCQ1')
            print(responseStudUpdateQuery)
            db.session.add(responseStudUpdateQuery)
            db.session.commit()
            if btn=='submitandnext':
                print('inside submitandnext')
                response_cap = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,question_id= last_q_id,student_id=studentRow.student_id).first()
                response_cap.answer_status = 239
                db.session.commit()
            if btn=='save':
                print('inside savebtn')
                response_cap = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,question_id= last_q_id,student_id=studentRow.student_id).first()
                response_cap.answer_status = 239
                db.session.commit()
        
        print('Question numbering')
        
        
        # if btn=='next':
        #     print('inside nextbtn')
        #     response_cap = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,question_id= last_q_id).first()
        #     response_cap.answer_status = 242
        #     db.session.add(response_cap)
        #     db.session.commit()
        
        ######################################################
    # correctOpt = ''
    # question = ''
    # questionOp = ''
    if int(qnum)<= int(totalQCount):
        print('qnum:'+str(qnum))
        print('totalQCount:'+str(totalQCount))
        print('###############q number LESS THAN TOTAL Q COUNT###############')
        question = QuestionDetails.query.filter_by(question_id=question_id, archive_status='N').first()
        questionOp = QuestionOptions.query.filter_by(question_id=question_id).order_by(QuestionOptions.option).all()
        print('this is the last q id#################:'+last_q_id)
        
        answerRes = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,student_id=studentRow.student_id).all()
        if len(answerRes)==0:
            insertData = ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
            question_id= last_q_id, teacher_id= teacherID,
            class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id, marks_scored= 0,answer_status=279,last_modified_date= now_local.strftime(format))
        # session['status'] = []
        answer_list = []
        print('response_session_id:'+str(resp_session_id))
        for row in answerRes:
            answer_pair = []
            answer_pair.append(row.question_id)
            answer_pair.append(row.answer_status)
            answer_list.append(answer_pair)
        # session['status'] = answer_list
        status = 0
        # print(session['status'])
        for row in answer_list:
            print(row[0])
            print(row[1])
        chooseOption = ResponseCapture.query.filter_by(resp_session_id = resp_session_id,question_id=question_id,student_id=studentRow.student_id).first()
        correctOpt = ''
        if chooseOption:
            correctOpt = chooseOption.response_option
        return render_template('_questionStud.html',uploadStatus=uploadStatus,correctOpt = correctOpt,duration=sessionDetailRow.test_time,btn=btn,answer_list=answer_list,question=question, questionOp=questionOp,qnum = int(qnum)+1,totalQCount = totalQCount, last_q_id=question_id)
    # else:
    #     return jsonify(['0'])
    #     print('###############q number MORE THAN TOTAL Q COUNT###############')
    #     # totalMarksQuery = "select sum(suggested_weightage) as total_marks, count(*) as num_of_questions  from question_details where question_id in "
    #     # totalMarksQuery =  totalMarksQuery +"(select distinct question_id from test_questions t1 inner join session_detail t2 on "
    #     # totalMarksQuery =  totalMarksQuery +"t1.test_id=t2.test_id and t2.resp_session_id='"+str(resp_session_id)+"') "

    #     totalMarksQuery = "select sum(marks_scored) as total_marks, count(*) as num_of_questions from response_capture where student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"'"
    #     print('Total Marks Query:'+totalMarksQuery)
    #     totalMarksVal = db.session.execute(text(totalMarksQuery)).first()
    #     neg_marks = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    #     # marksScoredQuery =  "select sum(suggested_weightage) as marks_scored, count(*) as correct_ans from question_details where question_id "
    #     # marksScoredQuery=marksScoredQuery+"in (select distinct question_id from response_capture where is_correct='Y' and "
    #     # marksScoredQuery=marksScoredQuery+"student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"')"
    #     incorrect_ques = "select count(*) as incorrect_ques from response_capture rc where is_correct = 'N' and resp_session_id = '"+str(resp_session_id)+"'"
    #     incorrect_ques = db.session.execute(text(incorrect_ques)).first()
    #     marksScoredQuery = "select sum(marks_scored) as marks_scored, count(*) as correct_ans from response_capture where is_correct='Y' and student_id="+str(studentRow.student_id)+" and resp_session_id='"+str(resp_session_id)+"'"
    #     marksScoredVal = db.session.execute(text(marksScoredQuery)).first()
    #     print('Marks Scored Query:'+marksScoredQuery)
    #     print('Marks Scored:'+str(marksScoredVal.marks_scored))
    #     print('Total Marks:'+str(totalMarksVal.total_marks))
    #     negative_marks = 0
    #     marks_scored = 0
    #     if neg_marks.incorrect_marks>0:
    #         negative_marks = int(neg_marks.incorrect_marks) * int(incorrect_ques.incorrect_ques)
    #     if negative_marks>0:
    #         marks_scored = int(marksScoredVal.marks_scored) - int(negative_marks)
    #     try:
    #         if marks_scored>0:
    #             marksPercentage = (marks_scored/totalMarksVal.total_marks) *100
    #         else:
    #             marksPercentage = 0
    #     except:
    #         marksPercentage=0        
        
    #     print('Marks Percentage:'+str(marksPercentage))
    #     if studentRow.points!=None and studentRow.points!="":
    #         studentRow.points = int(studentRow.points) + 1
    #         db.session.commit()
    #     return render_template('_feedbackReportIndiv.html',marksPercentage=marksPercentage,marksScoredVal=marksScoredVal , marks_scored= marks_scored,totalMarksVal =totalMarksVal, student_id=studentRow.student_id, student_name= studentRow.full_name, resp_session_id = resp_session_id )
    


#@app.route('/responseStudUpdate')
#@login_required
#def responseStudUpdate():        
#    question_id = request.args.get('question_id')
#    response_option = request.args.get('response_option')
#    resp_session_id = request.args.get('resp_session_id')
#    subject_id =  request.args.get('subject_id')
#
#    studentRow=StudentProfile.query.filter_by(user_id=current_user.id).first()
#
#    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id = resp_session_id).first()
#    teacherID = sessionDetailRow.teacher_id
#
#    optionCheckRow = QuestionOptions.query.filter_by(question_id=splitVal[0], option=response_option).first()                    
#
#    #print('this is optionCheckRow'+ str(optionCheckRow))
#    ansCheck = ''
#    if (optionCheckRow==None):
#        ansCheck='N'
#    elif (optionCheckRow.is_correct=='Y'):
#        ansCheck='Y'
#    else:
#        ansCheck='N'
#
#    responseStudUpdateQuery=ResponseCapture(school_id=studentRow.school_id,student_id=studentRow.student_id,
#        question_id= question_id, response_option=response_option, is_correct = ansCheck, teacher_id= teacherID,
#        class_sec_id=studentRow.class_sec_id, subject_id = subject_id, resp_session_id = resp_session_id,last_modified_date= date.today())
#    db.session.add(responseStudUpdateQuery)
#    db.session.commit()
#    return jsonify(['0'])
#

@app.route('/questionAllDetailsMob')
def questionAllDetailsMob():
    question_id = request.args.get('question_id')
    totalQCount = ''
    qnum= ''
    question = QuestionDetails.query.filter_by(question_id=question_id, archive_status='N').order_by(QuestionDetails.question_id).first()
    questionOp = QuestionOptions.query.filter_by(question_id=question_id).order_by(QuestionOptions.option_id).all()
    print('Question Id:'+str(question_id))
    print('Question Op:'+str(questionOp))
    return render_template('_questionMob.html',question=question, questionOp=questionOp,qnum = qnum,totalQCount = totalQCount,  )    


@app.route('/questionAllDetails')
def questionAllDetails():
    question_id = request.args.get('question_id')
    totalQCount = ''
    qnum= ''
    question = QuestionDetails.query.filter_by(question_id=question_id, archive_status='N').order_by(QuestionDetails.question_id).first()
    questionOp = QuestionOptions.query.filter_by(question_id=question_id).order_by(QuestionOptions.option_id).all()
    print('Question Id:'+str(question_id))
    print('Question Op:'+str(questionOp))
    return render_template('_question.html',question=question, questionOp=questionOp,qnum = qnum,totalQCount = totalQCount,  )    




#@app.route('/decodes', methods=['GET', 'POST'])
#def decodeAjax():
#    if request.method == 'POST':
#        decodedData = barCode.decode(request.form['imgBase64'])
#        if decodedData:
#            json_data = json.dumps(decodedData)
#            print(json_data)
#            return jsonify(json_data)
#        return jsonify(['NO BarCode Found'])

@app.route('/responseDBUpdate', methods=['POST'])
def responseDBUpdate():   
    print('Inside Response DB Update')     
    responseList=request.json
    responseSessionID=request.args.get('resp_session_id')
    responseArray = {}
    if responseList:
        #print(responseList)        
        for key, value in responseList.items():
            if key =="formdataVal":
                responseArray = value
                for val in responseArray:

                    #print('this is val: ' + str(val))
                    splitVal= re.split('[:]', val)
                    #print('this is splitVal'+ str(splitVal))
                    response = splitVal[1]
                    #print('this is response from SplitVal[1]' + response)
                    responseSplit = re.split('-|@',response)
                    #print('Response split is: '+ str(responseSplit ))
                    
                    teacherIDRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
                    print('Class Section Id in Response DB Update:'+str(responseSplit[0]))
                    studentDetailQuery = "select class_sec_id from student_profile where student_id=" + responseSplit[0]
                    studentDetailRow = db.session.execute(text(studentDetailQuery)).first()

                    #studentDetailQuery = "select class_sec_id from student_profile where student_id=" + responseSplit[0]
                    questionDetailRow = QuestionDetails.query.filter_by(question_id=splitVal[0], archive_status='N').first()
                    
                    dateVal= datetime.today().strftime("%d%m%Y")

                    # responseSessionID =  str(questionDetailRow.subject_id) + str(dateVal) + str(studentDetailRow.class_sec_id)
                    # print('Class section Id in response DB Update:'+str(studentDetailRow.class_sec_id))
                    print('this is the response session id: ' + str(responseSessionID))
                    #the response session id is a combination of today's date, subject id and the class section id

                    optionCheckRow = QuestionOptions.query.filter_by(question_id=splitVal[0], option=responseSplit[3]).first()
                    

                    #print('this is optionCheckRow'+ str(optionCheckRow))
                    ansCheck = ''
                    if (optionCheckRow==None):
                        ansCheck='N'
                    elif (optionCheckRow.is_correct=='Y'):
                        ansCheck='Y'
                    else:
                        ansCheck='N'

                    weightage = SessionDetail.query.filter_by(resp_session_id=responseSessionID).first()

                    responsesForQuest=ResponseCapture(school_id=teacherIDRow.school_id,student_id=responseSplit[0],
                    question_id= splitVal[0], response_option= responseSplit[3], is_correct = ansCheck, teacher_id= teacherIDRow.teacher_id,
                    class_sec_id=studentDetailRow.class_sec_id, subject_id = questionDetailRow.subject_id, resp_session_id = responseSessionID,last_modified_date= datetime.now(),marks_scored=weightage.correct_marks,answer_status=239)
                    db.session.add(responsesForQuest)
                    db.session.commit()
                #flash('Response Entered')
                return jsonify(['Data ready for entry'])
    return jsonify(['No records entered to DB'])
  
@app.route('/feedbackReport')
def feedbackReport():    
    fromClassPerf = request.args.get('fromClassPerf')
    responseSessionID=request.args.get('resp_session_id')
    print('Response Session Id in FeedBack Report route'+str(responseSessionID))
    responseDataQuery = ResponseCapture.query.filter_by(resp_session_id=responseSessionID).first()
    subjectId = ''

    if responseDataQuery:
        subjectId = responseDataQuery.subject_id
    testData = SessionDetail.query.filter_by(resp_session_id=responseSessionID).first()
    totalMarks = testData.total_marks
    testId = testData.test_id
    test = TestDetails.query.filter_by(test_id=testId).first()
    classDataQuery = ClassSection.query.filter_by(class_sec_id=testData.class_sec_id).first()
    classVal = classDataQuery.class_val
    section = classDataQuery.section
    subject = MessageDetails.query.filter_by(msg_id=test.subject_id).first()
    subjectName = ''
    if subject:
        subjectName = subject.description
    
    
    testType = test.test_type
    print('Class:'+str(classVal))
    print('Section:'+str(section))
    print('Subject:'+str(subjectName))
    print('testType:'+str(testType))
    if fromClassPerf!=None:
    #questionListJson=request.args.get('question_id')
        class_val=request.args.get('class_val')
        subject_id=request.args.get('subject_id')
    ##print('here is the class_val '+ str(class_val))
        section=request.args.get('section')
        section = section.strip()
        dateVal = request.args.get('date')
    print('Class Perf:'+str(fromClassPerf))
    
    #print('here is the section '+ str(section))
    #if (questionListJson != None) and (class_val != None) and (section != None):
    
    teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#
    if fromClassPerf!=None:
#
        classSecRow = ClassSection.query.filter_by(class_val=class_val, section=section, school_id=teacher.school_id).first()       
        print('here is the subject_id: '+ str(subject_id))
    #    #questionDetailRow = QuestionDetails.query.filter_by(question_id=questionListJson[1]).first()
    #    
        if dateVal == None or dateVal=="":
            dateVal= datetime.today().strftime("%d%m%Y")
        else:
            tempDate=dt.datetime.strptime(dateVal,'%Y-%m-%d').date()
            dateVal= tempDate.strftime("%d%m%Y")
            print('Date:'+str(dateVal)+'Subject Id:'+str(subject_id)+'Class Section Id:'+str(classSecRow.class_sec_id))
        responseSessionID = str(dateVal) + str(subject_id) + str(classSecRow.class_sec_id)

    if responseSessionID!=None:
        #
        testId = "select test_id from session_detail sd where resp_session_id = '"+str(responseSessionID)+"'"
        testId = db.session.execute(text(testId)).first()

        # print('Class Section Id in feedbackReport:'+str(classSecRow.class_sec_id))
        print('Here is response session id in feedback report: ' + responseSessionID)   
        # responseResultQuery = "with total_marks_cte as ( "
        # responseResultQuery = responseResultQuery + "select sum(suggested_weightage) as total_weightage, count(*) as num_of_questions  from question_details where question_id in "
        # responseResultQuery = responseResultQuery + "(select distinct question_id from test_questions t1 inner join session_detail t2 on "
        # responseResultQuery = responseResultQuery + "t1.test_id=t2.test_id and t2.resp_session_id='"+str(responseSessionID)+"') ) "
        # responseResultQuery = responseResultQuery + "select distinct sp.roll_number, sp.full_name, sp.student_id, "
        # responseResultQuery = responseResultQuery + "SUM(CASE WHEN rc.is_correct='Y' THEN qd.suggested_weightage ELSE 0 end) AS  points_scored , "
        # responseResultQuery = responseResultQuery + "total_marks_cte.total_weightage "
        # responseResultQuery = responseResultQuery + "from response_capture rc inner join student_profile sp on "
        # responseResultQuery = responseResultQuery + "rc.student_id=sp.student_id "
        # responseResultQuery = responseResultQuery + "inner join question_details qd on "
        # responseResultQuery = responseResultQuery + "qd.question_id=rc.question_id "
        # responseResultQuery = responseResultQuery + "and rc.resp_session_id='"+str(responseSessionID)+"', total_marks_cte "
        # responseResultQuery = responseResultQuery + "group by sp.roll_number, sp.full_name, sp.student_id, total_marks_cte.total_weightage "
        totQues = "select count(*) as totQues from test_questions tq where test_id = '"+str(testId.test_id)+"'"
        print(totQues)
        totQuesVal = db.session.execute(text(totQues)).first()
        print(totQuesVal)
        totMCQQues = "select count(*) as totMCQQues from test_questions tq inner join question_details qd on tq.question_id = qd.question_id where qd.question_type='MCQ1' and tq.test_id='"+str(testId.test_id)+"'"
        print(totMCQQues)
        totMCQQuesVal = db.session.execute(text(totMCQQues)).first()
        print(totMCQQuesVal)
        print('Total Ques:'+str(totQuesVal[0]))
        print('Total MCQ Ques:'+str(totMCQQuesVal[0]))
        if totQuesVal[0]> totMCQQuesVal[0]:
            print('If Subjective question is available')
            responseResultQuery = "select distinct sp.roll_number, sp.full_name, sp.student_id,SUM(CASE WHEN rc.resp_session_id = '"+str(responseSessionID)+"' and (rc.answer_status='239' or answer_status='241') and question_type ='MCQ1' THEN rc.marks_scored ELSE 0 end) as Objective_marks, SUM(CASE WHEN rc.resp_session_id = '"+str(responseSessionID)+"' and (rc.answer_status='239' or answer_status='241') and question_type = 'Subjective' THEN rc.marks_scored ELSE 0 end) AS Subjective_marks,SUM(case when rc.resp_session_id = '"+str(responseSessionID)+"' and (rc.answer_status='239' or answer_status='241') then rc.marks_scored else 0 end) as total_marks "
            responseResultQuery = responseResultQuery + "from response_capture rc inner join student_profile sp on sp.student_id = rc.student_id "
            responseResultQuery = responseResultQuery + "inner join session_detail sd on sd.resp_session_id = rc.resp_session_id where "
            responseResultQuery = responseResultQuery + "rc.resp_session_id = '"+str(responseSessionID)+"' "
            responseResultQuery = responseResultQuery + "group by sp.roll_number, sp.full_name, sp.student_id"
            print('Query:'+str(responseResultQuery))
            responseResultRow = db.session.execute(text(responseResultQuery)).fetchall()
            responseResultRowCount = 0
            total = SessionDetail.query.filter_by(resp_session_id=responseSessionID).first()
            incorrect_ques = "select count(*) as incorrect_ques from response_capture rc where is_correct = 'N' and resp_session_id = '"+str(responseSessionID)+"' and (answer_status='239' or answer_status='241')"
            incorrect_ques = db.session.execute(text(incorrect_ques)).first()
            # neg_marks = 0
            # if total.incorrect_marks>0:
            #     neg_marks = int(incorrect_ques.incorrect_ques)*int(total.incorrect_marks)
            if responseResultRow:
                allStudentsTotalMarks =  0
                totalPointsLimit = 0  
                grandTotal = 0 
                # print(responseResultRow)
                # print('ResultRow length:'+str(len(responseResultRow)))         
                # for row in responseResultRow:
                #     totalPointsScored = totalPointsScored + row.points_scored
                #     totalPointsLimit = totalPointsLimit + row.total_weightage
                totalPointsLimit = total.total_marks
                grandTotal = totalPointsLimit * len(responseResultRow)
                for row in responseResultRow:
                    print('Objective Question:'+str(row.objective_marks))
                    print('Subjective Question:'+str(row.subjective_marks))
                    allStudentsTotalMarks = allStudentsTotalMarks + row.total_marks
                classAverage = 0
                if totalPointsLimit !=0 and totalPointsLimit != None:
                    if allStudentsTotalMarks>0:
                        classAverage = (allStudentsTotalMarks/grandTotal) *100
                    else:
                        classAverage = 0
                else:
                    classAverage = 0
                    print("total Points limit is zero")

                responseResultRowCount = len(responseResultRow)
        else:
            responseResultQuery = "select distinct sp.roll_number, sp.full_name, sp.student_id,SUM(CASE WHEN rc.resp_session_id = '"+str(responseSessionID)+"' and (rc.answer_status='239' or answer_status='241') THEN rc.marks_scored ELSE 0 end) as marks_scored, (SUM(CASE WHEN rc.resp_session_id = '"+str(responseSessionID)+"' and (rc.answer_status='239' or answer_status='241') THEN rc.marks_scored ELSE 0 end)*100)/sd.total_marks AS percentage_marks from "
            responseResultQuery = responseResultQuery + "response_capture rc inner join student_profile sp on sp.student_id = rc.student_id "
            responseResultQuery = responseResultQuery + "inner join session_detail sd on sd.resp_session_id = rc.resp_session_id where "
            responseResultQuery = responseResultQuery + "rc.resp_session_id = '"+str(responseSessionID)+"' "
            responseResultQuery = responseResultQuery + "group by sp.roll_number, sp.full_name, sp.student_id,sd.total_marks"
            print('Query:'+str(responseResultQuery))
            responseResultRow = db.session.execute(text(responseResultQuery)).fetchall()
            responseResultRowCount = 0
            total = SessionDetail.query.filter_by(resp_session_id=responseSessionID).first()
            incorrect_ques = "select count(*) as incorrect_ques from response_capture rc where is_correct = 'N' and resp_session_id = '"+str(responseSessionID)+"' and (answer_status='239' or answer_status='241')"
            incorrect_ques = db.session.execute(text(incorrect_ques)).first()
            # neg_marks = 0
            # if total.incorrect_marks>0:
            #     neg_marks = int(incorrect_ques.incorrect_ques)*int(total.incorrect_marks)
            if responseResultRow:
                allStudentsTotalMarks =  0
                totalPointsLimit = 0  
                grandTotal = 0 
                # print(responseResultRow)
                # print('ResultRow length:'+str(len(responseResultRow)))         
                # for row in responseResultRow:
                #     totalPointsScored = totalPointsScored + row.points_scored
                #     totalPointsLimit = totalPointsLimit + row.total_weightage
                totalPointsLimit = total.total_marks
                grandTotal = totalPointsLimit * len(responseResultRow)
                for row in responseResultRow:
                    allStudentsTotalMarks = allStudentsTotalMarks + row.marks_scored
                classAverage = 0
                if totalPointsLimit !=0 and totalPointsLimit != None:
                    if allStudentsTotalMarks>0:
                        classAverage = (allStudentsTotalMarks/grandTotal) *100
                    else:
                        classAverage = 0
                else:
                    classAverage = 0
                    print("total Points limit is zero")

                responseResultRowCount = len(responseResultRow)

        #print('Here is the questionListJson: ' + str(questionListJson))
    else:
        print("Error collecting data from ajax request. Some values could be null")

    if responseResultRowCount>0:
        flag=''
        if totQuesVal[0]> totMCQQuesVal[0]:
            print('If Subjective Question included')
            flag=1
            return render_template('_feedbackReport.html',total_marks=totalMarks,class_val=classVal,section=section,subjectName=subjectName,testType=testType,flag=flag,totalPointsLimit=totalPointsLimit,classAverage=classAverage, classSecCheckVal=classSecCheck(),responseResultRow= responseResultRow,  responseResultRowCount = responseResultRowCount, resp_session_id = responseSessionID,exam_date=total.last_modified_date)
        else:
            return render_template('_feedbackReport.html',total_marks=totalMarks,class_val=classVal,section=section,subjectName=subjectName,testType=testType,flag=flag,totalPointsLimit=totalPointsLimit,classAverage=classAverage, classSecCheckVal=classSecCheck(),responseResultRow= responseResultRow,  responseResultRowCount = responseResultRowCount, resp_session_id = responseSessionID,exam_date=total.last_modified_date)
    else:
         return jsonify(['Test has not been attempted by any student.'])
         

@app.route('/reviewSubjective',methods=['GET','POST'])
def reviewSubjective():
    resp_sess_id = request.args.get('resp_session_id')
    student_id = request.args.get('student_id')
    print('Response_sess_id:'+str(resp_sess_id))
    print('Student_id:'+str(student_id))
    testId = SessionDetail.query.filter_by(resp_session_id=resp_sess_id).first()
    currentTestId = testId.test_id
    studentDet = StudentProfile.query.filter_by(student_id=student_id).first()
    questionDetailQuery = "select qd.question_description,qd.question_id,qd.reference_link,rc.response_option,rc.answer_type from test_questions tq inner join "
    questionDetailQuery = questionDetailQuery + "response_capture rc on tq.question_id=rc.question_id inner join "
    questionDetailQuery = questionDetailQuery + "question_details qd on tq.question_id=qd.question_id where rc.question_type='Subjective' and tq.test_id='"+str(currentTestId)+"' "
    questionDetailQuery = questionDetailQuery + "and rc.student_id='"+str(student_id)+"' and rc.resp_session_id='"+str(resp_sess_id)+"' and (rc.answer_status='239' or rc.answer_status='241')"
    print('Query:'+str(questionDetailQuery))
    questionDetailRow = db.session.execute(text(questionDetailQuery)).fetchall()
    TestDet = TestDetails.query.filter_by(test_id=currentTestId).first()
    testType = TestDet.test_type
    subject = MessageDetails.query.filter_by(msg_id=TestDet.subject_id).first()
    subjectName = subject.description
    return render_template('reviewPage.html',studentName=studentDet.full_name,questionDetailRow=questionDetailRow,testType=testType,subjectName=subjectName,resp_session_id=resp_sess_id)

@app.route('/studentDashboard',methods=['GET','POST'])
@login_required
def studentDashboard():
    print('inside student dashboard')
    # student_id = ''
    # studentDet = ''
    # if current_user.is_anonymous:
    #     student_id = request.args.get('student_id')
    #     studentDet = StudentProfile.query.filter_by(student_id=student_id).first()
    # else:
    studentDet = StudentProfile.query.filter_by(user_id=current_user.id).first()
    student_id = studentDet.student_id
    print('Student Id:'+str(student_id))
    testHistoryQuery = "SELECT fsprc.student_id,fsprc.subject,fsprc.topics,fsprc.test_date,fsprc.resp_session_id,fsprc.perf_percentage from fn_student_performance_response_capture("+str(student_id)+") fsprc order by test_date desc"
    # testHistoryQuery = testHistoryQuery + "inner join response_capture rc on fsprc.resp_session_id = rc.resp_session_id order by test_date desc limit 50"
    testHistory = db.session.execute(testHistoryQuery).fetchall()
    homeworkDetailQuery = "select sd.homework_id, homework_name, question_count, sd.last_modified_date,count(ssr.answer) as ans_count "
    homeworkDetailQuery = homeworkDetailQuery+ "from homework_detail sd left join student_homework_response ssr on ssr.homework_id =sd.homework_id "
    homeworkDetailQuery = homeworkDetailQuery+" where sd.school_id ="+str(studentDet.school_id)+ " and sd.is_archived='N' and sd.class_sec_id='"+str(studentDet.class_sec_id)+"' group by sd.homework_id,homework_name,question_count, sd.last_modified_date"
    homeworkDetailQuery = homeworkDetailQuery+" order by sd.last_modified_date desc limit 10"
    print(homeworkDetailQuery)
    homeworkData = db.session.execute(homeworkDetailQuery).fetchall()
    upcomingTestDetailQuery ="select md.description as subject,sd.test_due_date,sd.test_time, sd.total_marks,sd.test_link, sd.incorrect_marks, sd.test_id from session_detail sd "
    upcomingTestDetailQuery = upcomingTestDetailQuery + "inner join test_details td on sd.test_id = td.test_id "
    upcomingTestDetailQuery = upcomingTestDetailQuery + "inner join message_detail md on md.msg_id = td.subject_id "
    upcomingTestDetailQuery = upcomingTestDetailQuery + "where sd.resp_session_id not in (select distinct rc.resp_session_id from response_capture rc where student_id = '"+str(studentDet.student_id)+"') and sd.test_due_date > now() and sd.class_sec_id='"+str(studentDet.class_sec_id)+"'"
    print('fetch upcoming Test Query')
    print(upcomingTestDetailQuery)
    upcomigTestDetails = db.session.execute(upcomingTestDetailQuery).fetchall()
    print('Test Res Data:')
    print(testHistory)
    # Overall Performance
    overallSum = 0
    overallPerfValue = 0
    sumMarks = 0
    sum1 = 0
    sum2 = 0
    totalOfflineTestMarks = "select sum(marks_scored) as sum1 from result_upload ru where student_id = '"+str(student_id)+"'"
    print(totalOfflineTestMarks)
    totalOfflineTestMarks = db.session.execute(text(totalOfflineTestMarks)).first()
    if totalOfflineTestMarks.sum1:
        print(totalOfflineTestMarks.sum1)
        sum1 = totalOfflineTestMarks.sum1
    totalOnlineTestMarks = "select sum(student_score) as sum2 from performance_detail pd where student_id = '"+str(student_id)+"'"
    totalOnlineTestMarks = db.session.execute(text(totalOnlineTestMarks)).first()
    
    if totalOnlineTestMarks.sum2:
        print(totalOnlineTestMarks.sum2)
        sum2 = totalOnlineTestMarks.sum2
    sumMarks = int(sum1) + int(sum2)
    print('Total Marks:'+str(sumMarks))
    total1 = "select total_marks as offlineTotal from result_upload ru where student_id = '"+str(student_id)+"'"
    print(total1)
    total1 = db.session.execute(text(total1)).first()
    tot1 = 0
    if total1:
        print(total1.offlinetotal)
        tot1 = total1.offlinetotal
    total2 = "select count(*) as count from performance_detail pd where student_id = '"+str(student_id)+"'"
    total2 = db.session.execute(text(total2)).first()
    total3 = 0
    grandTotal = 0
    if total2.count:
        print(total2.count)
        total3 = total2.count*100
    grandTotal = int(tot1) + int(total3)
    print('Grand Total:'+str(grandTotal))
    # for rows in perfRows:
    #     overallSum = overallSum + int(rows.student_score)
        #print(overallSum)
    try:
        overallPerfValue = round(sumMarks/(grandTotal)*100,2)    
    except:
        overallPerfValue=0 
    # End
    # subjectPerf = ''
    try:
        subjectPerfQuery = "select subject,student_score from fn_leaderboard_responsecapture() where student_id='"+str(student_id)+"' "
        subjectPerf = db.session.execute(subjectPerfQuery).fetchall()
    except:
        subjectPerf = []
    topicTrackerQuery = "with cte_total_topics as "
    topicTrackerQuery = topicTrackerQuery + "(select subject_id,  "
    topicTrackerQuery = topicTrackerQuery +"count(is_covered) as total_topics , max(last_modified_Date) as last_updated_date "
    topicTrackerQuery = topicTrackerQuery +"  from topic_tracker where class_sec_id = '"+ str(studentDet.class_sec_id)+"' group by subject_id)  "
    topicTrackerQuery = topicTrackerQuery +"select c1.subject_id,  t2.description as subject_name, c1.last_updated_date, "
    topicTrackerQuery = topicTrackerQuery +"CASE WHEN COUNT(t1.subject_id) <> 0 THEN COUNT(c1.subject_id) ELSE 0 END "
    topicTrackerQuery = topicTrackerQuery +"topics_covered, c1.total_topics  "
    topicTrackerQuery = topicTrackerQuery +"from topic_tracker t1  "
    topicTrackerQuery = topicTrackerQuery +"right outer join cte_total_topics c1  "
    topicTrackerQuery = topicTrackerQuery +"on c1.subject_id=t1.subject_id and class_sec_id= '"+ str(studentDet.class_sec_id)+"'  "
    topicTrackerQuery = topicTrackerQuery +"and t1.is_covered='Y'  "
    topicTrackerQuery = topicTrackerQuery +"inner join   "
    topicTrackerQuery = topicTrackerQuery +"message_detail t2 on   "
    topicTrackerQuery = topicTrackerQuery +"c1.subject_id=t2.msg_id  "
    topicTrackerQuery = topicTrackerQuery +"group by c1.subject_id, t2.description, c1.total_topics,  c1.last_updated_date"         
    print(topicTrackerQuery)       
    topicRows  = db.session.execute(text(topicTrackerQuery)).fetchall()
    classQuery = ClassSection.query.filter_by(class_sec_id = studentDet.class_sec_id).first()
    qclass_val = classQuery.class_val
    writtenTestCountQuery = "SELECT fsprc.student_id,fsprc.subject,fsprc.topics,fsprc.test_date,fsprc.resp_session_id,fsprc.perf_percentage from fn_student_performance_response_capture("+str(student_id)+") fsprc order by test_date desc"
    writtenTestCountData = db.session.execute(text(writtenTestCountQuery)).fetchall()
    writtenTestCount = len(writtenTestCountData)
    pendingTestCountQuery = "select count(*) from session_detail sd "
    pendingTestCountQuery = pendingTestCountQuery + "inner join test_details td on sd.test_id = td.test_id "
    pendingTestCountQuery = pendingTestCountQuery + "inner join message_detail md on md.msg_id = td.subject_id "
    pendingTestCountQuery = pendingTestCountQuery + "where sd.resp_session_id not in (select distinct rc.resp_session_id from response_capture rc where student_id = '"+str(studentDet.student_id)+"') and sd.test_due_date > now() and sd.class_sec_id='"+str(studentDet.class_sec_id)+"'"
    pendingTestCount = db.session.execute(text(pendingTestCountQuery)).first()
    writtenHomeworkCountQuery = "select distinct homework_id from student_homework_response shr where student_id = '"+str(studentDet.student_id)+"'"
    writtenHomeworkCountData = db.session.execute(text(writtenHomeworkCountQuery)).fetchall()
    writtenHomeworkCount = len(writtenHomeworkCountData)
    pendingHomeworkCountQuery = "select distinct count(*) from homework_detail hd where homework_id not in "
    pendingHomeworkCountQuery = pendingHomeworkCountQuery + "(select distinct homework_id from student_homework_response shr where student_id = '"+str(studentDet.student_id)+"' ) and class_sec_id = '"+str(studentDet.class_sec_id)+"'"
    pendingHomeworkCount = db.session.execute(text(pendingHomeworkCountQuery)).first()
    topicCoveredCountQuery = "select distinct count(*) from topic_tracker tt where is_covered = 'Y' and school_id = '"+str(studentDet.school_id)+"' and is_archived = 'N'"
    topicCoveredCount = db.session.execute(text(topicCoveredCountQuery)).first()
    topicunCoveredCountQuery = "select distinct count(*) from topic_tracker tt where is_covered = 'N' and school_id = '"+str(studentDet.school_id)+"' and is_archived = 'N'"
    topicUncoveredCount = db.session.execute(text(topicunCoveredCountQuery)).first()
    return render_template('studentDashboard.html',subjectPerf=subjectPerf,topicUncoveredCount=topicUncoveredCount,topicCoveredCount=topicCoveredCount,pendingHomeworkCount=pendingHomeworkCount,writtenHomeworkCount=writtenHomeworkCount,pendingTestCount=pendingTestCount,writtenTestCount=writtenTestCount,qclass_val=qclass_val,topicRows=topicRows,overallPerfValue=overallPerfValue,upcomigTestDetails=upcomigTestDetails,homeworkData=homeworkData,testHistory=testHistory,studentDet=studentDet)

@app.route('/addSubjMarks',methods=['GET','POST'])
def addSubjMarks():
    marksList = request.form.getlist('marks')
    quesIdList = request.form.getlist('quesId')
    resp_session_id = request.args.get('resp_session_id')    
    isCorrect = request.form.getlist('isCorrect')
    remarksList = request.form.getlist('remarks')
    format = "%Y-%m-%d %H:%M:%S"
    # Current time in UTC
    now_utc = datetime.now(timezone('UTC'))
    print(now_utc.strftime(format))
    now_local = now_utc.astimezone(get_localzone())
    print('Time:')
    print(now_local.strftime(format))
    for i in range(len(quesIdList)):
        quesStatus = ResponseCapture.query.filter_by(resp_session_id=resp_session_id,question_id=quesIdList[i]).first()
        quesStatus.answer_status = 241
        quesStatus.last_modified_date = now_local.strftime(format)
        db.session.commit()
    for i in range(len(marksList)):
        print('Marks:'+str(marksList[i]))
        print('Remarks:'+str(remarksList[i]))
        print('QuesList:'+str(quesIdList[i]))
        questionDet = ResponseCapture.query.filter_by(resp_session_id=resp_session_id,question_id=quesIdList[i]).first()
        questionDet.marks_scored = marksList[i]
        questionDet.remark = remarksList[i]
        quesStatus.last_modified_date = now_local.strftime(format)
        questionDet.answer_status = 241
        if isCorrect[i]:
            questionDet.is_correct = 'Y'
        else:
            questionDet.is_correct = 'N'
        db.session.commit()
    db.session.execute(text('call sp_performance_detail_load_feedback()'))
    db.session.commit()
    print('After load sp_performance_detail_load_feedback()')
    return jsonify(['0'])

@app.route('/studentFeedbackReport')
@login_required
def studentFeedbackReport():
    student_id = request.args.get('student_id')  
    studentName = request.args.get('student_name') 
    resp_session_id = request.args.get('resp_session_id')
    student_id=student_id.strip()
    studentRow = ''
    if student_id!=None:
        studentRow=StudentProfile.query.filter_by(student_id=student_id).first()
    else:
        if student_id==None:
            student_id = current_user.id
        if studentName==None:
            studentName = str(current_user.first_name)+' '+str(current_user.last_name)
        if current_user.is_anonymous:        
            studentRow=StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()
        else:
            studentRow=StudentProfile.query.filter_by(user_id=current_user.id).first()   
            
    responseCaptureQuery = ''
    
    if studentRow:
        responseCaptureQuery = "select rc.student_id,qd.question_id, qd.question_description,rc.marks_scored,rc.answer_type, rc.response_option,rc.question_type,rc.is_correct, qo2.option_desc as option_desc,qo.option_desc as corr_option_desc, "   
        responseCaptureQuery = responseCaptureQuery +"qo.option as correct_option, rc.answer_status, "
        responseCaptureQuery = responseCaptureQuery +"CASE WHEN qo.option= response_option THEN 'Correct' ELSE 'Not Correct' END AS Result "
        responseCaptureQuery = responseCaptureQuery +"from response_capture rc  "
        responseCaptureQuery = responseCaptureQuery +"inner join question_Details qd on rc.question_id = qd.question_id  and qd.archive_status='N' "    
        responseCaptureQuery = responseCaptureQuery +"left join question_options qo on qo.question_id = rc.question_id and qo.is_correct='Y'  "
        responseCaptureQuery = responseCaptureQuery +"left join question_options qo2 on qo2.question_id = rc.question_id and qo2.option = rc.response_option "
        responseCaptureQuery = responseCaptureQuery +"where student_id='" +  str(student_id) + "' and rc.resp_session_id='"+str(resp_session_id)+ "'"
    else:
        responseCaptureQuery = "select rc.student_id,qd.question_id, qd.question_description,rc.answer_type, rc.response_option,rc.question_type,rc.is_correct, qo2.option_desc as option_desc,qo.option_desc as corr_option_desc, "   
        responseCaptureQuery = responseCaptureQuery +"qo.option as correct_option, rc.answer_status, "
        responseCaptureQuery = responseCaptureQuery +"CASE WHEN qo.option= response_option THEN 'Correct' ELSE 'Not Correct' END AS Result "
        responseCaptureQuery = responseCaptureQuery +"from response_capture rc  "
        responseCaptureQuery = responseCaptureQuery +"inner join question_Details qd on rc.question_id = qd.question_id  and qd.archive_status='N' "    
        responseCaptureQuery = responseCaptureQuery +"left join question_options qo on qo.question_id = rc.question_id and qo.is_correct='Y'  "
        responseCaptureQuery = responseCaptureQuery +"left join question_options qo2 on qo2.question_id = rc.question_id and qo2.option = rc.response_option "
        responseCaptureQuery = responseCaptureQuery +"where student_user_id='" +  str(student_id) + "' and rc.resp_session_id='"+str(resp_session_id)+ "'"
    print('Response Capture Query:'+str(responseCaptureQuery))
    responseCaptureRow = db.session.execute(text(responseCaptureQuery)).fetchall()
    marksScoredQuery = "select sum(marks_scored) as marks_scored from response_capture where student_id="+str(student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279'"
    print('Query for scored marks:'+str(marksScoredQuery))
    marksScoredVal = db.session.execute(text(marksScoredQuery)).first()
    sessionDetailRow = SessionDetail.query.filter_by(resp_session_id=resp_session_id).first()
    marks_scored = 0
        # if neg_marks.incorrect_marks>0:
        #     print('incorrect Ques:'+str(incorrect_ques.incorrect_ques))

        #     negative_marks = int(neg_marks.incorrect_marks) * int(incorrect_ques.incorrect_ques)
    if marksScoredVal.marks_scored!=None:
        print('inside marksscoredval is not empty')
        marks_scored = int(marksScoredVal.marks_scored)
        # if negative_marks>0:
        #     print('Negative Marks:'+str(negative_marks))
        #     marks_scored = int(marks_scored) - int(negative_marks)
        # else:
        #     marks_scored = int(marks_scored)
    try:
        if marks_scored>0:
            marksPercentage = (marks_scored/sessionDetailRow.total_marks) *100
        else:
            marksPercentage = 0
    except:
        marksPercentage=0        
        
    print('Marks Percentage:'+str(marksPercentage))
    SubjectiveMarks = "select sum(marks_scored) as marks_scored from response_capture where student_id="+str(student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279' and question_type='Subjective'"
    SubjectiveMarks = db.session.execute(text(SubjectiveMarks)).first()
    ObjectiveMarks = "select sum(marks_scored) as marks_scored from response_capture where student_id="+str(student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279' and question_type='MCQ1'"
    ObjectiveMarks = db.session.execute(text(ObjectiveMarks)).first()
    correctQuestions = "select count(*) as correctQues from response_capture where student_id="+str(student_id)+" and resp_session_id='"+str(resp_session_id)+"' and (answer_status='239' or answer_status='241') and answer_status<>'279' and is_correct='Y'"
    correctQuestions = db.session.execute(text(correctQuestions)).first()
    totalQuestions = "select count(*) as totalQues from test_questions where test_id='"+str(sessionDetailRow.test_id)+"'"
    totalQuestions = db.session.execute(text(totalQuestions)).first()
    objMarks = 0
    subjMarks = 0
    if ObjectiveMarks.marks_scored:
        objMarks = ObjectiveMarks.marks_scored
    if SubjectiveMarks.marks_scored:
        subjMarks = SubjectiveMarks.marks_scored
    x=''
    y=''
    for responseRow in responseCaptureRow:
        if responseRow.question_type=='Subjective':
            x=1
        else:
            y=1
    if x==1 and y=='':
        print('All are subjective')
    elif y==1 and x=='':
        print('All are objective')
    elif x==1 and y==1:
        print('Both')
    totalMarks = int(subjMarks) + int(objMarks)
    return render_template('studentFeedbackReport.html',classSecCheckVal=classSecCheck(),totalMarks=totalMarks,marksPercentage=marksPercentage,subjective_marks=SubjectiveMarks.marks_scored,objective_marks=ObjectiveMarks.marks_scored,correct_question=correctQuestions.correctques,total_questions=totalQuestions.totalques,studentName=studentName, student_id=student_id, resp_session_id = resp_session_id, responseCaptureRow = responseCaptureRow,disconn=1,x=x,y=y)

@app.route('/testPerformance')
@login_required
def testPerformance():
    user = User.query.filter_by(username=current_user.username).first_or_404()        
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    
    
    #setting up testperformance form
    form=testPerformanceForm()    
    test_date = "select distinct m.description as type,p.date as date,m.msg_id as id,c.class_val as class,c.section as section from performance_detail p,message_detail m,class_section c where c.school_id='"+str(teacher.school_id)+"' and p.school_id='"+str(teacher.school_id)+"' and p.test_type=m.msg_id and c.class_sec_id=p.class_sec_id order by date desc fetch next 10 rows only"  
    datelist = db.session.execute(text(test_date)).fetchall()
    for date in datelist:
        print(str(date.date.date())+' '+str(date.type))
    available_class=ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher.school_id).order_by(ClassSection.class_val).all()
    available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher.school_id).all()    
    available_test_type=MessageDetails.query.filter_by(category='Test type').all()

    class_list=[(str(i.class_val), "Class "+str(i.class_val)) for i in available_class]
    section_list=[(i.section,i.section) for i in available_section]    
    test_type_list=[(i.msg_id,i.description) for i in available_test_type]

    #selectfield choices
    form.class_val.choices = class_list
    form.subject_name.choices= ''
    form.section.choices = ''
    # section_list    
    form.test_type.choices=test_type_list

    #setting up studentperformance form
    form1=studentPerformanceForm()
    
    available_class=ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher.school_id).all()
    available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher.school_id).all()    
    available_test_type=MessageDetails.query.filter_by(category='Test type').all()
    available_student_list=StudentProfile.query.filter_by(school_id=teacher.school_id).all()


    class_list=[(str(i.class_val), "Class "+str(i.class_val)) for i in available_class]
    section_list=[(i.section,i.section) for i in available_section]    
    test_type_list=[(i.msg_id,i.description) for i in available_test_type]
    student_list=[(i.student_id,i.full_name) for i in available_student_list]
    resultSet = db.session.execute(text("Select * from fn_overall_performance_summary('"+str(teacher.school_id)+"') where class='All' and section='All'")).fetchall()
    avg_scores = []
    resultSetCount = 0
    # print('Length of Result Set:'+len(resultSet))
    for resultNum in resultSet:
        resultSetCount+=1
        print('#######################first  Resultset:'+str(resultSetCount))

    #for resultNum in resultSet:
    #    #resultSetCount+=1
    #    print('#######################second  Resultset:'+str(resultSetCount))

    #selectfield choices
    form1.class_val1.choices = class_list
    form1.section1.choices= ''
    # section_list    
    form1.test_type1.choices=test_type_list
    form1.student_name1.choices = ''
    findStudentData = "select p.full_name as full_name,p.student_id as student_id,c.class_val as class  from student_profile as p,class_section as c where p.school_id='"+str(teacher.school_id)+"' and c.school_id='"+str(teacher.school_id)+"' and p.class_sec_id=c.class_sec_id"
    print('Query:'+str(findStudentData))
    students = db.session.execute(text(findStudentData))
    print(students)
    print('Inside Test Performance')
    return render_template('testPerformance.html',classSecCheckVal=classSecCheck(),form=form,form1=form1,resultSetCount=resultSetCount,resultSet=resultSet,datelist=datelist,students=students)


@app.route('/testPerformanceGraph')
@login_required
def testPerformanceGraph():  
    print('Inside testPerformanceGraph')  
    class_val=request.args.get('class_val')
    section=request.args.get('section')
    section = section.strip()    
    test_type=request.args.get('test_type')
    print('class:'+class_val+'section:'+section+'test_type:'+test_type)
    #print('here is the class_val '+ str(class_val))
    dateVal = request.args.get('date')
    print('Date:'+dateVal)
    if dateVal ==None or dateVal=="":
        dateVal= datetime.today()

    teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    classSectionRows=ClassSection.query.filter_by(class_val=int(class_val),section=section, school_id=teacher.school_id).first()
        
    testPerformanceQuery = "select sp.full_name, pd.student_score, md.description "
    testPerformanceQuery = testPerformanceQuery + "from performance_detail pd "
    testPerformanceQuery = testPerformanceQuery + "inner join student_profile sp on "
    testPerformanceQuery = testPerformanceQuery + "sp.student_id=pd.student_id inner join "
    testPerformanceQuery = testPerformanceQuery + "message_detail md on md.msg_id=pd.subject_id "
    testPerformanceQuery = testPerformanceQuery + "where "
    testPerformanceQuery = testPerformanceQuery + "pd.class_sec_id='"+str(classSectionRows.class_sec_id) +"' "
    testPerformanceQuery = testPerformanceQuery + "and "
    testPerformanceQuery = testPerformanceQuery + "pd.date='"+ str(dateVal) +"' "
    testPerformanceQuery = testPerformanceQuery + "and test_type='"+ str(test_type)+ "'"  

    testPerformanceRecords = db.session.execute(text(testPerformanceQuery)).fetchall()
        
    if len(testPerformanceRecords) !=0:

        df = pd.DataFrame( [[ij for ij in i] for i in testPerformanceRecords])
        df.rename(columns={0: 'full_name', 1: 'student_score', 2: 'subject'}, inplace=True)

        student_names= list(df['full_name'])
        student_scores= list(df['student_score'])
        subject= list(df['subject'])    

        distinct_subjects= df['subject'].unique()
        
        subLevelData=[]
        i=0
        for subVal in distinct_subjects:
            filtered_df=df[df.subject==subVal]
            filtered_df_student = list(filtered_df['full_name'])
            filtered_df_student_scored = list(filtered_df['student_score'])            
            tempDict = dict(y=filtered_df_student,x=filtered_df_student_scored,type='bar',name=subVal,orientation='h')
            subLevelData.append(tempDict)
        print(str(subLevelData))

        graphData=[dict()]

        graphJSON = json.dumps(subLevelData, cls=plotly.utils.PlotlyJSONEncoder)        
        return render_template('_testPerformanceGraph.html',graphJSON=graphJSON)
    else:
        return jsonify(['No records found for the selected date'])
    

@app.route('/studentPerformanceGraph')
@login_required
def studentPerformanceGraph(): 
    print('Inside Student performance graph')
      
    class_val=request.args.get('class_val')
    section=request.args.get('section')
    fromPractice = request.args.get('fromPractice')
    print('##### this is the value of fromPractice: '+ str(fromPractice))
    if section!=None:
        section = section.strip()    
    test_type=request.args.get('test_type')
    student_id=request.args.get('student_id')      
    print("Student id is: "+str(student_id))  
    dateVal = request.args.get('date')
    
    if dateVal ==None or dateVal=="":
        dateVal= datetime.today()
    
    teacher=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    fromTestPerformance=0
    
    if fromPractice=='1':
        print('######using newer query')
        studPerformanceQuery = "select test_date as date, CAST(perf_percentage AS INTEGER) as student_score, subject as description"
        studPerformanceQuery = studPerformanceQuery + " from fn_student_performance_response_capture("+str(student_id)+") "
    else:
        print('#######using older query')
        studPerformanceQuery = "select pd.date, CAST(pd.student_score AS INTEGER) as student_score,md.description "
        studPerformanceQuery = studPerformanceQuery + "from performance_detail pd "
        studPerformanceQuery = studPerformanceQuery + "inner join "
        studPerformanceQuery = studPerformanceQuery + "message_detail md on md.msg_id=pd.subject_id "
        studPerformanceQuery = studPerformanceQuery + "and "
        studPerformanceQuery = studPerformanceQuery + "pd.student_id='"+ str(student_id) +"' order by date"            
    

    studPerformanceRecords = db.session.execute(text(studPerformanceQuery)).fetchall()
        
    if len(studPerformanceRecords) !=0:
        df = pd.DataFrame( [[ij for ij in i] for i in studPerformanceRecords])
        df.rename(columns={0: 'date', 1: 'student_score', 2: 'subject'}, inplace=True)

        dateRange= list(df['date'])
        student_scores= list(df['student_score'])
        subject= list(df['subject'])

        distinct_subjects= df['subject'].unique()
        subLevelData=[]
        i=0
        for subVal in distinct_subjects:
            filtered_df=df[df.subject==subVal]
            filtered_df_date = list(filtered_df['date'])
            filtered_df_student_scored = list(filtered_df['student_score'])
            tempDict = dict(y=filtered_df_student_scored,x=filtered_df_date,mode= 'lines+markers',type='scatter',name=subVal,line_shape="spline", smoothing= '1.5')
            subLevelData.append(tempDict)
        print(str(subLevelData))

        graphData=[dict()]

        graphJSON = json.dumps(subLevelData, cls=plotly.utils.PlotlyJSONEncoder)        
        return render_template('_studentPerformanceGraph.html',graphJSON=graphJSON)
    else:
        return jsonify(['No performance data found'])
  


@app.route('/classPerformance')
@login_required
def classPerformance():
    form=feedbackReportForm()

    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()

    available_class=ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()
    available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher_id.school_id).all()    
    available_subject=MessageDetails.query.filter_by(category='Subject').all()


    class_list=[(str(i.class_val), "Class "+str(i.class_val)) for i in available_class]
    section_list=[(i.section,i.section) for i in available_section]    
    subject_name_list=[(i.msg_id,i.description) for i in available_subject]

    #selectfield choices
    form.class_val.choices = class_list
    form.section.choices= ''
    # section_list    
    form.subject_name.choices= ''
    # subject_name_list

    testDetailQuery = "select distinct t1.resp_session_id, t1.last_modified_Date as test_date, t5.teacher_name as conducted_by,t4.class_val, t4.section,t1.session_id, "
    testDetailQuery = testDetailQuery+ "t3.description as subject "
    testDetailQuery = testDetailQuery+ " from session_detail t1 "
    testDetailQuery = testDetailQuery+ " inner join test_details t2 on t2.test_id=t1.test_id "
    testDetailQuery = testDetailQuery+ " inner join message_detail t3 on t2.subject_id=t3.msg_id "
    testDetailQuery = testDetailQuery+ " inner join class_section t4 on t1.class_Sec_id=t4.class_sec_id "
    testDetailQuery = testDetailQuery+ " inner join teacher_profile t5 on t5.teacher_id=t1.teacher_id  and t5.school_id='"+str(teacher_id.school_id)+"' order by test_date desc "
    print(testDetailQuery)
    testDetailRows= db.session.execute(text(testDetailQuery)).fetchall()
    indic='DashBoard'
    return render_template('classPerformance.html',indic=indic,title='Online Test Reports',classSecCheckVal=classSecCheck(),form=form, school_id=teacher_id.school_id, testDetailRows=testDetailRows,user_type_val=str(current_user.user_type))



@app.route('/resultUpload',methods=['POST','GET'])
@login_required
def resultUpload():
    #selectfield choices list
    qclass_val = request.args.get('class_val')
    qsection=request.args.get('section')
    print('Section:'+str(qsection))

    user = User.query.filter_by(username=current_user.username).first_or_404()    
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first()     
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    board_id=SchoolProfile.query.with_entities(SchoolProfile.board_id).filter_by(school_id=teacher_id.school_id).first()
    distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher.school_id)+" GROUP BY class_val order by s")).fetchall()        
    classSections=ClassSection.query.filter_by(school_id=teacher.school_id).all()

    count = 0
    #this section is to load the page for the first class section if no query value has been provided
    if qclass_val==None:
        for section in classSections:                
            if count==0:
                getClassVal = section.class_val
                getSection = section.section          
                print('set class and section values for first load')      
        #if no value has been passed for class and section in query string then use the values fetched from db
                qclass_val = getClassVal
                qsection=getSection
                count+=1
            
    test_type = MessageDetails.query.filter_by(category='Test type').all()
    test_details = "select test_id,date_of_creation,description as subject_name from test_details td inner join message_detail md on md.msg_id=td.subject_id where class_val='"+str(qclass_val)+"' and school_id='"+str(teacher_id.school_id)+"'"
    test_details = db.session.execute(test_details).fetchall()
    for section in classSections:
            print("Class Section:"+section.section)
    subject_name = []
    if current_user.is_authenticated:
        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
        print('Class value:'+str(qclass_val))
        print('school_id:'+str(teacher_id.school_id))
        class_sec_id=ClassSection.query.filter_by(class_val=str(qclass_val),school_id=teacher_id.school_id,section=qsection).all()
        
        print(teacher_id.school_id)
        class_value = str(qclass_val)
        print('Class Value:'+str(class_value))
        for section in class_sec_id:
            print('Section Id:'+str(section.class_sec_id))
        student_list = []
        for section in class_sec_id:
            student_list = StudentProfile.query.filter_by(class_sec_id=section.class_sec_id,school_id=teacher_id.school_id).all()
        queryForSubjectName = "select description,msg_id from message_detail where msg_id in (select distinct subject_id from topic_detail where class_val='"+ str(class_value) +"' and board_id='"+ str(board_id[0]) +"')"
        subject_name = db.session.execute(queryForSubjectName).fetchall()
        print('No of Subjects:'+str(len(subject_name)))
        for subjects in subject_name:
            print('Subjects Name:'+subjects[0])  
        indic='DashBoard'      
        return render_template('resultUpload.html',indic=indic,title='Result Upload',classSecCheckVal=classSecCheck(),test_details=test_details,test_type=test_type,qclass_val=qclass_val,subject_name=subject_name,qsection=qsection, distinctClasses=distinctClasses, classsections=classSections,student_list=student_list,user_type_val=str(current_user.user_type))
        

@app.route('/resultUpload/<class_val>')
def section(class_val):
    if class_val!='All':

        teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
        sections = ClassSection.query.distinct().filter_by(class_val=class_val,school_id=teacher_id.school_id).all()
        sectionArray = []

        for section in sections:
            sectionObj = {}
            sectionObj['section_id'] = section.class_sec_id
            sectionObj['section_val'] = section.section
            sectionArray.append(sectionObj)
        return jsonify({'sections' : sectionArray})
    else:
        return "return"



@app.route('/fetchStudentsName')
def fetchStudentsName():
    print('Inside fetchStudentsName')
    class_val = request.args.get('class_val')
    section = request.args.get('section')
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
    print('class_val:'+class_val)
    query = "select p.full_name as name,p.student_id as id  from student_profile as p,class_section as c where p.school_id='"+str(teacher_id.school_id)+"' and c.school_id='"+str(teacher_id.school_id)+"' and p.class_sec_id="
    if class_val!='All' and section=='All': 
        query =query + "(select class_sec_id from class_section where school_id='"+str(teacher_id.school_id)+"' and class_val='"+str(class_val)+"')"
    if section!='All' and class_val=='All':
        query = query +"(select class_sec_id from class_section where school_id='"+str(teacher_id.school_id)+"' and section='"+str(section)+"')"
    if class_val=='All' and section=='All':
        query = query + "c.class_sec_id"
    if class_val!='All' and section!='All':
        query = query + "(select class_sec_id from class_section where school_id='"+str(teacher_id.school_id)+"' and class_val='"+str(class_val)+"' and section='"+str(section)+"')"
    print('Fetch Student Query:'+query)
    studentList = db.session.execute(text(query))
    
    studentArray = []
    for student in studentList:
        studentObject = {}
        studentObject['student_name'] = str(student[0])
        studentObject['student_id'] = str(student[1])
        studentArray.append(studentObject)
    return jsonify({'studentData':studentArray})




@app.route('/fetchTestDates')
def fetchTestDates():
    print('Inside fetchTestDates')
    class_val = request.args.get('class_val')
    section = request.args.get('section')
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
    query = "select distinct m.description as type,p.date as date,m.msg_id as id,c.class_val as class,c.section as section from performance_detail p,message_detail m,class_section c where c.school_id='"+str(teacher_id.school_id)+"' and p.school_id='"+str(teacher_id.school_id)+"'" 
    print(class_val+' '+section)
    if(class_val!='All'):
        query = query + "and c.class_val='"+str(class_val)+"'"
    if(section!='All'): 
        query = query + "and c.section='"+str(section)+"'" 
    query = query + "and p.test_type=m.msg_id and c.class_sec_id=p.class_sec_id order by date desc fetch next 10 rows only"
    print(query)
    dateData = db.session.execute(text(query))
    dateArray = []    
    for data in dateData:
        print('Inside for')
        dateObject = {}
        dateObject['type'] = str(data[0])
        dateObject['date'] = str(data[1].date())
        dateObject['id'] = str(data[2])
        dateObject['class'] = str(data[3])
        dateObject['section'] = str(data[4])
        print(str(data[0])+' '+str(data[1].date())+' '+str(data[2])+' '+str(data[3])+' '+str(data[4]))
        dateArray.append(dateObject)
    return jsonify({'dateData':dateArray})


@app.route('/scoreGraph')
def scoreGraph():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    class_value = request.args.get('class_val')
    section = request.args.get('section')
    query = "select *from fn_overall_attendance_summary('"+str(teacher_id.school_id)+"') where class='"+str(class_value)+"' and section='"+str(section)+"'"
    attendance = db.session.execute(text(query))
    resultArray = []

    for result in attendance:
        Array = {}
        Array['present'] = str(result.no_of_student_present)
        Array['absent'] = str(result.no_of_student_absent)
        resultArray.append(Array)
    return jsonify({'result':resultArray})

@app.route('/linkWithImpact/<int:impactSchoolID>', methods=["GET","POST"])
@login_required
def linkWithImpact(impactSchoolID):
    teacherProfileData = TeacherProfile.query.filter_by(user_id = current_user.id).first()
    if teacherProfileData!=None:
        schoolProfileData = SchoolProfile.query.filter_by(school_id=teacherProfileData.school_id).first()
        if schoolProfileData!=None:
            schoolProfileData.impact_school_id=impactSchoolID
            db.session.commit()
            status = 0
        else:
            status = 1
    else:
        status = 1
    urlForImpact = app.config['IMPACT_HOST']
    return render_template('linkWithImpact.html',urlForImpact=urlForImpact,status=status,schoolId=schoolProfileData.school_id, impactSchoolID = impactSchoolID)


@app.route('/api/schoolRegistration/<int:impactSchoolID>',methods=["GET","POST"])
def apiSchoolRegistration(impactSchoolID,):
    #Firstly register the school using the data from impact form
    #When inserting a record in schoolProfile ensure that you add the impactSchoolID to it
    ###Use this section to insert into school registration, class section, address and update user as admin
    #
    #
    #
    #
    # if the insert is successful, set status =1
    ###status =1
    ###message = "School registered successfully"
    ###else
    ###status = 0
    ###message = "Error registering school"

    #Secondly return the school_id that has newly been created back to the impact system
    #schoolData = SchoolProfile.query.filter_by(impact_school_id = request.form('impact_school_id')).first()
    #

    #postData= {
    #    "status" : status,
    #    "message" : message
    #    "perf_eval_school_id" : schoolData.school_id
    #}
    return {'result' : postData}    

@app.route('/api/userRegistration',methods=["POST"])
def apiUserRegistration():
    if request.form['internalKey'] == app.config['ALLLEARN_INTERNAL_KEY']:
        userData = User.query.filter_by(email = request.form['email']).first()
        if userData==None:
            user = User(username=request.form['email'], email=request.form['email'], user_type='140', access_status='144', phone=request.form['phone'],
                    first_name = request.form['firstName'],last_name= request.form['lastName'],password_hash=request.form['key'])        
            db.session.add(user)
            db.session.commit()
            print("########## no existing user of the same email")
            return {'result' : 'User Registered successfully', 'value':'0'}  
        else:
            print(str(userData))
            print(str(userData.email))
            print('######### user already present')
            return {'result' : 'User already present','value':'err1'}  
    else:
        return {'result' : 'Key mismatch','value':'err2'}


   


@app.route('/api/performanceSummary/<int:schoolID>')
def apiPerformanceSummary(schoolID):
    query = "Select * from fn_overall_performance_summary("+str(schoolID)+") where class='All'and section='All' and subject='All'"
    
    resultSet = db.session.execute(text(query))
    
    resultArray = []

    for result in resultSet:
        Array = {}
        Array['avg_score'] = str(round(result.avg_score,2))
        Array['highest_mark'] = str(result.highest_mark)
        Array['lowest_mark'] = str(result.lowest_mark)
        Array['no_of_students_above_90'] = str(result.no_of_students_above_90)
        Array['no_of_students_80_90'] = str(result.no_of_students_80_90)
        Array['no_of_students_70_80'] = str(result.no_of_students_70_80)
        Array['no_of_students_50_70'] = str(result.no_of_students_50_70)
        Array['no_of_students_below_50'] = str(result.no_of_students_below_50)
        Array['no_of_students_cross_50'] = str(result.no_of_students_cross_50)
        resultArray.append(Array)
    return {'result' : resultArray}    


@app.route('/testDateSearch/<class_val>')
def testDate(class_val):
    print('Inside TestDate Search')
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    class_sec_id=ClassSection.query.filter_by(class_val=int(class_val),school_id=teacher_id.school_id).first()
    testdates = ResultUpload.query.with_entities(ResultUpload.exam_date).distinct().filter_by(class_sec_id=class_sec_id.class_sec_id,school_id=teacher_id.school_id).all()
    print('class value:'+str(class_val)+"School Id:"+str(teacher_id.school_id)+"Teacher_id:"+str(teacher_id.teacher_id)+"userId:"+str(current_user.id))
    
    dates = []
    for test in testdates:
        print("Test Dates:"+str(test.exam_date))
        test = test.exam_date.date().strftime("%d-%B-%Y")
        print('Dates:'+str(test))
        dates.append(test)
    testdateArray = []
    print(dates)
    for testdate in dates:
        testdateObj = {}
        testdateObj['date'] = testdate
        testdateArray.append(testdateObj)
    return jsonify({'testdates' : testdateArray})

@app.route('/resultUploadHistory')
def resultUploadHistory():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()

    uploadHistoryQuery = "select distinct upload_id, cs.class_val, cs.section, "
    uploadHistoryQuery = uploadHistoryQuery + "md.description as test_type, md2.description as subject, date(ru.last_modified_date) as upload_date, date(ru.exam_date) as exam_date "
    uploadHistoryQuery = uploadHistoryQuery +"from result_upload ru inner join  class_section cs on  cs.class_sec_id=ru.class_sec_id and cs.school_id='"+str(teacher_id.school_id)+"' "
    uploadHistoryQuery = uploadHistoryQuery +"inner join message_detail md on md.msg_id=ru.test_type inner join message_detail md2 on md2.msg_id=ru.subject_id order by exam_date desc"
    
    uploadHistoryRecords = db.session.execute(text(uploadHistoryQuery)).fetchall()
    indic='DashBoard'
    return render_template('resultUploadHistory.html',indic=indic,title='Result History',uploadHistoryRecords=uploadHistoryRecords,user_type_val=str(current_user.user_type))

@app.route('/generateOTP',methods=['POST','GET'])
def generateOTP():
    print('inside generateOTP')
    phone = request.args.get('phone')
    print('phone',phone)
    message = 'new testing message'
    send_sms(phone,message)
    return jsonify([0])

@app.route('/uploadHistoryDetail',methods=['POST','GET'])
def uploadHistoryDetail():
    upload_id=request.args.get('upload_id')
    resultDetailQuery = "select distinct sp.full_name, sp.profile_picture, ru.total_marks, ru.marks_scored as marks_scored, md.description as test_type, ru.exam_date,cs.class_val, cs.section, ru.question_paper_ref "
    resultDetailQuery = resultDetailQuery + "from result_upload ru inner join student_profile sp on sp.student_id=ru.student_id "
    resultDetailQuery = resultDetailQuery + "inner join message_detail md on md.msg_id=ru.test_type "
    resultDetailQuery = resultDetailQuery + "and ru.upload_id='"+ str(upload_id) +"' inner join class_section cs on cs.class_sec_id=ru.class_sec_id order by marks_scored desc" 
    resultUploadRows = db.session.execute(text(resultDetailQuery)).fetchall()
    # paper = TestDetails.query.filter_by(test).first()
    runcount=0
    class_val_record = ""    
    section_record=""
    test_type_record=""
    exam_date_record=""
    question_paper_ref = ""
    for value in resultUploadRows:
        if runcount==0:        
            class_val_record = value.class_val
            section_record = value.section
            test_type_record=value.test_type
            exam_date_record= value.exam_date
            question_paper_ref = value.question_paper_ref
        runcount+1
    print('Upload Id:'+str(upload_id))
    print('Question Paper link:'+str(question_paper_ref))
    return render_template('_uploadHistoryDetail.html',resultUploadRows=resultUploadRows, class_val_record=class_val_record,section_record=section_record, test_type_record=test_type_record,exam_date_record=exam_date_record,question_paper_ref=question_paper_ref)

@app.route('/studentList/<class_val>/<section>/')
def studentList(class_val,section):
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    classSecRow = ClassSection.query.filter_by(class_val=class_val, section=section,school_id=teacher_id.school_id).first()
    students = StudentProfile.query.distinct().filter_by(class_sec_id=classSecRow.class_sec_id).all()
    studentArray = []

    

    for student in students:
        print('Student Name:'+str(student.full_name))
        studentObj = {}
        studentObj['student_id'] = student.student_id
        studentObj['student_name'] = student.full_name
        studentObj['class_value'] = classSecRow.class_val
        studentArray.append(studentObj)

    print(str(studentArray))
    return jsonify({'students' : studentArray})

@app.route('/quesFileUpload',methods=['POST','GET'])
def quesFileUpload():
    print('inside quesFileUpload')
    data = request.get_json()
    fileData = data[0]
    print('print File Content')
    print(fileData)
    imgData = data[1]
    print('print image list')
    print(imgData)
    class_val = request.args.get('classValue')
    subject_name = request.args.get('subjectValue')
    # print(subject_name)
    # print(fileData)
    # print(imgData)
    # print(class_val)
    k=0
    # print(len(fileData))
    # print(fileData[0])
    # print(fileData[0][0])
    # print('after first row')
    # print(fileData[1][0])
    # print('after second row')
    for i in range(0,len(fileData)):
        print('inside for loop:'+str(k))
        print(fileData[k][0])
        if fileData[k][2]=='MCQ1':
            print("Inside MCQ")
            print('loop index:'+str(k))
            print(k)
            print(fileData[k][0])
            print(fileData[k][1])
            print(fileData[k][2])
            
            print(imgData[k])
            question=QuestionDetails(class_val=str(class_val),subject_id=str(subject_name),question_description=fileData[k][0],
            topic_id=fileData[k][1],question_type=fileData[k][2],reference_link=str(imgData[k]),archive_status=str('N'),suggested_weightage=fileData[k][3])
            db.session.add(question)
            db.session.commit()
            question_id=question.question_id
            for j in range(4,8):
                if fileData[k][8]==fileData[k][j]:
                    correct='Y'
                    weightage=fileData[k][3]
                else:
                    correct='N'
                    weightage='0'
                if j==4:
                    option_val='A'
                elif j==5:
                    option_val='B'
                elif j==6:
                    option_val='C'
                else:
                    option_val='D'
                options=QuestionOptions(option_desc=fileData[k][j],question_id=question_id,is_correct=correct,option=option_val,weightage=int(fileData[k][3]))
                db.session.add(options)
                db.session.commit()
        k=k+1
    return jsonify(['1'])

@app.route('/createQuestion',methods=['POST','GET'])
def createQuestion():
    weightage = request.args.get('weightage')
    questionDesc = request.args.get('questionDesc')
    questionTypeValue = request.args.get('questionTypeValue')
    topicValue = request.args.get('topicValue')
    subjectValue = request.args.get('subjectValue')
    classValue = request.args.get('classValue')
    reference = request.args.get('reference')
    correct = request.args.get('correct')
    options = request.get_json()
    for op in options:
        print('Options:'+str(op))
    print('correct:'+str(correct))
    print('weightage:'+str(weightage))
    print('questionDesc:'+str(questionDesc))
    print('questionTypeValue:'+str(questionTypeValue))
    print('topicValue:'+str(topicValue))
    print('subjectValue:'+str(subjectValue))
    print('classValue:'+str(classValue))
    print('reference:'+str(reference))
    question=QuestionDetails(class_val=classValue,subject_id=str(subjectValue),question_description=str(questionDesc),
        reference_link=str(reference),topic_id=str(topicValue),question_type=str(questionTypeValue),suggested_weightage=str(weightage),archive_status=str('N'))
    print(question)
    db.session.add(question)
    if questionTypeValue=='Subjective':
        print('question is subjetive')
        question_id=db.session.query(QuestionDetails).filter_by(class_val=classValue,topic_id=str(topicValue),question_description=str(questionDesc)).first()
        options=QuestionOptions(question_id=question_id.question_id,weightage=str(weightage))
        print('Options Desc:'+str(options))
        db.session.add(options)
        db.session.commit()
    else:
        print('question is MCQ')
        question_id=db.session.query(QuestionDetails).filter_by(class_val=classValue,topic_id=str(topicValue),question_description=str(questionDesc)).first()
        i=1
        answer = ''
        weigh = ''
        for op in options:
            if int(correct)==int(i):
                answer='Y'
                weigh=str(weightage)
            else:
                weigh=0
                answer='N'
            if i==1:
                opt='A'
            elif i==2:
                opt='B'
            elif i==3:
                opt='C'
            else:
                opt='D'
            i=i+1   
            options=QuestionOptions(option_desc=op,question_id=question.question_id,is_correct=answer,weightage=weigh,option=opt)
            print("Options in question Builder:"+str(options))
            db.session.add(options)
        db.session.commit()
    return jsonify(['1'])

@app.route('/questionBuilder',methods=['POST','GET'])
@login_required
def questionBuilder():
    form=QuestionBuilderQueryForm()
    print("Inside Question Builder")
    if request.method=='POST':
        if form.submit.data:
            print('Question Image:'+str(request.form['reference']))
            question=QuestionDetails(class_val=str(request.form['class_val']),subject_id=int(request.form['subject_name']),question_description=request.form['question_desc'],
            reference_link=request.form['reference'],topic_id=int(request.form['topics']),question_type=form.question_type.data,suggested_weightage=int(request.form['weightage']),archive_status=str('N'))
            print(question)
            db.session.add(question)
            if form.question_type.data=='Subjective':
                question_id=db.session.query(QuestionDetails).filter_by(class_val=str(request.form['class_val']),topic_id=int(request.form['topics']),question_description=request.form['question_desc']).first()
                options=QuestionOptions(question_id=question_id.question_id,weightage=request.form['weightage'])
                print('Options Desc:'+str(options))
                db.session.add(options)
                db.session.commit()
                flash('Success')
                return render_template('questionBuilder.html',user_type_val=str(current_user.user_type))
            else:
                option_list=request.form.getlist('option_desc')
                question_id=db.session.query(QuestionDetails).filter_by(class_val=str(request.form['class_val']),topic_id=int(request.form['topics']),question_description=request.form['question_desc']).first()
                if request.form['correct']=='':
                    flash('Correct option not seleted !')
                    return render_template('questionBuilder.html',user_type_val=str(current_user.user_type))
                for i in range(len(option_list)):
                    print('option value:'+str(request.form['option']))
                    if int(request.form['option'])==i+1:
                        correct='Y'
                        weightage=int(request.form['weightage'])
                    else:
                        weightage=0
                        correct='N'
                    if i+1==1:
                        option='A'
                    elif i+1==2:
                        option='B'
                    elif i+1==3:
                        option='C'
                    else:
                        option='D'
                        
                    options=QuestionOptions(option_desc=option_list[i],question_id=question.question_id,is_correct=correct,weightage=weightage,option=option)
                    print("Options in question Builder:"+str(options))
                    db.session.add(options)
                db.session.commit()
                flash('Success')
                return render_template('questionBuilder.html',user_type_val=str(current_user.user_type))
        else:
            csv_file=request.files['file-input']
            print('if question bulk upload')
            print(csv_file)
            print(type(csv_file))
            df1=pd.read_csv(csv_file)
            for index ,row in df1.iterrows():
                print('index:'+str(index))
                print('Ques Desc:'+str(row['Question Description']))
                if str(row['Question Description'])!="nan" or str(row['Question Description'])!='nan':
                    print('Ques Desc'+str(row['Question Description']))
                    if index<20:
                        if row['Question Type']=='MCQ1':
                            print("Inside MCQ")
                            print('Image Url:'+str(request.form['reference-url'+str(index+1)]))
                            question=QuestionDetails(class_val=str(request.form['class_val']),subject_id=int(request.form['subject_name']),question_description=row['Question Description'],
                            topic_id=row['Topic Id'],question_type='MCQ1',reference_link=request.form['reference-url'+str(index+1)],archive_status=str('N'),suggested_weightage=row['Suggested Weightage'])
                            db.session.add(question)
                            question_id=db.session.query(QuestionDetails).filter_by(class_val=str(request.form['class_val']),topic_id=row['Topic Id'],question_description=row['Question Description']).first()
                            for i in range(1,5):
                                option_no=str(i)
                                option_name='Option'+option_no

                                print(row[option_name])
                                if row['CorrectAnswer']==row['Option'+option_no]:
                                    correct='Y'
                                    weightage=row['Suggested Weightage']
                                else:
                                    correct='N'
                                    weightage='0'
                                if i==1:
                                        option_val='A'
                                elif i==2:
                                        option_val='B'
                                elif i==3:
                                        option_val='C'
                                else:
                                    option_val='D'
                                print(row[option_name])
                                print(question_id.question_id)
                                print(correct)
                                print(option_val)
                                options=QuestionOptions(option_desc=row[option_name],question_id=question.question_id,is_correct=correct,option=option_val,weightage=int(weightage))
                                print(options)
                                db.session.add(options)
                                db.session.commit()
                        else:
                            print("Inside Subjective")
                            question=QuestionDetails(class_val=str(request.form['class_val']),subject_id=int(request.form['subject_name']),question_description=row['Question Description'],
                            topic_id=row['Topic Id'],question_type='Subjective',reference_link=request.form['reference-url'+str(index+1)],archive_status=str('N'),suggested_weightage=row['Suggested Weightage'])
                            db.session.add(question)
            db.session.commit()
            flash('Successfully Uploaded !')
            return render_template('questionBuilder.html',user_type_val=str(current_user.user_type))
    return render_template('questionBuilder.html',user_type_val=str(current_user.user_type))



@app.route('/questionUpload',methods=['GET'])
def questionUpload():
    flag = False
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form=QuestionBuilderQueryForm()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
    form.subject_name.choices= ''
    # [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
    form.topics.choices= ''
    # [(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]
    if request.method=='POST':
        print('Inside if question Upload')
        topic_list=Topic.query.filter_by(class_val=str(form.class_val.data),subject_id=str(form.subject_name.data),chapter_num=str(form.chapter_num.data)).all()
        return render_template('questionUpload.html',form=form, flag=flag,topic_list=topic_list)
    else:
        return render_template('questionUpload.html',form=form,flag=flag)

@app.route('/uploadMarks',methods=['GET'])
def uploadMarks():
    user = User.query.filter_by(username=current_user.username).first_or_404()    
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first()     
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    board_id=SchoolProfile.query.with_entities(SchoolProfile.board_id).filter_by(school_id=teacher_id.school_id).first()
    classValue = request.args.get('class_val')
    class_section = request.args.get('class_section')
    class_sec_id=ClassSection.query.filter_by(class_val=str(classValue),school_id=teacher_id.school_id,section=class_section).first()
    student_list=StudentProfile.query.filter_by(class_sec_id=class_sec_id.class_sec_id,school_id=teacher_id.school_id).all()
    paperUrl = request.args.get('paperUrl')
    subject_id = request.args.get('subject_id')
    marks = request.args.get('marks')
    testdate = request.args.get('testdate')
    Tmarks = request.args.get('Tmarks')
    testId = request.args.get('testId')
    test_type = request.args.get('test_type')
    marks = marks.split(",")
    count = 0
    now = datetime.now()
 
    print("now =", now)

    dt_string = now.strftime("%d/%m/%Y/%H:%M:%S")
    num = 1
    list_id = []
    for student_id in student_list:
        list_id.append(student_id.student_id)
    paper = ''
    print('Coming Paper Url :'+str(paperUrl))
    if testId!='':
        paper = TestDetails.query.filter_by(test_id=testId).first()
        print('Paper:'+str(paper.test_paper_link))
    else:
        testType = MessageDetails.query.filter_by(msg_id=test_type).first()
        testPaper = TestDetails(test_type=testType.description,total_marks=Tmarks,
        last_modified_date=datetime.today(),board_id=board_id.board_id,subject_id=subject_id,class_val=str(classValue),date_of_creation=datetime.today(),school_id=teacher_id.school_id,teacher_id=teacher_id.teacher_id,test_paper_link=paperUrl)
        db.session.add(testPaper)
    for marksSubjectWise in marks:
        if marksSubjectWise=='-1':
            marksSubjectWise=0
            is_present=MessageDetails.query.filter_by(description='Not Present').first()
        else:
            is_present=MessageDetails.query.filter_by(description='Present').first()
        #print('Marks:'+marksSubjectWise)
        upload_id=str(teacher_id.school_id)+str(class_sec_id.class_sec_id)+str(subject_id) + str(test_type) + dt_string
        upload_id=upload_id.replace('-','')
        print('Test Id:'+testId)
        
        if testId=='':
            print('If test id is null')
            print('Test Date:'+str(testdate))
            print('upload Date:'+str(datetime.today()))
            Marks=ResultUpload(school_id=teacher_id.school_id,student_id=list_id[count],
            exam_date=testdate,marks_scored=marksSubjectWise,class_sec_id=class_sec_id.class_sec_id,
            test_type=test_type,subject_id=subject_id,is_present=is_present.msg_id,total_marks=Tmarks,
            uploaded_by=teacher_id.teacher_id, upload_id=upload_id,last_modified_date=datetime.today(),question_paper_ref=paperUrl
            )
        else:
            print('If test id not null')
            print('Test Date:'+str(testdate))
            print('upload Date:'+str(datetime.today()))
            
            Marks=ResultUpload(school_id=teacher_id.school_id,student_id=list_id[count],
            exam_date=testdate,marks_scored=marksSubjectWise,class_sec_id=class_sec_id.class_sec_id,
            test_type=test_type,subject_id=subject_id,is_present=is_present.msg_id,total_marks=Tmarks,test_id=testId,
            uploaded_by=teacher_id.teacher_id, upload_id=upload_id,last_modified_date=datetime.today(),question_paper_ref=str(paper.test_paper_link)
            )
        
        db.session.add(Marks)
        count = count + 1
    db.session.execute(text('call sp_performance_detail_load()'))
    db.session.commit()
    print('after call performance_detail_load function()')
    flash('Login required !')
    print('Class_val:'+str(classValue)+'subject_id:'+str(subject_id)+'classSection:'+class_section+"testdate:"+testdate+"Total marks:"+Tmarks+"TestId:"+testId+"Test type:"+test_type)
    

    flash('Marks successfully uploaded')
    return render_template('resultUpload.html')


@app.route('/questionFile',methods=['GET']) 
def questionFile():
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form=QuestionBuilderQueryForm()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()]
    form.subject_name.choices= ''
    # [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
    # form.chapter_num.choices= ''
    form.topics.choices= ''
    # [(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]
    return render_template('questionFile.html',form=form)

@app.route('/addChapterTopics',methods=['GET','POST'])
def addChapterTopics():
    class_val = request.args.get('class_val')
    subject_id = request.args.get('subject_id')
    if current_user.is_anonymous:
        studentData = StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()
        school_id = studentData.school_id
    else:
        if current_user.user_type==134 or current_user.user_type==234:
            studentData = StudentProfile.query.filter_by(user_id=current_user.id).first()
            school_id = studentData.school_id            
        else:
            teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
            school_id = teacherData.school_id

    query = "select distinct bd.book_name ,topic_name, chapter_name, td.topic_id, td.chapter_num from topic_tracker tt "
    query = query + "inner join topic_detail td on td.topic_id = tt.topic_id "
    query = query + "inner join book_details bd on td.book_id = bd.book_id "
    query = query + "where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id)+"' and tt.school_id='"+str(school_id)+"' and tt.is_archived='N' order by td.chapter_num "
    print('Query:'+str(query))
    chapters = db.session.execute(text(query)).fetchall()
    chaptersArray = []
    i=1
    book = ''
    ch = ''
    print(chapters)
    for chapter in chapters:
        if len(chapters)>1:
            book = chapter.book_name
            ch = chapter.topic_id
            if i==1:
                ch = str(ch)+"/"
            elif i==len(chapters):
                book = "/"+str(book)
            else:
                book = "/"+str(book)
                ch = str(ch)+"/"
            i=i+1
        else:
            book = chapter.book_name
            ch = chapter.topic_id
        chaptersArray.append(str(book)+"@"+str(chapter.topic_name)+"@"+str(chapter.chapter_name)+"@"+str(ch))
    
    if chaptersArray:
        return jsonify([chaptersArray])
    else:
        return ""


@app.route('/addClass',methods=['GET','POST'])
def addClass():
    class_val = request.args.get('class_val')
    print('inside add class')
    ##########
    if current_user.is_anonymous:
        studentData = StudentProfile.query.filter_by(user_id=app.config['ANONYMOUS_USERID']).first()
        school_id = studentData.school_id
    else:
        if current_user.user_type==134 or current_user.user_type==234:
            studentData = StudentProfile.query.filter_by(user_id=current_user.id).first()
            school_id = studentData.school_id            
        else:
            teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
            school_id = teacherData.school_id
    ##########
    #teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
    #board_id = SchoolProfile.query.filter_by(school_id=school_id).first()
    subjects = BoardClassSubject.query.filter_by(class_val=str(class_val),school_id=school_id).all()
    subjectArray = []
    print(subjects)
    for subject in subjects:
        subject_name = "select description as subject_name from message_detail where msg_id='"+str(subject.subject_id)+"'"
        subject_name = db.session.execute(text(subject_name)).first()
        print(str(subject.subject_id)+":"+str(subject_name.subject_name))
        subjectArray.append(str(subject.subject_id)+":"+str(subject_name.subject_name))

    if subjectArray:
        return jsonify([subjectArray])
    else:
        return ""

@app.route('/topperList')
def topperList():
    classValue = request.args.get('class_val')
    # subject_id = request.args.get('subject_id')
    #test_type = request.args.get('test_type')
    # section_val = request.args.get('section_val')
    #test_date = request.args.get('test_date')
    #print('Class Value:'+classValue+"Subject Value:"+subject_id+"Test Type:"+test_type+"Section value:"+section_val+"Test date:"+test_date)
    user = User.query.filter_by(username=current_user.username).first_or_404()
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first() 
    query = "select  * from fn_performance_leaderboard_detail('"+ str(teacher.school_id) +"') where class='"+classValue+"' order by marks desc"
    #print('Query topperList:'+query)
    leaderBoardData = db.session.execute(text(query)).fetchall()
    return render_template('_leaderBoardTable.html',leaderBoardData=leaderBoardData)

@app.route('/topperListForAll')
def topperListBySubject():
    classValue = request.args.get('class_val')
    subjectValue = request.args.get('subject_id')
    test_type = request.args.get('test_type')
    section_val = request.args.get('section_val')
    user = User.query.filter_by(username=current_user.username).first_or_404()
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first() 
    subjectName = ''
    print('Subject id:'+subjectValue)
    if subjectValue!='All':
        subjectName = MessageDetails.query.filter_by(msg_id=subjectValue).first()
    query = "select *from public.fn_performance_leaderboard('"+ str(teacher.school_id) +"') where "
    print('Subject Value:'+subjectValue)
    if subjectValue=='All':
        sub = 'All'
    else:
        sub = subjectName.description
    if classValue!='':
        query = query + "class='"+classValue+"' and subjects='"+sub+"' and section='"+section_val+"' order by marks desc"
    else:
        query = query + "subjects='"+sub+"' and section='"+section_val+"' order by marks desc"
    print('Query topperList:'+query)
    leaderBoardData = db.session.execute(text(query)).fetchall()
    return render_template('_leaderBoardTable.html',leaderBoardData=leaderBoardData)

@app.route('/questionBuilder/<class_val>')
def subject_list(class_val):
    if class_val!='All':
        user_type_val = current_user.user_type
        teacher_id = ''
        if user_type_val==134:
            teacher_id = StudentProfile.query.filter_by(user_id=current_user.id).first()
        else:
            teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
        cl = class_val.replace("-","/")
        print('New Class value:'+str(cl))
        board_id=SchoolProfile.query.with_entities(SchoolProfile.board_id).filter_by(school_id=teacher_id.school_id).first()
        print(cl)
        print(board_id)
        print(teacher_id.school_id)
        subject_id=BoardClassSubject.query.with_entities(BoardClassSubject.subject_id).distinct().filter_by(class_val=str(cl),board_id=board_id.board_id,school_id=teacher_id.school_id).all()
        subject_name_list=[]
        print(subject_id)
        for id in subject_id:

            subject_name=MessageDetails.query.filter_by(msg_id=id).first()
            if subject_name in subject_name_list:
                continue
            subject_name_list.append(subject_name)
        subjectArray = []
        print('Subject Array:')
        print(subjectArray)
        for subject in subject_name_list:
            subjectObj = {}
            subjectObj['subject_id'] = subject.msg_id
            subjectObj['subject_name'] = subject.description
            subjectArray.append(subjectObj)

        return jsonify({'subjects' : subjectArray})
    else:
        return "return"

# topic list generation dynamically
# @app.route('/questionBuilder/<class_val>/<subject_id>')
# def topic_list(class_val,subject_id):
#     cl = class_val.replace("-","/")
#     topic_list="select td.topic_id,td.topic_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val = '"+str(cl)+"' and td.subject_id = '"+str(subject_id)+"' and tt.is_archived = 'N'"
#     topic_list = db.session.execute(text(topic_list))
#     topicArray=[]

#     for topic in topic_list:
#         topicObj={}
#         topicObj['topic_id']=topic.topic_id
#         topicObj['topic_name']=topic.topic_name
#         topicArray.append(topicObj)
    
#     return jsonify({'topics':topicArray})

# @app.route('/questionTopicPicker')
# def questionTopicPicker():
#     print('Inside topic picker')
#     class_val = request.args.get('class_val')
#     subject_id = request.args.get('subject_id')
#     topic_list="select td.topic_id,td.chapter_num,td.topic_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val = '"+str(class_val)+"' and td.subject_id = '"+str(subject_id)+"' and tt.is_archived = 'N'"
#     topic_list = db.session.execute(text(topic_list)).fetchall()
#     for topic in topic_list:
#         print(topic.topic_id)
#         print(topic.topic_name)
#         print(topic.chapter_num)
    
#     return render_template('_topics.html',topic_list=topic_list)

# # @app.route('/coveredTopic',methods=['GET','POST'])
# # def coveredTopic():
# #     print('covered Topics')
# #     class_v = request.args.get('class_val')
# #     section = request.args.get('section')
# #     query = "select *from topic_details where "
# #     return render_template('_topicCovered.html')

# @app.route('/questionChapterpicker/<class_val>/<subject_id>')
# def chapter_list(class_val,subject_id):
#     cl = class_val.replace('-','/')
#     chapter_num = "select distinct td.chapter_num,td.chapter_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val='"+cl+"' and td.subject_id='"+subject_id+"' and tt.is_archived='N' order by td.chapter_num,td.chapter_name"
#     print(chapter_num)
#     print('Inside chapterPicker')
    
#     chapter_num_list = db.session.execute(text(chapter_num))
#     chapter_num_array=[]
#     for chapterno in chapter_num_list:
#         chapterNo = {}
#         chapterNo['chapter_num']=chapterno.chapter_num
#         chapterNo['chapter_name']=chapterno.chapter_name
#         chapter_num_array.append(chapterNo)
#     return jsonify({'chapterNum':chapter_num_array})

# @app.route('/addEvent', methods = ["GET","POST"])
# @login_required
# def addEvent():        
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     form = addEventForm()
#     if form.validate_on_submit():
#         dataForEntry = EventDetail(event_name=form.eventName.data, event_duration=form.duration.data,event_date=form.eventDate.data,event_start=form.startDate.data,event_end=form.endDate.data,event_category=form.category.data,school_id=teacher_id.school_id, last_modified_date=datetime.today())                
#         db.session.add(dataForEntry)
#         db.session.commit()
#         flash('Event Added!')
#     indic='DashBoard'
#     return render_template('addEvent.html',indic=indic, form=form,title='Add Event')

# @app.route('/studentProfileOld')
# @login_required
# def studentProfileOld():
#     return render_template('studentProfile.html')


# @app.route('/allocateStudentToSponsor')
# def allocateStudentToSponsor():
#     student_id = request.args.get('student_id')
#     sponsor_id = request.args.get('sponsor_id')
#     sponsor_name = request.args.get('sponsor_name')
#     amount = request.args.get('amount')

#     studentData = StudentProfile.query.filter_by(student_id=student_id).first()
#     studentData.sponsor_id = sponsor_id
#     studentData.sponsor_name = sponsor_name
#     studentData.sponsored_amount = amount
#     studentData.sponsored_on = datetime.today()
#     studentData.sponsored_status='Y'
#     studentData.last_modified_date = datetime.today()
#     db.session.commit()
    
#     return jsonify(['0'])

# @app.route('/indivStudentProfile')
# @login_required
# def indivStudentProfile():    
#     student_id=request.args.get('student_id')
#     flag = request.args.get('flag')
#     #spnsor data check
#     sponsor_id = request.args.get('sponsor_id')
#     sponsor_name = request.args.get('sponsor_name')
#     amount = request.args.get('amount')
    
#     #New updated query
#     studentProfileQuery = "select full_name, email, sponsored_status,sponsored_on,sponsor_name,phone,sp.school_id as school_id, dob, md.description as gender,class_val, section, "
#     studentProfileQuery = studentProfileQuery + "roll_number,school_adm_number,profile_picture,student_id from student_profile sp inner join "
#     studentProfileQuery = studentProfileQuery + "class_section cs on sp.class_sec_id= cs.class_sec_id and sp.student_id='"+str(student_id)+"'" 
#     studentProfileQuery = studentProfileQuery + "inner join message_detail md on md.msg_id =sp.gender "
#     studentProfileQuery = studentProfileQuery + "left join address_detail ad on ad.address_id=sp.address_id"

#     studentProfileRow = db.session.execute(text(studentProfileQuery)).first()  
    
#     #performanceData
#     performanceQuery = "SELECT * from vw_leaderboard WHERE student_id = '"+str(student_id)+ "'"    

#     perfRows = db.session.execute(text(performanceQuery)).fetchall()
    
    # testCountQuery = "select count(distinct resp_session_id) as testcountval from response_capture where student_id='"+str(student_id)+ "'"

#     testCount = db.session.execute(text(testCountQuery)).first() 

#     totalofflineTestQuery = "select  count(*) as count from result_upload ru where student_id  = '"+str(student_id)+ "'"

#     totalofflineTestCount = db.session.execute(text(totalofflineTestQuery)).first()

#     totalTestCount = int(testCount.testcountval) + int(totalofflineTestCount.count)
#     print('totalTestCount:'+str(totalTestCount))

#     testResultQuery = "select exam_date, t2.description as test_type, test_id, t3.description as subject, marks_scored, total_marks "
#     testResultQuery = testResultQuery+ "from result_upload t1 inner join message_detail t2 on t1.test_type=t2.msg_id "
#     testResultQuery = testResultQuery + "inner join message_detail t3 on t3.msg_id=t1.subject_id "
#     testResultQuery = testResultQuery + " where student_id=%s order by exam_date desc" % student_id
#     print(testResultQuery)
#     testResultRows = db.session.execute(text(testResultQuery)).fetchall()
    
#     onlineTestResultQuery = "select sd.last_modified_date,rc.resp_session_id,td.test_type,md.description , sum(marks_scored) as marks_scored,sd.total_marks from response_capture rc "
#     onlineTestResultQuery = onlineTestResultQuery + "inner join session_detail sd on rc.resp_session_id=sd.resp_session_id "
#     onlineTestResultQuery = onlineTestResultQuery + "inner join test_details td on sd.test_id = td.test_id "
#     onlineTestResultQuery = onlineTestResultQuery + "inner join message_detail md on md.msg_id = rc.subject_id where rc.student_id='"+str(student_id)+"' "
#     onlineTestResultQuery = onlineTestResultQuery + "group by rc.resp_session_id,sd.last_modified_date,sd.total_marks,td.test_type,md.description order by sd.last_modified_date desc "
#     onlineTestResultRows = db.session.execute(text(onlineTestResultQuery)).fetchall()
#     #Remarks info
#     studentRemarksQuery = "select student_id, tp.teacher_id, teacher_name, profile_picture, remark_desc, sr.last_modified_date as last_modified_date"
#     studentRemarksQuery= studentRemarksQuery+ " from student_remarks sr inner join teacher_profile tp on sr.teacher_id=tp.teacher_id and student_id="+str(student_id) + " "
#     studentRemarkRows = db.session.execute(text(studentRemarksQuery)).fetchall()
#     #studentRemarkRows = StudentRemarks.query.filter_by(student_id=student_id).order_by(StudentRemarks.last_modified_date.desc()).limit(5).all()

#     #Sponsor allocation
#     urlForAllocationComplete = str(app.config['IMPACT_HOST']) + '/responseStudentAllocate'
#     overallSum = 0
#     overallPerfValue = 0
#     sumMarks = 0
#     sum1 = 0
#     sum2 = 0
#     totalOfflineTestMarks = "select sum(marks_scored) as sum1 from result_upload ru where student_id = '"+str(student_id)+"'"
#     print(totalOfflineTestMarks)
#     totalOfflineTestMarks = db.session.execute(text(totalOfflineTestMarks)).first()
#     if totalOfflineTestMarks.sum1:
#         print(totalOfflineTestMarks.sum1)
#         sum1 = totalOfflineTestMarks.sum1
#     totalOnlineTestMarks = "select sum(rc2.marks_scored) as sum2,rc2.resp_session_id from response_capture rc2 where student_id = '"+str(student_id)+"' group by rc2.resp_session_id "
#     totalOnlineTestMarks = db.session.execute(text(totalOnlineTestMarks)).fetchall()
    
#     for row in totalOnlineTestMarks:
#         print(row.sum2)
#         sum2 = sum2 + row.sum2
#     sumMarks = int(sum1) + int(sum2)
#     print('Total Marks:'+str(sumMarks))
#     total1 = "select total_marks as offlineTotal from result_upload ru where student_id = '"+str(student_id)+"'"
#     print(total1)
#     total1 = db.session.execute(text(total1)).first()
#     tot1 = 0
#     if total1:
#         print(total1.offlinetotal)
#         tot1 = total1.offlinetotal
#     total2 = "select sd.total_marks,sd.resp_session_id from session_detail sd inner join response_capture rc on sd.resp_session_id = rc.resp_session_id where rc.student_id = '"+str(student_id)+"' group by sd.total_marks,sd.resp_session_id"
#     total2 = db.session.execute(text(total2)).fetchall()
#     total3 = 0
#     grandTotal = 0
#     for row in total2:
#         print(row.total_marks)
#         total3 = total3 + row.total_marks
#     grandTotal = int(tot1) + int(total3)
#     print(tot1)
#     print(total3)
#     print('Grand Total:'+str(grandTotal))
#     print('Sum Marks:'+str(sumMarks))
#     for rows in perfRows:
#         overallSum = overallSum + int(rows.student_score)
#         #print(overallSum)
#     try:
#         overallPerfValue = round(sumMarks/(grandTotal)*100,2)    
#     except:
#         overallPerfValue=0    
#     guardianRows = GuardianProfile.query.filter_by(student_id=student_id).all()
#     qrRows = studentQROptions.query.filter_by(student_id=student_id).all()
#     qrAPIURL = "https://api.qrserver.com/v1/create-qr-code/?size=150x150&data="    
#     qrArray=[]
#     x = range(4)    

#     #section for fetching surveys
#     try:
#         surveyRows = SurveyDetail.query.filter_by(school_id=studentProfileRow.school_id,is_archived='N').all()
#     except:
#         surveyRows=[]
#         print('survey error')

#     for n in x:               
#         if studentProfileRow!=None and studentProfileRow!="":
#             optionURL = qrAPIURL+str(student_id)+ '-'+str(studentProfileRow.roll_number)+'-'+ studentProfileRow.full_name.replace(" ", "%20")+'@'+string.ascii_uppercase[n]
#         else:
#             optionURL=""
#         qrArray.append(optionURL) 
#         #print(optionURL)
#     return render_template('_indivStudentProfile.html',surveyRows=surveyRows, studentRemarkRows=studentRemarkRows, urlForAllocationComplete=urlForAllocationComplete, studentProfileRow=studentProfileRow,guardianRows=guardianRows, 
#         qrArray=qrArray,totalTestCount=totalTestCount,perfRows=perfRows,overallPerfValue=overallPerfValue,student_id=student_id,testCount=testCount,
#         testResultRows = testResultRows,onlineTestResultRows=onlineTestResultRows,disconn=1, sponsor_name=sponsor_name, sponsor_id=sponsor_id,amount=amount,flag=flag)


# @app.route('/attendanceReport',methods=["GET","POST"])
# def attendanceReport():
#     class_sec_id = request.args.get('class_sec_id')
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
#     attendanceData = "select full_name, is_present from attendance a inner join student_profile sp on a.class_sec_id = sp.class_sec_id where "
#     attendanceData = attendanceData + "sp.school_id = '"+str(teacherDataRow.school_id)+"' and sp.class_sec_id = '"+str(class_sec_id)+"'"
#     attendanceData = db.session.execute(text(attendanceData)).fetchall()
#     file_name='Attendance'+str(class_sec_id)+str(teacherDataRow.school_id)+'.csv'
#     file_name = file_name.replace(" ", "")
#     if not os.path.exists('tempdocx'):
#         os.mkdir('tempdocx')
#     # document.save('tempdocx/'+file_name)
#     with open('tempdocx/'+file_name, 'w', newline='') as file:
#     #uploading to s3 bucket
        
#         # with open(filePath, 'w', newline='') as file:
            
#         print('file')
            
#         file.write('Introduction \n')
#         writer = csv.writer(file)
#         writer.writerow(["Student Name ", "Present"])
#         for data in attendanceData:
#             writer.writerow([data.full_name,data.is_present])
#     client = boto3.client('s3', region_name='ap-south-1')
#     client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'attendance_report/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
#         #deleting file from temporary location after upload to s3
#     os.remove('tempdocx/'+file_name)
#     filePath = 'https://'+str(os.environ.get('S3_BUCKET_NAME'))+'.s3.ap-south-1.amazonaws.com/attendance_report/'+str(file_name)
#     print(filePath)
#     return jsonify([filePath])


# @app.route('/addAttendence',methods=["GET","POST"])
# def addAttendence():
#     class_sec_id = request.args.get('class_sec_id')
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
#     students = StudentProfile.query.filter_by(class_sec_id=class_sec_id,school_id=teacherDataRow.school_id).all()
#     print('class_Sec_id:'+str(class_sec_id))
#     print('school_id'+str(teacherDataRow.school_id))
#     student_ids = request.get_json()
#     print(students)
#     print(student_ids)
#     data = Attendance.query.filter_by(class_sec_id=class_sec_id,school_id=teacherDataRow.school_id).first()
#     if data:
#         for student_id in students:
#             sel = 0
            
#             for selected in student_ids:
#                 if str(student_id.student_id) == str(selected):
#                     print('if student id is selected:'+str(student_id.student_id))
#                     sel =1
#                     break
#             if sel==1:
#                 print('set Y for selected ')
#                 update = Attendance.query.filter_by(class_sec_id=class_sec_id,school_id=teacherDataRow.school_id,student_id=student_id.student_id).first()
#                 update.is_present = 'Y'
#             else:
#                 update = Attendance.query.filter_by(class_sec_id=class_sec_id,school_id=teacherDataRow.school_id,student_id=student_id.student_id).first()
#                 update.is_present = 'N'
#             db.session.commit()    
#     else:
#         for student_id in students:
#             sel = 0
            
#             for selected in student_ids:
#                 if str(student_id.student_id) == str(selected):
#                     print('if student id is selected:'+str(student_id.student_id))
#                     sel =1
#                     break
#             if sel==1:
#                 print('set Y for selected ')
#                 addData = Attendance(school_id=teacherDataRow.school_id,teacher_id=teacherDataRow.teacher_id,attendance_date=datetime.today(),last_modified_date=datetime.today(),
#                 is_present='Y',class_sec_id=class_sec_id,student_id=student_id.student_id)
#             else:
#                 addData = Attendance(school_id=teacherDataRow.school_id,teacher_id=teacherDataRow.teacher_id,attendance_date=datetime.today(),last_modified_date=datetime.today(),
#                 is_present='N',class_sec_id=class_sec_id,student_id=student_id.student_id)
#             db.session.add(addData)
#             db.session.commit()
#     return jsonify(['1'])


# @app.route('/fetchAttendenceList',methods=["GET","POST"])
# def fetchAttendenceList():
#     class_sec_id = request.args.get('class_sec_id')
#     date = request.args.get('date')
#     print('Date:'+str(date))
#     teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first() 
#     if date:
#         fetchData = "select sd.student_id,full_name,sd.profile_picture,a.is_present "
#         fetchData = fetchData + "from attendance a right join student_profile sd on a.student_id=sd.student_id "
#         fetchData = fetchData + "where sd.class_sec_id="+str(class_sec_id)+" and sd.school_id='"+str(teacherDataRow.school_id)+"' and date(attendance_date)='"+str(date)+"' order by full_name "
#     else:
#         fetchData = "select sd.student_id,full_name,sd.profile_picture,a.is_present "
#         fetchData = fetchData + "from attendance a right join student_profile sd on a.student_id=sd.student_id "
#         fetchData = fetchData + "where sd.class_sec_id="+str(class_sec_id)+" and sd.school_id='"+str(teacherDataRow.school_id)+"' order by full_name "
#     print('Query:'+str(fetchData))
#     studAttendeceList = db.session.execute(fetchData).fetchall()
#     print('class_sec_id:'+str(class_sec_id))
#     for row in studAttendeceList:
#         print('Fetch Data:'+str(row.is_present))
#     return render_template('_attendanceTable.html',studAttendeceList=studAttendeceList)

@app.route('/addStudentRemarks',methods = ["GET","POST"])
def addStudentRemarks():
    remark_desc=request.form.get('remark')
    student_id = request.form.get('student_id')
    teacherDataRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()        
    try:
        studentRemarkAdd = StudentRemarks(student_id=student_id, remark_desc=remark_desc, teacher_id=teacherDataRow.teacher_id,
            is_archived = 'N', last_modified_date = datetime.today())
        db.session.add(studentRemarkAdd)
        db.session.commit()
        return jsonify(['0'])
    except:
        return jsonify(['1'])

def classChecker(available_class):
    class_list = []
    for k in available_class:
        if k.class_val == '99':
            class_list.append((str('LKG')+"-"+str(k.section),str('LKG')+"-"+str(k.section)))
            print('inside available classes')
        elif k.class_val == '100':
            class_list.append((str('UKG')+"-"+str(k.section),str('UKG')+"-"+str(k.section)))
        else:
            class_list.append((str(k.class_val)+"-"+str(k.section),str(k.class_val)+"-"+str(k.section)))
    return class_list

# @app.route('/studentProfile') 
# @login_required
# def studentProfile():    
#     qstudent_id=request.args.get('student_id')
#     #####Section for sponsor data fetch
#     qsponsor_name = request.args.get('sponsor_name')
#     qsponsor_id = request.args.get('sponsor_id')
#     qamount = request.args.get('amount')
#     studentDetails = StudentProfile.query.filter_by(user_id = current_user.id).first()

#     if qstudent_id==None or qstudent_id=='':
#         form=studentDirectoryForm()
#         user = User.query.filter_by(username=current_user.username).first_or_404()        
#         teacher= TeacherProfile.query.filter_by(user_id=user.id).first()    

#         available_class=ClassSection.query.with_entities(ClassSection.class_val,ClassSection.section).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher.school_id).all()
#         available_section=ClassSection.query.with_entities(ClassSection.section).distinct().filter_by(school_id=teacher.school_id).all()    
#         available_test_type=MessageDetails.query.filter_by(category='Test type').all()
#         # available_student_list=StudentProfile.query.filter_by(school_id=teacher.school_id).all()
#         available_student_list = "select student_id,full_name,profile_picture,class_val, section from student_profile sp inner join class_section cs on sp.class_sec_id = cs.class_sec_id where cs.school_id ='"+str(teacher.school_id)+"'"
#         available_student_list = db.session.execute(available_student_list).fetchall()
        
#         class_list = classChecker(available_class)
        
#         section_list=[(i.section,i.section) for i in available_section]    
#         test_type_list=[(i.msg_id,i.description) for i in available_test_type]
#         # student_list=[(i.student_id,i.full_name) for i in available_student_list]

#         #selectfield choices
#         print(class_list)
#         form.class_section.choices = class_list
#         # form.section1.choices= ''
#         # section_list    
#         # form.test_type1.choices=test_type_list
#         form.student_name.choices = ''
#         flag = 1
#         indic='DashBoard'
#         return render_template('studentProfileNew.html',indic=indic,title='Student Profile',form=form, sponsor_name=qsponsor_name, sponsor_id = qsponsor_id, amount = qamount,available_student_list=available_student_list,flag=flag,user_type_val=str(current_user.user_type),studentDetails=studentDetails)
#     else:
#         value=0
#         flag = 0
#         if current_user.user_type==72:
#             value=1
#         #print(qstudent_id)
#         indic='DashBoard'
#         return render_template('studentProfileNew.html',indic=indic,title='Student Profile',qstudent_id=qstudent_id,disconn=value, sponsor_name=qsponsor_name, sponsor_id = qsponsor_id, amount = qamount,flag=flag,user_type_val=str(current_user.user_type),studentDetails=studentDetails)
#         flag = 0       
#         #print(qstudent_id)
#         indic='DashBoard'
#         return render_template('studentProfileNew.html',indic=indic,title='Student Profile',qstudent_id=qstudent_id,disconn=disconn, sponsor_name=qsponsor_name, sponsor_id = qsponsor_id, amount = qamount,flag=flag, user_type_val=str(current_user.user_type),studentDetails=studentDetails)

#Addition of new section to conduct student surveys
# @app.route('/studentSurveys')
# def studentSurveys():
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     surveyDetailQuery = "select sd.survey_id, survey_name, question_count, count(ssr.student_id ) as student_responses, question_count, sd.last_modified_date "
#     surveyDetailQuery = surveyDetailQuery+ "from survey_detail sd left join student_survey_response ssr on ssr.survey_id =sd.survey_id "
#     surveyDetailQuery = surveyDetailQuery+" where sd.school_id ="+str(teacherRow.school_id)+ " and sd.is_archived='N' group by sd.survey_id,survey_name, question_count,question_count, sd.last_modified_date"
#     surveyDetailRow = db.session.execute(surveyDetailQuery).fetchall()
#     #surveyDetailRow = SurveyDetail.query.filter_by(school_id=teacherRow.school_id).all()
#     indic='DashBoard'
#     return render_template('studentSurveys.html',indic=indic,title='Student Surveys', surveyDetailRow=surveyDetailRow,user_type_val=str(current_user.user_type))

# @app.route('/indivSurveyDetail/')
# def indivSurveyDetail():
#     survey_id = request.args.get('survey_id')
#     survey_id = survey_id.split('_')[1]
#     student_id = request.args.get('student_id')    
#     studSurveyData = " select sq.sq_id as sq_id, question,sq.survey_id,survey_response_id , sp.student_id, answer from student_survey_response ssr "
#     studSurveyData = studSurveyData  + " right join survey_questions sq on "
#     studSurveyData = studSurveyData  +  " sq.survey_id =ssr.survey_id and "
#     studSurveyData = studSurveyData  +  " sq.sq_id =ssr.sq_id and ssr.student_id ="+ str(student_id)
#     studSurveyData = studSurveyData  +  " left join student_profile sp "
#     studSurveyData = studSurveyData  +  " on sp.student_id =ssr.student_id "
#     studSurveyData = studSurveyData  +  " where sq.survey_id =" + str(survey_id)
#     surveyQuestions = db.session.execute(text(studSurveyData)).fetchall()
#     return render_template('_indivSurveyDetail.html',surveyQuestions=surveyQuestions,student_id=student_id,survey_id=survey_id)

# @app.route('/updateSurveyAnswer',methods=["GET","POST"])
# def updateSurveyAnswer():
#     sq_id_list = request.form.getlist('sq_id')
#     survey_response_id_list = request.form.getlist('survey_response_id')
#     answer_list = request.form.getlist('answer')
#     survey_id = request.form.get('survey_id')
#     student_id = request.form.get('student_id')
#     for i in range(len(sq_id_list)):
#         if survey_response_id_list[i]!='None':
#             studentSurveyAnsUpdate = StudentSurveyResponse.query.filter_by(sq_id=sq_id_list[i], survey_response_id=survey_response_id_list[i]).first()
#             studentSurveyAnsUpdate.answer = answer_list[i]
#             surveyDetailRow = SurveyDetail.query.filter_by(survey_id=survey_id).first()            
#         else:
#             addNewSurveyResponse = StudentSurveyResponse(survey_id=survey_id, sq_id=sq_id_list[i], 
#                 student_id=student_id, answer=answer_list[i], last_modified_date=datetime.today())
#             db.session.add(addNewSurveyResponse)
#     db.session.commit()
#     return jsonify(['0'])

# @app.route('/addNewSurvey',methods=["GET","POST"])
# def addNewSurvey():    
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     questions = request.form.getlist('questionInput')
#     questionCount = len(questions)
#     newSurveyRow = SurveyDetail(survey_name=request.form.get('surveyName'),teacher_id= teacherRow.teacher_id, 
#         school_id=teacherRow.school_id, question_count = questionCount, is_archived='N',last_modified_date=datetime.today())
#     db.session.add(newSurveyRow)
#     db.session.commit()
#     currentSurvey = SurveyDetail.query.filter_by(teacher_id=teacherRow.teacher_id).order_by(SurveyDetail.last_modified_date.desc()).first()
#     for i in range(questionCount):
#         newSurveyQuestion= SurveyQuestions(survey_id=currentSurvey.survey_id, question=questions[i], is_archived='N',last_modified_date=datetime.today())
#         db.session.add(newSurveyQuestion)
#     db.session.commit()
#     return jsonify(['0'])


# @app.route('/archiveSurvey')
# def archiveSurvey():
#     survey_id = request.args.get('survey_id')
#     surveyData = SurveyDetail.query.filter_by(survey_id=survey_id).first()
#     surveyData.is_archived='Y'
#     db.session.commit()
#     return jsonify(['0'])
# #End of student survey section


#Start of inventory pages
# @app.route('/inventoryManagement')
# def inventoryManagement():
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     print('#####################'+str(teacherRow))
#     inventoryDetailRow = InventoryDetail.query.filter_by(school_id = teacherRow.school_id, is_archived='N').all()
    
#     class_list=ClassSection.query.distinct().order_by(ClassSection.class_val).filter_by(school_id=teacherRow.school_id).all()    
#     return render_template('inventoryManagement.html',inventoryDetailRow=inventoryDetailRow,class_list=class_list)

# @app.route('/addInventoryItem', methods=["GET","POST"])
# def addInventoryItem():
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     inventoryName = request.form.get('inventoryName')    
#     newInventoryRow = InventoryDetail(inv_name=request.form.get('invName'),teacher_id= teacherRow.teacher_id, 
#         inv_description = request.form.get('invDescription'), inv_category = 225, total_stock = request.form.get('totalStock'),
#         item_rate = request.form.get('itemRate'), total_cost = request.form.get('totalCost'), stock_out=0, 
#         school_id=teacherRow.school_id, is_archived='N',last_modified_date=datetime.today())
#     db.session.add(newInventoryRow)
#     db.session.commit()        
#     addedInventory = InventoryDetail.query.filter_by(teacher_id = teacherRow.teacher_id, is_archived='N').order_by(InventoryDetail.last_modified_date.desc()).first()
#     return jsonify([addedInventory.inv_id])


# @app.route('/archiveInventory')
# def archiveInventory():
#     inv_id = request.args.get('inv_id')
#     InventoryData = InventoryDetail.query.filter_by(inv_id=inv_id).first()
#     InventoryData.is_archived='Y'
#     db.session.commit()
#     return jsonify(['0'])



# @app.route('/studentInventoryAlloc')
# def studentInventoryAlloc():
#     class_sec_id = request.args.get('class_sec_id')
#     inv_id = request.args.get('inv_id')
#     inv_id = inv_id.split('_')[1]
#     #studentInventoryQuery = "select sp.student_id , sp.full_name from student_profile sp  where sp.class_Sec_id="+ str(class_sec_id)    
#     #studentInventoryData = db.session.execute(studentInventoryQuery).fetchall()    
#     studentInventoryData = StudentProfile.query.filter_by(class_sec_id = str(class_sec_id)).all()
#     return render_template('_studentInventoryAlloc.html',studentInventoryData=studentInventoryData)


# @app.route('/updateInventoryAllocation')
# def updateInventoryAllocation():
#     ##last bit of changes required here to save inventory allocation
#     return jsonify(['0'])
# #End of inventory pages



@app.route('/performance')
def performance():
    df = pd.read_csv('data.csv').drop('Open', axis=1)
    chart_data = df.to_dict(orient='records')
    chart_data = json.dumps(chart_data, indent=2)
    data = {'chart_data': chart_data}
    return render_template("_performance.html", data=data)

@app.route('/help')
@login_required
def help():
    indic='DashBoard'
    return render_template('help.html',indic=indic,title='Help',user_type_val=str(current_user.user_type))


@app.route('/search')
def search():
    return render_template('search.html',title='Search',home=1)



@app.route('/checkout')
@login_required
def checkout():
    MERCHANT_KEY = "jS0QCS0O"
    key = ""
    SALT = "m4UzSP4umv"
    PAYU_BASE_URL = "https://sandboxsecure.payu.in/_payment"
    action = ''
    posted={}
    # Merchant Key and Salt provided y the PayU.
    for i in request.form:
    	posted[i]=request.form[i]
    hash_object = hashlib.sha256(b'randint(0,20)')
    txnid=hash_object.hexdigest()[0:20]
    hashh = ''
    posted['txnid']=txnid
    hashSequence = "key|txnid|amount|productinfo|firstname|email|udf1|udf2|udf3|udf4|udf5|udf6|udf7|udf8|udf9|udf10"
    posted['key']=key
    hash_string=''
    hashVarsSeq=hashSequence.split('|')
    for i in hashVarsSeq:
    	try:
    		hash_string+=str(posted[i])
    	except Exception:
    		hash_string+=''
    	hash_string+='|'
    hash_string+=SALT
    hash_string=hash_string.encode('utf-8')
    hashh=hashlib.sha512(hash_string).hexdigest().lower()
    action =PAYU_BASE_URL
    if(posted.get("key")!=None and posted.get("txnid")!=None and posted.get("productinfo")!=None and posted.get("firstname")!=None and posted.get("email")!=None):
    	return render_template('checkout.html',posted=posted,hashh=hashh,MERCHANT_KEY=MERCHANT_KEY,txnid=txnid,hash_string=hash_string,action='https://test.payu.in/_payment')
    else:
    	return render_template('checkout.html',posted=posted,hashh=hashh,MERCHANT_KEY=MERCHANT_KEY,txnid=txnid,hash_string=hash_string,action='.')   


@app.route('/paymentSuccess')
@login_required
def paymentSuccess():
    return render_template('paymentSuccess.html')

@app.route('/paymentFailure')
@login_required
def paymentFailure():
    return render_template('paymentFailure.html')


@app.route('/createSubscription',methods = ["GET","POST"])
@login_required
def createSubscription():
    form = createSubscriptionForm()
    if form.validate_on_submit():
        subcription_data=SubscriptionDetail(sub_name=form.sub_name.data,
            monthly_charge=form.monthly_charge.data,start_date=form.start_date.data,end_date=form.end_date.data,
            student_limit=form.student_limit.data,teacher_limit=form.teacher_limit.data,test_limit=form.test_limit.data, 
            sub_desc= form.sub_desc.data, last_modified_date= datetime.now(), archive_status='N', sub_duration_months=form.sub_duration.data)
        db.session.add(subcription_data)
        db.session.commit()
        flash('New subscription plan created.')
    return render_template('createSubscription.html',form=form)


# Routes for New Task

# @app.route('/studentHomeWork')
# @login_required
# def studentHomeWork():
#     user_type = current_user.user_type
#     if user_type==134:
#         user_id = User.query.filter_by(id=current_user.id).first()
#         student_id = StudentProfile.query.filter_by(user_id=user_id.id).first()
#         print('class_sec_id:'+str(student_id.class_sec_id))
#         homeworkDetailQuery = "select sd.homework_id, homework_name, question_count, sd.last_modified_date,count(ssr.answer) as ans_count "
#         homeworkDetailQuery = homeworkDetailQuery+ "from homework_detail sd left join student_homework_response ssr on ssr.homework_id =sd.homework_id "
#         homeworkDetailQuery = homeworkDetailQuery+" where sd.school_id ="+str(student_id.school_id)+ " and sd.is_archived='N' and sd.class_sec_id='"+str(student_id.class_sec_id)+"' group by sd.homework_id,homework_name,question_count, sd.last_modified_date"
#         homeworkDetailQuery = homeworkDetailQuery+" order by sd.last_modified_date desc"
#         print(homeworkDetailQuery)
#         homeworkData = db.session.execute(homeworkDetailQuery).fetchall()
#         print('student_id:'+str(student_id.student_id))
#         studentDetails = StudentProfile.query.filter_by(user_id=current_user.id).first()  
#         indic='homework'
#     return render_template('studentHomeWork.html',indic=indic,student_id=student_id.student_id,homeworkData=homeworkData,user_type_val=str(current_user.user_type), studentDetails=studentDetails)

# @app.route('/HomeWork')
# @login_required
# def HomeWork():
#     qclass_val = request.args.get('class_val')
#     qsection=request.args.get('section')
#     teacherRow = ''
#     if current_user.user_type==134:
#         teacherRow = StudentProfile.query.filter_by(user_id=current_user.id).first()        
#     else:
#         teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     classSections=ClassSection.query.filter_by(school_id=teacherRow.school_id).all()
#     count = 0
#     for section in classSections:
#         print("Class Section:"+section.section)
#             #this section is to load the page for the first class section if no query value has been provided
#         if count==0:
#             getClassVal = section.class_val
#             getSection = section.section
#             count+=1
#     if qclass_val is None:
#         qclass_val = getClassVal
#         qsection = getSection
    
#     class_sec_id = ClassSection.query.filter_by(school_id=teacherRow.school_id,class_val=qclass_val,section=qsection).first()
#     homeworkDetailQuery = "select sd.homework_id, homework_name, question_count, count(ssr.student_id ) as student_responses, question_count, sd.last_modified_date "
#     homeworkDetailQuery = homeworkDetailQuery+ "from homework_detail sd left join student_homework_response ssr on ssr.homework_id =sd.homework_id "
#     homeworkDetailQuery = homeworkDetailQuery+" where sd.school_id ="+str(teacherRow.school_id)+ " and sd.is_archived='N' and sd.class_sec_id='"+str(class_sec_id.class_sec_id)+"' group by sd.homework_id,homework_name, question_count,question_count, sd.last_modified_date"
#     homeworkDetailQuery = homeworkDetailQuery+" order by sd.last_modified_date desc"
#     print(homeworkDetailQuery)
#     homeworkDetailRow = db.session.execute(homeworkDetailQuery).fetchall()
#     #surveyDetailRow = SurveyDetail.query.filter_by(school_id=teacherRow.school_id).all()
#     distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacherRow.school_id)+" GROUP BY class_val order by s")).fetchall() 
#     classSections=ClassSection.query.filter_by(school_id=teacherRow.school_id).all()
#     indic='DashBoard'
#     return render_template('HomeWork.html',indic=indic,title='Homework', homeworkDetailRow=homeworkDetailRow,distinctClasses=distinctClasses,classSections=classSections,qclass_val=qclass_val,qsection=qsection,user_type_val=str(current_user.user_type))

# @app.route('/homeworkReview')
# @login_required
# def homeworkReview():
#     homework_id = request.args.get('homework_id')
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #homeworkRevData = "select *from fn_student_homework_status("+str(teacherRow.school_id)+","+str(homework_id)+")"
#     homeworkRevData = "select sp.full_name as student_name, sp.student_id ,count(answer) as ans_count,hd.question_count as qcount,hd.homework_id from student_homework_response shr inner join student_profile sp "
#     homeworkRevData = homeworkRevData + "on sp.student_id = shr.student_id inner join homework_detail hd on hd.homework_id = shr.homework_id "
#     homeworkRevData = homeworkRevData + "where sp.school_id = '"+str(teacherRow.school_id)+"' and shr.homework_id='"+str(homework_id)+"' group by student_name , qcount, sp.student_id, hd.homework_id"
#     homeworkRevData = db.session.execute(text(homeworkRevData)).fetchall()    
#     homework_name = HomeWorkDetail.query.filter_by(homework_id=homework_id).first()
#     classSections=ClassSection.query.filter_by(school_id=teacherRow.school_id,class_sec_id=homework_name.class_sec_id ).first()
#     return render_template('homeworkReview.html',homeworkRevData=homeworkRevData,class_val=classSections.class_val,section=classSections.section,homework_name=homework_name.homework_name,homework_id=homework_id)

# @app.route('/indivHomeworkReview',methods=['GET','POST'])
# @login_required
# def indivHomeworkReview():
#     homework_id = request.args.get('homework_id') 
#     student_id = request.args.get('student_id')
#     #homework_id = HomeWorkDetail.query.filter_by(homework_id=homework_id).first()
#     #reviewData = "select  hq.sq_id as sq_id, hq.question,hq.ref_type,hq.ref_url,shr.answer,shr.teacher_remark as teacher_remark from homework_questions hq left join student_homework_response shr "
#     #reviewData = reviewData + "on hq.homework_id = shr.homework_id and hq.sq_id =shr.sq_id where hq.homework_id = '"+str(homework_id.homework_id)+"'"
#     reviewData = "select  hq.sq_id as sq_id, hq.question,hq.ref_type,hq.ref_url,shr.answer,shr.teacher_remark as teacher_remark "
#     #from homework_questions hq left join student_homework_response shr "
#     #reviewData = reviewData + "on hq.homework_id = shr.homework_id and hq.sq_id =shr.sq_id where hq.homework_id = '"+str(homework_id.homework_id)+"'"
#     reviewData = reviewData + " from homework_detail hd inner join homework_questions hq on "
#     reviewData = reviewData + " hd.homework_id = hq.homework_id and hd.homework_id =" + str(homework_id)
#     reviewData = reviewData + " left join student_homework_response shr on "
#     reviewData = reviewData + " hq.sq_id =shr.sq_id and shr.student_id = " + str(student_id)
#     reviewData = reviewData + " and shr.homework_response_id in (select min(homework_response_id )from student_homework_response shr "
#     reviewData = reviewData + " where student_id ="+ str(student_id) +" and homework_id ="+ str(homework_id) +" group by sq_id ) "
#     #print(reviewData)
#     reviewData = db.session.execute(text(reviewData)).fetchall()    
#     return render_template('_indivHomeWorkReview.html',reviewData=reviewData,student_id=student_id)

# @app.route('/indivHomeworkDetail',methods=['GET','POST'])
# @login_required
# def indivHomeWorkDetail():
#     homework_id = request.args.get('homework_id') 
#     user_id = User.query.filter_by(id=current_user.id).first()
#     student_id = StudentProfile.query.filter_by(user_id=user_id.id).first()
#     homework_name = HomeWorkDetail.query.filter_by(homework_id=homework_id).first()
#     #homeworkQuestions = HomeWorkQuestions.query.filter_by(homework_id=homework_id).all()

#     #homeworkDataQQuery = " select distinct hq.sq_id as sq_id, question,hq.homework_id ,ref_type, ref_url, homework_response_id , sp.student_id, answer,teacher_remark from student_homework_response shr "
#     #homeworkDataQQuery = homeworkDataQQuery + "right join homework_questions hq on "
#     #homeworkDataQQuery = homeworkDataQQuery +  "hq.homework_id =shr.homework_id and "
#     #homeworkDataQQuery = homeworkDataQQuery +  "hq.sq_id =shr.sq_id and shr.student_id = "+ str(student_id.student_id)
#     #homeworkDataQQuery = homeworkDataQQuery +  " left join student_profile sp "
#     #homeworkDataQQuery = homeworkDataQQuery +  "on sp.student_id =shr.student_id where hq.homework_id ="+ str(homework_id)
#     #homeworkDataQQuery = homeworkDataQQuery +  " and homework_response_id in "
#     #homeworkDataQQuery = homeworkDataQQuery +  " (select min(homework_response_id )from student_homework_response shr "
#     #homeworkDataQQuery = homeworkDataQQuery +  " where student_id ="+ str(student_id.student_id) + " and homework_id ="+ str(homework_id) +" group by sq_id ) "
#     homeworkDataQQuery = "select distinct hq.sq_id as sq_id, question,hq.homework_id ,ref_type, ref_url, homework_response_id , shr.student_id, answer,teacher_remark "
#     homeworkDataQQuery = homeworkDataQQuery +  " from homework_detail hd inner join homework_questions hq "
#     homeworkDataQQuery = homeworkDataQQuery +  " on hd.homework_id = hq.homework_id and hd.homework_id =" + str(homework_id)
#     homeworkDataQQuery = homeworkDataQQuery +  " left join student_homework_response shr on "
#     homeworkDataQQuery = homeworkDataQQuery +  " hq.sq_id =shr.sq_id and shr.student_id =" + str(student_id.student_id)
#     homeworkDataQQuery = homeworkDataQQuery +  " and shr.homework_response_id in (select min(homework_response_id )from student_homework_response shr "
#     homeworkDataQQuery = homeworkDataQQuery +  " where student_id ="+ str(student_id.student_id) +" and homework_id ="+ str(homework_id) +" group by sq_id )"

#     print(homeworkDataQQuery)
#     homeworkDataRows = db.session.execute(text(homeworkDataQQuery)).fetchall()
#     homeworkAttach = db.session.execute(text("select attachment from homework_detail where homework_id='"+str(homework_id)+"'")).first()
#     return render_template('_indivHomeWorkDetail.html',homeworkDataRows=homeworkDataRows,homework_name=homework_name,homework_id=homework_id,student_id=student_id,homeworkAttach=homeworkAttach)

# @app.route('/addAnswerRemark',methods=["GET","POST"])
# def addAnswerRemark():
#     remark = request.form.getlist('remark')
#     student_id = request.args.get('student_id')
#     sq_id_list = request.form.getlist('sq_id')
#     print('######'+str(len(sq_id_list) ))
#     for i in range(len(sq_id_list)):
#         remarkData = StudentHomeWorkResponse.query.filter_by(student_id=student_id,sq_id= sq_id_list[i]).first()  

#         print('################################e   entered remark section')
#         print(str(StudentHomeWorkResponse.query.filter_by(student_id=student_id,sq_id= sq_id_list[i])))
#         if remarkData!=None:            
#             remarkData.teacher_remark = remark[i]
#     db.session.commit()
#     return jsonify(['0'])

# checkValue = ''
# @app.route('/addHomeworkAnswer',methods=["GET","POST"])
# def addHomeworkAnswer():
#     sq_id_list = request.form.getlist('sq_id')
#     answer_list = request.form.getlist('answer')
#     homework_id = request.form.get('homework_id')
#     print('add homework answer')
#     user_id = User.query.filter_by(id=current_user.id).first()
#     student_id = StudentProfile.query.filter_by(user_id=user_id.id).first()
#     for i in range(len(sq_id_list)):
#         checkStudentReponse = StudentHomeWorkResponse.query.filter_by(student_id=student_id.student_id,sq_id=sq_id_list[i]).first()
#         if checkStudentReponse==None or checkStudentReponse=="":
#             addNewHomeWorkResponse = StudentHomeWorkResponse(homework_id=homework_id, sq_id=sq_id_list[i], 
#                 student_id=student_id.student_id, answer=answer_list[i], last_modified_date=datetime.today())
#             db.session.add(addNewHomeWorkResponse)
#             print("Not present")
#         else:
#             return jsonify(['1'])
#         #    checkStudentReponse.answer=answer_list[i]
#         #    print("Already present")
#     db.session.commit()
#     return jsonify(['0'])


# def get_yt_video_id(url):
#     """Returns Video_ID extracting from the given url of Youtube
    
#     Examples of URLs:
#       Valid:
#         'http://youtu.be/_lOT2p_FCvA',
#         'www.youtube.com/watch?v=_lOT2p_FCvA&feature=feedu',
#         'http://www.youtube.com/embed/_lOT2p_FCvA',
#         'http://www.youtube.com/v/_lOT2p_FCvA?version=3&amp;hl=en_US',
#         'https://www.youtube.com/watch?v=rTHlyTphWP0&index=6&list=PLjeDyYvG6-40qawYNR4juzvSOg-ezZ2a6',
#         'youtube.com/watch?v=_lOT2p_FCvA',
      
#       Invalid:
#         'youtu.be/watch?v=_lOT2p_FCvA',
#     """
    

#     if url.startswith(('youtu', 'www')):
#         url = 'http://' + url
#     embeddingURL = "https://www.youtube.com/embed/"
#     query = urlparse(url)
    
#     if 'youtube' in query.hostname:
#         if query.path == '/watch':
#             return "96", embeddingURL + parse_qs(query.query)['v'][0]
#         elif query.path.startswith(('/embed/', '/v/')):
#             return "96",embeddingURL + query.path.split('/')[2]
#     elif 'youtu.be' in query.hostname:
#         return "96",embeddingURL + query.path[1:]
#     else:
#         return "97",url


# def checkContentType(contentName):
#     with urlopen(contentName) as response:
#         info = response.info()
#         contentTypeVal = info.get_content_type()
#         splittedContentType = contentTypeVal.split('/')
#         if splittedContentType[1]=='pdf' or splittedContentType[1]=='msword':
#             return "99"
#         elif splittedContentType[0]=='audio':
#             return "97"
#         elif splittedContentType[1]=='image':
#             return "98"
#         else:
#             return "227"


# @app.route('/addNewHomeWork',methods=["GET","POST"])
# def addNewHomeWork():     
#     teacherRow=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     questions = request.form.getlist('questionInput')
#     #contentType = request.form.getlist('contentType')
#     contentName = request.form.getlist('contentName')
#     homeworkContent = request.form.get('homeworkContent')
#     print('inside addNew Homework')
#     print(homeworkContent)
#     for i in range(len(contentName)):
#         print(contentName[i])
#     #for i in range(len(contentType)):
#     #    print('content type:'+str(contentType[i]))
#     questionCount = len(questions)
#     class_val = request.form.get('class')
#     section = request.form.get('section')
#     class_sec_id = ClassSection.query.filter_by(school_id=teacherRow.school_id,class_val=class_val,section=section).first()
#     newHomeWorkRow = HomeWorkDetail(homework_name=request.form.get('homeworkName'),teacher_id= teacherRow.teacher_id, 
#         school_id=teacherRow.school_id, question_count = questionCount, is_archived='N',class_sec_id=class_sec_id.class_sec_id,last_modified_date=datetime.today(),attachment=homeworkContent)
#     db.session.add(newHomeWorkRow)
#     db.session.commit()
#     currentHomeWork = HomeWorkDetail.query.filter_by(teacher_id=teacherRow.teacher_id).order_by(HomeWorkDetail.last_modified_date.desc()).first()
        
#     for i in range(questionCount):           
#         if contentName[i] !='':               
#             refType ,contentName[i] = get_yt_video_id(contentName[i])
#             if refType!=96:
#                 refType= checkContentType(contentName[i])                
#         else:
#             refType=226
#         newHomeWorkQuestion= HomeWorkQuestions(homework_id=currentHomeWork.homework_id, question=questions[i], is_archived='N',last_modified_date=datetime.today(),ref_type=int(refType),ref_url=contentName[i])
#         db.session.add(newHomeWorkQuestion)
#     db.session.commit()
#     return jsonify(['0:'+ str(currentHomeWork.homework_id)])






# @app.route('/archiveHomeWork')
# def archiveHomeWork():
#     homework_id = request.args.get('homework_id')
#     homeworkData = HomeWorkDetail.query.filter_by(homework_id=homework_id).first()
#     homeworkData.is_archived='Y'
#     db.session.commit()
#     return jsonify(['0'])


# End

#??start
def conversion(num):
    strings = []
    for integ in num:
        strings.append(str(integ))
    a_string = "".join(strings)
    an_integer = int(a_string)
    return(an_integer)
      
def convList(left):
    res = [int(x) for x in str(left)]
    return res
def cal(num): 
    digit = len(num) 
    print(digit)
    powTen = pow(10, digit - 1) 
    res = []

    
    an_integer = conversion(num)
    res.append(num)
    for i in range(digit - 1):
        firstDigit = an_integer // powTen
        left = (an_integer * 10 + firstDigit - (firstDigit * powTen * 10))
        an_integer = left
        out = convList(an_integer)
        res.append(out)
    return res 

# Route for Schedule or Time Table
# @app.route('/updateSchedule',methods=['POST','GET'])
# def updateSchedule():
#     slots = request.form.get('slots')
#     class_value = request.args.get('class_val')
#     print('No of slots:'+str(slots))
#     print('Class_val:'+str(class_value))
#     # start = request.form.getlist('start')
#     # end = request.form.getlist('end')
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     sections = ClassSection.query.filter_by(school_id=teacher.school_id,class_val=class_value).all()
#     batches = len(sections)
#     slotTime = []
#     # for (s,e) in itertools.zip_longest(start,end):
#     #     slotTime.append(s+'-'+e)
#     print('inside schedule')
#     # nDays = request.form.get('nDays')
#     noTeachers = request.form.get('noTeachers')
#     nameTeacher = request.form.getlist('nameTeacher')
#     nameSubject = request.form.getlist('nameSubject')
#     subject = request.form.getlist('subject')
#     time = request.form.getlist('time')
#     day = request.form.getlist('day')
#         # batches = request.form.get('batches')
#     print('No of Days:'+str(len(day)))
#     nDays = len(day)
#     totalTime = 0
#     for ti in time:
#             # print('Time:'+str(ti))
#         if ti:
#             totalTime = totalTime + int(ti)
#     totalSlots = int(nDays)*int(slots)
#         # print('Total slots:'+str(totalSlots))
#         # print('Total Time:'+str(totalTime))
#     perSlots = []
#     for s in range(int(slots)):
                
#         perSlots.append(s+1)
#     slotsoutput = cal(perSlots)

#         # for slot in range(len(slotsoutput)):
#             # print('slot'+str(slot))
#             # print(slotsoutput[slot])
#     finalSlot = []
#     indexSlot = []
#         # print('outside while loop')
#     p=0
#         # print('length of slots output')
#         # print(len(slotsoutput))
#     while p<len(slotsoutput):
#             # print('inside while loop')
#         r = randint(0,len(slotsoutput)-1)
#         if r not in indexSlot:
#             p=p+1
#             indexSlot.append(r)
#         # print(indexSlot)
#     for index in indexSlot:
#         finalSlot.append(slotsoutput[index])
#         # print(finalSlot)
        

#         # print('slot output print')
#         # print(slotsoutput)
#     t=1
#     z=0
#     class_sec_ids = ClassSection.query.filter_by(class_val = class_value, school_id = teacher.school_id).all() 
#     for class_sec_id in class_sec_ids:
#         updateTable = "update schedule_detail set is_archived = 'Y' where class_sec_id='"+str(class_sec_id.class_sec_id)+"' and school_id='"+str(teacher.school_id)+"'"
#         updateTable = db.session.execute(text(updateTable))
#     for l in subject:
#             # print(l)
#         if(l):
#             z=z+1
#         # print('No of Subjects:'+str(z))
#     if totalSlots>=totalTime:

#         for arr1 in finalSlot:
#             if t<=int(batches):    
                            
                    
#                 for i in range(0,z):

#                     for j in range(0,int(time[i])):
                            
#                         class_sec_id = "select class_sec_id from class_section where class_val='"+str(class_value)+"' and section='"+chr(ord('@')+t)+"' and school_id='"+str(teacher.school_id)+"'"
                            
#                         class_sec_id = db.session.execute(text(class_sec_id)).first()
#                         subject_id = "select msg_id from message_detail where description='"+str(subject[i])+"'"
#                         subject_id = db.session.execute(text(subject_id)).first()
#                         teacher_id = TeacherSubjectClass.query.filter_by(class_sec_id=class_sec_id.class_sec_id,subject_id=subject_id.msg_id,school_id=teacher.school_id,is_archived='N').first()
                           
#                         if teacher_id:
#                             insertData = ScheduleDetail(class_sec_id=class_sec_id.class_sec_id ,school_id=teacher.school_id, days_name=day[j] ,subject_id=subject_id.msg_id , teacher_id=teacher_id.teacher_id ,slot_no=arr1[i] , last_modified_date=dt.datetime.now(), is_archived= 'N')
#                         else:
#                             insertData = ScheduleDetail(class_sec_id=class_sec_id.class_sec_id ,school_id=teacher.school_id, days_name=day[j] ,subject_id=subject_id.msg_id  ,slot_no=arr1[i] , last_modified_date=dt.datetime.now(), is_archived= 'N')
#                         db.session.add(insertData)
#             t=t+1
#     else:
#         return jsonify(['1'])
#     db.session.commit()
#     print('Data is Submitted') 
#     return jsonify(['0'])

# @app.route('/schedule')
# def schedule():
#     # slots = request.form.get('slots')
#     # print('inside schedule function')
#     # print(slots)
#     qclass_val = request.args.get("class_val")
#     qsection = request.args.get("section")
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     classSections=ClassSection.query.filter_by(school_id=teacher.school_id).all()
#     available_class_section = "select distinct class_val,section from class_section where school_id='"+str(teacher.school_id)+"'"
#     available_class_section = db.session.execute(text(available_class_section)).fetchall()
#     distinctClasses = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher.school_id)+" GROUP BY class_val order by s")).fetchall()  
#     if qclass_val==None and qsection == None:
#         qclass_val = db.session.execute(text("SELECT  distinct class_val,sum(class_sec_id),count(section) as s FROM class_section cs where school_id="+ str(teacher.school_id)+" GROUP BY class_val order by s")).first()  
#         qclass_val = qclass_val.class_val
#         qsection = ClassSection.query.filter_by(school_id=teacher.school_id).first()
#         qsection = qsection.section
#     indic='DashBoard'
#     return render_template('schedule.html',indic=indic,classsections=classSections,distinctClasses=distinctClasses,available_class_section=available_class_section,qclass_val=qclass_val,qsection=qsection)
    
# @app.route('/fetchTimeTable',methods=['GET','POST'])
# def fetchTimeTable():
#     class_val = request.args.get('class_value')
#     section = request.args.get('section')
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     print(class_val)
#     print(section)
#     print(teacher.school_id)
#     class_section = ClassSection.query.filter_by(class_val=str(class_val),section = section, school_id= teacher.school_id).first()

#     query = "select *from fn_time_table("+str(teacher.school_id)+","+str(class_section.class_sec_id)+")"
#     print(query)
#     timeTableData = db.session.execute(text(query)).fetchall()
#     print(timeTableData)
#     for data in timeTableData:
#         print(data)
#     fetchTeacher = "select *from fn_teacher_allocation("+str(teacher.school_id)+","+str(class_section.class_sec_id)+")"
#     print(fetchTeacher)
#     fetchTeacher = db.session.execute(text(fetchTeacher)).fetchall()
#     print(fetchTeacher)
#     return render_template('_timeTable.html',timeTableData=timeTableData,fetchTeacher=fetchTeacher)

# @app.route('/downloadTimeTable',methods=['GET','POST'])
# def downloadTimeTable():
#     print('inside download time table')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     # fetchTeacher = "select *from fn_teacher_allocation("+str(teacher_id.school_id)+","+str(class_section.class_sec_id)+")"
#     # print(fetchTeacher)
#     # fetchTeacher = db.session.execute(text(fetchTeacher)).fetchall()
#     board_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     class_sec_ids = ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
#     # filepath = 'static/images/'+str(teacher_id.school_id)+str(dt.datetime.now())+'TimeTable.csv'
#     file_name = 'Timetable'+str(teacher_id.school_id)+'.csv'
#     with open('tempdocx/'+file_name, 'w', newline='') as file:
        
    
        
#         print('file')
        
#         file.write('Introduction \n')
#         writer = csv.writer(file)
#         for class_sec_id in class_sec_ids:
#             fetchTeacher = "select *from fn_teacher_allocation("+str(teacher_id.school_id)+","+str(class_sec_id.class_sec_id)+")"
#             print(fetchTeacher)
            
#             fetchTeacher = db.session.execute(text(fetchTeacher)).fetchall()
#             writer.writerow(["Subject Name", "Teacher Name"])
#             # writer.writerow(["", "Name", "Contribution"])
#             for teacher in fetchTeacher:
#                 print(teacher.subject_name)
#                 print(teacher.teacher_name)
#                 writer.writerow([teacher.subject_name,teacher.teacher_name])
#                 # os.remove(filepath)
#                 print('File path')
#             query = "select *from fn_time_table("+str(teacher_id.school_id)+","+str(class_sec_id.class_sec_id)+")"
#             print(query)
#             timeTableData = db.session.execute(text(query)).fetchall()
#             writer.writerow(["Periods", "Monday","Tuesday","Wednesday","Thursday","Friday","Saturday"])
            
#             for table in timeTableData:
#                 writer.writerow([table.period_no,table.monday,table.tuesday,table.wednesday,table.thursday,table.friday,table.saturday])
#             writer.writerow(["","","","","","",""])
#     client = boto3.client('s3', region_name='ap-south-1')
#     client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'time_table/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
#     os.remove('tempdocx/'+file_name)
#     filepath ='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/time_table/'+file_name
#     print('filepath:'+str(filepath))
#     return jsonify([filepath])

# @app.route('/allSubjects',methods=['GET','POST'])
# def allSubjects():
#     print('inside all Subjects')
#     teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     school_id = SchoolProfile.query.filter_by(school_id=teacher_id.school_id).first()
#     class_val = request.args.get('class_value') 
#     print(class_val)
#     print(teacher_id.school_id)
#     subjects = BoardClassSubject.query.filter_by(class_val = str(class_val),school_id=teacher_id.school_id,is_archived='N').all()
#     subjectList = []
#     for subject in subjects:
#         subject_name = "select description from message_detail where msg_id='"+str(subject.subject_id)+"'"
#         subject_name = db.session.execute(text(subject_name)).first()
        
#         subjectList.append(str(subject_name.description))
#     return jsonify([subjectList])

#End
#Start

# @app.route('/addTeacherClassSubject',methods=['GET','POST'])
# def addTeacherClassSubject():
#     subjectNames = request.get_json()
#     print('Inside add Teacher class Subject')
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_id = request.args.get('class_sec_id')
#     teacher_id = request.args.get('teacher_id')
#     remdata = "update teacher_subject_class set is_archived = 'Y' where class_sec_id='"+str(class_sec_id)+"' and school_id = '"+str(teacher.school_id)+"'  and teacher_id = '"+str(teacher_id)+"'"
#     remdata = db.session.execute(text(remdata))
#     for subjects in subjectNames:
#         print('inside for')
#         print(subjects)
#         subject_id = "select msg_id from message_detail where description='"+str(subjects)+"'"
#         subject_id = db.session.execute(text(subject_id)).first()
        
        
#         addTeacherClass = TeacherSubjectClass(teacher_id=teacher_id,class_sec_id=class_sec_id,subject_id=subject_id.msg_id,is_archived='N',school_id=teacher.school_id,last_modified_date=dt.datetime.now())
#         db.session.add(addTeacherClass)
#         db.session.commit()
#     return ""

# @app.route('/loadSubject',methods=['GET','POST'])
# def loadSubject():
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     teacher_id = request.args.get('teacher_id')
#     class_sec_id = request.args.get('class_sec_id')
#     class_val = ClassSection.query.filter_by(class_sec_id=class_sec_id,school_id=teacher.school_id).first()
#     allSubjects = "select distinct subject_id from board_class_subject bcs where school_id = '"+str(teacher.school_id)+"' and class_val = '"+str(class_val.class_val)+"' and is_archived = 'N' and subject_id "
#     allSubjects = allSubjects + "not in (select distinct subject_id from teacher_subject_class where school_id= '"+str(teacher.school_id)+"' and is_archived = 'N' and class_sec_id = '"+str(class_sec_id)+"' and teacher_id='"+str(teacher_id)+"')"
#     allSubjects = db.session.execute(text(allSubjects)).fetchall()

#     selSubjects = "select distinct subject_id from teacher_subject_class where school_id= '"+str(teacher.school_id)+"' and is_archived = 'N' and class_sec_id = '"+str(class_sec_id)+"' and teacher_id='"+str(teacher_id)+"'"
#     selSubjects = db.session.execute(text(selSubjects)).fetchall()
#     selArray = []
#     for subjects in allSubjects:
#         subject_name = "select description from message_detail where msg_id='"+str(subjects.subject_id)+"'"
#         subject_name = db.session.execute(text(subject_name)).first()
#         selArray.append(str(subject_name.description)+':'+str('false'))
#     for sub in selSubjects:
#         subject_name = "select description from message_detail where msg_id='"+str(sub.subject_id)+"'"
#         subject_name = db.session.execute(text(subject_name)).first()
#         selArray.append(str(subject_name.description)+':'+str('true'))
#     if selArray:
#         return jsonify([selArray])
#     else:
#         return ""

# @app.route('/loadClasses',methods=['GET','POST'])
# def loadClasses():
#     teacher_id = request.args.get('teacher_id')
#     teacher = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_val = ClassSection.query.filter_by(school_id=teacher.school_id).all()
#     classes = []
#     print(teacher.school_id)
#     for clas in class_val: 
#         teacher = TeacherSubjectClass.query.filter_by(teacher_id=teacher_id,class_sec_id=clas.class_sec_id).first()
#         if teacher:
#             classes.append(str(clas.class_val)+':'+str(clas.class_sec_id)+':'+str(clas.section)+':'+str('true'))
#         else:
#             classes.append(str(clas.class_val)+':'+str(clas.class_sec_id)+':'+str(clas.section)+':'+str('false'))
        
#     return jsonify([classes])

# @app.route('/teacherAllocation',methods=['GET','POST'])
# def teacherAllocation():
#     teacher_id = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     class_sec_ids = ClassSection.query.filter_by(school_id=teacher_id.school_id).all()
#     fetchTeacherNameQuery = "select distinct tp.teacher_name,tp.teacher_id from teacher_profile tp "
#     fetchTeacherNameQuery = fetchTeacherNameQuery + "left join teacher_subject_class tsc on tsc.teacher_id=tp.teacher_id "
#     fetchTeacherNameQuery = fetchTeacherNameQuery + "where tp.school_id='"+str(teacher_id.school_id)+"'"
#     print(fetchTeacherNameQuery)
#     teacherNames = db.session.execute(text(fetchTeacherNameQuery)).fetchall()
#     print('Inside teacher class subject allocation')
#     return render_template('_teacherAllocation.html',teacherNames=teacherNames,class_sec_ids=class_sec_ids)

#End

@app.route('/subscriptionPlans')
def subscriptionPlans():
    subscriptionRow = SubscriptionDetail.query.filter_by(archive_status='N').order_by(SubscriptionDetail.sub_duration_months).all()    
    distinctSubsQuery = db.session.execute(text("select distinct group_name, sub_desc, student_limit, teacher_limit, test_limit from subscription_detail where archive_status='N' order by student_limit ")).fetchall()
    return render_template('/subscriptionPlans.html', subscriptionRow=subscriptionRow, distinctSubsQuery=distinctSubsQuery)

# @app.route('/studTC',methods=["GET","POST"])
# @login_required
# def studTC():
#     teacherData = TeacherProfile.query.filter_by(user_id=current_user.id).first()
#     #tcdata = TransferCerts.query.filter_by(is_archived='N', school_id=teacherData.school_id).all()
#     tcDataQuery = "select sp.student_id, sp.school_adm_number, sp.first_name, sp.last_name, cs.class_val, cs.section, tc.tc_url, tc.tc_id "
#     tcDataQuery = tcDataQuery + " from transfer_certs tc left join student_profile sp on tc.student_id=sp.student_id "
#     tcDataQuery = tcDataQuery + " inner join class_Section cs on cs.class_Sec_id=sp.class_sec_id"
#     tcDataQuery = tcDataQuery + " where tc.is_archived='N' and tc.school_id="+ str(teacherData.school_id)
#     tcData = db.session.execute(tcDataQuery).fetchall()
#     if request.method=="POST":
#         student_id = request.form.get('student_id')
#         school_adm_number = request.form.get('school_adm_number')
#         teacher_id = teacherData.teacher_id
#         school_id = teacherData.school_id
#         tc_url =  request.form.get('pdfURL')
#         #student_id = request.form.get('student_id')
#         if school_adm_number:
#             chkStudentData = StudentProfile.query.filter_by(school_adm_number=school_adm_number,school_id=school_id).first()
#             if chkStudentData!=None:
#                 transferDataAdd = TransferCerts(student_id=chkStudentData.student_id, school_adm_number=school_adm_number, teacher_id=teacher_id
#                     ,school_id=school_id, tc_url=tc_url,is_archived='N', last_modified_date=datetime.today())
#                 db.session.add(transferDataAdd)
#                 db.session.commit()
#                 flash('TC Uploaded Successfully!')
#             else:
#                 flash(Markup('<span class="red-text"> Student ID invalid. Please try again.</span>'))
#         elif student_id:
#             chkStudentData = StudentProfile.query.filter_by(student_id=student_id,school_id=school_id).first()
#             if chkStudentData!=None:
#                 transferDataAdd = TransferCerts(student_id=student_id, school_adm_number=chkStudentData.school_adm_number, teacher_id=teacher_id
#                     ,school_id=school_id, tc_url=tc_url,is_archived='N', last_modified_date=datetime.today())
#                 db.session.add(transferDataAdd)
#                 db.session.commit()
#                 flash('TC Uploaded Successfully!')
#             else:
#                 flash(Markup('<span class="red-text"> Student ID invalid. Please try again.</span>'))
#         else:
#             flash(Markup('<span class="red-text"> Please Enter Student Id or School Admission Number. Please try again.</span>'))
#     indic='DashBoard'
#     return render_template('studTC.html',indic=indic,tcData=tcData,title='Student TC')

# @app.route('/archiveTCClass',methods=["GET","POST"])
# @login_required
# def archiveTCClass():
#     tc_id = request.args.get('tc_id')
#     tcData = TransferCerts.query.filter_by(tc_id=tc_id).first()
#     tcData.is_archived='Y'
#     db.session.commit()
#     return jsonify(['0'])

# @app.route('/accessStudTC')
# def accessStudTC():
#     school_id = request.args.get('school_id')
#     if school_id!=None and school_id!="":
#         schoolData = SchoolProfile.query.filter_by(school_id=str(school_id)).first()
#         if schoolData!=None:
#             return render_template('accessStudTC.html',schoolData=schoolData)
#         else:
#             return jsonify(['Invalid School Data'])
#     #if ("alllearn" in str(request.url)) or  ("localhost" in str(request.url)) ("tc" in str(request.url)) or ("wix" in str(request.url)) :        
    
#     #else:
#     #    return jsonify(['Invalid Call'])

# @app.route('/fetchStudTC')
# def fetchStudTC():
#     # student_id = request.args.get('student_id')
#     school = current_user.school_id
#     print('School id:'+str(school))
#     school_adm_number = request.args.get('school_adm_number')
#     student = StudentProfile.query.filter_by(school_adm_number = school_adm_number,school_id=school).first()
#     tcData = TransferCerts.query.filter_by(student_id=student.student_id,is_archived='N'  ).first()    
#     if tcData!=None:
#         if tcData.tc_url!=None and tcData.tc_url!="":
#             return jsonify([tcData.tc_url])
#         else:
#             return jsonify(['NA'])    
#     else:
#         return jsonify(['NA'])

def format_currency(value):
    return "???{:,.2f}".format(value)


if __name__=="__main__":
    app.debug=True  
    #app.use_reloader=False  
    app.jinja_env.filters['zip'] = zip
    #app.run(host=os.getenv('IP', '127.0.0.1'), 
    #        port=int(os.getenv('PORT', 8000)))
    app.run(host=os.getenv('IP', '0.0.0.0'),         
        port=int(os.getenv('PORT', 8000))
        # ssl_context='adhoc'
        )
    #app.run()
