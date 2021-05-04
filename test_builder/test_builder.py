from flask import Flask, Blueprint, Markup, render_template, request, flash, redirect, url_for, Response,session,jsonify
from applicationDB import *
from test_builder.utils import *
from forms import LoginForm, RegistrationForm,ContentManager,LeaderBoardQueryForm, EditProfileForm, ResetPasswordRequestForm, ResetPasswordForm,ResultQueryForm,MarksForm, TestBuilderQueryForm,SchoolRegistrationForm, PaymentDetailsForm, addEventForm,QuestionBuilderQueryForm, SingleStudentRegistration, SchoolTeacherForm, feedbackReportForm, testPerformanceForm, studentPerformanceForm, QuestionUpdaterQueryForm,  QuestionBankQueryForm,studentDirectoryForm, promoteStudentForm
from flask_login import LoginManager, current_user, login_user, logout_user, login_required
import os
from datetime import datetime
from pytz import timezone
from tzlocal import get_localzone
from flask import g, jsonify
import json, boto3
from sqlalchemy import func, distinct, text, update
from miscFunctions import subjects,topics,subjectPerformance,signs3Folder,chapters
from docx import Document
from docx.shared import Inches
from io import StringIO, BytesIO
from random import randint
import requests as rq
import urllib
from urllib.parse import quote,urlparse, parse_qs
from google.auth.transport import requests
import json


test_builder= Blueprint('test_builder',__name__)

@test_builder.route('/testBuilder',methods=['POST','GET'])
@login_required
def testBuilder():
    topic_list=None
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form=TestBuilderQueryForm()
    print(teacher_id.school_id)
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()]
    form.subject_name.choices= ''
    form.chapter_num.choices= ''
    # [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
    form.test_type.choices= [(i.description,i.description) for i in MessageDetails.query.filter_by(category='Test type').all()]
    test_papers = MessageDetails.query.filter_by(category='Test type').all()
    # print(request.form['class_val'])
    # print(request.form['subject_id'])
    available_class = "select distinct class_val from class_section where school_id='"+str(teacher_id.school_id)+"'"
    available_class = db.session.execute(text(available_class)).fetchall()
    if request.method=='POST':
        if request.form['test_date']=='':
            # flash('Select Date')
            # form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(int(form.class_val.data))]
            indic='DashBoard'
            return render_template('testBuilder.html',indic=indic,form=form)
        topic_list=Topic.query.filter_by(class_val=str(form.class_val.data),subject_id=int(form.subject_name.data),chapter_num=int(form.chapter_num.data)).all()
        subject=MessageDetails.query.filter_by(msg_id=int(form.subject_name.data)).first()
        session['class_val']=form.class_val.data
        session['date']=request.form['test_date']
        session['sub_name']=subject.description
        session['sub_id']=form.subject_name.data
        session['test_type_val']=form.test_type.data
        session['chapter_num']=form.chapter_num.data 
        form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(str(form.class_val.data))]
        form.chapter_num.choices= [(int(i['chapter_num']), str(i['chapter_num'])+' - '+str(i['chapter_name'])) for i in chapters(str(form.class_val.data),int(form.subject_name.data))]
        indic='DashBoard'
        return render_template('testBuilder.html',indic=indic,title='Test Builder',form=form,topics=topic_list,user_type_val=str(current_user.user_type))
    indic='DashBoard'
    return render_template('testBuilder.html',indic=indic,title='Test Builder',form=form,available_class=available_class,test_papers=test_papers,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

@test_builder.route('/filterQuestionsfromTopic',methods=['GET','POST'])
def filterQuestionsfromTopic():
    topics = request.get_json()
    print('topics:'+str(topics))
    
    for topic in topics:
        print('topic:'+str(topic))
    class_val = request.args.get('class_val')
    subject_id = request.args.get('subject_id')
    test_type = request.args.get('test_type')
    print('Class feedback:'+str(test_type))
    questions = []
    questionList = ''
    if topics:
        print('if topics available')
        for topic in topics:
            # if test_type == 'Class Feedback':
            if class_val!=None:
                if subject_id!=None:
                    if test_type!=None:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                    else:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                else:
                    if test_type!=None:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                    else:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()           
            else:
                if subject_id!=None:
                    if test_type!=None:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                    else:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                else:
                    if test_type!=None:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
                    else:
                        questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
            # else:
            #     if class_val!=None:
            #         if subject_id!=None:
            #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
            #         else:
            #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all()
            #     else:
            #         if subject_id!=None:
            #             questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id,topic_id=topic).order_by(QuestionDetails.question_id).all()
            #         else:
            #             questionList = QuestionDetails.query.filter_by(archive_status='N',topic_id=topic).order_by(QuestionDetails.question_id).all() 
            if questionList:  
                questions.append(questionList)
    else:
        print('if topics not available')
        # if test_type == 'Class Feedback':
        if class_val!=None:
            print('if class_val available:'+str(class_val))
            if subject_id!=None:
                print('if subject_id available:'+str(subject_id))
                if test_type:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
                else:
                    print('if test type is empty')
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
            else:
                if test_type!=None:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
                else:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
        else:
            if subject_id!=None:
                if test_type!=None:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
                else:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
            else:
                if test_type!=None:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
                else:
                    questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
        # else:
        #     if class_val!=None:
        #         if subject_id!=None:
        #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').order_by(QuestionDetails.question_id).all()
        #         else:
        #             questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N').order_by(QuestionDetails.question_id).all()
        #     else:
        #         if subject_id!=None:
        #             questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id).order_by(QuestionDetails.question_id).all()
        #         else:
        #             questionList = QuestionDetails.query.filter_by(archive_status='N').order_by(QuestionDetails.question_id).all() 
        print('QuestionList:'+str(questionList))
        for ques in questionList:
            print('inside for of QuestionList')
            print(ques.question_id)
        if len(questionList)==0:
            print('returning 1')
            return jsonify(['1']) 
        else:
            return render_template('testBuilderQuestions.html',questions=questionList)
    if len(questions)==0:
        print('returning 1')
        return jsonify(['1']) 
    else:
        return render_template('testBuilderQuestions.html',questions=questions,flagTopic = 'true')


@test_builder.route('/fetchRequiredQues',methods=['GET','POST'])
def fetchRequiredQues():
    class_val = request.args.get('class_val')
    subject_id = request.args.get('subject_id')
    
    if class_val!=None:
        if subject_id!=None:
            questionList = QuestionDetails.query.filter_by(class_val = str(class_val),subject_id=subject_id,archive_status='N').all()
        else:
            questionList = QuestionDetails.query.filter_by(class_val = str(class_val),archive_status='N').all()
    else:
        if subject_id!=None:
            questionList = QuestionDetails.query.filter_by(archive_status='N',subject_id=subject_id).all()
        else:
            questionList = QuestionDetails.query.filter_by(archive_status='N').all()
    if len(questionList)==0:
        print('returning 1')
        return jsonify(['1'])
    else:
        print('returning template'+ str(questionList))
        return render_template('testBuilderQuestions.html',questions=questionList)

@test_builder.route('/testBuilderQuestions',methods=['GET','POST'])  
def testBuilderQuestions():
    questions=[]
    topicList=request.get_json()
    for topic in topicList:
        # questionList = QuestionDetails.query.join(QuestionOptions, QuestionDetails.question_id==QuestionOptions.question_id).add_columns(QuestionDetails.question_id, QuestionDetails.question_description, QuestionDetails.question_type, QuestionOptions.weightage).filter(QuestionDetails.topic_id == int(topic),QuestionDetails.archive_status=='N' ).filter(QuestionOptions.is_correct=='Y').all()
        questionList = QuestionDetails.query.filter_by(topic_id = int(topic),archive_status='N').order_by(QuestionDetails.question_id).all()
        questions.append(questionList)
    if len(questionList)==0:
        print('returning 1')
        return jsonify(['1'])
    else:
        print('returning template'+ str(questionList))
        return render_template('testBuilderQuestions.html',questions=questions)



@test_builder.route('/testBuilderFileUpload',methods=['GET','POST'])
def testBuilderFileUpload():
    class_val = request.args.get('class_val')
    test_type = request.args.get('test_type')
    subject_id = request.args.get('subject_id')
    subject_name = MessageDetails.query.filter_by(msg_id=subject_id).first()
    date = request.args.get('date')
    print('class_val:'+str(class_val))
    print('test_type:'+str(test_type))
    print('subject_id:'+str(subject_id))
    print('Date:'+str(date))
    print('Inside Test builder file upload Test Type value:'+str(test_type))
    #question_list=request.get_json()
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    board_id = SchoolProfile.query.filter_by(school_id = teacher_id.school_id).first()
    data=request.get_json()
    question_list=data[0]
    count_marks=data[1]
    
    document = Document()
    print('Date')
    print(date)
    document.add_heading(schoolNameVal(), 0)
    document.add_heading('Class '+str(class_val)+" - "+str(test_type)+" - "+str(date) , 1)
    document.add_heading("Subject : "+str(subject_name.description),2)
    document.add_heading("Total Marks : "+str(count_marks),3)
    p = document.add_paragraph()
    #For every selected question add the question description
    for question in question_list:
        data=QuestionDetails.query.filter_by(question_id=int(question), archive_status='N').first()
        #for every question add it's options
        options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
        #add question desc
        document.add_paragraph(
            data.question_description, style='List Number'
        )    
    #Add the image associated with the question
        if data.reference_link!='' and data.reference_link!=None:
            try:
                response = requests.get(data.reference_link, stream=True)
                image = BytesIO(response.content)
                document.add_picture(image, width=Inches(1.25))
            except:
                pass

        for option in options:
            if option.option_desc is not None:
                document.add_paragraph(
                    option.option+". "+option.option_desc)     
    #document.add_page_break()
    #naming file here
    cl = class_val.replace("/","-")
    file_name=str(teacher_id.school_id)+str(cl)+str(subject_name.description)+str(test_type)+str(datetime.today().strftime("%Y%m%d"))+str(count_marks)+'.docx'
    file_name = file_name.replace(" ", "")
    if not os.path.exists('tempdocx'):
        os.mkdir('tempdocx')
    document.save('tempdocx/'+file_name)
    #uploading to s3 bucket
    client = boto3.client('s3', region_name='ap-south-1')
    client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
    #deleting file from temporary location after upload to s3
    os.remove('tempdocx/'+file_name)

    ###### Inserting record in the Test Detail table
    file_name_val='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name

    # Test create date
    format = "%Y-%m-%d %H:%M:%S"
    # Current time in UTC
    now_utc = datetime.now(timezone('UTC'))
    print(now_utc.strftime(format))
    # Convert to local time zone
    now_local = now_utc.astimezone(get_localzone())
    print('Date of test creation:'+str(now_local.strftime(format)))
    # date_utc = date.now(timezone('UTC'))
    # date_local = date_utc.astimezone(get_localzone())
    # print('Date of Test:'+str(date_local.strftime(format)))
    # Test end date
    testDetailsUpd = TestDetails(test_type=str(test_type), total_marks=str(count_marks),last_modified_date= datetime.now(),
        board_id=str(board_id.board_id), subject_id=int(subject_id),class_val=str(class_val),date_of_creation=now_local.strftime(format),
        date_of_test=str(date), school_id=teacher_id.school_id,test_paper_link=file_name_val, teacher_id=teacher_id.teacher_id)
    db.session.add(testDetailsUpd)
    db.session.commit()

    ##### This section to insert values into test questions table #####
    #try:
    createdTestID = TestDetails.query.filter_by(teacher_id=teacher_id.teacher_id).order_by(TestDetails.last_modified_date.desc()).first()
    for questionVal in question_list:
        testQuestionInsert= TestQuestions(test_id=createdTestID.test_id, question_id=questionVal, last_modified_date=datetime.now(),is_archived='N')
        db.session.add(testQuestionInsert)
    db.session.commit()
    #except:
    #    print('error inserting values into the test questions table')
    #### End of section ####
    testPaperData= TestDetails.query.filter_by(school_id=teacher_id.school_id,teacher_id=teacher_id.teacher_id).order_by(TestDetails.date_of_creation.desc()).first()
    sections = ClassSection.query.filter_by(school_id=teacher_id.school_id,class_val=testPaperData.class_val).all()
    return render_template('testPaperDisplay.html',file_name=file_name_val,testPaperData=testPaperData,sections=sections)

@test_builder.route('/testPapers')
@login_required
def testPapers():
    indic='DashBoard'
    return render_template('testPapers.html',indic=indic,title='Test Papers')

@test_builder.route('/testPaperTable')
def testPaperTable():
    paper_count = request.args.get('paper_count')
    if paper_count=="all":
        paper_count=500
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()    
    #testPaperData= TestDetails.query.filter_by(school_id=teacher_id.school_id).order_by(TestDetails.date_of_creation.desc()).all()
    testPaperQuery = "select *from test_details where school_id="+ str(teacher_id.school_id)
    testPaperQuery = testPaperQuery  +" order by date_of_creation desc fetch first "+str(paper_count)+" rows only"
    print('Query:'+str(testPaperQuery))
    testPaperData = db.session.execute(testPaperQuery).fetchall()
    subjectNames=MessageDetails.query.filter_by(category='Subject')
    return render_template('_testPaperTable.html',testPaperData=testPaperData,subjectNames=subjectNames,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

@test_builder.route('/findSection',methods=['GET','POST'])
def findSection():
    class_val=request.args.get('class_val')
    school_id = User.query.filter_by(id=current_user.id).first()
    print('Class:'+str(class_val))
    print('School Id:'+str(school_id.school_id))
    sections = ClassSection.query.filter_by(school_id=school_id.school_id,class_val=class_val)
    sec = []
    for section in sections:
        sec.append(section.section)
    return jsonify(sec)

@test_builder.route('/getQuestionDetails')
def getQuestionDetails():
    print('inside getQuestionDetails')
    qtest_id = request.args.get('test_id')
    print('Test Id:'+str(qtest_id)) 
    getQuestionQuery = "select tq.question_id,qd.question_type,qd.question_description,td.total_marks as weightage from test_questions tq "
    getQuestionQuery = getQuestionQuery + "inner join question_details qd on tq.question_id = qd.question_id "
    getQuestionQuery = getQuestionQuery + "inner join test_details td on td.test_id = tq.test_id where td.test_id = '"+str(qtest_id)+"'"
    testQuestionsDetails = db.session.execute(text(getQuestionQuery)).fetchall()
    for questions in testQuestionsDetails:
        print(questions.question_description)
    totalQues = len(testQuestionsDetails)
    return render_template('_getQuestionDetails.html',totalQues=totalQues,testQuestionsDetails=testQuestionsDetails)

@test_builder.route('/getChapterDetails')
def getChapterDetails():
    qtest_id=request.args.get('test_id')
    getChapterQuery = "select distinct topic_name, chapter_name, chapter_num from test_questions tq "
    getChapterQuery= getChapterQuery+ "inner join question_details qd  on "
    getChapterQuery= getChapterQuery+ " qd.question_id=tq.question_id inner join "
    getChapterQuery= getChapterQuery+ " topic_detail td on td.topic_id=qd.topic_id "
    getChapterQuery= getChapterQuery+ "where tq.test_id='"+str(qtest_id)+"'"

    getChapterRows = db.session.execute(text(getChapterQuery)).fetchall()

    return render_template('_getChapterDetails.html', getChapterRows=getChapterRows)

    return jsonify({'data':finalLink})
    
@test_builder.route('/getTestPaperLink',methods=['POST','GET'])
def getTestPaperLink():
    if request.method == 'POST':
        print('inside getTestPaperLink')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
        a = json.dumps(jsonExamData)
      
        z = json.loads(a)
        
        
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for data in z['results'].values():
            
            paramList.append(data)
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)
        print(conList[2])
        # Test for topic
        print('Testing for topic')
        # print(type(paramList[1]))
        # print(int(paramList[1]))
        # 
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)
            classList.append(classVal)
            j=j+1
        
        selClass = ''
        print('Selected Class option:')
        print(paramList[4])
        for className in classList:
            num = className.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[4]))
            if int(num) == int(paramList[4]):
                print(className)
                selClass = className.split('-')[1]
                print('selClass:'+str(selClass))
        print('class')
        selClass = selClass.strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[2]))
            if int(num) == int(paramList[2]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
                
        print('Subject:')
        selSubject = selSubject.strip()
        # Start for topic
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        selChapter = ''
        for chapterName in chapterDetList:
            num = chapterName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(chapterName)
                selChapter = chapterName.split('-')[1]
                print('selChapter:'+str(selChapter))
        #End topic
        selChapter = selChapter.strip()
        print('Chapter'+str(selChapter))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        fetchQuesIdsQuery = "select qd.question_id from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name = '"+str(selChapter)+"' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit '"+str(paramList[3])+"'"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))        
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        oldQuesIds = []
        for ques in fetchQuesIds:
            if ques:
                oldQuesIds.append(ques.question_id)
        testPaperQuery = "select test_id,test_paper_link from test_details order by test_id desc limit 1"
        print(testPaperQuery)
        testPaperData = db.session.execute(text(testPaperQuery)).first()
        fetchLastPaperQuestionIds = TestQuestions.query.filter_by(test_id=testPaperData.test_id).all()
        newQuesIds = []
        for ques in fetchLastPaperQuestionIds:
            if ques:
                newQuesIds.append(ques.question_id) 
        testPaperLink = ''
        if  oldQuesIds ==  newQuesIds:   
            testPaperLink = str("Here's the test paper link:\n")+str(testPaperData.test_paper_link)
            print('testPaperLink:'+str(testPaperLink))
        else:
            testPaperLink = 'No testpaper available'
        return jsonify({'TestPaperLink':testPaperLink})

    
@test_builder.route('/testLinkWhatsappBot', methods=['POST','GET'])
def testLinkWhatsappBot(): 
    phone = request.args.get('phone') 
    user = User.query.filter_by(phone=phone).first()
    teacher= TeacherProfile.query.filter_by(user_id=user.id).first() 
    student = StudentProfile.query.filter_by(user_id=user.id).first()
    subject_id = request.args.get('subjectId')
    print('inside testlinkwhatsappbot')
    print(subject_id)
    subjectQuery = MessageDetails.query.filter_by(msg_id=subject_id).first()
    subjectName = subjectQuery.description
    classVal = request.args.get('classVal')
    emailDet = ''
    if student:
      emailDet = StudentProfile.query.filter_by(student_id=student.student_id).first()
    user = ''
    
    if emailDet:
        user = User.query.filter_by(email=teacher.email).first()
    if user:
        login_user(user,remember='Y')
    clasVal = classVal.replace('@','_')
    respsessionid = request.args.get('respsessionid')
    testQuery = SessionDetail.query.filter_by(resp_session_id=respsessionid).first()
    testId = testQuery.test_id
    section = request.args.get('section')
    fetchQuesQuery = "select question_id from test_questions where test_id='"+str(testId)+"'"
    fetchQuesIds = db.session.execute(fetchQuesQuery).fetchall()
    quesIds = []
    for fetchIds in fetchQuesIds:
        quesIds.append(fetchIds.question_id)
    questions = QuestionDetails.query.filter(QuestionDetails.question_id.in_(quesIds)).all()  
    for ques in questions:
        print('question description:')
        print(ques.question_id)
        print(ques.question_description)
    # questions = QuestionDetails.query.filter(QuestionDetails.question_id.in_(fetchQuesIds)).all()
    questionListSize = len(fetchQuesIds)
    respsessionid = request.args.get('respsessionid')
    total_marks = request.args.get('totalMarks')
    weightage = request.args.get('weightage')
    test_type = request.args.get('testType')
    test_type = test_type.replace('@','_')
    uploadStatus = request.args.get('uploadStatus')
    resultStatus = request.args.get('resultStatus')
    advance = request.args.get('advance')
    print('inside testLinkWhatsappBot')
    print('Subject Id:'+str(subject_id))
    studId = None
    if current_user.is_anonymous:
        print('user id student')
        return redirect(url_for('feedbackCollectionStudDev',student_id=studId,resp_session_id=respsessionid,school_id=teacher.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance,_external=True))
        # return render_template('feedbackCollectionStudDev.html',resp_session_id=str(respsessionid),studId=studId,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)
    else:
        print('user is teacher') 
        url = "http://www.school.alllearn.in/feedbackCollectionStudDev?resp_session_id="+str(respsessionid)+"&school_id="+str(teacher.school_id)
        responseSessionIDQRCode = "https://api.qrserver.com/v1/create-qr-code/?size=150x150&data="+url
        return render_template('feedbackCollectionTeachDev.html',classSecCheckVal='Y', subject_id=subject_id, 
            class_val = clasVal, section = section,questions=questions, questionListSize = questionListSize, resp_session_id = respsessionid,responseSessionIDQRCode=responseSessionIDQRCode,
            subjectName = subjectName, totalMarks=total_marks,weightage=weightage, 
            batch_test=0,testType=test_type,school_id=teacher.school_id,uploadStatus=uploadStatus,resultStatus=resultStatus,advance=advance)

@test_builder.route('/newTestLinkGenerate',methods=['POST','GET'])
def newTestLinkGenerate():
    if request.method == 'POST':
        print('newTestLinkGenerate')
        jsonExamData = request.json
        # jsonExamData = {"results": {"weightage": "10","topics": "1","subject": "1","question_count": "10","class_val": "3","uploadStatus":"Y","duration":"0","resultStatus":"Y","instructions":"","advance":"Y","negativeMarking":"0","test_type":"Class Feedback"},"custom_key": "custom_value","contact": {"phone": "9008262739"}}
        
        a = json.dumps(jsonExamData)
      
        z = json.loads(a)
        
        
        paramList = []
        conList = []
        print('data:')
        # print(z['result'].class_val)
        # print(z['result'])
        for data in z['results'].values():
            
            paramList.append(data)
        for con in z['contact'].values():
            conList.append(con)
        print(paramList)
        print(conList[2])
        # Test for topic
        print('Testing for topic')
        # print(type(paramList[1]))
        # print(int(paramList[1]))
        # 
        print('Data Contact')
        # print(conList[2])
        contactNo = conList[2][-10:]
        print(contactNo)
        userId = User.query.filter_by(phone=contactNo).first()
        teacher_id = TeacherProfile.query.filter_by(user_id=userId.id).first()
        classesListData = ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).all()
        classList = [] 
        j=1
        for classlist in classesListData:
            classVal = str(j)+str(' - ')+str(classlist.class_val)
            classList.append(classVal)
            j=j+1
        
        selClass = ''
        print('Selected Class option:')
        print(paramList[4])
        for className in classList:
            num = className.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[4]))
            if int(num) == int(paramList[4]):
                print(className)
                selClass = className.split('-')[1]
                print('selClass:'+str(selClass))
        print('class')
        selClass = selClass.strip()
        print(selClass)
        subQuery = "select md.description as subject,md.msg_id from board_class_subject bcs inner join message_detail md on bcs.subject_id = md.msg_id where school_id='"+str(teacher_id.school_id)+"' and class_val = '"+str(selClass)+"'"
        print(subQuery)
        subjectData = db.session.execute(text(subQuery)).fetchall()
        print(subjectData)
        subjectList = []
        k=1
        subId = ''
        for subj in subjectData:
            sub = str(k)+str('-')+str(subj.subject)
            subjectList.append(sub)
            k=k+1
        for subjectName in subjectList:
            num = subjectName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[2]))
            if int(num) == int(paramList[2]):
                print(subjectName)
                selSubject = subjectName.split('-')[1]
                print('selSubject:'+str(selSubject))
                
        print('Subject:')
        selSubject = selSubject.strip()
        # Start for topic
        subQuery = MessageDetails.query.filter_by(description=selSubject).first()
        subId = subQuery.msg_id
        print(selSubject)
        print('SubId:'+str(subId))
        extractChapterQuery = "select td.chapter_name ,td.chapter_num ,bd.book_name from topic_detail td inner join book_details bd on td.book_id = bd.book_id where td.class_val = '"+str(selClass)+"' and td.subject_id = '"+str(subId)+"'"
        print('Query:'+str(extractChapterQuery))
        extractChapterData = db.session.execute(text(extractChapterQuery)).fetchall()
        print(extractChapterData)
        c=1
        chapterDetList = []
        for chapterDet in extractChapterData:
            chap = str(c)+str('-')+str(chapterDet.chapter_name)+str('-')+str(chapterDet.book_name)+str("\n")
            chapterDetList.append(chap)
            c=c+1
        selChapter = ''
        for chapterName in chapterDetList:
            num = chapterName.split('-')[0]
            print('num:'+str(num))
            print('class:'+str(paramList[1]))
            if int(num) == int(paramList[1]):
                print(chapterName)
                selChapter = chapterName.split('-')[1]
                print('selChapter:'+str(selChapter))
        #End topic
        selChapter = selChapter.strip()
        print('Chapter'+str(selChapter))
        dateVal= datetime.today().strftime("%d%m%Y%H%M%S")
        fetchQuesIdsQuery = "select td.board_id,qd.suggested_weightage,qd.question_type,qd.question_id,qd.question_description,td.subject_id,td.topic_id from question_details qd "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join topic_detail td on qd.topic_id = td.topic_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "inner join message_detail md on md.msg_id = td.subject_id "
        fetchQuesIdsQuery = fetchQuesIdsQuery + "where td.chapter_name = '"+str(selChapter)+"' and md.description = '"+str(selSubject)+"' and td.class_val = '"+str(selClass)+"' limit '"+str(paramList[3])+"'"
        print('fetchQuesIds Query:'+str(fetchQuesIdsQuery))
        fetchQuesIds = db.session.execute(fetchQuesIdsQuery).fetchall()
        msg = 'no questions available'
        print('fetchQuesIds:'+str(fetchQuesIds))
        if len(fetchQuesIds)==0 or fetchQuesIds=='':
            return jsonify({'onlineTestLink':msg})
        listLength = len(fetchQuesIds)
        count_marks = int(paramList[0]) * int(listLength)
        
        subjId = ''
        topicID = ''
        boardID = ''
        for det in fetchQuesIds:
            subjId = det.subject_id
            topicID = det.topic_id
            boardID = det.board_id
            break
        print('subjId:'+str(subjId))
        print(fetchQuesIds)
        currClassSecRow=ClassSection.query.filter_by(school_id=str(teacher_id.school_id),class_val=str(selClass).strip()).first()
        resp_session_id = str(subId).strip()+ str(dateVal).strip() + str(randint(10,99)).strip()
        # threadUse(currClassSecRow.class_sec_id,resp_session_id,fetchQuesIds,paramList[11],count_marks,selClass,teacher_id.teacher_id,teacher_id.school_id)
        # task = exampleData.delay(10,20)
        task = insertData.delay(currClassSecRow.class_sec_id,resp_session_id,fetchQuesIds,paramList[11],count_marks,selClass,teacher_id.teacher_id,teacher_id.school_id)
        clasVal = selClass.replace('_','@')
        testType = paramList[11].replace('_','@')
        linkForTeacher=url_for('testLinkWhatsappBot',testType=paramList[11],totalMarks=count_marks,respsessionid=resp_session_id,fetchQuesIds=fetchQuesIds,weightage=10,negativeMarking=paramList[10],uploadStatus=paramList[5],resultStatus=paramList[7],advance=paramList[9],instructions=paramList[8],duration=paramList[6],classVal=clasVal,section=currClassSecRow.section,subjectId=subjId,phone=contactNo, _external=True)
        key = '265e29e3968fc62f68da76a373e5af775fa60'
        url = urllib.parse.quote(linkForTeacher)
        name  = ''
        r = rq.get('http://cutt.ly/api/api.php?key={}&short={}&name={}'.format(key, url, name))
        print('New Link')
        print(r.text)
        print(type(r.text))
        linkList = []
        jsonLink = json.dumps(r.text)
        newData = json.loads(r.text)
        print(type(newData))
        for linkData in newData['url'].values():
            linkList.append(linkData)
        finalLink = linkList[3]
        newLink = str('Here is the link to the online test:\n')+finalLink+str('\nDo you want to download the question paper?\n1 - Yes\n2 - No')
        print('newLink'+str(newLink))
        return jsonify({'onlineTestLink':newLink})
