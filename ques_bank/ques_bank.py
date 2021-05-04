from ques_bank.utils import *
from applicationDB import *
from flask import Flask, Blueprint, Markup, render_template, request, redirect, url_for, Response,session
from applicationDB import *
from forms import RegistrationForm,ContentManager, EditProfileForm, ResetPasswordRequestForm, ResetPasswordForm,ResultQueryForm,MarksForm, TestBuilderQueryForm,SchoolRegistrationForm, PaymentDetailsForm, \
    QuestionBuilderQueryForm, SingleStudentRegistration, SchoolTeacherForm, testPerformanceForm, studentPerformanceForm, QuestionUpdaterQueryForm,  QuestionBankQueryForm, \
    promoteStudentForm
from flask_login import current_user, login_user, login_required
import os
from datetime import datetime
from flask import g, jsonify
import json, boto3
from sqlalchemy import distinct, text, update
from miscFunctions import subjects,topics,subjectPerformance, chapters
from docx import Document


ques_bank= Blueprint('ques_bank',__name__)

@ques_bank.route('/questionBank',methods=['POST','GET'])
@login_required
def questionBank():
    topic_list=None
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form=QuestionBankQueryForm()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().order_by(ClassSection.class_val).filter_by(school_id=teacher_id.school_id).all()]
    form.subject_name.choices= ''
    form.chapter_num.choices= ''
    form.test_type.choices= [(i.description,i.description) for i in MessageDetails.query.filter_by(category='Test type').all()]
    if request.method=='POST':
        topic_list="select td.topic_id,td.topic_name from topic_detail td inner join topic_tracker tt on td.topic_id = tt.topic_id where td.class_val = '"+str(form.class_val.data)+"' and td.subject_id = '"+str(form.subject_name.data)+"' and td.chapter_num='"+str(form.chapter_num.data)+"' and tt.is_archived = 'N'"
        topic_list = db.session.execute(text(topic_list)).fetchall()
        subject=MessageDetails.query.filter_by(msg_id=int(form.subject_name.data)).first()
        session['class_val']=form.class_val.data
        session['sub_name']=subject.description
        session['test_type_val']=form.test_type.data
        session['chapter_num']=form.chapter_num.data  
        print('Class value:'+str(form.class_val.data))
        form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(str(form.class_val.data))]
        print('Class value:'+str(form.class_val.data))
        form.chapter_num.choices= [(int(i['chapter_num']), str(i['chapter_num'])+' - '+str(i['chapter_name'])) for i in chapters(str(form.class_val.data),int(form.subject_name.data))]
        indic='DashBoard'
        return render_template('questionBank.html',indic=indic,title='Question Bank',form=form,topics=topic_list,user_type_val=str(current_user.user_type))
    indic='DashBoard'
    return render_template('questionBank.html',indic=indic,title='Question Bank',form=form,classSecCheckVal=classSecCheck(),user_type_val=str(current_user.user_type))

@ques_bank.route('/visitedQuestions',methods=['GET','POST'])
def visitedQuestions():
    retake = request.args.get('retake')
    questions=[]
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    topicList=request.get_json() 
    for topic in topicList:
        print(str(retake)+'Retake')
        topicFromTracker = TopicTracker.query.filter_by(school_id = teacher_id.school_id, topic_id=int(topic)).first()
        topicFromTracker.is_covered='Y'
        if topicFromTracker.reteach_count:
            topicFromTracker.reteach_count=topicFromTracker.reteach_count+1
        db.session.commit()

    return jsonify(['1'])
    
@ques_bank.route('/questionBankQuestions',methods=['GET','POST'])
def questionBankQuestions():
    questions=[]
    topicList=request.get_json()
    for topic in topicList:
        # question_Details=QuestionDetails.query.filter_by(QuestionDetails.topic_id == int(topic)).first()
        # questionList = QuestionDetails.query.join(QuestionOptions, QuestionDetails.question_id==QuestionOptions.question_id).add_columns(QuestionDetails.question_id, QuestionDetails.question_description, QuestionDetails.question_type, QuestionDetails.suggested_weightage).filter(QuestionDetails.topic_id == int(topic)).filter(QuestionOptions.is_correct=='Y').all()
        questionList = QuestionDetails.query.filter_by(topic_id=int(topic),archive_status='N').order_by(QuestionDetails.question_id).all()
        questions.append(questionList)
        for q in questionList:
            print("Question List"+str(q))    
    if len(questionList)==0:
        print('returning 1')
        return jsonify(['1'])
    else:
        print('returning template'+ str(questionList))
        return render_template('questionBankQuestions.html',questions=questions)

@ques_bank.route('/questionBankFileUpload',methods=['GET','POST'])
def questionBankFileUpload():
    #question_list=request.get_json()
    data=request.get_json()
    question_list=data[0]
    count_marks=data[1]
    document = Document()
    document.add_heading(schoolNameVal(), 0)
    document.add_heading('Class '+session.get('class_val',None)+" - "+session.get('test_type_val',None)+" - "+str(session.get('date',None)) , 1)
    document.add_heading("Subject : "+session.get('sub_name',None),2)
    document.add_heading("Total Marks : "+str(count_marks),3)
    p = document.add_paragraph()
    for question in question_list:
        data=QuestionDetails.query.filter_by(question_id=int(question), archive_status='N').first()
        document.add_paragraph(
            data.question_description, style='List Number'
        )    
        options=QuestionOptions.query.filter_by(question_id=data.question_id).all()
        for option in options:
            if option.option_desc is not None:
                document.add_paragraph(
                    option.option+". "+option.option_desc)     
    #document.add_page_break()
    file_name='S'+'1'+'C'+session.get('class_val',"0")+session.get('sub_name',"0")+session.get('test_type_val',"0")+str(datetime.today().strftime("%d%m%Y"))+'.docx'
    if not os.path.exists('tempdocx'):
        os.mkdir('tempdocx')
    document.save('tempdocx/'+file_name)
    client = boto3.client('s3', region_name='ap-south-1')
    client.upload_file('tempdocx/'+file_name , os.environ.get('S3_BUCKET_NAME'), 'test_papers/{}'.format(file_name),ExtraArgs={'ACL':'public-read'})
    os.remove('tempdocx/'+file_name)

    return render_template('testPaperDisplay.html',file_name='https://'+os.environ.get('S3_BUCKET_NAME')+'.s3.ap-south-1.amazonaws.com/test_papers/'+file_name)

@ques_bank.route('/updateQuestion')
def updateQuestion():
    question_id = request.args.get('question_id')
    updatedCV = request.args.get('updatedCV')
    topicId = request.args.get('topicName')
    subId = request.args.get('subName')
    qType = request.args.get('qType')
    qDesc = request.args.get('qDesc')
    corrans = request.args.get('corrans')
    weightage = request.args.get('weightage')
    print('Weightage:'+str(weightage))
    print('Correct option:'+str(str(corrans)))
    preview = request.args.get('preview')
    options = request.args.get('options')
    op1 = request.args.get('op1')
    op2 = request.args.get('op2')
    op3 = request.args.get('op3')
    op4 = request.args.get('op4')
    print(op1)
    print(op2)
    print(op3)
    print(op4)
    form = QuestionBuilderQueryForm()
    print("Updated class Value+:"+updatedCV)
    print(str(updatedCV)+" "+str(topicId)+" "+str(subId)+" "+str(qType)+" "+str(qDesc)+" "+str(preview)+" "+str(corrans)+" "+str(weightage))
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
    form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
    form.topics.choices=[(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]
    flag = False
    # updateQuery = "update question_details t1 set topic_id='" + str(topicId) + "' where question_id='" + question_id + "'"
    updateQuery = "update question_details set class_val='" + str(updatedCV) +  "',topic_id='"+ str(topicId) + "',subject_id='"+ str(subId) + "',question_type='" + str(qType) + "',question_description='"+ str(qDesc) + "',reference_link='"+ str(preview) +"' where question_id='" + str(question_id) + "'"

    queryOneExe = db.session.execute(text(updateQuery))
    updateWeightage = "update question_details set suggested_weightage='" + str(weightage) + "' where question_id='" + str(question_id) + "'" 
    querytwoExe = db.session.execute(text(updateWeightage))
    

    option_id_list = QuestionOptions.query.filter_by(question_id=question_id).order_by(QuestionOptions.option_id).all()
    if corrans:
        print(option_id_list)
        i=0
        opId1=''
        opId2=''
        opId3=''
        opId4=''
        for opt in option_id_list:
            if i==0:
                opId1 = opt.option_id
            elif i==1:
                opId2 = opt.option_id
            elif i==2:
                opId3 = opt.option_id
            else:
                opId4 = opt.option_id
            i=i+1
        print("Options order of id:"+str(opId1)+' '+str(opId2)+' '+str(opId3)+' '+str(opId4))
        # print(option_id) 
        # option = "select option_id from question_options where question_id='"+ str(question_id) + "'"
        # opt = db.session.execute(text(option))
        # print("Updated options "+str(opt))
        if opId1 and opId2 and opId3 and opId4:
            updateOption1 = "update question_options set option_desc='"+str(op1)+"' where option_id='"+ str(opId1) + "'"
            print(updateOption1)
            updateOpt1Exe = db.session.execute(text(updateOption1))
            updateOption2 = "update question_options set option_desc='"+str(op2)+"' where option_id='"+ str(opId2) + "'"
            print(updateOption2)
            updateOpt2Exe = db.session.execute(text(updateOption2))
            updateOption3 = "update question_options set option_desc='"+str(op3)+"' where option_id='"+ str(opId3) + "'"
            print(updateOption3)
            updateOpt3Exe = db.session.execute(text(updateOption3))
            updateOption4 = "update question_options set option_desc='"+str(op4)+"' where option_id='"+ str(opId4) + "'"
            print(updateOption4)
            updateOpt4Exe = db.session.execute(text(updateOption4))
            if str(corrans)!='':
                updatequery1 = "update question_options set is_correct='N' where is_correct='Y' and question_id='" +str(question_id)+"'"
                print(updatequery1)
                update1 = db.session.execute(text(updatequery1))
                updateCorrectOption = "update question_options set is_correct='Y' where option_desc='"+str(corrans)+"' and question_id='"+str(question_id)+"'"
                print(updateCorrectOption)
                updateOp = db.session.execute(text(updateCorrectOption))
        else:
            optionlist = []
            optionlist.append(op1)
            optionlist.append(op2)
            optionlist.append(op3)
            optionlist.append(op4)
            corrAns = 'Y'
            for optionDesc in optionlist:
                if optionDesc==corrans:
                    query = "insert into question_options(option_desc,question_id,weightage,is_correct,option) values('"+optionDesc+"','"+question_id+"','"+weightage+"','Y','A')"
                else:
                    query = "insert into question_options(option_desc,question_id,weightage,option) values('"+optionDesc+"','"+question_id+"','"+weightage+"','A')"
                    db.session.execute(query)

    print('Inside Update Questions')
    db.session.commit()
    print(updateQuery)
    # updateSecondQuery = "update question_options set weightage='" + str(weightage) +"' where question_id='" + str(question_id) + "'"
    # querySecondExe = db.session.execute(text(updateSecondQuery)) 
    # db.session.commit()
    print("Question Id in update Question:"+question_id)
    # print(updatedData)
    return render_template('questionUpload.html', form=form, flag=flag)


@ques_bank.route('/questionOptions')
def questionOptions():
    question_id_arg=request.args.get('question_id')
    questionOptionResults = QuestionOptions.query.filter_by(question_id=question_id_arg).all()
    questionOptionsList=[]
    for value in questionOptionResults:
        print("This is the value: "+str(value))        
        questionOptionsList.append(value.option+". "+value.option_desc)

    print(questionOptionsList)

    return jsonify([questionOptionsList])


@ques_bank.route('/deleteQuestion')
def deleteQuestion():
    question_id = request.args.get('question_id')
    print('Delete Question Id:'+question_id)
    print("Question Id:-"+question_id)

    updateQuery = "update question_details set archive_status='Y' where question_id='"+question_id+"'"
    print(updateQuery)
    db.session.execute(updateQuery)
    db.session.commit()
    return "text" 




@ques_bank.route('/questionDetails')
def questionDetails():
    flag = True
    question_id = request.args.get('question_id')
    print("Question Id-:"+question_id)
    form = QuestionBuilderQueryForm()
    teacher_id=TeacherProfile.query.filter_by(user_id=current_user.id).first()
    form.class_val.choices = [(str(i.class_val), "Class "+str(i.class_val)) for i in ClassSection.query.with_entities(ClassSection.class_val).distinct().filter_by(school_id=teacher_id.school_id).order_by(ClassSection.class_val).all()]
    form.subject_name.choices= [(str(i['subject_id']), str(i['subject_name'])) for i in subjects(1)]
    form.topics.choices=[(str(i['topic_id']), str(i['topic_name'])) for i in topics(1,54)]

    questionDetailsQuery = "select t2.class_val, t1.question_id, t2.subject_id, t1.reference_link, t1.suggested_weightage, t2.topic_name, t2.topic_id, t1.question_type, t1.question_description, t4.description from question_details t1 "
    questionDetailsQuery = questionDetailsQuery + "inner join topic_detail t2 on t1.topic_id=t2.topic_id "
    questionDetailsQuery = questionDetailsQuery + "inner join message_detail t4 on t1.subject_id = t4.msg_id"
    questionDetailsQuery = questionDetailsQuery + " where t1.question_id ='" + question_id + "' order by t1.question_id"

    questionUpdateUploadSubjective = db.session.execute(text(questionDetailsQuery)).first()   
    question_desc = questionUpdateUploadSubjective.question_description.replace('\n', ' ').replace('\r', '')
    print(questionUpdateUploadSubjective)
    questionUpdateUpload=questionUpdateUploadSubjective
    if questionUpdateUpload.question_type=='MCQ1':
       
        query = "select option_desc from question_options where question_id='" + question_id + "' order by option_id"
        #print(query)
        avail_options = db.session.execute(text(query)).fetchall()
        queryCorrectoption = "select option_desc from question_options where is_correct='Y' and question_id='" + question_id + "'"  
        #print(queryCorrectoption)
        correctoption = db.session.execute(text(queryCorrectoption)).fetchall()
        print(correctoption)
        correctOption = ''
        for c in correctoption:
            print(c.option_desc)
            correctOption = c.option_desc
        print('Correct Option:'+correctOption)
        for q in questionUpdateUploadSubjective:
            print('this is check for MCQ ' + str(q))
        for a in avail_options:
            print(a)
        print('Correct Option Again:'+correctOption)
        return render_template('questionUpload.html', question_id=question_id, questionUpdateUpload=questionUpdateUpload, form=form, flag=flag, avail_options=avail_options, correctOption=correctOption,question_desc=question_desc)
        # return render_template('questionUpload.html',question_id=question_id, questionUpdateUploadSubjective=questionUpdateUploadSubjective,form=form,flag=flag,avail_options=avail_options,correctOption=correctOption)

    for q in questionUpdateUpload:
        print('this is check for Subjective ' + str(q))
    
    return render_template('questionUpload.html', question_id=question_id, questionUpdateUpload=questionUpdateUpload, form=form, flag=flag,question_desc=question_desc)
